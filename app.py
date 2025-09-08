from flask import (
    Flask,
    render_template,
    redirect,
    url_for,
    flash,
    request,
    send_file,
    session,
    jsonify,
    g,
)
from flask_compress import Compress
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate  # Added for database migrations
from sqlalchemy.orm import joinedload
import sqlalchemy as sa  # Adăugat pentru server_default=sa.text()
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from flask_login import (
    LoginManager,
    UserMixin,
    login_user,
    logout_user,
    current_user,
    login_required,
)
from werkzeug.security import generate_password_hash, check_password_hash
import bcrypt
import os
import secrets
from datetime import datetime, date, time, timedelta
import time as _time
from sqlalchemy import (
    func,
    or_,
    and_,
    inspect,
    text,
)  # Am adăugat text aici explicit, deși sa.text ar fi funcționat cu importul de mai sus
import re
from unidecode import unidecode
import json
import pytz  # Adăugat pentru fusuri orare
from functools import wraps

# Inițializare aplicație Flask
app = Flask(__name__)
try:
    Compress(app)  # HTTP compression for faster responses
except Exception as _e:
    print("Flask-Compress not active:", _e)

# Definire fus orar pentru România
EUROPE_BUCHAREST = pytz.timezone("Europe/Bucharest")


# Funcție helper pentru a obține ora curentă localizată
def get_localized_now():
    return datetime.now(EUROPE_BUCHAREST)


# Filtru Jinja pentru a formata datetime-uri cu fusul orar local
@app.template_filter("localdatetime")
def localdatetime_filter(dt, fmt="%d-%m-%Y %H:%M:%S"):
    if not dt:
        return ""
    if dt.tzinfo is None:
        # Presupunem că datetime-urile naive sunt în ora serverului (care ar trebui să fie Europe/Bucharest)
        # sau sunt UTC și trebuie convertite. Pentru siguranță, dacă e naive, îl localizăm ca UTC apoi convertim.
        # Dar majoritatea datetime-urilor create cu datetime.now() fără tz sunt deja în ora sistemului.
        # Cel mai sigur este să verificăm dacă provin din datetime.utcnow() sau datetime.now()
        # Pentru ActionLog.timestamp care e default=datetime.utcnow, va fi conștient de fus (UTC)
        # Pentru celelalte (ex: Permission.start_datetime), sunt stocate ca naive.
        # Le vom considera ca fiind în ora serverului și le vom localiza.
        # Totuși, o practică mai bună ar fi să stocăm totul ca UTC.
        # Având în vedere structura actuală, vom considera datetime-urile naive ca fiind în ora serverului.
        localized_dt = EUROPE_BUCHAREST.localize(
            dt, is_dst=None
        )  # is_dst=None pentru a gestiona tranzițiile DST
    else:
        localized_dt = dt.astimezone(EUROPE_BUCHAREST)
    return localized_dt.strftime(fmt)


@app.template_filter("localtime")
def localtime_filter(t, fmt="%H:%M"):
    if not t:
        return ""
    # Ora (time object) nu are fus orar, deci o formatăm direct.
    # Dacă ar fi nevoie de conversie bazată pe dată, ar fi mai complex.
    return t.strftime(fmt)


@app.template_filter("localdate")
def localdate_filter(d, fmt="%d-%m-%Y"):
    if not d:
        return ""
    # Data (date object) nu are fus orar.
    return d.strftime(fmt)


@app.context_processor
def inject_global_vars():
    # This context processor makes the get_localized_now function available
    # to all Jinja templates. If 'get_localized_now' is undefined in a template,
    # ensure this processor is correctly registered and the Flask app is restarted.
    return dict(get_localized_now=get_localized_now)


@app.context_processor
def inject_app_config():
    # Makes the Flask app object available as 'app' in all templates
    return dict(app=app)


# IMPORTANT: Change this to a strong, unique, and static secret key in a real environment!
# Using a static key is crucial for session persistence across app restarts.
# For production, load from an environment variable.
app.config["SECRET_KEY"] = os.environ.get(
    "FLASK_SECRET_KEY", "dev_fallback_super_secret_key_123!@#"
)
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///site.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(
    days=30
)  # Sesiune permanentă de 30 de zile
app.config["SEND_FILE_MAX_AGE_DEFAULT"] = timedelta(
    days=30
)  # Cache static files for 30 days
# Cookie hardening (non-breaking defaults)
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["REMEMBER_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SECURE"] = os.environ.get(
    "FLASK_COOKIE_SECURE", "0"
) in [
    "1",
    "true",
    "True",
]
# Keep users logged in unless they explicitly logout
app.config["REMEMBER_COOKIE_DURATION"] = timedelta(days=30)
app.config["REMEMBER_COOKIE_REFRESH_EACH_REQUEST"] = True

db = SQLAlchemy(app)
migrate = Migrate(app, db)  # Initialize Flask-Migrate
login_manager = LoginManager(app)
login_manager.login_view = "user_login"
login_manager.login_message_category = "info"
login_manager.login_message = (
    "Te rugăm să te autentifici pentru a accesa această pagină."
)


# Security and caching headers
@app.after_request
def add_default_headers(response):
    try:
        path = request.path or ""
        if path.startswith("/static/"):
            response.headers["Cache-Control"] = (
                "public, max-age=2592000, immutable"
            )
        else:
            # Avoid caching dynamic content to keep data fresh
            response.headers["Cache-Control"] = "no-store"

        # Helpful security headers (non-breaking)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["X-Frame-Options"] = "SAMEORIGIN"
        response.headers["Permissions-Policy"] = (
            "geolocation=(), camera=(), microphone=()"
        )

        # Content Security Policy tuned to current external CDNs and inline usage
        csp = (
            "default-src 'self'; "
            "script-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; "
            "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; "
            "img-src 'self' data: blob: https:; "
            "font-src 'self' data: https://fonts.gstatic.com https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; "
            "connect-src 'self'; "
            "object-src 'none'; "
            "frame-ancestors 'self'; "
            "base-uri 'self'"
        )
        response.headers.setdefault("Content-Security-Policy", csp)

        # HSTS only when HTTPS
        try:
            if request.is_secure:
                response.headers.setdefault(
                    "Strict-Transport-Security",
                    "max-age=31536000; includeSubDomains",
                )
        except Exception:
            pass
    except Exception:
        pass
    return response


# Friendly error pages (non-intrusive)
@app.errorhandler(404)
def handle_404(_e):
    try:
        return render_template("errors/404.html"), 404
    except Exception:
        return "Pagina nu a fost găsită.", 404


@app.errorhandler(403)
def handle_403(_e):
    try:
        return render_template("errors/403.html"), 403
    except Exception:
        return "Acces interzis.", 403


@app.errorhandler(500)
def handle_500(_e):
    try:
        return render_template("errors/500.html"), 500
    except Exception:
        return "Eroare internă a serverului.", 500


SERVICE_TYPES = [
    "GSS",
    "SVM",
    "Planton 1",
    "Planton 2",
    "Planton 3",
    "Intervenție",
    "Altul",
]
GENDERS = ["Nespecificat", "M", "F"]
KNOWN_RANK_PATTERNS = [
    re.compile(r"^(Mm V)\s+", re.IGNORECASE),
    re.compile(r"^(Sd cap)\s+", re.IGNORECASE),
    re.compile(r"^(Sg Maj)\s+", re.IGNORECASE),
    re.compile(r"^(Mm IV)\s+", re.IGNORECASE),
    re.compile(r"^(Sdt\.?)\s+", re.IGNORECASE),
    re.compile(r"^(Sd\.?)\s+", re.IGNORECASE),
    re.compile(r"^(Cap\.?)\s+", re.IGNORECASE),
    re.compile(r"^(Sg\.?)\s+", re.IGNORECASE),
    re.compile(r"^(Frt\.?)\s+", re.IGNORECASE),
    re.compile(r"^(Plt\.? Adj\.?)\s+", re.IGNORECASE),
    re.compile(r"^(Plt\.? Maj\.?)\s+", re.IGNORECASE),
    re.compile(r"^(Plt\.?)\s+", re.IGNORECASE),
]

# --- Simple in-memory rate limiting (per-process) ---
_RATE_LIMIT_BUCKETS = {}


def _is_rate_limited(bucket_key: str, limit: int, window_seconds: int) -> bool:
    window_start = _time.time() - window_seconds
    timestamps = _RATE_LIMIT_BUCKETS.get(bucket_key, [])
    # prune
    timestamps = [ts for ts in timestamps if ts >= window_start]
    if len(timestamps) >= limit:
        _RATE_LIMIT_BUCKETS[bucket_key] = timestamps
        return True
    timestamps.append(_time.time())
    _RATE_LIMIT_BUCKETS[bucket_key] = timestamps
    return False


class User(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=True)
    role = db.Column(db.String(50), nullable=False)
    unique_code = db.Column(db.String(100), unique=True, nullable=True)
    personal_code_hash = db.Column(db.String(256), nullable=True)
    is_first_login = db.Column(db.Boolean, default=True)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return (
            check_password_hash(self.password_hash, password)
            if self.password_hash
            else False
        )

    def set_personal_code(self, code):
        self.personal_code_hash = bcrypt.hashpw(
            code.encode("utf-8"), bcrypt.gensalt()
        ).decode("utf-8")
        self.is_first_login = False

    def check_personal_code(self, code):
        return (
            bcrypt.checkpw(
                code.encode("utf-8"), self.personal_code_hash.encode("utf-8")
            )
            if self.personal_code_hash
            else False
        )

    def can_login_with_personal_code(self):
        return self.role != "admin" and self.personal_code_hash is not None

    def __repr__(self):
        return f"<User {self.username} ({self.role})>"


class Student(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nume = db.Column(db.String(100), nullable=False)
    prenume = db.Column(db.String(100), nullable=False)
    grad_militar = db.Column(db.String(50), nullable=False)
    id_unic_student = db.Column(db.String(50), unique=True, nullable=True)
    pluton = db.Column(db.String(50), nullable=False)
    companie = db.Column(db.String(50), nullable=False)
    batalion = db.Column(db.String(50), nullable=False)
    gender = db.Column(db.String(10), default="Nespecificat", nullable=False)
    volunteer_points = db.Column(db.Integer, default=0, nullable=False)
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=False
    )
    creator = db.relationship(
        "User", backref=db.backref("students_created", lazy=True)
    )
    is_platoon_graded_duty = db.Column(
        db.Boolean, default=False, nullable=False
    )  # True if leader in own platoon
    assigned_graded_platoon = db.Column(
        db.String(50), nullable=True
    )  # Platoon ID they are assigned to lead, if different from their own or to be explicit
    is_smt = db.Column(
        db.Boolean, default=False, server_default=sa.text("0"), nullable=False
    )  # Total Medical Exemption
    exemption_details = db.Column(
        db.String(255), nullable=True
    )  # Details for other/partial exemptions

    def __repr__(self):
        details = []
        if self.is_platoon_graded_duty:
            details.append("Gradat Pluton Propriu")
        if self.assigned_graded_platoon:
            details.append(f"Gradat la Pl.{self.assigned_graded_platoon}")
        if self.is_smt:
            details.append("SMT")
        if self.exemption_details:
            details.append(f"Scutire: {self.exemption_details[:30]}")

        detail_str = ""
        if details:
            detail_str = " (" + ", ".join(details) + ")"

        return f"<Student {self.grad_militar} {self.nume} {self.prenume} - Pl.{self.pluton}{detail_str}>"


class Permission(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(
        db.Integer,
        db.ForeignKey("student.id", ondelete="CASCADE"),
        nullable=False,
    )
    start_datetime = db.Column(db.DateTime, nullable=False)
    end_datetime = db.Column(db.DateTime, nullable=False)
    reason = db.Column(db.String(255), nullable=True)
    status = db.Column(db.String(50), default="Aprobată", nullable=False)
    destination = db.Column(db.String(255), nullable=True)
    transport_mode = db.Column(db.String(100), nullable=True)
    student = db.relationship(
        "Student",
        backref=db.backref(
            "permissions", lazy=True, cascade="all, delete-orphan"
        ),
    )
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=True
    )  # Made nullable
    creator = db.relationship(
        "User", backref=db.backref("permissions_created", lazy=True)
    )

    @property
    def is_active(self):
        now = get_localized_now()
        start_dt_aware = (
            EUROPE_BUCHAREST.localize(self.start_datetime)
            if self.start_datetime.tzinfo is None
            else self.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        end_dt_aware = (
            EUROPE_BUCHAREST.localize(self.end_datetime)
            if self.end_datetime.tzinfo is None
            else self.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        return (
            start_dt_aware <= now <= end_dt_aware and self.status == "Aprobată"
        )

    @property
    def is_upcoming(self):
        now = get_localized_now()
        start_dt_aware = (
            EUROPE_BUCHAREST.localize(self.start_datetime)
            if self.start_datetime.tzinfo is None
            else self.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        return start_dt_aware > now and self.status == "Aprobată"

    @property
    def is_past(self):
        now = get_localized_now()
        end_dt_aware = (
            EUROPE_BUCHAREST.localize(self.end_datetime)
            if self.end_datetime.tzinfo is None
            else self.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        return end_dt_aware < now or self.status == "Anulată"


class DailyLeave(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(
        db.Integer,
        db.ForeignKey("student.id", ondelete="CASCADE"),
        nullable=False,
    )
    leave_date = db.Column(db.Date, nullable=False)
    start_time = db.Column(db.Time, nullable=False)
    end_time = db.Column(db.Time, nullable=False)
    reason = db.Column(db.String(255), nullable=True)
    status = db.Column(db.String(50), default="Aprobată", nullable=False)
    student = db.relationship(
        "Student",
        backref=db.backref(
            "daily_leaves", lazy=True, cascade="all, delete-orphan"
        ),
    )
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=True
    )  # Made nullable
    creator = db.relationship(
        "User", backref=db.backref("daily_leaves_created", lazy=True)
    )

    @property
    def start_datetime(self):
        return datetime.combine(self.leave_date, self.start_time)

    @property
    def end_datetime(self):
        effective_end_date = self.leave_date
        if self.end_time < self.start_time:
            effective_end_date += timedelta(days=1)
        return datetime.combine(
            effective_end_date, self.end_time
        )  # Returns naive datetime

    @property
    def is_active(self):
        now = get_localized_now()
        # self.start_datetime and self.end_datetime are properties returning naive datetimes
        start_dt_aware = EUROPE_BUCHAREST.localize(self.start_datetime)
        end_dt_aware = EUROPE_BUCHAREST.localize(self.end_datetime)
        return (
            start_dt_aware <= now <= end_dt_aware and self.status == "Aprobată"
        )

    @property
    def is_upcoming(self):
        now = get_localized_now()
        start_dt_aware = EUROPE_BUCHAREST.localize(self.start_datetime)
        return start_dt_aware > now and self.status == "Aprobată"

    @property
    def is_past(self):
        now = get_localized_now()
        end_dt_aware = EUROPE_BUCHAREST.localize(self.end_datetime)
        return end_dt_aware < now or self.status == "Anulată"

    @property
    def leave_type_display(self):
        in_program_start, in_program_end = time(7, 0), time(14, 20)
        out_program_evening_start, out_program_morning_end = time(22, 0), time(
            7, 0
        )
        if (
            in_program_start <= self.start_time <= in_program_end
            and in_program_start <= self.end_time <= in_program_end
            and self.start_time < self.end_time
        ):
            return "În program"
        elif (
            (
                self.start_time >= out_program_evening_start
                or self.start_time < out_program_morning_end
            )
            and (
                self.end_time <= out_program_morning_end
                or self.end_time > self.start_time
                or self.start_time > self.end_time
            )
            and not (
                in_program_start <= self.start_time <= in_program_end
                and in_program_start <= self.end_time <= in_program_end
            )
        ):
            return "Afară program"
        return "Nespecificat"


class WeekendLeave(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(
        db.Integer,
        db.ForeignKey("student.id", ondelete="CASCADE"),
        nullable=False,
    )
    weekend_start_date = db.Column(db.Date, nullable=False)
    day1_selected = db.Column(db.String(10), nullable=True)
    day1_date = db.Column(db.Date, nullable=True)
    day1_start_time = db.Column(db.Time, nullable=True)
    day1_end_time = db.Column(db.Time, nullable=True)
    day2_selected = db.Column(db.String(10), nullable=True)
    day2_date = db.Column(db.Date, nullable=True)
    day2_start_time = db.Column(db.Time, nullable=True)
    day2_end_time = db.Column(db.Time, nullable=True)
    day3_selected = db.Column(db.String(10), nullable=True)
    day3_date = db.Column(db.Date, nullable=True)
    day3_start_time = db.Column(db.Time, nullable=True)
    day3_end_time = db.Column(db.Time, nullable=True)
    duminica_biserica = db.Column(
        db.Boolean, default=False, nullable=False
    )  # New field for church attendance
    reason = db.Column(db.String(255), nullable=True)
    status = db.Column(db.String(50), default="Aprobată", nullable=False)
    student = db.relationship(
        "Student",
        backref=db.backref(
            "weekend_leaves", lazy=True, cascade="all, delete-orphan"
        ),
    )
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=True
    )  # Made nullable
    creator = db.relationship(
        "User", backref=db.backref("weekend_leaves_created", lazy=True)
    )

    def get_intervals(self):
        intervals = []
        days_info = [
            (
                self.day1_date,
                self.day1_start_time,
                self.day1_end_time,
                self.day1_selected,
            ),
            (
                self.day2_date,
                self.day2_start_time,
                self.day2_end_time,
                self.day2_selected,
            ),
            (
                self.day3_date,
                self.day3_start_time,
                self.day3_end_time,
                self.day3_selected,
            ),
        ]
        for d_date, s_time, e_time, d_name in days_info:
            if d_date and s_time and e_time:
                # Create naive datetime objects first
                s_dt_naive = datetime.combine(d_date, s_time)
                e_dt_naive = datetime.combine(d_date, e_time)

                # Handle overnight case for naive end datetime
                if e_dt_naive < s_dt_naive:  # e.g. start 22:00, end 02:00
                    e_dt_naive += timedelta(days=1)

                # Localize to make them aware
                s_dt_aware = EUROPE_BUCHAREST.localize(s_dt_naive)
                e_dt_aware = EUROPE_BUCHAREST.localize(e_dt_naive)

                intervals.append(
                    {
                        "day_name": d_name,
                        "start": s_dt_aware,
                        "end": e_dt_aware,
                    }
                )
        return sorted(intervals, key=lambda x: x["start"])

    @property
    def is_overall_active_or_upcoming(self):
        now = get_localized_now()
        if self.status != "Aprobată":
            return False
        return any(interval["end"] >= now for interval in self.get_intervals())

    @property
    def is_any_interval_active_now(self):
        if self.status != "Aprobată":
            return False
        # This code is only reached if status is 'Aprobată'
        now = get_localized_now()
        return any(
            interval["start"] <= now <= interval["end"]
            for interval in self.get_intervals()
        )

    @property
    def is_overall_past(self):
        now = get_localized_now()
        return (
            True
            if self.status == "Anulată"
            else not self.is_overall_active_or_upcoming
        )

    @property
    def display_days_and_times(self):
        return (
            "; ".join(
                [
                    f"{i['day_name']} ({i['start'].strftime('%d.%m')}) {i['start'].strftime('%H:%M')}-{i['end'].strftime('%H:%M')}"
                    for i in self.get_intervals()
                ]
            )
            or "Nespecificat"
        )


class VolunteerActivity(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=True)
    activity_date = db.Column(db.Date, nullable=False)
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=True
    )  # Made nullable
    creator = db.relationship(
        "User", backref=db.backref("volunteer_activities_created", lazy=True)
    )
    participants = db.relationship(
        "ActivityParticipant",
        backref="activity",
        lazy="dynamic",
        cascade="all, delete-orphan",
    )


class ActivityParticipant(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    activity_id = db.Column(
        db.Integer,
        db.ForeignKey("volunteer_activity.id", ondelete="CASCADE"),
        nullable=False,
    )
    student_id = db.Column(
        db.Integer,
        db.ForeignKey("student.id", ondelete="CASCADE"),
        nullable=False,
    )
    points_awarded = db.Column(db.Integer, default=0)
    student = db.relationship(
        "Student",
        backref=db.backref(
            "participations", lazy=True, cascade="all, delete-orphan"
        ),
    )


# --- Models for "Save for Later" Volunteer Generation ---
volunteer_session_participants = db.Table(
    "volunteer_session_participants",
    db.Column(
        "volunteer_session_id",
        db.Integer,
        db.ForeignKey("volunteer_session.id"),
        primary_key=True,
    ),
    db.Column(
        "student_id", db.Integer, db.ForeignKey("student.id"), primary_key=True
    ),
)


class VolunteerSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(150), nullable=False)
    created_at = db.Column(
        db.DateTime, default=datetime.utcnow, nullable=False
    )
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=False
    )
    creator = db.relationship(
        "User", backref=db.backref("volunteer_sessions_created", lazy=True)
    )
    students = db.relationship(
        "Student",
        secondary=volunteer_session_participants,
        lazy="dynamic",
        backref=db.backref("volunteer_sessions", lazy="dynamic"),
    )

    def __repr__(self):
        return f"<VolunteerSession {self.name} (Created by User ID: {self.created_by_user_id})>"


class ServiceAssignment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(
        db.Integer,
        db.ForeignKey("student.id", ondelete="CASCADE"),
        nullable=False,
    )
    service_type = db.Column(db.String(100), nullable=False)
    service_date = db.Column(db.Date, nullable=False)
    start_datetime = db.Column(db.DateTime, nullable=False)
    end_datetime = db.Column(db.DateTime, nullable=False)
    participates_in_roll_call = db.Column(db.Boolean, default=True)
    notes = db.Column(db.Text, nullable=True)
    student = db.relationship(
        "Student",
        backref=db.backref(
            "service_assignments", lazy=True, cascade="all, delete-orphan"
        ),
    )
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=True
    )  # Made nullable
    creator = db.relationship(
        "User", backref=db.backref("service_assignments_created", lazy=True)
    )

    @property
    def is_active(self):
        now = get_localized_now()
        start_dt_aware = (
            EUROPE_BUCHAREST.localize(self.start_datetime)
            if self.start_datetime.tzinfo is None
            else self.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        end_dt_aware = (
            EUROPE_BUCHAREST.localize(self.end_datetime)
            if self.end_datetime.tzinfo is None
            else self.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        return start_dt_aware <= now <= end_dt_aware

    @property
    def is_upcoming(self):
        now = get_localized_now()
        start_dt_aware = (
            EUROPE_BUCHAREST.localize(self.start_datetime)
            if self.start_datetime.tzinfo is None
            else self.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        return start_dt_aware > now

    @property
    def is_past(self):
        now = get_localized_now()
        end_dt_aware = (
            EUROPE_BUCHAREST.localize(self.end_datetime)
            if self.end_datetime.tzinfo is None
            else self.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        return end_dt_aware < now


class ActionLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=True
    )  # Nullable if action can be system-initiated or by non-logged-in user
    timestamp = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    action_type = db.Column(
        db.String(50), nullable=False
    )  # e.g., CREATE, UPDATE, DELETE, LOGIN, LOGOUT, RESET_CODE
    target_model = db.Column(
        db.String(100), nullable=True
    )  # e.g., "Student", "Permission", "User"
    target_id = db.Column(
        db.Integer, nullable=True
    )  # ID of the affected record
    details_before = db.Column(
        db.Text, nullable=True
    )  # JSON string of relevant fields before change
    details_after = db.Column(
        db.Text, nullable=True
    )  # JSON string of relevant fields after change
    description = db.Column(
        db.Text, nullable=True
    )  # General description (e.g., IP address for login, or summary)

    user = db.relationship(
        "User", backref=db.backref("action_logs", lazy="dynamic")
    )  # Changed to lazy='dynamic'

    def __repr__(self):
        user_desc = (
            f"User {self.user_id}" if self.user_id else "System/UnknownUser"
        )
        target_desc = (
            f" on {self.target_model}({self.target_id})"
            if self.target_model and self.target_id
            else ""
        )
        description_desc = (
            f" - {self.description[:50]}..." if self.description else ""
        )
        return f'<ActionLog {self.timestamp.strftime("%Y-%m-%d %H:%M:%S")} - {user_desc} - {self.action_type}{target_desc}{description_desc}>'


class AuditLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(
        db.Integer,
        db.ForeignKey("student.id", ondelete="CASCADE"),
        nullable=False,
    )
    user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=True
    )  # Who performed the action
    timestamp = db.Column(
        db.DateTime, default=get_localized_now, nullable=False
    )
    action = db.Column(
        db.String(255), nullable=False
    )  # e.g., "CREATE_STUDENT", "UPDATE_LEAVE", "DELETE_SERVICE"
    details = db.Column(db.Text, nullable=True)  # Human-readable description

    student = db.relationship(
        "Student",
        backref=db.backref(
            "audit_logs", lazy="dynamic", cascade="all, delete-orphan"
        ),
    )
    user = db.relationship(
        "User", backref=db.backref("performed_student_actions", lazy="dynamic")
    )

    def __repr__(self):
        return f"<AuditLog for Student {self.student_id} at {self.timestamp}>"


def log_student_action(student_id, action, details):
    """Helper function to log an action related to a specific student."""
    try:
        user_id = current_user.id if current_user.is_authenticated else None
        log_entry = AuditLog(
            student_id=student_id,
            user_id=user_id,
            action=action,
            details=details,
        )
        db.session.add(log_entry)
    except Exception as e:
        app.logger.error(
            f"Failed to log student action for student {student_id}: {e}"
        )


class UpdateTopic(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(
        db.DateTime, default=datetime.utcnow, nullable=False
    )
    updated_at = db.Column(
        db.DateTime,
        default=datetime.utcnow,
        onupdate=datetime.utcnow,
        nullable=False,
    )
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    author = db.relationship(
        "User", backref=db.backref("update_topics_authored", lazy=True)
    )
    is_pinned = db.Column(db.Boolean, default=False, nullable=False)
    # status_color can map to Bootstrap alert/badge classes e.g., 'primary', 'success', 'warning', 'danger', 'info', 'light', 'dark'
    status_color = db.Column(db.String(20), nullable=True)
    is_visible = db.Column(
        db.Boolean, default=True, nullable=False
    )  # For soft delete or drafts

    def __repr__(self):
        return f"<UpdateTopic {self.id}: {self.title[:50]}>"


class SiteSetting(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    key = db.Column(
        db.String(50), unique=True, nullable=False
    )  # e.g., 'home_page_title', 'home_page_badge_text'
    value = db.Column(db.Text, nullable=True)

    def __repr__(self):
        return f'<SiteSetting {self.key}: {self.value[:20] if self.value else "None"}>'


class PublicViewCode(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(16), unique=True, nullable=False)
    scope_type = db.Column(
        db.String(20), nullable=False
    )  # 'company' or 'battalion'
    scope_id = db.Column(
        db.String(50), nullable=False
    )  # The ID of the company or battalion
    expires_at = db.Column(db.DateTime, nullable=False)
    created_at = db.Column(
        db.DateTime, default=datetime.utcnow, nullable=False
    )
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=False
    )
    creator = db.relationship(
        "User", backref=db.backref("public_view_codes_created", lazy=True)
    )
    is_active = db.Column(db.Boolean, default=True, nullable=False)

    def __repr__(self):
        return f"<PublicViewCode {self.code} for {self.scope_type} {self.scope_id}>"


class ScopedAccessCode(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(16), unique=True, nullable=False)
    description = db.Column(
        db.String(255), nullable=True
    )  # e.g., "Acces Voluntariat pentru Sdt. Popescu"
    # This JSON field will store a list of route prefixes, e.g., '["/volunteer", "/gradat/students"]'
    permissions = db.Column(db.Text, nullable=False)
    expires_at = db.Column(db.DateTime, nullable=False)
    created_at = db.Column(
        db.DateTime, default=datetime.utcnow, nullable=False
    )
    # The gradat who created this code and whose data will be accessed
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=False
    )
    creator = db.relationship(
        "User", backref=db.backref("scoped_access_codes_created", lazy=True)
    )
    is_active = db.Column(db.Boolean, default=True, nullable=False)

    def get_permissions_list(self):
        try:
            return json.loads(self.permissions)
        except (json.JSONDecodeError, TypeError):
            return []

    def __repr__(self):
        return f"<ScopedAccessCode {self.code} for User ID {self.created_by_user_id}>"


class LeaveTemplate(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(150), nullable=False)
    template_type = db.Column(
        db.String(50), nullable=False
    )  # 'permission', 'daily_leave', 'service'
    created_by_user_id = db.Column(
        db.Integer, db.ForeignKey("user.id"), nullable=False
    )
    # JSON column to store the template data
    data = db.Column(db.Text, nullable=False)

    creator = db.relationship(
        "User", backref=db.backref("leave_templates", lazy=True)
    )

    def get_data(self):
        try:
            return json.loads(self.data)
        except (json.JSONDecodeError, TypeError):
            return {}

    def __repr__(self):
        return f"<LeaveTemplate {self.id}: {self.name} ({self.template_type})>"


@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


def init_db():
    with app.app_context():
        db.create_all()  # This will create UpdateTopic table if it doesn't exist
        if not User.query.filter_by(username="admin").first():
            admin = User(username="admin", role="admin", is_first_login=False)
            admin.set_password("admin123")
            db.session.add(admin)
            db.session.commit()
            print("Admin user created.")
        else:
            print("Admin user already exists.")
        print("DB initialized.")


# --- Insights / Operational Overview ---
@app.route("/gradat/insights")
@login_required
def gradat_insights():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    # Scope to students managed by current gradat
    students_managed = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .with_entities(Student.id, Student.nume, Student.prenume)
        .all()
    )
    student_ids = [sid for sid, *_ in students_managed]
    if not student_ids:
        return render_template(
            "insights.html",
            scope_label="Plutonul meu",
            stats={},
            upcoming=[],
            conflicts=[],
        )

    now = get_localized_now()
    horizon = now + timedelta(days=7)

    # Upcoming items
    dl_up = (
        DailyLeave.query.options(joinedload(DailyLeave.student))
        .filter(
            DailyLeave.student_id.in_(student_ids),
            DailyLeave.start_time >= (now - timedelta(days=1)).time(),
        )
        .order_by(DailyLeave.start_time.asc())
        .all()
    )

    # Weekend leaves are stored per-day; build intervals via get_intervals()
    wl_all = (
        WeekendLeave.query.options(joinedload(WeekendLeave.student))
        .filter(
            WeekendLeave.student_id.in_(student_ids),
            WeekendLeave.status == "Aprobată",
        )
        .all()
    )

    perm_up = (
        Permission.query.options(joinedload(Permission.student))
        .filter(
            Permission.student_id.in_(student_ids),
            Permission.end_datetime >= now,
        )
        .order_by(Permission.start_datetime.asc())
        .all()
    )

    serv_up = (
        ServiceAssignment.query.options(joinedload(ServiceAssignment.student))
        .filter(
            ServiceAssignment.student_id.in_(student_ids),
            ServiceAssignment.end_datetime >= now,
        )
        .order_by(ServiceAssignment.start_datetime.asc())
        .all()
    )

    def student_name(st):
        try:
            return (
                f"{st.grad_militar} {st.nume} {st.prenume}"
                if hasattr(st, "grad_militar")
                else f"{st.nume} {st.prenume}"
            )
        except Exception:
            return "—"

    upcoming = []
    for p in perm_up:
        start_aware = (
            EUROPE_BUCHAREST.localize(p.start_datetime)
            if p.start_datetime.tzinfo is None
            else p.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        end_aware = (
            EUROPE_BUCHAREST.localize(p.end_datetime)
            if p.end_datetime.tzinfo is None
            else p.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        upcoming.append(
            dict(
                kind="Permisie",
                student=student_name(p.student),
                start=start_aware,
                end=end_aware,
                meta=p.status,
            )
        )
    for d in dl_up:
        start_dt = datetime.combine(now.date(), d.start_time)
        end_dt = datetime.combine(now.date(), d.end_time)
        start_aware = EUROPE_BUCHAREST.localize(start_dt)
        end_aware = EUROPE_BUCHAREST.localize(end_dt)
        upcoming.append(
            dict(
                kind="Învoire Zilnică",
                student=student_name(d.student),
                start=start_aware,
                end=end_aware,
                meta=d.reason if hasattr(d, "reason") else "",
            )
        )
    for w in wl_all:
        for it in w.get_intervals():
            # it['start'] and it['end'] already aware
            if it["end"] >= now:
                upcoming.append(
                    dict(
                        kind="Învoire Weekend",
                        student=student_name(w.student),
                        start=it["start"],
                        end=it["end"],
                        meta=w.status,
                    )
                )
    for s in serv_up:
        start_aware = (
            EUROPE_BUCHAREST.localize(s.start_datetime)
            if s.start_datetime.tzinfo is None
            else s.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        end_aware = (
            EUROPE_BUCHAREST.localize(s.end_datetime)
            if s.end_datetime.tzinfo is None
            else s.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        upcoming.append(
            dict(
                kind="Serviciu",
                student=student_name(s.student),
                start=start_aware,
                end=end_aware,
                meta=s.service_type,
            )
        )

    # Limit upcoming to next 7 days for view brevity
    upcoming = [u for u in upcoming if u["start"] <= horizon]
    upcoming.sort(key=lambda x: (x["start"], x["student"]))

    # Quick stats
    stats = {
        "students_count": len(student_ids),
        "permissions_active": sum(
            1
            for p in perm_up
            if (
                (
                    EUROPE_BUCHAREST.localize(p.start_datetime)
                    if p.start_datetime.tzinfo is None
                    else p.start_datetime.astimezone(EUROPE_BUCHAREST)
                )
                <= now
                <= (
                    EUROPE_BUCHAREST.localize(p.end_datetime)
                    if p.end_datetime.tzinfo is None
                    else p.end_datetime.astimezone(EUROPE_BUCHAREST)
                )
            )
        ),
        "services_active": sum(
            1
            for s in serv_up
            if (
                (
                    EUROPE_BUCHAREST.localize(s.start_datetime)
                    if s.start_datetime.tzinfo is None
                    else s.start_datetime.astimezone(EUROPE_BUCHAREST)
                )
                <= now
                <= (
                    EUROPE_BUCHAREST.localize(s.end_datetime)
                    if s.end_datetime.tzinfo is None
                    else s.end_datetime.astimezone(EUROPE_BUCHAREST)
                )
            )
        ),
        "leaves_today": sum(
            1 for d in dl_up if d.start_time <= now.time() <= d.end_time
        ),
        "weekend_leaves_upcoming": sum(
            1
            for w in wl_all
            for it in w.get_intervals()
            if it["start"].date() <= horizon.date() and it["end"] >= now
        ),
    }

    # Conflict detection (overlaps per student across all types)
    per_student = {}

    def add_evt(sid, label, start, end):
        # Ensure all events are aware
        start_aware = (
            EUROPE_BUCHAREST.localize(start)
            if start.tzinfo is None
            else start.astimezone(EUROPE_BUCHAREST)
        )
        end_aware = (
            EUROPE_BUCHAREST.localize(end)
            if end.tzinfo is None
            else end.astimezone(EUROPE_BUCHAREST)
        )
        per_student.setdefault(sid, []).append((start_aware, end_aware, label))

    for p in perm_up:
        add_evt(
            p.student_id,
            f"Permisie ({p.status})",
            p.start_datetime,
            p.end_datetime,
        )
    for s in serv_up:
        add_evt(
            s.student_id,
            f"Serviciu ({s.service_type})",
            s.start_datetime,
            s.end_datetime,
        )
    for w in wl_all:
        for it in w.get_intervals():
            add_evt(
                w.student_id, f"Weekend ({w.status})", it["start"], it["end"]
            )
    # Daily leave as same day window
    for d in dl_up:
        sd = datetime.combine(now.date(), d.start_time)
        ed = datetime.combine(now.date(), d.end_time)
        sd_aware = EUROPE_BUCHAREST.localize(sd)
        ed_aware = EUROPE_BUCHAREST.localize(ed)
        add_evt(d.student_id, "Învoire Zilnică", sd_aware, ed_aware)

    conflicts = []
    for sid, events in per_student.items():
        ev = sorted(events, key=lambda t: t[0])
        for i in range(len(ev)):
            for j in range(i + 1, len(ev)):
                a_start, a_end, a_lbl = ev[i]
                b_start, b_end, b_lbl = ev[j]
                if a_end > b_start and a_start < b_end:
                    st = db.session.get(Student, sid)
                    conflicts.append(
                        dict(
                            student=student_name(st),
                            a=a_lbl,
                            b=b_lbl,
                            start=max(a_start, b_start),
                            end=min(a_end, b_end),
                        )
                    )
    conflicts.sort(key=lambda x: (x["student"], x["start"]))

    return render_template(
        "insights.html",
        scope_label="Plutonul meu",
        stats=stats,
        upcoming=upcoming,
        conflicts=conflicts,
    )


@app.route("/admin/insights")
@login_required
def admin_insights():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    # Reuse gradat_insights logic but for all students (capped)
    now = get_localized_now()
    horizon = now + timedelta(days=7)

    perm_up = (
        Permission.query.options(joinedload(Permission.student))
        .filter(Permission.end_datetime >= now)
        .order_by(Permission.start_datetime.asc())
        .limit(500)
        .all()
    )
    dl_up = (
        DailyLeave.query.options(joinedload(DailyLeave.student))
        .order_by(DailyLeave.start_time.asc())
        .limit(500)
        .all()
    )
    wl_all = (
        WeekendLeave.query.options(joinedload(WeekendLeave.student))
        .filter(WeekendLeave.status == "Aprobată")
        .limit(500)
        .all()
    )
    serv_up = (
        ServiceAssignment.query.options(joinedload(ServiceAssignment.student))
        .filter(ServiceAssignment.end_datetime >= now)
        .order_by(ServiceAssignment.start_datetime.asc())
        .limit(500)
        .all()
    )

    def student_name(st):
        try:
            return (
                f"{st.grad_militar} {st.nume} {st.prenume}"
                if hasattr(st, "grad_militar")
                else f"{st.nume} {st.prenume}"
            )
        except Exception:
            return "—"

    upcoming = []
    for p in perm_up:
        start_aware = (
            EUROPE_BUCHAREST.localize(p.start_datetime)
            if p.start_datetime.tzinfo is None
            else p.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        end_aware = (
            EUROPE_BUCHAREST.localize(p.end_datetime)
            if p.end_datetime.tzinfo is None
            else p.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        upcoming.append(
            dict(
                kind="Permisie",
                student=student_name(p.student),
                start=start_aware,
                end=end_aware,
                meta=p.status,
            )
        )
    for d in dl_up:
        # For admin, assume today context for daily
        sd = datetime.combine(now.date(), d.start_time)
        ed = datetime.combine(now.date(), d.end_time)
        sd_aware = EUROPE_BUCHAREST.localize(sd)
        ed_aware = EUROPE_BUCHAREST.localize(ed)
        upcoming.append(
            dict(
                kind="Învoire Zilnică",
                student=student_name(d.student),
                start=sd_aware,
                end=ed_aware,
                meta=getattr(d, "reason", ""),
            )
        )
    for w in wl_all:
        for it in w.get_intervals():
            if it["end"] >= now:
                upcoming.append(
                    dict(
                        kind="Învoire Weekend",
                        student=student_name(w.student),
                        start=it["start"],
                        end=it["end"],
                        meta=w.status,
                    )
                )
    for s in serv_up:
        start_aware = (
            EUROPE_BUCHAREST.localize(s.start_datetime)
            if s.start_datetime.tzinfo is None
            else s.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        end_aware = (
            EUROPE_BUCHAREST.localize(s.end_datetime)
            if s.end_datetime.tzinfo is None
            else s.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        upcoming.append(
            dict(
                kind="Serviciu",
                student=student_name(s.student),
                start=start_aware,
                end=end_aware,
                meta=s.service_type,
            )
        )
    upcoming = [u for u in upcoming if u["start"] <= horizon]
    upcoming.sort(key=lambda x: (x["start"], x["student"]))

    stats = {
        "permissions_active": sum(
            1
            for p in perm_up
            if (
                (
                    EUROPE_BUCHAREST.localize(p.start_datetime)
                    if p.start_datetime.tzinfo is None
                    else p.start_datetime.astimezone(EUROPE_BUCHAREST)
                )
                <= now
                <= (
                    EUROPE_BUCHAREST.localize(p.end_datetime)
                    if p.end_datetime.tzinfo is None
                    else p.end_datetime.astimezone(EUROPE_BUCHAREST)
                )
            )
        ),
        "services_active": sum(
            1
            for s in serv_up
            if (
                (
                    EUROPE_BUCHAREST.localize(s.start_datetime)
                    if s.start_datetime.tzinfo is None
                    else s.start_datetime.astimezone(EUROPE_BUCHAREST)
                )
                <= now
                <= (
                    EUROPE_BUCHAREST.localize(s.end_datetime)
                    if s.end_datetime.tzinfo is None
                    else s.end_datetime.astimezone(EUROPE_BUCHAREST)
                )
            )
        ),
        "leaves_today": sum(
            1 for d in dl_up if d.start_time <= now.time() <= d.end_time
        ),
        "weekend_leaves_upcoming": sum(
            1
            for w in wl_all
            for it in w.get_intervals()
            if it["start"].date() <= horizon.date() and it["end"] >= now
        ),
    }

    # Conflicts as above (limited)
    per_student = {}

    def add_evt(sid, label, start, end):
        per_student.setdefault(sid, []).append((start, end, label))

    for p in perm_up:
        add_evt(
            p.student_id,
            f"Permisie ({p.status})",
            p.start_datetime,
            p.end_datetime,
        )
    for s in serv_up:
        add_evt(
            s.student_id,
            f"Serviciu ({s.service_type})",
            s.start_datetime,
            s.end_datetime,
        )
    for w in wl_all:
        for it in w.get_intervals():
            if it["end"] >= now:  # Only consider active/upcoming for conflicts
                add_evt(
                    w.student_id,
                    f"Weekend ({w.status})",
                    it["start"],
                    it["end"],
                )
    for d in dl_up:
        sd = datetime.combine(now.date(), d.start_time)
        ed = datetime.combine(now.date(), d.end_time)
        sd_aware = EUROPE_BUCHAREST.localize(sd)
        ed_aware = EUROPE_BUCHAREST.localize(ed)
        add_evt(d.student_id, "Învoire Zilnică", sd_aware, ed_aware)
    conflicts = []
    for sid, events in per_student.items():
        ev = sorted(events, key=lambda t: t[0])
        for i in range(len(ev)):
            for j in range(i + 1, len(ev)):
                a_start, a_end, a_lbl = ev[i]
                b_start, b_end, b_lbl = ev[j]
                if a_end > b_start and a_start < b_end:
                    st = db.session.get(Student, sid)
                    conflicts.append(
                        dict(
                            student=student_name(st),
                            a=a_lbl,
                            b=b_lbl,
                            start=max(a_start, b_start),
                            end=min(a_end, b_end),
                        )
                    )
    conflicts.sort(key=lambda x: (x["student"], x["start"]))

    return render_template(
        "insights.html",
        scope_label="Sistem (Admin)",
        stats=stats,
        upcoming=upcoming,
        conflicts=conflicts,
    )


# --- Daily Brief (today/tomorrow overview) ---
def _collect_brief(scope_student_ids=None):
    now = get_localized_now()
    today = now.date()
    tomorrow = today + timedelta(days=1)

    def within_scope(q):
        if scope_student_ids is None:
            return q
        return q.filter(
            getattr(q.column_descriptions[0]["entity"], "student_id").in_(
                scope_student_ids
            )
        )

    # Permissions
    perm_today = Permission.query.options(
        joinedload(Permission.student)
    ).filter(
        Permission.start_datetime <= datetime.combine(today, time(23, 59, 59)),
        Permission.end_datetime >= datetime.combine(today, time(0, 0, 0)),
    )
    perm_tom = Permission.query.options(joinedload(Permission.student)).filter(
        Permission.start_datetime
        <= datetime.combine(tomorrow, time(23, 59, 59)),
        Permission.end_datetime >= datetime.combine(tomorrow, time(0, 0, 0)),
    )
    if scope_student_ids is not None:
        perm_today = perm_today.filter(
            Permission.student_id.in_(scope_student_ids)
        )
        perm_tom = perm_tom.filter(
            Permission.student_id.in_(scope_student_ids)
        )

    # Services
    serv_today = ServiceAssignment.query.options(
        joinedload(ServiceAssignment.student)
    ).filter(
        ServiceAssignment.start_datetime
        <= datetime.combine(today, time(23, 59, 59)),
        ServiceAssignment.end_datetime
        >= datetime.combine(today, time(0, 0, 0)),
    )
    serv_tom = ServiceAssignment.query.options(
        joinedload(ServiceAssignment.student)
    ).filter(
        ServiceAssignment.start_datetime
        <= datetime.combine(tomorrow, time(23, 59, 59)),
        ServiceAssignment.end_datetime
        >= datetime.combine(tomorrow, time(0, 0, 0)),
    )
    if scope_student_ids is not None:
        serv_today = serv_today.filter(
            ServiceAssignment.student_id.in_(scope_student_ids)
        )
        serv_tom = serv_tom.filter(
            ServiceAssignment.student_id.in_(scope_student_ids)
        )

    # Daily leaves (today only by definition)
    dl_today = DailyLeave.query.options(joinedload(DailyLeave.student))
    if scope_student_ids is not None:
        dl_today = dl_today.filter(
            DailyLeave.student_id.in_(scope_student_ids)
        )

    # Weekend leaves via intervals
    wl_all = WeekendLeave.query.options(joinedload(WeekendLeave.student))
    if scope_student_ids is not None:
        wl_all = wl_all.filter(WeekendLeave.student_id.in_(scope_student_ids))
    wl_all = wl_all.filter(WeekendLeave.status == "Aprobată").all()

    return dict(
        today=dict(
            permissions=perm_today.order_by(
                Permission.start_datetime.asc()
            ).all(),
            services=serv_today.order_by(
                ServiceAssignment.start_datetime.asc()
            ).all(),
            daily_leaves=dl_today.all(),
            weekend_leaves=[
                w
                for w in wl_all
                if any(
                    it["start"].date() <= today <= it["end"].date()
                    for it in w.get_intervals()
                )
            ],
        ),
        tomorrow=dict(
            permissions=perm_tom.order_by(
                Permission.start_datetime.asc()
            ).all(),
            services=serv_tom.order_by(
                ServiceAssignment.start_datetime.asc()
            ).all(),
            weekend_leaves=[
                w
                for w in wl_all
                if any(
                    it["start"].date() <= tomorrow <= it["end"].date()
                    for it in w.get_intervals()
                )
            ],
        ),
    )


@app.route("/gradat/brief")
@login_required
def gradat_brief():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    students = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .with_entities(Student.id)
        .all()
    )
    student_ids = [sid for sid, in students]
    data = _collect_brief(scope_student_ids=student_ids)
    return render_template(
        "daily_brief.html", scope_label="Plutonul meu", data=data
    )


@app.route("/admin/brief")
@login_required
def admin_brief():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    data = _collect_brief(scope_student_ids=None)
    return render_template(
        "daily_brief.html", scope_label="Sistem (Admin)", data=data
    )


@app.route("/brief/ack", methods=["POST"])
@login_required
def brief_ack():
    try:
        log_action(
            "BRIEF_ACK",
            description=f"Brief acknowledged by user {current_user.id}",
        )
        db.session.commit()
        flash("Brief marcat ca citit.", "success")
    except Exception:
        db.session.rollback()
        flash("Eroare la marcarea brief-ului.", "danger")
    ref = request.referrer or url_for("dashboard")
    return redirect(ref)


# --- Conflict Center ---
def _collect_conflicts(scope_student_ids=None):
    now = get_localized_now()
    horizon = now + timedelta(days=14)

    # Query relevant records
    perm_up = Permission.query.options(joinedload(Permission.student)).filter(
        Permission.end_datetime >= now
    )
    wl_all = WeekendLeave.query.options(
        joinedload(WeekendLeave.student)
    ).filter(WeekendLeave.status == "Aprobată")
    serv_up = ServiceAssignment.query.options(
        joinedload(ServiceAssignment.student)
    ).filter(ServiceAssignment.end_datetime >= now)
    dl_up = DailyLeave.query.options(joinedload(DailyLeave.student)).filter(
        DailyLeave.status == "Aprobată"
    )

    if scope_student_ids is not None:
        perm_up = perm_up.filter(Permission.student_id.in_(scope_student_ids))
        wl_all = wl_all.filter(WeekendLeave.student_id.in_(scope_student_ids))
        serv_up = serv_up.filter(
            ServiceAssignment.student_id.in_(scope_student_ids)
        )
        dl_up = dl_up.filter(DailyLeave.student_id.in_(scope_student_ids))

    perm_up = (
        perm_up.order_by(Permission.start_datetime.asc()).limit(1000).all()
    )
    wl_all = wl_all.limit(1000).all()
    serv_up = (
        serv_up.order_by(ServiceAssignment.start_datetime.asc())
        .limit(1000)
        .all()
    )
    dl_up = (
        dl_up.order_by(
            DailyLeave.leave_date.asc(), DailyLeave.start_time.asc()
        )
        .limit(1000)
        .all()
    )

    per_student = {}

    def add_evt(store, sid, label, start, end):
        start_aware = (
            EUROPE_BUCHAREST.localize(start)
            if start.tzinfo is None
            else start.astimezone(EUROPE_BUCHAREST)
        )
        end_aware = (
            EUROPE_BUCHAREST.localize(end)
            if end.tzinfo is None
            else end.astimezone(EUROPE_BUCHAREST)
        )
        store.setdefault(sid, []).append((start_aware, end_aware, label))

    # Add permissions
    for p in perm_up:
        add_evt(
            per_student,
            p.student_id,
            f"Permisie ({p.status})",
            p.start_datetime,
            p.end_datetime,
        )

    # Add services
    for s in serv_up:
        add_evt(
            per_student,
            s.student_id,
            f"Serviciu ({s.service_type})",
            s.start_datetime,
            s.end_datetime,
        )

    # Add approved weekend leaves
    for w in wl_all:
        for it in w.get_intervals():
            if it["end"] >= now:
                add_evt(
                    per_student,
                    w.student_id,
                    f"Weekend ({w.status})",
                    it["start"],
                    it["end"],
                )

    # Add daily leaves (use model properties for accurate intervals)
    for d in dl_up:
        add_evt(
            per_student,
            d.student_id,
            "Învoire Zilnică",
            d.start_datetime,
            d.end_datetime,
        )

    # Compute conflicts
    conflicts = []
    for sid, events in per_student.items():
        ev = sorted(events, key=lambda t: t[0])
        for i in range(len(ev)):
            for j in range(i + 1, len(ev)):
                a_start, a_end, a_lbl = ev[i]
                b_start, b_end, b_lbl = ev[j]
                overlap_start = max(a_start, b_start)
                overlap_end = min(a_end, b_end)
                if (
                    a_end > b_start
                    and a_start < b_end
                    and overlap_start <= horizon
                ):
                    st = db.session.get(Student, sid)
                    conflicts.append(
                        dict(
                            student=st,
                            a=a_lbl,
                            b=b_lbl,
                            start=overlap_start,
                            end=overlap_end,
                        )
                    )

    conflicts.sort(
        key=lambda x: (x["student"].nume if x["student"] else "", x["start"])
    )
    return conflicts


@app.route("/gradat/conflicts")
@login_required
def gradat_conflicts():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    students = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .with_entities(Student.id)
        .all()
    )
    student_ids = [sid for sid, in students]
    conflicts = _collect_conflicts(student_ids)
    return render_template(
        "conflicts.html", scope_label="Plutonul meu", conflicts=conflicts
    )


@app.route("/admin/conflicts")
@login_required
def admin_conflicts():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    conflicts = _collect_conflicts(None)
    return render_template(
        "conflicts.html", scope_label="Sistem (Admin)", conflicts=conflicts
    )


def get_upcoming_fridays(num_fridays=5):
    """
    Returnează o listă de până la `num_fridays` vinări (YYYY-MM-DD + etichetă),
    începând cu vinerea săptămânii curente. Dacă azi este Luni/Marti/Miercuri,
    include și vinerea anterioară ca prim element pentru a acoperi weekendul curent.
    """
    today = get_localized_now().date()
    weekday = today.weekday()  # Monday=0 ... Sunday=6

    # Găsește vinerea săptămânii curente (poate fi în trecut dacă e sâmbătă/duminică)
    if weekday <= 4:
        current_friday = today + timedelta(days=(4 - weekday))
    else:
        current_friday = today - timedelta(days=(weekday - 4))

    fridays = [current_friday + timedelta(weeks=i) for i in range(num_fridays)]

    # Dacă e Luni(0)/Marți(1)/Miercuri(2), inserează vinerea precedentă
    if weekday < 3:
        prev_friday = current_friday - timedelta(weeks=1)
        fridays.insert(0, prev_friday)
        fridays = fridays[:num_fridays]

    return [
        {"value": d.strftime("%Y-%m-%d"), "display": f"{d.strftime('%d %B %Y')} (Vineri)"}
        for d in fridays
    ]


def validate_daily_leave_times(start_time_obj, end_time_obj, leave_date_obj):
    if leave_date_obj.weekday() > 3:
        return False, "Învoirile zilnice sunt permise doar de Luni până Joi."
    if start_time_obj == end_time_obj:
        return False, "Ora de început și de sfârșit nu pot fi identice."
    return True, "Valid"


# Funcția get_student_status a fost integrată și optimizată în _calculate_presence_data
# și nu mai este necesară separat dacă _calculate_presence_data este singurul apelant.


def check_leave_conflict(
    student_id,
    leave_start_dt,
    leave_end_dt,
    existing_leave_id=None,
    leave_type=None,
):
    """
    Verifică conflictele pentru un student într-un interval dat, excluzând opțional o învoire existentă.
    """
    # 1. Verifică conflicte cu servicii blocante
    blocking_services = ["GSS", "Intervenție"]
    conflicting_service_query = ServiceAssignment.query.filter(
        ServiceAssignment.student_id == student_id,
        ServiceAssignment.service_type.in_(blocking_services),
        ServiceAssignment.start_datetime < leave_end_dt,
        ServiceAssignment.end_datetime > leave_start_dt,
    )
    conflicting_service = conflicting_service_query.first()
    if conflicting_service:
        return f"serviciu ({conflicting_service.service_type}) pe {conflicting_service.service_date.strftime('%d-%m-%Y')}"

    # 2. Verifică conflicte cu Permisii
    perm_query = Permission.query.filter(
        Permission.student_id == student_id,
        Permission.status == "Aprobată",
        Permission.start_datetime < leave_end_dt,
        Permission.end_datetime > leave_start_dt,
    )
    if leave_type == "permission" and existing_leave_id:
        perm_query = perm_query.filter(Permission.id != existing_leave_id)
    conflicting_permission = perm_query.first()
    if conflicting_permission:
        return f"o permisie existentă ({conflicting_permission.start_datetime.strftime('%d.%m %H:%M')} - {conflicting_permission.end_datetime.strftime('%d.%m %H:%M')})"

    # 3. Verifică conflicte cu Învoiri Zilnice
    daily_leaves_query = DailyLeave.query.filter(
        DailyLeave.student_id == student_id, DailyLeave.status == "Aprobată"
    )
    if leave_type == "daily_leave" and existing_leave_id:
        daily_leaves_query = daily_leaves_query.filter(
            DailyLeave.id != existing_leave_id
        )

    for dl in daily_leaves_query.all():
        if (
            dl.start_datetime < leave_end_dt
            and dl.end_datetime > leave_start_dt
        ):
            return f"o învoire zilnică pe {dl.leave_date.strftime('%d.%m')}"

    # 4. Verifică conflicte cu Învoiri de Weekend
    weekend_leaves_query = WeekendLeave.query.filter(
        WeekendLeave.student_id == student_id,
        WeekendLeave.status == "Aprobată",
    )
    if leave_type == "weekend_leave" and existing_leave_id:
        weekend_leaves_query = weekend_leaves_query.filter(
            WeekendLeave.id != existing_leave_id
        )

    # Make the incoming naive datetimes aware for comparison with aware datetimes from get_intervals()
    try:
        leave_start_dt_aware = EUROPE_BUCHAREST.localize(leave_start_dt)
        leave_end_dt_aware = EUROPE_BUCHAREST.localize(leave_end_dt)
    except ValueError:  # Already aware
        leave_start_dt_aware = leave_start_dt.astimezone(EUROPE_BUCHAREST)
        leave_end_dt_aware = leave_end_dt.astimezone(EUROPE_BUCHAREST)

    for wl in weekend_leaves_query.all():
        for interval in wl.get_intervals():
            if (
                interval["start"] < leave_end_dt_aware
                and interval["end"] > leave_start_dt_aware
            ):
                return f"o învoire de weekend ({interval['day_name']})"

    return None


def check_service_conflict_for_student(
    student_id,
    service_start_dt,
    service_end_dt,
    service_type,
    current_service_id=None,
):
    if service_type in ["Intervenție", "GSS"]:
        conflicting_permission = Permission.query.filter(
            Permission.student_id == student_id,
            Permission.status == "Aprobată",
            Permission.start_datetime < service_end_dt,
            Permission.end_datetime > service_start_dt,
        ).first()
        if conflicting_permission:
            return f"permisie ({conflicting_permission.start_datetime.strftime('%d.%m %H:%M')} - {conflicting_permission.end_datetime.strftime('%d.%m %H:%M')})"
        daily_leaves = DailyLeave.query.filter(
            DailyLeave.student_id == student_id,
            DailyLeave.status == "Aprobată",
        ).all()
        for dl in daily_leaves:
            if (
                dl.start_datetime < service_end_dt
                and dl.end_datetime > service_start_dt
            ):
                return f"învoire zilnică ({dl.leave_date.strftime('%d.%m')} {dl.start_time.strftime('%H:%M')}-{dl.end_time.strftime('%H:%M')})"
        weekend_leaves = WeekendLeave.query.filter(
            WeekendLeave.student_id == student_id,
            WeekendLeave.status == "Aprobată",
        ).all()
        for wl in weekend_leaves:
            for interval in wl.get_intervals():
                if (
                    interval["start"] < service_end_dt
                    and interval["end"] > service_start_dt
                ):
                    return f"învoire de weekend ({interval['day_name']} {interval['start'].strftime('%H:%M')}-{interval['end'].strftime('%H:%M')})"
    query_other_services = ServiceAssignment.query.filter(
        ServiceAssignment.student_id == student_id,
        ServiceAssignment.start_datetime < service_end_dt,
        ServiceAssignment.end_datetime > service_start_dt,
    )
    if current_service_id:
        query_other_services = query_other_services.filter(
            ServiceAssignment.id != current_service_id
        )
    conflicting_other_service = query_other_services.first()
    if conflicting_other_service:
        return f"alt serviciu ({conflicting_other_service.service_type} pe {conflicting_other_service.service_date.strftime('%d.%m')})"
    return None


# --- Action Logging Utilities ---
def model_to_dict(instance, exclude_fields=None):
    """Converts a SQLAlchemy model instance to a dictionary, excluding specified fields."""
    if not instance:
        return {}

    default_exclude = [
        "_sa_instance_state",
        "password_hash",
        "personal_code_hash",
        "unique_code",
    ]  # Standard + sensitive user fields
    fields_to_exclude = set(default_exclude)
    if exclude_fields:
        fields_to_exclude.update(exclude_fields)

    data = {}
    for c in inspect(instance).mapper.column_attrs:
        if c.key not in fields_to_exclude:
            val = getattr(instance, c.key)
            # Convert datetime/date/time objects to ISO format string
            if isinstance(val, (datetime, date, time)):
                data[c.key] = val.isoformat()
            else:
                data[c.key] = val
    return data


def log_action(
    action_type,
    target_model_name=None,
    target_id=None,
    details_before_dict=None,
    details_after_dict=None,
    description=None,
    user_override=None,
):
    """
    Logs an action to the ActionLog table.
    - action_type: Verb describing the action (e.g., "USER_LOGIN", "CREATE_STUDENT").
    - target_model_name: String name of the model class being affected (e.g., "Student").
    - target_id: Integer ID of the specific record affected.
    - details_before_dict: Dictionary representing the state of the object before changes.
    - details_after_dict: Dictionary representing the state of the object after changes.
    - description: A textual description of the event (e.g., IP address, summary of bulk operation).
    - user_override: Optionally provide a User object if current_user is not appropriate (e.g. system actions).
    """
    try:
        log = ActionLog(action_type=action_type)

        acting_user = user_override
        if (
            not acting_user
            and hasattr(current_user, "is_authenticated")
            and current_user.is_authenticated
        ):
            acting_user = current_user

        if acting_user:
            log.user_id = acting_user.id

        log.target_model = target_model_name
        log.target_id = (
            int(target_id) if target_id is not None else None
        )  # Ensure int or None

        if details_before_dict:
            log.details_before = json.dumps(
                details_before_dict, ensure_ascii=False, default=str
            )
        if details_after_dict:
            log.details_after = json.dumps(
                details_after_dict, ensure_ascii=False, default=str
            )

        log.description = description

        db.session.add(log)
        # Committing is handled by the calling route after the main operation succeeds.
    except Exception as e:
        app.logger.error(
            f"AUDIT LOGGING FAILED for action '{action_type}': {str(e)}"
        )
        # Avoid rollback here, as it might interfere with the main operation's transaction.
        # The main operation should handle its own rollback on failure.


# --- Rute Comune ---
@app.route("/")
def home():
    total_students = 0
    total_users = 0
    total_volunteer_activities = 0
    try:
        total_students = Student.query.count()
        total_users = User.query.filter(User.role != "admin").count()
        total_volunteer_activities = VolunteerActivity.query.count()
    except Exception as e:
        pass

    # Fetch homepage settings
    default_title = "UNAP User Panel"
    default_badge_text = "Beta v2.5"  # Original badge text

    title_setting = SiteSetting.query.filter_by(key="home_page_title").first()
    badge_setting = SiteSetting.query.filter_by(
        key="home_page_badge_text"
    ).first()

    display_title = (
        title_setting.value
        if title_setting and title_setting.value
        else default_title
    )
    display_badge_text = (
        badge_setting.value
        if badge_setting and badge_setting.value
        else default_badge_text
    )

    # If badge text is explicitly set to "None" or empty string by admin, don't show badge.
    # The template will handle the logic of not rendering the span if display_badge_text is empty.

    return render_template(
        "home.html",
        total_students=total_students,
        total_users=total_users,
        total_volunteer_activities=total_volunteer_activities,
        home_page_title=display_title,
        home_page_badge_text=display_badge_text,
    )


@app.route("/updates")
@login_required  # Sau eliminați @login_required dacă pagina este publică
def public_updates_page():
    # Aici veți adăuga logica pentru a prelua anunțurile/actualizările din baza de date
    # De exemplu: updates = UpdateTopic.query.order_by(UpdateTopic.created_at.desc()).all()
    # Momentan, vom presupune că există un template public_updates.html
    # return render_template('public_updates.html', updates=updates)
    # Pentru a evita o nouă eroare dacă template-ul nu există, redăm un text simplu
    # return "Pagina de Anunțuri Publice va fi aici."
    # Sau, dacă aveți un template generic pentru "în construcție"
    # return render_template('placeholder.html', page_title="Anunțuri")
    # Să presupunem că aveți un template 'public_updates.html' și vom trece o listă goală deocamdată
    page = request.args.get("page", 1, type=int)
    per_page = 10  # Sau orice alt număr de elemente pe pagină doriți
    updates_pagination = (
        UpdateTopic.query.filter_by(is_visible=True)
        .order_by(UpdateTopic.is_pinned.desc(), UpdateTopic.updated_at.desc())
        .paginate(page=page, per_page=per_page, error_out=False)
    )
    return render_template(
        "public_updates.html",
        updates_pagination=updates_pagination,
        title="Anunțuri",
    )


@app.route("/user_login", methods=["GET", "POST"])
def user_login():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    if request.method == "POST":
        # Rate limiting: 5 attempts / 5 minutes per IP
        client_ip = (
            request.headers.get(
                "X-Forwarded-For", request.remote_addr or "unknown"
            )
            .split(",")[0]
            .strip()
        )
        bucket_key = f"login:user:{client_ip}"
        if _is_rate_limited(bucket_key, limit=5, window_seconds=300):
            flash(
                "Prea multe încercări. Încearcă din nou în câteva minute.",
                "warning",
            )
            return render_template("errors/429.html"), 429
        login_code = request.form.get("login_code")
        user_by_unique_code = User.query.filter_by(
            unique_code=login_code
        ).first()

        if user_by_unique_code:
            if user_by_unique_code.is_first_login:
                login_user(user_by_unique_code, remember=True)
                try:
                    session.permanent = True
                except Exception:
                    pass
                log_action(
                    "USER_FIRST_LOGIN_SUCCESS",
                    target_model_name="User",
                    target_id=user_by_unique_code.id,
                    description=f"User {user_by_unique_code.username} first login with unique code. IP: {request.remote_addr}",
                    user_override=user_by_unique_code,
                )
                db.session.commit()
                flash(
                    "Autentificare reușită! Setează-ți codul personal.", "info"
                )
                return redirect(url_for("set_personal_code"))
            else:
                log_action(
                    "USER_LOGIN_FAIL_UNIQUE_CODE_USED",
                    target_model_name="User",
                    target_id=user_by_unique_code.id,
                    description=f"Attempt to use already used unique code for user {user_by_unique_code.username}. IP: {request.remote_addr}",
                )
                db.session.commit()
                flash(
                    "Acest cod unic a fost deja folosit pentru prima autentificare. Te rugăm folosește codul personal setat.",
                    "warning",
                )
                return redirect(url_for("user_login"))

        # Narrow search to non-admin users who have a personal code set
        users_non_admin = User.query.filter(
            User.role != "admin", User.personal_code_hash.isnot(None)
        ).all()
        user_by_personal_code = next(
            (u for u in users_non_admin if u.check_personal_code(login_code)),
            None,
        )

        if user_by_personal_code:
            if (
                user_by_personal_code.is_first_login
            ):  # Should not happen if logic is correct
                log_action(
                    "USER_LOGIN_FAIL_CONFIG_ERROR",
                    target_model_name="User",
                    target_id=user_by_personal_code.id,
                    description=f"User {user_by_personal_code.username} attempted login with personal code but is_first_login is true. IP: {request.remote_addr}",
                )
                db.session.commit()
                flash(
                    "Eroare de configurare cont. Contactează administratorul.",
                    "danger",
                )
                return redirect(url_for("user_login"))

            login_user(user_by_personal_code, remember=True)
            try:
                session.permanent = True
            except Exception:
                pass
            log_action(
                "USER_LOGIN_SUCCESS",
                target_model_name="User",
                target_id=user_by_personal_code.id,
                description=f"User {user_by_personal_code.username} login with personal code. IP: {request.remote_addr}",
                user_override=user_by_personal_code,
            )
            db.session.commit()
            flash("Autentificare reușită!", "success")
            return redirect(url_for("dashboard"))

        log_action(
            "USER_LOGIN_FAIL_INVALID_CODE",
            description=f"Invalid/Expired login code provided: '{login_code[:20]}...'. IP: {request.remote_addr}",
        )
        db.session.commit()
        flash("Cod de autentificare invalid sau expirat.", "danger")
        return redirect(url_for("user_login"))
    return render_template("user_login.html")


@app.route("/admin_login", methods=["GET", "POST"])
def admin_login():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    if request.method == "POST":
        # Rate limiting: 5 attempts / 5 minutes per IP
        client_ip = (
            request.headers.get(
                "X-Forwarded-For", request.remote_addr or "unknown"
            )
            .split(",")[0]
            .strip()
        )
        bucket_key = f"login:admin:{client_ip}"
        if _is_rate_limited(bucket_key, limit=5, window_seconds=300):
            flash(
                "Prea multe încercări. Încearcă din nou în câteva minute.",
                "warning",
            )
            return render_template("errors/429.html"), 429
        username = request.form.get("username")
        password = request.form.get("password")
        user = User.query.filter_by(username=username, role="admin").first()
        if user and user.check_password(password):
            login_user(user, remember=True)
            try:
                session.permanent = True
            except Exception:
                pass
            log_action(
                "ADMIN_LOGIN_SUCCESS",
                target_model_name="User",
                target_id=user.id,
                description=f"Admin user {user.username} logged in. IP: {request.remote_addr}",
                user_override=user,
            )
            db.session.commit()
            flash("Autentificare admin reușită!", "success")
            return redirect(url_for("dashboard"))
        else:
            log_action(
                "ADMIN_LOGIN_FAIL",
                description=f"Failed admin login attempt for username '{username}'. IP: {request.remote_addr}",
            )
            db.session.commit()
            flash("Nume de utilizator sau parolă admin incorecte.", "danger")
            return redirect(url_for("admin_login"))
    return render_template("admin_login.html")


# Lightweight health endpoint
@app.route("/healthz")
def healthz():
    try:
        # DB quick check
        db.session.execute(sa.text("SELECT 1"))
        db_ok = True
    except Exception:
        db_ok = False
    status = 200 if db_ok else 500
    return (
        jsonify(
            {
                "status": "ok" if db_ok else "degraded",
                "db": "ok" if db_ok else "error",
                "time": get_localized_now().isoformat(),
            }
        ),
        status,
    )


@app.route("/logout")
@login_required
def logout():
    user_id_logged_out = current_user.id
    username_logged_out = current_user.username
    logout_user()
    log_action(
        "USER_LOGOUT",
        target_model_name="User",
        target_id=user_id_logged_out,
        description=f"User {username_logged_out} logged out. IP: {request.remote_addr}",
    )
    db.session.commit()
    flash("Ai fost deconectat.", "success")
    return redirect(url_for("home"))


@app.route("/set_personal_code", methods=["GET", "POST"])
@login_required
def set_personal_code():
    if not current_user.is_first_login:
        flash("Codul personal a fost deja setat.", "info")
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        personal_code = request.form.get("personal_code")
        confirm_personal_code = request.form.get("confirm_personal_code")

        if not personal_code or len(personal_code) < 4:
            flash(
                "Codul personal trebuie să aibă minim 4 caractere.", "warning"
            )
            return redirect(url_for("set_personal_code"))
        if personal_code != confirm_personal_code:
            flash("Codurile personale nu se potrivesc.", "warning")
            return redirect(url_for("set_personal_code"))

        # Security fix: Check if the personal code is already in use by another user
        other_users = User.query.filter(
            User.id != current_user.id,
            User.role != "admin",
            User.personal_code_hash.isnot(None),
        ).all()
        for other_user in other_users:
            if other_user.check_personal_code(personal_code):
                flash(
                    "Acest cod personal este deja utilizat de altcineva. Vă rugăm alegeți un alt cod.",
                    "danger",
                )
                return redirect(url_for("set_personal_code"))

        try:
            details_before = {
                "is_first_login": current_user.is_first_login,
                "personal_code_hash_exists": current_user.personal_code_hash
                is not None,
            }

            current_user.set_personal_code(
                personal_code
            )  # This also sets is_first_login = False

            details_after = model_to_dict(
                current_user, exclude_fields=["password_hash", "unique_code"]
            )  # Get updated state
            details_after["personal_code_hash_exists"] = (
                True  # Explicitly state it now exists
            )

            log_action(
                "USER_SET_PERSONAL_CODE_SUCCESS",
                target_model_name="User",
                target_id=current_user.id,
                details_before_dict=details_before,
                details_after_dict=details_after,
                description=f"User {current_user.username} set personal code. IP: {request.remote_addr}",
            )

            db.session.commit()  # Commit user changes and log entry

            flash_message = "Codul personal a fost setat cu succes. Te rugăm să te autentifici din nou folosind noul cod."

            # Log automatic logout after setting personal code
            user_id_logged_out = (
                current_user.id
            )  # Capture before logout_user() invalidates current_user
            username_logged_out = current_user.username
            logout_user()
            log_action(
                "USER_LOGOUT_POST_SET_CODE",
                target_model_name="User",
                target_id=user_id_logged_out,
                description=f"User {username_logged_out} automatically logged out after setting personal code. IP: {request.remote_addr}",
            )
            db.session.commit()  # Commit logout log

            flash(flash_message, "success")
            return redirect(url_for("user_login"))

        except Exception as e:
            db.session.rollback()
            flash_msg = (
                f"A apărut o eroare la setarea codului personal: {str(e)}"
            )
            flash(flash_msg, "danger")
            try:
                log_action(
                    "USER_SET_PERSONAL_CODE_FAIL",
                    target_model_name="User",
                    target_id=current_user.id,
                    description=f"Failed to set personal code for user {current_user.username}. Error: {str(e)}. IP: {request.remote_addr}",
                )
                db.session.commit()
            except Exception as log_e:
                app.logger.error(
                    f"CRITICAL: Failed to commit failure log for USER_SET_PERSONAL_CODE_FAIL: {str(log_e)}"
                )
            return redirect(url_for("set_personal_code"))

    return render_template("set_personal_code.html")


@app.route("/admin/dashboard")
@login_required
def admin_dashboard_route():
    if current_user.role != "admin":
        flash("Acces neautorizat la panoul de administrare.", "danger")
        return redirect(url_for("dashboard"))

    users_to_display = (
        User.query.filter(User.role != "admin").order_by(User.username).all()
    )
    total_user_count = User.query.count()
    total_students_count = Student.query.count()

    # Fetch all active public codes for the admin view
    now = get_localized_now()
    active_public_codes = (
        PublicViewCode.query.options(joinedload(PublicViewCode.creator))
        .filter(
            PublicViewCode.is_active == True, PublicViewCode.expires_at > now
        )
        .order_by(PublicViewCode.created_at.desc())
        .all()
    )

    return render_template(
        "admin_dashboard.html",
        users=users_to_display,
        total_users=total_user_count,
        total_students=total_students_count,
        active_public_codes=active_public_codes,
    )


@app.route("/dashboard")
@login_required
def dashboard():
    if current_user.role == "admin":
        return redirect(url_for("admin_dashboard_route"))
    elif current_user.role == "gradat":
        student_ids_managed = [
            sid
            for (sid,) in Student.query.filter_by(
                created_by_user_id=current_user.id
            )
            .with_entities(Student.id)
            .all()
        ]
        student_count = len(student_ids_managed)

        today_localized = get_localized_now().date()
        today_start = datetime.combine(today_localized, time.min)
        today_end = datetime.combine(today_localized, time.max)

        # Permisii active AZI (se suprapun cu ziua de azi)
        permissions_today_count = Permission.query.filter(
            Permission.student_id.in_(student_ids_managed),
            Permission.status == "Aprobată",
            Permission.start_datetime <= today_end,
            Permission.end_datetime >= today_start,
        ).count()

        # Învoiri zilnice active AZI
        daily_leaves_today_count = DailyLeave.query.filter(
            DailyLeave.student_id.in_(student_ids_managed),
            DailyLeave.status == "Aprobată",
            DailyLeave.leave_date
            == today_localized,  # Folosim data localizată
        ).count()

        # Învoiri weekend active AZI
        weekend_leaves_active_today = 0
        all_wl_gradat = WeekendLeave.query.filter(
            WeekendLeave.student_id.in_(student_ids_managed),
            WeekendLeave.status == "Aprobată",
        ).all()
        for wl in all_wl_gradat:
            for interval in wl.get_intervals():
                if (
                    interval["start"].date() == today_localized
                    or interval["end"].date() == today_localized
                    or (
                        interval["start"].date() < today_localized
                        and interval["end"].date() > today_localized
                    )
                ):
                    weekend_leaves_active_today += 1
                    break

        # Servicii active AZI
        services_today_count = ServiceAssignment.query.filter(
            ServiceAssignment.student_id.in_(student_ids_managed),
            ServiceAssignment.start_datetime <= today_end,
            ServiceAssignment.end_datetime >= today_start,
        ).count()

        total_volunteer_activities = VolunteerActivity.query.filter_by(
            created_by_user_id=current_user.id
        ).count()

        # Data for "Mini Situație Pluton (ACUM)"
        now_for_dashboard = get_localized_now()
        students_managed_by_gradat = Student.query.filter_by(
            created_by_user_id=current_user.id
        ).all()

        # Use _calculate_presence_data for current situation
        # Ensure _calculate_presence_data correctly identifies "present in formation"
        # _calculate_presence_data returns:
        # "efectiv_control", "efectiv_prezent_total", "efectiv_absent_total",
        # "in_formation_count", "on_duty_count", "platoon_graded_duty_count", "absent_students_details"

        # "Prezenți" for mini-dashboard = "in_formation_count" (cei care nu sunt in permisie/serviciu/etc.)
        # "Învoiți" = "efectiv_absent_total" (cei care sunt in permisie, invoire zilnica, invoire weekend)
        # "În Serviciu" = "on_duty_count"
        # "Gradat Pluton Prezent" = "platoon_graded_duty_count" (dacă e considerat separat de "Prezenți")

        current_platoon_situation = {}
        if students_managed_by_gradat:
            current_platoon_situation = _calculate_presence_data(
                students_managed_by_gradat, now_for_dashboard
            )
            # Consolidate "absent" categories for the "Învoiți" count for simplicity on dashboard
            # _calculate_presence_data already gives 'efectiv_absent_total' which is this sum.
        else:  # Default values if no students
            current_platoon_situation = {
                "efectiv_control": 0,
                "in_formation_count": 0,  # Prezenți în formație
                "efectiv_absent_total": 0,  # Total învoiți/absenți motivat
                "on_duty_count": 0,  # În Serviciu
                "platoon_graded_duty_count": 0,  # Gradat pluton (dacă e separat)
            }

        # Serviciile de azi pentru pluton
        todays_services = (
            ServiceAssignment.query.options(
                joinedload(ServiceAssignment.student)
            )
            .filter(
                ServiceAssignment.student_id.in_(student_ids_managed),
                ServiceAssignment.start_datetime <= today_end,
                ServiceAssignment.end_datetime >= today_start,
            )
            .order_by(ServiceAssignment.start_datetime.asc())
            .all()
        )

        return render_template(
            "gradat_dashboard.html",
            student_count=student_count,
            permissions_today_count=permissions_today_count,  # Statistică veche "azi"
            daily_leaves_today_count=daily_leaves_today_count,  # Statistică veche "azi"
            weekend_leaves_today_count=weekend_leaves_active_today,  # Statistică veche "azi"
            services_today_count=services_today_count,  # Statistică veche "azi"
            total_volunteer_activities=total_volunteer_activities,
            current_platoon_situation=current_platoon_situation,
            # Date noi pentru "Mini Situație ACUM"
            sit_total_studenti=current_platoon_situation.get(
                "efectiv_control", 0
            ),
            sit_prezenti_formatie=current_platoon_situation.get(
                "in_formation_count", 0
            ),
            sit_total_invoiti_acum=current_platoon_situation.get(
                "efectiv_absent_total", 0
            ),
            sit_in_serviciu_acum=current_platoon_situation.get(
                "on_duty_count", 0
            ),
            sit_gradat_pluton_prezent_acum=current_platoon_situation.get(
                "platoon_graded_duty_count", 0
            ),
            # Statistici Servicii
            total_services_count=ServiceAssignment.query.filter(
                ServiceAssignment.student_id.in_(student_ids_managed)
            ).count(),
            upcoming_services_count=ServiceAssignment.query.filter(
                ServiceAssignment.student_id.in_(student_ids_managed),
                ServiceAssignment.start_datetime > get_localized_now(),
                ServiceAssignment.start_datetime
                <= get_localized_now() + timedelta(days=7),
            ).count(),
            todays_services=todays_services,
            # Quick Stats
            students_with_high_points=Student.query.filter(
                Student.created_by_user_id == current_user.id,
                Student.volunteer_points > 10,
            ).count(),
            active_scoped_codes=ScopedAccessCode.query.filter_by(
                created_by_user_id=current_user.id, is_active=True
            )
            .filter(ScopedAccessCode.expires_at > get_localized_now())
            .all(),
        )
    elif current_user.role == "comandant_companie":
        return redirect(url_for("company_commander_dashboard"))
    elif current_user.role == "comandant_batalion":
        return redirect(url_for("battalion_commander_dashboard"))

    return render_template("dashboard.html", name=current_user.username)


# --- Admin User Management ---
@app.route(
    "/admin/users/create", methods=["POST"], endpoint="admin_create_user"
)  # Form is on admin_dashboard, so this handles POST
@login_required
def admin_create_user():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("home"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        role = request.form.get("role")
        valid_roles = ["gradat", "comandant_companie", "comandant_batalion"]

        if not username:
            flash("Numele de utilizator este obligatoriu.", "warning")
            return redirect(url_for("admin_dashboard_route"))
        if User.query.filter_by(username=username).first():
            flash(f'Numele de utilizator "{username}" există deja.', "warning")
            return redirect(url_for("admin_dashboard_route"))
        if not role or role not in valid_roles:
            flash("Rolul selectat este invalid.", "warning")
            return redirect(url_for("admin_dashboard_route"))

        unique_code = secrets.token_hex(8)
        while User.query.filter_by(
            unique_code=unique_code
        ).first():  # Ensure unique_code is truly unique
            unique_code = secrets.token_hex(8)

        new_user = User(
            username=username,
            role=role,
            unique_code=unique_code,
            is_first_login=True,
            # password_hash is not set, user will set personal_code
        )
        db.session.add(new_user)
        try:
            db.session.add(new_user)  # Add to session first
            db.session.flush()  # Assign an ID to new_user so we can log it

            log_action(
                "ADMIN_CREATE_USER_SUCCESS",
                target_model_name="User",
                target_id=new_user.id,
                details_after_dict=model_to_dict(new_user),
                description=f"Admin {current_user.username} created user {new_user.username} ({new_user.role}). Unique code: {unique_code}. IP: {request.remote_addr}",
            )
            db.session.commit()
            flash(
                f'Utilizatorul "{username}" ({role}) a fost creat cu succes! Cod unic de autentificare: {unique_code}',
                "success",
            )
        except Exception as e:
            db.session.rollback()
            flash_msg = "Eroare la crearea utilizatorului."
            flash(flash_msg, "danger")
            try:
                log_action(
                    "ADMIN_CREATE_USER_FAIL",
                    description=f"Admin {current_user.username} failed to create user {username} ({role}). Error: {str(e)}. IP: {request.remote_addr}",
                )
                db.session.commit()
            except Exception as log_e:
                app.logger.error(
                    f"CRITICAL: Failed to commit failure log for ADMIN_CREATE_USER_FAIL: {str(log_e)}"
                )
        return redirect(url_for("admin_dashboard_route"))

    return redirect(url_for("admin_dashboard_route"))


@app.route(
    "/admin/users/reset_code/<int:user_id>",
    methods=["POST"],
    endpoint="admin_reset_user_code",
)
@login_required
def admin_reset_user_code(user_id):
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("home"))

    user_to_reset = db.session.get(User, user_id)
    if not user_to_reset:
        flash("Utilizatorul nu a fost găsit.", "danger")
        return redirect(url_for("admin_dashboard_route"))

    if user_to_reset.role == "admin":
        flash(
            "Codul utilizatorului admin nu poate fi resetat prin această metodă.",
            "warning",
        )
        return redirect(url_for("admin_dashboard_route"))

    details_before = model_to_dict(user_to_reset)
    old_unique_code = user_to_reset.unique_code  # For logging

    new_unique_code = secrets.token_hex(8)
    while User.query.filter_by(unique_code=new_unique_code).first():
        new_unique_code = secrets.token_hex(8)

    user_to_reset.unique_code = new_unique_code
    user_to_reset.is_first_login = True
    user_to_reset.password_hash = None
    user_to_reset.personal_code_hash = None

    try:
        details_after = model_to_dict(user_to_reset)
        log_action(
            "ADMIN_RESET_USER_CODE_SUCCESS",
            target_model_name="User",
            target_id=user_to_reset.id,
            details_before_dict=details_before,
            details_after_dict=details_after,
            description=f"Admin {current_user.username} reset code for user {user_to_reset.username}. Old unique code: {old_unique_code}, New unique code: {new_unique_code}. IP: {request.remote_addr}",
        )
        db.session.commit()
        flash(
            f'Codul pentru utilizatorul "{user_to_reset.username}" a fost resetat. Noul cod unic de autentificare: {new_unique_code}',
            "success",
        )
    except Exception as e:
        db.session.rollback()
        flash_msg = f"Eroare la resetarea codului: {str(e)}"
        flash(flash_msg, "danger")
        try:
            log_action(
                "ADMIN_RESET_USER_CODE_FAIL",
                target_model_name="User",
                target_id=user_to_reset.id,
                details_before_dict=details_before,  # Log what it was before the attempt
                description=f"Admin {current_user.username} failed to reset code for user {user_to_reset.username}. Error: {str(e)}. IP: {request.remote_addr}",
            )
            db.session.commit()
        except Exception as log_e:
            app.logger.error(
                f"CRITICAL: Failed to commit failure log for ADMIN_RESET_USER_CODE_FAIL: {str(log_e)}"
            )
    return redirect(url_for("admin_dashboard_route"))


@app.route(
    "/admin/users/delete/<int:user_id>",
    methods=["POST"],
    endpoint="admin_delete_user",
)
@login_required
def admin_delete_user(user_id):
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("home"))

    user_to_delete = db.session.get(User, user_id)
    if not user_to_delete:
        flash("Utilizatorul nu a fost găsit.", "danger")
        # Log attempt to delete non-existent user?
        log_action(
            "ADMIN_DELETE_USER_FAIL_NOT_FOUND",
            target_id=user_id,
            description=f"Admin {current_user.username} attempt to delete non-existent user ID {user_id}. IP: {request.remote_addr}",
        )
        db.session.commit()
        return redirect(url_for("admin_dashboard_route"))

    if user_to_delete.role == "admin":
        flash(
            "Conturile de administrator nu pot fi șterse prin această interfață.",
            "warning",
        )
        log_action(
            "ADMIN_DELETE_USER_FAIL_IS_ADMIN",
            target_model_name="User",
            target_id=user_to_delete.id,
            description=f"Admin {current_user.username} attempt to delete admin user {user_to_delete.username}. IP: {request.remote_addr}",
        )
        db.session.commit()
        return redirect(url_for("admin_dashboard_route"))

    username_deleted_log = user_to_delete.username
    role_deleted_log = user_to_delete.role
    details_before = model_to_dict(user_to_delete)
    num_students_deleted = 0
    num_activities_orphaned = 0
    num_permissions_orphaned = 0
    num_daily_leaves_orphaned = 0
    num_weekend_leaves_orphaned = 0
    num_services_orphaned = 0

    try:
        # Handle VolunteerActivities created by this user
        activities_to_orphan = VolunteerActivity.query.filter_by(
            created_by_user_id=user_to_delete.id
        ).all()
        num_activities_orphaned = len(activities_to_orphan)
        for activity in activities_to_orphan:
            activity.created_by_user_id = None

        # Handle Permissions created by this user
        permissions_to_orphan = Permission.query.filter_by(
            created_by_user_id=user_to_delete.id
        ).all()
        num_permissions_orphaned = len(permissions_to_orphan)
        for perm in permissions_to_orphan:
            perm.created_by_user_id = None

        # Handle DailyLeaves created by this user
        daily_leaves_to_orphan = DailyLeave.query.filter_by(
            created_by_user_id=user_to_delete.id
        ).all()
        num_daily_leaves_orphaned = len(daily_leaves_to_orphan)
        for dl in daily_leaves_to_orphan:
            dl.created_by_user_id = None

        # Handle WeekendLeaves created by this user
        weekend_leaves_to_orphan = WeekendLeave.query.filter_by(
            created_by_user_id=user_to_delete.id
        ).all()
        num_weekend_leaves_orphaned = len(weekend_leaves_to_orphan)
        for wl in weekend_leaves_to_orphan:
            wl.created_by_user_id = None

        # Handle ServiceAssignments created by this user
        services_to_orphan = ServiceAssignment.query.filter_by(
            created_by_user_id=user_to_delete.id
        ).all()
        num_services_orphaned = len(services_to_orphan)
        for service_assignment in services_to_orphan:
            service_assignment.created_by_user_id = None

        # Handle VolunteerSession created by this user
        volunteer_sessions_to_orphan = VolunteerSession.query.filter_by(
            created_by_user_id=user_to_delete.id
        ).all()
        num_volunteer_sessions_orphaned = len(volunteer_sessions_to_orphan)
        for vs in volunteer_sessions_to_orphan:
            vs.created_by_user_id = (
                None  # Or set to a default admin ID if preferred
            )

        if user_to_delete.role == "gradat":
            students_to_delete = Student.query.filter_by(
                created_by_user_id=user_to_delete.id
            ).all()
            num_students_deleted = len(students_to_delete)
            for student in students_to_delete:
                db.session.delete(
                    student
                )  # Cascading deletes for student's permissions, leaves etc. should be handled by DB/ORM settings
            if num_students_deleted > 0:
                flash(
                    f"Toți studenții ({num_students_deleted}) și datele asociate direct LOR pentru gradatul {username_deleted_log} au fost șterse.",
                    "info",
                )

        db.session.delete(user_to_delete)

        log_description = (
            f"Admin {current_user.username} deleted user {username_deleted_log} ({role_deleted_log}). "
            f"{num_students_deleted} students (and their direct data) also deleted if user was gradat. "
            f"Orphaned records: Activities({num_activities_orphaned}), Permissions({num_permissions_orphaned}), "
            f"DailyLeaves({num_daily_leaves_orphaned}), WeekendLeaves({num_weekend_leaves_orphaned}), Services({num_services_orphaned}), "
            f"VolunteerSessions({num_volunteer_sessions_orphaned}). "
            f"IP: {request.remote_addr}"
        )

        log_action(
            "ADMIN_DELETE_USER_SUCCESS",
            target_model_name="User",
            target_id=user_id,
            details_before_dict=details_before,
            description=log_description,
        )
        db.session.commit()
        flash(
            f'Utilizatorul "{username_deleted_log}" a fost șters cu succes. Înregistrările create de acesta au fost disociate (păstrate fără creator).',
            "success",
        )

    except Exception as e:
        db.session.rollback()
        flash_msg = f"Eroare la ștergerea utilizatorului {username_deleted_log}: {str(e)}"
        flash(flash_msg, "danger")
        try:
            log_action(
                "ADMIN_DELETE_USER_FAIL",
                target_model_name="User",
                target_id=user_id,
                details_before_dict=details_before,  # Log what it was before the attempt
                description=f"Admin {current_user.username} failed to delete user {username_deleted_log}. Error: {str(e)}. IP: {request.remote_addr}",
            )
            db.session.commit()
        except Exception as log_e:
            app.logger.error(
                f"CRITICAL: Failed to commit failure log for ADMIN_DELETE_USER_FAIL: {str(log_e)}"
            )

    return redirect(url_for("admin_dashboard_route"))


@app.route(
    "/admin/user/edit/<int:user_id>",
    methods=["GET", "POST"],
    endpoint="admin_edit_user",
)
@login_required
def admin_edit_user(user_id):
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("home"))

    user_to_edit = db.session.get(User, user_id)
    if not user_to_edit:
        flash("Utilizatorul nu a fost găsit.", "danger")
        return redirect(url_for("admin_dashboard_route"))

    if (
        user_to_edit.role == "admin"
    ):  # Prevent admin from editing their own or other admin usernames via this form
        flash(
            "Numele de utilizator pentru conturile de admin nu poate fi modificat prin această interfață.",
            "warning",
        )
        return redirect(url_for("admin_dashboard_route"))

    original_username = user_to_edit.username  # For logging and comparison
    details_before_edit = model_to_dict(user_to_edit)

    if request.method == "POST":
        new_username = request.form.get("new_username", "").strip()

        if not new_username:
            flash("Noul nume de utilizator nu poate fi gol.", "warning")
            return render_template(
                "admin_edit_user.html", user_to_edit=user_to_edit
            )  # Re-render with error

        if new_username != original_username:
            existing_user = User.query.filter(
                User.username == new_username
            ).first()
            if existing_user:
                flash(
                    f"Numele de utilizator '{new_username}' există deja. Te rugăm alege altul.",
                    "warning",
                )
                return render_template(
                    "admin_edit_user.html", user_to_edit=user_to_edit
                )

            user_to_edit.username = new_username

            try:
                details_after_edit = model_to_dict(user_to_edit)
                log_action(
                    "ADMIN_UPDATE_USERNAME_SUCCESS",
                    target_model_name="User",
                    target_id=user_to_edit.id,
                    details_before_dict=details_before_edit,
                    details_after_dict=details_after_edit,
                    description=f"Admin {current_user.username} changed username for user ID {user_to_edit.id} from '{original_username}' to '{new_username}'.",
                )
                db.session.commit()
                flash(
                    f"Numele de utilizator pentru '{original_username}' a fost schimbat în '{new_username}'.",
                    "success",
                )

                if user_to_edit.role in [
                    "comandant_companie",
                    "comandant_batalion",
                ]:
                    # Check if new username still allows unit ID derivation
                    derived_id_after_change = _get_commander_unit_id(
                        new_username,
                        (
                            "CmdC"
                            if user_to_edit.role == "comandant_companie"
                            else "CmdB"
                        ),
                    )
                    if not derived_id_after_change:
                        flash(
                            "Atenție: Noul nume de utilizator '{new_username}' pentru comandantul {new_username} nu mai corespunde modelului pentru extragerea automată a ID-ului unității. Funcționalitatea specifică rolului său ar putea fi afectată.",
                            "warning",
                        )
                    else:
                        # Optionally, confirm the new derived ID if it's part of the username structure
                        flash(
                            "Verificare ID unitate pentru '{new_username}': ID derivat este '{derived_id_after_change}'. Asigurați-vă că este corect.",
                            "info",
                        )

                return redirect(url_for("admin_dashboard_route"))
            except Exception as e:
                db.session.rollback()
                user_to_edit.username = (
                    original_username  # Revert optimistic change
                )
                flash(
                    f"Eroare la salvarea noului nume de utilizator: {str(e)}",
                    "danger",
                )
                try:
                    log_action(
                        "ADMIN_UPDATE_USERNAME_FAIL",
                        target_model_name="User",
                        target_id=user_to_edit.id,
                        details_before_dict=details_before_edit,
                        description=f"Admin {current_user.username} failed to change username for user ID {user_to_edit.id} from '{original_username}' to '{new_username}'. Error: {str(e)}",
                    )
                    db.session.commit()
                except Exception as log_e:
                    app.logger.error(
                        f"CRITICAL: Failed to commit failure log for ADMIN_UPDATE_USERNAME_FAIL: {str(log_e)}"
                    )
        else:
            flash(
                "Numele de utilizator nu a fost schimbat.", "info"
            )  # No change made
            return redirect(url_for("admin_dashboard_route"))

    # For GET request
    return render_template("admin_edit_user.html", user_to_edit=user_to_edit)


@app.route("/admin/student/edit/<int:student_id>", methods=["GET", "POST"])
@login_required
def admin_edit_student(student_id):
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))  # Or home

    student_to_edit = db.session.get(Student, student_id)
    if not student_to_edit:
        flash("Studentul nu a fost găsit.", "danger")
        return redirect(
            url_for("list_students", _anchor="admin_view_params_if_any")
        )  # Assuming admin view of list_students

    details_before_edit = model_to_dict(student_to_edit)
    original_creator_username = (
        student_to_edit.creator.username if student_to_edit.creator else "N/A"
    )

    if request.method == "POST":
        form = request.form
        # Capture values from form
        student_to_edit.nume = form.get("nume", "").strip()
        student_to_edit.prenume = form.get("prenume", "").strip()
        student_to_edit.grad_militar = form.get("grad_militar", "").strip()
        student_to_edit.pluton = form.get("pluton", "").strip()
        student_to_edit.companie = form.get("companie", "").strip()
        student_to_edit.batalion = form.get("batalion", "").strip()
        student_to_edit.gender = form.get("gender")
        student_to_edit.is_platoon_graded_duty = (
            "is_platoon_graded_duty" in request.form
        )
        student_to_edit.assigned_graded_platoon = (
            form.get("assigned_graded_platoon", "").strip() or None
        )
        student_to_edit.is_smt = "is_smt" in request.form
        student_to_edit.exemption_details = (
            form.get("exemption_details", "").strip() or None
        )
        new_id_unic = form.get("id_unic_student", "").strip() or None

        # Potentially allow changing created_by_user_id by admin? For now, no.
        # student_to_edit.created_by_user_id = form.get('created_by_user_id', student_to_edit.created_by_user_id)

        if not all(
            [
                student_to_edit.nume,
                student_to_edit.prenume,
                student_to_edit.grad_militar,
                student_to_edit.companie,
                student_to_edit.batalion,
                student_to_edit.gender,
            ]
        ):
            flash(
                "Toate câmpurile marcate cu * (cu excepția plutonului pentru gradați) sunt obligatorii.",
                "warning",
            )
            # Pass original creator info for display even on error
            return render_template(
                "admin_edit_student.html",
                form_title=f"Editare Student (Admin): {details_before_edit.get('grad_militar','')} {details_before_edit.get('nume','')}",
                student=student_to_edit,
                genders=GENDERS,
                original_creator_username=original_creator_username,
                form_data=request.form,
            )

        if student_to_edit.gender not in GENDERS:
            flash("Valoare invalidă pentru gen.", "warning")
            return render_template(
                "admin_edit_student.html",
                form_title=f"Editare Student (Admin): {details_before_edit.get('grad_militar','')} {details_before_edit.get('nume','')}",
                student=student_to_edit,
                genders=GENDERS,
                original_creator_username=original_creator_username,
                form_data=request.form,
            )

        # Check for ID Unic conflict if changed
        if new_id_unic and new_id_unic != details_before_edit.get(
            "id_unic_student"
        ):  # Check against original value from details_before
            existing_student_with_id = Student.query.filter(
                Student.id_unic_student == new_id_unic,
                Student.id != student_to_edit.id,
            ).first()
            if existing_student_with_id:
                flash(
                    f"Alt student (ID: {existing_student_with_id.id}) cu ID unic '{new_id_unic}' există deja.",
                    "warning",
                )
                return render_template(
                    "admin_edit_student.html",
                    form_title=f"Editare Student (Admin): {details_before_edit.get('grad_militar','')} {details_before_edit.get('nume','')}",
                    student=student_to_edit,
                    genders=GENDERS,
                    original_creator_username=original_creator_username,
                    form_data=request.form,
                )
        student_to_edit.id_unic_student = new_id_unic

        try:
            details_after_edit = model_to_dict(student_to_edit)
            log_action(
                "ADMIN_UPDATE_STUDENT_SUCCESS",
                target_model_name="Student",
                target_id=student_to_edit.id,
                details_before_dict=details_before_edit,
                details_after_dict=details_after_edit,
                description=f"Admin {current_user.username} updated student {student_to_edit.grad_militar} {student_to_edit.nume} (ID: {student_to_edit.id}).",
            )
            db.session.commit()
            flash(
                f"Studentul {student_to_edit.grad_militar} {student_to_edit.nume} {student_to_edit.prenume} a fost actualizat de admin!",
                "success",
            )
            return redirect(
                url_for("list_students")
            )  # Redirects to admin view due to session/role context
        except Exception as e:
            db.session.rollback()
            flash_msg = (
                f"Eroare la actualizarea studentului de către admin: {str(e)}"
            )
            flash(flash_msg, "danger")
            try:
                log_action(
                    "ADMIN_UPDATE_STUDENT_FAIL",
                    target_model_name="Student",
                    target_id=student_to_edit.id,
                    details_before_dict=details_before_edit,
                    description=f"Admin {current_user.username} failed to update student ID {student_to_edit.id}. Error: {str(e)}",
                )
                db.session.commit()
            except Exception as log_e:
                app.logger.error(
                    f"CRITICAL: Failed to commit failure log for ADMIN_UPDATE_STUDENT_FAIL: {str(log_e)}"
                )
            return render_template(
                "admin_edit_student.html",
                form_title=f"Editare Student (Admin): {details_before_edit.get('grad_militar','')} {details_before_edit.get('nume','')}",
                student=student_to_edit,
                genders=GENDERS,
                original_creator_username=original_creator_username,
                form_data=request.form,
            )

    # For GET request, populate form with student's current data
    form_data_for_get = model_to_dict(student_to_edit)
    # Ensure boolean is_platoon_graded_duty is correctly passed for checkbox
    form_data_for_get["is_platoon_graded_duty"] = (
        student_to_edit.is_platoon_graded_duty
    )

    return render_template(
        "admin_edit_student.html",
        form_title=f"Editare Student (Admin): {student_to_edit.grad_militar} {student_to_edit.nume} {student_to_edit.prenume}",
        student=student_to_edit,
        genders=GENDERS,
        original_creator_username=original_creator_username,
        form_data=form_data_for_get,
    )


# --- Admin List View for Permissions ---
@app.route("/admin/permissions", endpoint="admin_list_permissions")
@login_required
def admin_list_permissions():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    page = request.args.get("page", 1, type=int)
    per_page = 20

    query = Permission.query.options(
        joinedload(Permission.student).joinedload(
            Student.creator
        ),  # Creator of student
        joinedload(Permission.creator),  # Creator of permission
    ).order_by(Permission.start_datetime.desc())

    # Basic Filtering (can be expanded)
    search_student_name = request.args.get("search_student_name", "").strip()
    filter_status = request.args.get("filter_status", "").strip()
    # TODO: Add date range filters if needed

    if search_student_name:
        # Search in student's nume or prenume
        search_pattern = f"%{unidecode(search_student_name.lower())}%"
        query = query.join(Permission.student).filter(
            or_(
                func.lower(unidecode(Student.nume)).like(search_pattern),
                func.lower(unidecode(Student.prenume)).like(search_pattern),
            )
        )

    if filter_status:
        query = query.filter(Permission.status == filter_status)

    permissions_pagination = query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    # For status filter dropdown
    statuses = sorted(
        list(
            set(
                item[0]
                for item in db.session.query(Permission.status)
                .distinct()
                .all()
                if item[0]
            )
        )
    )

    return render_template(
        "admin_list_permissions.html",
        permissions_pagination=permissions_pagination,
        search_student_name=search_student_name,
        filter_status=filter_status,
        statuses=statuses,
        title="Listă Generală Permisii (Admin)",
    )


# --- Admin List View for Daily Leaves ---
@app.route("/admin/daily_leaves", endpoint="admin_list_daily_leaves")
@login_required
def admin_list_daily_leaves():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    page = request.args.get("page", 1, type=int)
    per_page = 20

    query = DailyLeave.query.options(
        joinedload(DailyLeave.student), joinedload(DailyLeave.creator)
    ).order_by(DailyLeave.leave_date.desc(), DailyLeave.start_time.desc())

    # Filtering
    search_student_name = request.args.get("search_student_name", "").strip()
    filter_status = request.args.get("filter_status", "").strip()
    filter_date = request.args.get("filter_date", "").strip()

    if search_student_name:
        search_pattern = f"%{unidecode(search_student_name.lower())}%"
        query = query.join(DailyLeave.student).filter(
            or_(
                func.lower(unidecode(Student.nume)).like(search_pattern),
                func.lower(unidecode(Student.prenume)).like(search_pattern),
            )
        )

    if filter_status:
        query = query.filter(DailyLeave.status == filter_status)

    if filter_date:
        try:
            date_obj = datetime.strptime(filter_date, "%Y-%m-%d").date()
            query = query.filter(DailyLeave.leave_date == date_obj)
        except ValueError:
            flash(
                "Format dată invalid pentru filtrare. Folosiți YYYY-MM-DD.",
                "warning",
            )

    daily_leaves_pagination = query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    statuses = sorted(
        list(
            set(
                item[0]
                for item in db.session.query(DailyLeave.status)
                .distinct()
                .all()
                if item[0]
            )
        )
    )

    return render_template(
        "admin_list_daily_leaves.html",
        daily_leaves_pagination=daily_leaves_pagination,
        search_student_name=search_student_name,
        filter_status=filter_status,
        filter_date=filter_date,
        statuses=statuses,
        title="Listă Generală Învoiri Zilnice (Admin)",
    )


# --- Admin List View for Weekend Leaves ---
@app.route("/admin/weekend_leaves", endpoint="admin_list_weekend_leaves")
@login_required
def admin_list_weekend_leaves():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    page = request.args.get("page", 1, type=int)
    per_page = 20

    query = WeekendLeave.query.options(
        joinedload(WeekendLeave.student), joinedload(WeekendLeave.creator)
    ).order_by(WeekendLeave.weekend_start_date.desc())

    # Filtering
    search_student_name = request.args.get("search_student_name", "").strip()
    filter_status = request.args.get("filter_status", "").strip()
    filter_weekend_start_date = request.args.get(
        "filter_weekend_start_date", ""
    ).strip()  # This is the Friday

    if search_student_name:
        search_pattern = f"%{unidecode(search_student_name.lower())}%"
        query = query.join(WeekendLeave.student).filter(
            or_(
                func.lower(unidecode(Student.nume)).like(search_pattern),
                func.lower(unidecode(Student.prenume)).like(search_pattern),
            )
        )

    if filter_status:
        query = query.filter(WeekendLeave.status == filter_status)

    if filter_weekend_start_date:
        try:
            date_obj = datetime.strptime(
                filter_weekend_start_date, "%Y-%m-%d"
            ).date()
            query = query.filter(WeekendLeave.weekend_start_date == date_obj)
        except ValueError:
            flash(
                "Format dată invalid pentru filtrare (Vineri weekend). Folosiți YYYY-MM-DD.",
                "warning",
            )

    weekend_leaves_pagination = query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    statuses = sorted(
        list(
            set(
                item[0]
                for item in db.session.query(WeekendLeave.status)
                .distinct()
                .all()
                if item[0]
            )
        )
    )

    return render_template(
        "admin_list_weekend_leaves.html",
        weekend_leaves_pagination=weekend_leaves_pagination,
        search_student_name=search_student_name,
        filter_status=filter_status,
        filter_weekend_start_date=filter_weekend_start_date,
        statuses=statuses,
        title="Listă Generală Învoiri Weekend (Admin)",
    )


# --- Admin List View for Service Assignments ---
@app.route("/admin/services", endpoint="admin_list_services")
@login_required
def admin_list_services():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    page = request.args.get("page", 1, type=int)
    per_page = 20

    query = ServiceAssignment.query.options(
        joinedload(ServiceAssignment.student),
        joinedload(ServiceAssignment.creator),
    ).order_by(ServiceAssignment.start_datetime.desc())

    # Filtering
    search_student_name = request.args.get("search_student_name", "").strip()
    filter_service_type = request.args.get("filter_service_type", "").strip()
    filter_service_date = request.args.get("filter_service_date", "").strip()

    if search_student_name:
        search_pattern = f"%{unidecode(search_student_name.lower())}%"
        query = query.join(ServiceAssignment.student).filter(
            or_(
                func.lower(unidecode(Student.nume)).like(search_pattern),
                func.lower(unidecode(Student.prenume)).like(search_pattern),
            )
        )

    if filter_service_type:
        query = query.filter(
            ServiceAssignment.service_type == filter_service_type
        )

    if filter_service_date:
        try:
            date_obj = datetime.strptime(
                filter_service_date, "%Y-%m-%d"
            ).date()
            # ServiceAssignment has start_datetime and end_datetime, and service_date
            # Filter by service_date for simplicity, or choose start/end date part.
            query = query.filter(ServiceAssignment.service_date == date_obj)
        except ValueError:
            flash(
                "Format dată invalid pentru filtrare (Data serviciu). Folosiți YYYY-MM-DD.",
                "warning",
            )

    services_pagination = query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    # For service type filter dropdown
    service_types_for_filter = sorted(
        list(
            set(
                item[0]
                for item in db.session.query(ServiceAssignment.service_type)
                .distinct()
                .all()
                if item[0]
            )
        )
    )

    return render_template(
        "admin_list_services.html",
        services_pagination=services_pagination,
        search_student_name=search_student_name,
        filter_service_type=filter_service_type,
        filter_service_date=filter_service_date,
        service_types_for_filter=service_types_for_filter,  # Pass this to template
        title="Listă Generală Servicii (Admin)",
    )


# --- Helper function to parse commander's unit ID ---
def _get_commander_unit_id(username, role_prefix):
    # role_prefix could be "CmdC" or "CmdB"
    # We are looking for patterns like "CmdC1", "Vasile_CmdC1", "CmdB12", "Popescu_CmdB123"
    # The ID is the numeric part immediately following the role_prefix in the username string.

    # Escape role_prefix in case it contains special regex characters (though "CmdC" and "CmdB" don't)
    # Then look for one or more digits (\d+) captured in a group.
    match = re.search(f"{re.escape(role_prefix)}(\\d+)", username)
    if match:
        unit_id_part = match.group(
            1
        )  # Get the captured digits (e.g., "1", "123")
        # The \d+ ensures unit_id_part is not empty and contains only digits.
        return unit_id_part
    return None  # Return None if prefix followed by digits is not found


# --- Helper function to determine standard roll call time ---
def get_standard_roll_call_datetime(for_date=None):
    target_date = (
        for_date if for_date else get_localized_now().date()
    )  # Folosim data localizată
    weekday = target_date.weekday()  # Monday is 0 and Sunday is 6

    if 0 <= weekday <= 3:  # Monday to Thursday
        roll_call_time = time(20, 0)
    else:  # Friday to Sunday
        roll_call_time = time(22, 0)

    return datetime.combine(target_date, roll_call_time)


# --- Helper function to calculate presence data for a list of students (Optimized) ---
def _calculate_presence_data(student_list, check_datetime):
    if not student_list:
        return {
            "efectiv_control": 0,
            "efectiv_prezent_total": 0,
            "efectiv_absent_total": 0,
            "in_formation_count": 0,
            "on_duty_count": 0,
            "platoon_graded_duty_count": 0,
            "all_present_details": [],
            "all_absent_details": [],
            "in_formation_students_details": [],
            "absent_students_details": [],
            "smt_students_details": [],
            "exempt_other_students_details": [],
            "present_exempt_not_in_formation_details": [],
            "present_exempt_not_in_formation_count": 0,
        }

    # --- Step 1: Bulk Data Fetch ---
    student_ids = [s.id for s in student_list]
    now_naive = check_datetime.replace(tzinfo=None)

    active_services = ServiceAssignment.query.filter(
        ServiceAssignment.student_id.in_(student_ids),
        ServiceAssignment.start_datetime <= now_naive,
        ServiceAssignment.end_datetime >= now_naive,
    ).all()
    active_permissions = Permission.query.filter(
        Permission.student_id.in_(student_ids),
        Permission.status == "Aprobată",
        Permission.start_datetime <= now_naive,
        Permission.end_datetime >= now_naive,
    ).all()
    all_daily_leaves = DailyLeave.query.filter(
        DailyLeave.student_id.in_(student_ids), DailyLeave.status == "Aprobată"
    ).all()
    all_weekend_leaves = WeekendLeave.query.filter(
        WeekendLeave.student_id.in_(student_ids),
        WeekendLeave.status == "Aprobată",
    ).all()
    # Volunteers active on the specific date (treated as absent on that day)
    vol_entries = (
        db.session.query(ActivityParticipant, VolunteerActivity)
        .join(
            VolunteerActivity,
            ActivityParticipant.activity_id == VolunteerActivity.id,
        )
        .filter(
            ActivityParticipant.student_id.in_(student_ids),
            VolunteerActivity.activity_date == check_datetime.date(),
        )
        .all()
    )

    # --- Step 2: Process into Fast-Lookup Maps ---
    active_services_map = {sa.student_id: sa for sa in active_services}
    active_permissions_map = {p.student_id: p for p in active_permissions}
    active_daily_leaves_map = {
        dl.student_id: dl for dl in all_daily_leaves if dl.is_active
    }
    active_weekend_leaves_map = {}
    for wl in all_weekend_leaves:
        if wl.is_any_interval_active_now:
            active_interval = next(
                (
                    interval
                    for interval in wl.get_intervals()
                    if interval["start"] <= check_datetime <= interval["end"]
                ),
                None,
            )
            if active_interval:
                active_weekend_leaves_map[wl.student_id] = {
                    "leave": wl,
                    "interval": active_interval,
                }
    active_volunteer_map = {}
    for ap, act in vol_entries:
        # If multiple activities in a day, keep the first name encountered
        active_volunteer_map.setdefault(ap.student_id, act.name)

    # --- Step 3: Categorize Students ---
    efectiv_control = len(student_list)
    present_in_formation = []
    present_on_duty = []
    present_graded_staff = (
        []
    )  # Combined list for platoon leaders, assigned to other platoons, and company/battalion staff
    present_exempt_not_in_formation = (
        []
    )  # Temporary exemptions: present physically, not in formation
    all_absent_details = []
    general_absent_details = (
        []
    )  # Leaves/permits/weekend/daily/volunteer (excludes SMT and other exemptions now treated as present)
    smt_students_details = []
    exempt_other_students_details = (
        []
    )  # Kept for backward compatibility if templates use it anywhere

    for s in student_list:
        student_display_name = f"{s.grad_militar} {s.nume} {s.prenume}"
        status_found = False

        # SMT first (still considered absent motivated)
        if getattr(s, "is_smt", False):
            smt_students_details.append(f"{student_display_name} - SMT")
            all_absent_details.append(f"{student_display_name} - SMT")
            status_found = True
        else:
            # Temporary exemption (scutire temporară) is independent of presence.
            # We record it for reporting, but DO NOT force a presence category here.
            ex_det = (getattr(s, "exemption_details", "") or "").strip()
            ex_det_lower = ex_det.lower()
            if ex_det and ex_det_lower not in {
                "none",
                "null",
                "-",
                "n/a",
                "na",
            }:
                label = f"{student_display_name} - Scutire temporară"
                if ex_det:
                    label += f": {ex_det}"
                present_exempt_not_in_formation.append(label)
                # Backward-compatible bucket for any other exemption aggregations
                exempt_other_students_details.append(
                    f"{student_display_name} - Scutire temporară: {ex_det}"
                    if ex_det
                    else f"{student_display_name} - Scutire temporară"
                )
                # Do not set status_found here.

        # Service/leave/volunteer checks
        if not status_found and s.id in active_services_map:
            present_on_duty.append(
                f"{student_display_name} - Serviciu ({active_services_map[s.id].service_type})"
            )
            status_found = True
        elif not status_found and s.id in active_permissions_map:
            general_absent_details.append(f"{student_display_name} - Permisie")
            all_absent_details.append(f"{student_display_name} - Permisie")
            status_found = True
        elif not status_found and s.id in active_weekend_leaves_map:
            wl_data = active_weekend_leaves_map[s.id]
            general_absent_details.append(
                f"{student_display_name} - Învoire Weekend ({wl_data['interval']['day_name']})"
            )
            all_absent_details.append(
                f"{student_display_name} - Învoire Weekend ({wl_data['interval']['day_name']})"
            )
            status_found = True
        elif not status_found and s.id in active_daily_leaves_map:
            general_absent_details.append(
                f"{student_display_name} - Învoire Zilnică ({active_daily_leaves_map[s.id].leave_type_display})"
            )
            all_absent_details.append(
                f"{student_display_name} - Învoire Zilnică ({active_daily_leaves_map[s.id].leave_type_display})"
            )
            status_found = True
        elif not status_found and s.id in active_volunteer_map:
            general_absent_details.append(
                f"{student_display_name} - Voluntariat ({active_volunteer_map[s.id]})"
            )
            all_absent_details.append(
                f"{student_display_name} - Voluntariat ({active_volunteer_map[s.id]})"
            )
            status_found = True

        if not status_found:
            # Student is present at the unit; determine their presence category.
            # is_platoon_graded_duty is for platoon leaders (own platoon)
            # assigned_graded_platoon indicates assigned to lead another platoon or unit (e.g., '99' for battalion/company level)
            # pluton == '0' is used for company/battalion staff
            assigned_other_pl = (
                getattr(s, "assigned_graded_platoon", None) or ""
            ).strip()
            if assigned_other_pl:
                if assigned_other_pl == "99":
                    present_graded_staff.append(
                        f"{student_display_name} - Gradat (Comp./Bat.)"
                    )
                else:
                    present_graded_staff.append(
                        f"{student_display_name} - Gradat la Pl.{assigned_other_pl}"
                    )
            elif getattr(s, "is_platoon_graded_duty", False):
                present_graded_staff.append(
                    f"{student_display_name} - Gradat Pluton"
                )
            elif str(getattr(s, "pluton", "")).strip() == "0":
                present_graded_staff.append(
                    f"{student_display_name} - Gradat (Comp./Bat.)"
                )
            else:
                present_in_formation.append(student_display_name)

    # --- Step 4: Compile Results ---
    in_formation_count = len(present_in_formation)
    on_duty_count = len(present_on_duty)
    graded_staff_count = len(present_graded_staff)
    present_exempt_count = len(present_exempt_not_in_formation)
    absent_total_count = len(all_absent_details)

    # Exemptions are independent of presence; do not add them to totals.
    efectiv_prezent_total = (
        in_formation_count + on_duty_count + graded_staff_count
    )
    efectiv_absent_total = absent_total_count

    # Consistency check
    if efectiv_control != efectiv_prezent_total + efectiv_absent_total:
        app.logger.warning(
            f"Presence data discrepancy: EC({efectiv_control}) != EP({efectiv_prezent_total}) + EA({efectiv_absent_total})"
        )

    return {
        "efectiv_control": efectiv_control,
        "efectiv_prezent_total": efectiv_prezent_total,
        "efectiv_absent_total": efectiv_absent_total,
        "in_formation_count": in_formation_count,
        "on_duty_count": on_duty_count,
        "platoon_graded_duty_count": graded_staff_count,
        # Details
        "all_present_details": sorted(present_in_formation),
        "in_formation_students_details": sorted(
            present_in_formation
        ),  # Backward compatible key for templates
        "on_duty_students_details": sorted(present_on_duty),
        "platoon_graded_duty_students_details": sorted(present_graded_staff),
        "present_exempt_not_in_formation_details": sorted(
            present_exempt_not_in_formation
        ),
        "present_exempt_not_in_formation_count": present_exempt_count,
        "all_absent_details": sorted(all_absent_details),
        "absent_students_details": sorted(
            general_absent_details
        ),  # Excludes SMT/other exemptions
        "smt_students_details": sorted(smt_students_details),
        "exempt_other_students_details": sorted(exempt_other_students_details),
        # Convenience counters
        "smt_count": len(smt_students_details),
        "exempt_other_count": len(exempt_other_students_details),
    }


# --- Commander Dashboards ---
@app.route("/dashboard/company")
@login_required
def company_commander_dashboard():
    if current_user.role != "comandant_companie":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    company_id_str = _get_commander_unit_id(current_user.username, "CmdC")
    if not company_id_str:
        # Avoid redirect loop; render dashboard with empty datasets and guidance
        flash(
            "ID-ul companiei nu a putut fi determinat din numele de utilizator. Numele trebuie să conțină modelul „CmdC<NUMĂR>”, ex: „Popescu_CmdC1”.",
            "warning",
        )
        roll_call_datetime = get_standard_roll_call_datetime()
        roll_call_time_str = roll_call_datetime.strftime("%d %B %Y, %H:%M")
        empty_presence = {
            "efectiv_control": 0,
            "efectiv_prezent_total": 0,
            "efectiv_absent_total": 0,
            "in_formation_count": 0,
            "on_duty_count": 0,
            "platoon_graded_duty_count": 0,
            "in_formation_students_details": [],
            "on_duty_students_details": [],
            "platoon_graded_duty_students_details": [],
            "all_absent_details": [],
        }
        return render_template(
            "company_commander_dashboard.html",
            company_id="N/A",
            roll_call_time_str=roll_call_time_str,
            total_company_presence=empty_presence,
            platoons_data={},
            platoons_data_now={},
            permissions_today_count=0,
            daily_leaves_today_company=0,
            weekend_leaves_today_company=0,
            total_leaves_today_count=0,
            services_today_count=0,
            total_students_company=0,
            permissions_active_now_company=0,
            daily_leaves_active_now_company=0,
            weekend_leaves_active_now_company=0,
            total_on_leave_now_company=0,
            services_active_now_company=0,
            current_time_for_display=get_localized_now(),
            active_public_codes=[],
            services_last_7_days_count=0,
            services_today_breakdown={},
            todays_services=[],
        )

    roll_call_datetime = get_standard_roll_call_datetime()
    roll_call_time_str = roll_call_datetime.strftime("%d %B %Y, %H:%M")

    # Fetch all students in this specific company
    students_in_company_all = Student.query.filter_by(
        companie=company_id_str
    ).all()
    student_ids_in_company = [s.id for s in students_in_company_all]

    total_company_presence_roll_call = _calculate_presence_data(
        students_in_company_all, roll_call_datetime
    )  # For existing roll call report

    # New stats for "today"
    today_localized_company = get_localized_now().date()
    today_start = datetime.combine(today_localized_company, time.min)
    today_end = datetime.combine(today_localized_company, time.max)

    permissions_today_company_count = Permission.query.filter(
        Permission.student_id.in_(student_ids_in_company),
        Permission.status == "Aprobată",
        Permission.start_datetime <= today_end,
        Permission.end_datetime >= today_start,
    ).count()

    daily_leaves_today_company_count = DailyLeave.query.filter(
        DailyLeave.student_id.in_(student_ids_in_company),
        DailyLeave.status == "Aprobată",
        DailyLeave.leave_date == today_localized_company,
    ).count()

    weekend_leaves_today_company_count = 0
    all_wl_company_for_today_stats = WeekendLeave.query.filter(
        WeekendLeave.student_id.in_(student_ids_in_company),
        WeekendLeave.status == "Aprobată",
    ).all()
    for wl_today in all_wl_company_for_today_stats:
        for interval_today in wl_today.get_intervals():
            if (
                interval_today["start"].date() == today_localized_company
                or interval_today["end"].date() == today_localized_company
                or (
                    interval_today["start"].date() < today_localized_company
                    and interval_today["end"].date() > today_localized_company
                )
            ):
                weekend_leaves_today_company_count += 1
                break

    total_leaves_today_company_count = (
        daily_leaves_today_company_count + weekend_leaves_today_company_count
    )

    services_today_company_count = ServiceAssignment.query.filter(
        ServiceAssignment.student_id.in_(student_ids_in_company),
        ServiceAssignment.start_datetime <= today_end,
        ServiceAssignment.end_datetime >= today_start,
    ).count()

    # Stats for "NOW"
    now_localized = get_localized_now()
    permissions_active_now_company = Permission.query.filter(
        Permission.student_id.in_(student_ids_in_company),
        Permission.status == "Aprobată",
        Permission.start_datetime <= now_localized,
        Permission.end_datetime >= now_localized,
    ).count()

    daily_leaves_active_now_company = 0
    all_dl_company = DailyLeave.query.filter(
        DailyLeave.student_id.in_(student_ids_in_company),
        DailyLeave.status == "Aprobată",
    ).all()
    for dl_now in all_dl_company:
        if dl_now.is_active:  # is_active already uses get_localized_now()
            daily_leaves_active_now_company += 1

    weekend_leaves_active_now_company = 0
    all_wl_company_for_now_stats = WeekendLeave.query.filter(
        WeekendLeave.student_id.in_(student_ids_in_company),
        WeekendLeave.status == "Aprobată",
    ).all()
    for wl_now in all_wl_company_for_now_stats:
        if (
            wl_now.is_any_interval_active_now
        ):  # is_any_interval_active_now uses get_localized_now()
            weekend_leaves_active_now_company += 1

    total_on_leave_now_company = (
        permissions_active_now_company
        + daily_leaves_active_now_company
        + weekend_leaves_active_now_company
    )

    services_active_now_company = ServiceAssignment.query.filter(
        ServiceAssignment.student_id.in_(student_ids_in_company),
        ServiceAssignment.start_datetime <= now_localized,
        ServiceAssignment.end_datetime >= now_localized,
    ).count()

    # New service stats for company
    seven_days_ago = get_localized_now().date() - timedelta(days=7)
    services_last_7_days_count = ServiceAssignment.query.filter(
        ServiceAssignment.student_id.in_(student_ids_in_company),
        ServiceAssignment.service_date > seven_days_ago,
        ServiceAssignment.service_date <= get_localized_now().date(),
    ).count()

    services_today_breakdown = dict(
        db.session.query(
            ServiceAssignment.service_type, func.count(ServiceAssignment.id)
        )
        .filter(
            ServiceAssignment.student_id.in_(student_ids_in_company),
            ServiceAssignment.start_datetime <= today_end,
            ServiceAssignment.end_datetime >= today_start,
        )
        .group_by(ServiceAssignment.service_type)
        .all()
    )

    platoons_data_roll_call = {}
    platoons_data_now = {}  # New dictionary for current data
    platoons_in_company = sorted(
        list(set(s.pluton for s in students_in_company_all if s.pluton))
    )

    students_without_platoon = [
        s for s in students_in_company_all if not s.pluton
    ]
    if students_without_platoon:
        platoon_name = "Gradati / Personal neîncadrat"
        platoons_data_roll_call[platoon_name] = _calculate_presence_data(
            students_without_platoon, roll_call_datetime
        )
        platoons_data_now[platoon_name] = _calculate_presence_data(
            students_without_platoon, now_localized
        )

    for pluton_id_str in platoons_in_company:
        students_in_pluton = [
            s for s in students_in_company_all if s.pluton == pluton_id_str
        ]
        platoon_name = f"Plutonul {pluton_id_str}"
        platoons_data_roll_call[platoon_name] = _calculate_presence_data(
            students_in_pluton, roll_call_datetime
        )
        platoons_data_now[platoon_name] = _calculate_presence_data(
            students_in_pluton, now_localized
        )  # Calculate for now

    active_public_codes = (
        PublicViewCode.query.filter_by(
            created_by_user_id=current_user.id,
            is_active=True,
            scope_type="company",
            scope_id=company_id_str,
        )
        .filter(PublicViewCode.expires_at > get_localized_now())
        .all()
    )

    # Serviciile de azi pentru companie (listă)
    todays_services_company = (
        ServiceAssignment.query.options(joinedload(ServiceAssignment.student))
        .filter(
            ServiceAssignment.student_id.in_(student_ids_in_company),
            ServiceAssignment.start_datetime <= today_end,
            ServiceAssignment.end_datetime >= today_start,
        )
        .order_by(ServiceAssignment.start_datetime.asc())
        .all()
    )

    return render_template(
        "company_commander_dashboard.html",
        company_id=company_id_str,
        roll_call_time_str=roll_call_time_str,
        total_company_presence=total_company_presence_roll_call,
        platoons_data=platoons_data_roll_call,
        platoons_data_now=platoons_data_now,  # Pass new data to template
        # Statistici "Astăzi"
        permissions_today_count=permissions_today_company_count,
        daily_leaves_today_company=daily_leaves_today_company_count,  # Nume nou pentru claritate
        weekend_leaves_today_company=weekend_leaves_today_company_count,  # Nume nou
        total_leaves_today_count=total_leaves_today_company_count,
        services_today_count=services_today_company_count,
        total_students_company=len(student_ids_in_company),
        # Statistici "ACUM"
        permissions_active_now_company=permissions_active_now_company,
        daily_leaves_active_now_company=daily_leaves_active_now_company,
        weekend_leaves_active_now_company=weekend_leaves_active_now_company,
        total_on_leave_now_company=total_on_leave_now_company,
        services_active_now_company=services_active_now_company,
        current_time_for_display=now_localized,
        active_public_codes=active_public_codes,
        # New service stats
        services_last_7_days_count=services_last_7_days_count,
        services_today_breakdown=services_today_breakdown,
        todays_services=todays_services_company,
    )


@app.route("/company_commander/logs", endpoint="company_commander_logs")
@login_required
def company_commander_logs():
    if current_user.role != "comandant_companie":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    company_id_str = _get_commander_unit_id(current_user.username, "CmdC")
    if not company_id_str:
        flash("ID-ul companiei nu a putut fi determinat.", "warning")
        return redirect(url_for("company_commander_dashboard"))

    page = request.args.get("page", 1, type=int)
    per_page = 25

    # Subquery pentru a găsi ID-urile studenților din compania comandantului
    student_ids_in_company = (
        db.session.query(Student.id)
        .filter(Student.companie == company_id_str)
        .scalar_subquery()
    )

    # Interogare principală pentru ActionLog
    # Filtrăm logurile unde target_id (dacă modelul este Student, Permission, etc.) corespunde unui student din companie.
    # Acest query este complex și ar putea necesita optimizări sau o abordare diferită pentru performanță pe volume mari.

    # Interogare principală pentru ActionLog, optimizată
    base_query = ActionLog.query.join(
        User, ActionLog.user_id == User.id, isouter=True
    ).options(joinedload(ActionLog.user))

    conditions = []

    # Condiție pentru Studenți
    student_subquery = (
        db.session.query(Student.id)
        .filter(Student.companie == company_id_str)
        .scalar_subquery()
    )
    conditions.append(
        and_(
            ActionLog.target_model == "Student",
            ActionLog.target_id.in_(student_subquery),
        )
    )

    # Condiție pentru Permisii
    permission_subquery = (
        db.session.query(Permission.id)
        .join(Student, Permission.student_id == Student.id)
        .filter(Student.companie == company_id_str)
        .scalar_subquery()
    )
    conditions.append(
        and_(
            ActionLog.target_model == "Permission",
            ActionLog.target_id.in_(permission_subquery),
        )
    )

    # Condiție pentru DailyLeave
    daily_leave_subquery = (
        db.session.query(DailyLeave.id)
        .join(Student, DailyLeave.student_id == Student.id)
        .filter(Student.companie == company_id_str)
        .scalar_subquery()
    )
    conditions.append(
        and_(
            ActionLog.target_model == "DailyLeave",
            ActionLog.target_id.in_(daily_leave_subquery),
        )
    )

    # Condiție pentru WeekendLeave
    weekend_leave_subquery = (
        db.session.query(WeekendLeave.id)
        .join(Student, WeekendLeave.student_id == Student.id)
        .filter(Student.companie == company_id_str)
        .scalar_subquery()
    )
    conditions.append(
        and_(
            ActionLog.target_model == "WeekendLeave",
            ActionLog.target_id.in_(weekend_leave_subquery),
        )
    )

    # Condiție pentru ServiceAssignment
    service_subquery = (
        db.session.query(ServiceAssignment.id)
        .join(Student, ServiceAssignment.student_id == Student.id)
        .filter(Student.companie == company_id_str)
        .scalar_subquery()
    )
    conditions.append(
        and_(
            ActionLog.target_model == "ServiceAssignment",
            ActionLog.target_id.in_(service_subquery),
        )
    )

    # TODO: Adaugă condiții pentru alte modele relevante dacă este necesar (ex. VolunteerActivity, ActivityParticipant)
    # Exemplu pentru ActivityParticipant (dacă s-ar loga target_id-ul participantului)
    # participant_subquery = db.session.query(ActivityParticipant.id).join(Student, ActivityParticipant.student_id == Student.id).filter(Student.companie == company_id_str).scalar_subquery()
    # conditions.append(and_(ActionLog.target_model == 'ActivityParticipant', ActionLog.target_id.in_(participant_subquery)))

    # Aplică condițiile cu OR
    if conditions:
        final_query = base_query.filter(or_(*conditions)).order_by(
            ActionLog.timestamp.desc()
        )
    else:  # În caz că nu există condiții (deși e improbabil aici)
        final_query = base_query.order_by(
            ActionLog.timestamp.desc()
        )  # Sau returnează o listă goală direct

    # Filters
    filter_action_type = request.args.get("action_type_filter_val", "").strip()
    filter_target_model = request.args.get(
        "target_model_filter_val", ""
    ).strip()
    filter_performed_by_user_id_str = request.args.get(
        "performed_by_user_filter_val", ""
    ).strip()
    filter_date_from_str = request.args.get("filter_date_from", "").strip()
    filter_date_to_str = request.args.get("filter_date_to", "").strip()

    if filter_action_type:
        final_query = final_query.filter(
            ActionLog.action_type.ilike(f"%{filter_action_type}%")
        )
    if filter_target_model:
        final_query = final_query.filter(
            ActionLog.target_model.ilike(f"%{filter_target_model}%")
        )
    if filter_performed_by_user_id_str:
        try:
            user_id_val = int(filter_performed_by_user_id_str)
            final_query = final_query.filter(ActionLog.user_id == user_id_val)
        except ValueError:
            flash(
                "ID Utilizator (efectuat de) invalid pentru filtrare.",
                "warning",
            )
    if filter_date_from_str:
        try:
            date_from = datetime.strptime(
                filter_date_from_str, "%Y-%m-%d"
            ).date()
            final_query = final_query.filter(
                ActionLog.timestamp
                >= datetime.combine(date_from, time.min).replace(
                    tzinfo=pytz.UTC
                )
            )
        except ValueError:
            flash("Format dată 'De la' invalid.", "warning")
    if filter_date_to_str:
        try:
            date_to = datetime.strptime(filter_date_to_str, "%Y-%m-%d").date()
            final_query = final_query.filter(
                ActionLog.timestamp
                <= datetime.combine(date_to, time.max).replace(tzinfo=pytz.UTC)
            )
        except ValueError:
            flash("Format dată 'Până la' invalid.", "warning")

    logs_pagination = final_query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    return render_template(
        "commander_action_logs.html",
        logs_pagination=logs_pagination,
        title=f"Jurnal Acțiuni Compania {company_id_str}",
        unit_id=company_id_str,
        unit_type="Compania",
        action_type_filter_val=filter_action_type,
        target_model_filter_val=filter_target_model,
        performed_by_user_filter_val=filter_performed_by_user_id_str,
        filter_date_from=filter_date_from_str,
        filter_date_to=filter_date_to_str,
    )


@app.route("/battalion_commander/logs", endpoint="battalion_commander_logs")
@login_required
def battalion_commander_logs():
    if current_user.role != "comandant_batalion":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    battalion_id_str = _get_commander_unit_id(current_user.username, "CmdB")
    if not battalion_id_str:
        flash("ID-ul batalionului nu a putut fi determinat.", "warning")
        return redirect(url_for("battalion_commander_dashboard"))

    page = request.args.get("page", 1, type=int)
    per_page = 25

    # Similar cu company_commander_logs, dar filtrăm pe battalion_id_str
    base_query_battalion = ActionLog.query.join(
        User, ActionLog.user_id == User.id, isouter=True
    ).options(joinedload(ActionLog.user))
    conditions_battalion = []

    student_subquery_b = (
        db.session.query(Student.id)
        .filter(Student.batalion == battalion_id_str)
        .scalar_subquery()
    )
    conditions_battalion.append(
        and_(
            ActionLog.target_model == "Student",
            ActionLog.target_id.in_(student_subquery_b),
        )
    )

    permission_subquery_b = (
        db.session.query(Permission.id)
        .join(Student, Permission.student_id == Student.id)
        .filter(Student.batalion == battalion_id_str)
        .scalar_subquery()
    )
    conditions_battalion.append(
        and_(
            ActionLog.target_model == "Permission",
            ActionLog.target_id.in_(permission_subquery_b),
        )
    )

    daily_leave_subquery_b = (
        db.session.query(DailyLeave.id)
        .join(Student, DailyLeave.student_id == Student.id)
        .filter(Student.batalion == battalion_id_str)
        .scalar_subquery()
    )
    conditions_battalion.append(
        and_(
            ActionLog.target_model == "DailyLeave",
            ActionLog.target_id.in_(daily_leave_subquery_b),
        )
    )

    weekend_leave_subquery_b = (
        db.session.query(WeekendLeave.id)
        .join(Student, WeekendLeave.student_id == Student.id)
        .filter(Student.batalion == battalion_id_str)
        .scalar_subquery()
    )
    conditions_battalion.append(
        and_(
            ActionLog.target_model == "WeekendLeave",
            ActionLog.target_id.in_(weekend_leave_subquery_b),
        )
    )

    service_subquery_b = (
        db.session.query(ServiceAssignment.id)
        .join(Student, ServiceAssignment.student_id == Student.id)
        .filter(Student.batalion == battalion_id_str)
        .scalar_subquery()
    )
    conditions_battalion.append(
        and_(
            ActionLog.target_model == "ServiceAssignment",
            ActionLog.target_id.in_(service_subquery_b),
        )
    )

    if conditions_battalion:
        final_query_battalion = base_query_battalion.filter(
            or_(*conditions_battalion)
        ).order_by(ActionLog.timestamp.desc())
    else:
        final_query_battalion = base_query_battalion.order_by(
            ActionLog.timestamp.desc()
        )

    # Filters
    filter_action_type = request.args.get("action_type_filter_val", "").strip()
    filter_target_model = request.args.get(
        "target_model_filter_val", ""
    ).strip()
    filter_performed_by_user_id_str = request.args.get(
        "performed_by_user_filter_val", ""
    ).strip()
    filter_date_from_str = request.args.get("filter_date_from", "").strip()
    filter_date_to_str = request.args.get("filter_date_to", "").strip()

    if filter_action_type:
        final_query_battalion = final_query_battalion.filter(
            ActionLog.action_type.ilike(f"%{filter_action_type}%")
        )
    if filter_target_model:
        final_query_battalion = final_query_battalion.filter(
            ActionLog.target_model.ilike(f"%{filter_target_model}%")
        )
    if filter_performed_by_user_id_str:
        try:
            user_id_val = int(filter_performed_by_user_id_str)
            final_query_battalion = final_query_battalion.filter(
                ActionLog.user_id == user_id_val
            )
        except ValueError:
            flash(
                "ID Utilizator (efectuat de) invalid pentru filtrare.",
                "warning",
            )
    if filter_date_from_str:
        try:
            date_from = datetime.strptime(
                filter_date_from_str, "%Y-%m-%d"
            ).date()
            final_query_battalion = final_query_battalion.filter(
                ActionLog.timestamp
                >= datetime.combine(date_from, time.min).replace(
                    tzinfo=pytz.UTC
                )
            )
        except ValueError:
            flash("Format dată 'De la' invalid.", "warning")
    if filter_date_to_str:
        try:
            date_to = datetime.strptime(filter_date_to_str, "%Y-%m-%d").date()
            final_query_battalion = final_query_battalion.filter(
                ActionLog.timestamp
                <= datetime.combine(date_to, time.max).replace(tzinfo=pytz.UTC)
            )
        except ValueError:
            flash("Format dată 'Până la' invalid.", "warning")

    logs_pagination = final_query_battalion.paginate(
        page=page, per_page=per_page, error_out=False
    )

    return render_template(
        "commander_action_logs.html",
        logs_pagination=logs_pagination,
        title=f"Jurnal Acțiuni Batalionul {battalion_id_str}",
        unit_id=battalion_id_str,
        unit_type="Batalionul",
        action_type_filter_val=filter_action_type,
        target_model_filter_val=filter_target_model,
        performed_by_user_filter_val=filter_performed_by_user_id_str,
        filter_date_from=filter_date_from_str,
        filter_date_to=filter_date_to_str,
    )


@app.route("/dashboard/battalion")
@login_required
def battalion_commander_dashboard():
    if current_user.role != "comandant_batalion":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    battalion_id_str = _get_commander_unit_id(current_user.username, "CmdB")
    if not battalion_id_str:
        # Avoid redirect loop; render dashboard with empty datasets and guidance
        flash(
            "ID-ul batalionului nu a putut fi determinat din numele de utilizator. Numele trebuie să conțină modelul „CmdB<NUMĂR>”, ex: „Ionescu_CmdB2”.",
            "warning",
        )
        roll_call_datetime = get_standard_roll_call_datetime()
        roll_call_time_str = roll_call_datetime.strftime("%d %B %Y, %H:%M")
        empty_presence = {
            "efectiv_control": 0,
            "efectiv_prezent_total": 0,
            "efectiv_absent_total": 0,
            "in_formation_count": 0,
            "on_duty_count": 0,
            "platoon_graded_duty_count": 0,
            "in_formation_students_details": [],
            "on_duty_students_details": [],
            "platoon_graded_duty_students_details": [],
            "all_absent_details": [],
        }
        return render_template(
            "battalion_commander_dashboard.html",
            battalion_id="N/A",
            roll_call_time_str=roll_call_time_str,
            total_battalion_presence=empty_presence,
            companies_data={},
            companies_data_now={},
            permissions_today_count=0,
            daily_leaves_today_battalion=0,
            weekend_leaves_today_battalion=0,
            total_leaves_today_count=0,
            services_today_count=0,
            total_students_battalion=0,
            permissions_active_now_battalion=0,
            daily_leaves_active_now_battalion=0,
            weekend_leaves_active_now_battalion=0,
            total_on_leave_now_battalion=0,
            services_active_now_battalion=0,
            current_time_for_display=get_localized_now(),
            active_public_codes=[],
            services_last_7_days_count=0,
            services_today_breakdown={},
            todays_services=[],
        )

    roll_call_datetime = get_standard_roll_call_datetime()
    roll_call_time_str = roll_call_datetime.strftime("%d %B %Y, %H:%M")

    students_in_battalion_all = Student.query.filter_by(
        batalion=battalion_id_str
    ).all()
    student_ids_in_battalion = [s.id for s in students_in_battalion_all]

    total_battalion_presence_roll_call = _calculate_presence_data(
        students_in_battalion_all, roll_call_datetime
    )  # For existing roll call report

    # New stats for "today"
    today_localized_battalion = get_localized_now().date()
    today_start = datetime.combine(today_localized_battalion, time.min)
    today_end = datetime.combine(today_localized_battalion, time.max)

    permissions_today_battalion_count = Permission.query.filter(
        Permission.student_id.in_(student_ids_in_battalion),
        Permission.status == "Aprobată",
        Permission.start_datetime <= today_end,
        Permission.end_datetime >= today_start,
    ).count()

    daily_leaves_today_battalion_count = DailyLeave.query.filter(
        DailyLeave.student_id.in_(student_ids_in_battalion),
        DailyLeave.status == "Aprobată",
        DailyLeave.leave_date == today_localized_battalion,
    ).count()

    weekend_leaves_today_battalion_count = 0
    all_wl_battalion_for_today_stats = WeekendLeave.query.filter(
        WeekendLeave.student_id.in_(student_ids_in_battalion),
        WeekendLeave.status == "Aprobată",
    ).all()
    for wl_today in all_wl_battalion_for_today_stats:
        for interval_today in wl_today.get_intervals():
            if (
                interval_today["start"].date() == today_localized_battalion
                or interval_today["end"].date() == today_localized_battalion
                or (
                    interval_today["start"].date() < today_localized_battalion
                    and interval_today["end"].date()
                    > today_localized_battalion
                )
            ):
                weekend_leaves_today_battalion_count += 1
                break

    total_leaves_today_battalion_count = (
        daily_leaves_today_battalion_count
        + weekend_leaves_today_battalion_count
    )

    services_today_battalion_count = ServiceAssignment.query.filter(
        ServiceAssignment.student_id.in_(student_ids_in_battalion),
        ServiceAssignment.start_datetime <= today_end,
        ServiceAssignment.end_datetime >= today_start,
    ).count()

    # Stats for "NOW"
    now_localized_b = get_localized_now()
    permissions_active_now_battalion = Permission.query.filter(
        Permission.student_id.in_(student_ids_in_battalion),
        Permission.status == "Aprobată",
        Permission.start_datetime <= now_localized_b,
        Permission.end_datetime >= now_localized_b,
    ).count()

    daily_leaves_active_now_battalion = 0
    all_dl_battalion = DailyLeave.query.filter(
        DailyLeave.student_id.in_(student_ids_in_battalion),
        DailyLeave.status == "Aprobată",
    ).all()
    for dl_now_b in all_dl_battalion:
        if dl_now_b.is_active:
            daily_leaves_active_now_battalion += 1

    weekend_leaves_active_now_battalion = 0
    all_wl_battalion_for_now_stats = WeekendLeave.query.filter(
        WeekendLeave.student_id.in_(student_ids_in_battalion),
        WeekendLeave.status == "Aprobată",
    ).all()
    for wl_now_b in all_wl_battalion_for_now_stats:
        if wl_now_b.is_any_interval_active_now:
            weekend_leaves_active_now_battalion += 1

    total_on_leave_now_battalion = (
        permissions_active_now_battalion
        + daily_leaves_active_now_battalion
        + weekend_leaves_active_now_battalion
    )

    services_active_now_battalion = ServiceAssignment.query.filter(
        ServiceAssignment.student_id.in_(student_ids_in_battalion),
        ServiceAssignment.start_datetime <= now_localized_b,
        ServiceAssignment.end_datetime >= now_localized_b,
    ).count()

    # New service stats for battalion
    seven_days_ago = get_localized_now().date() - timedelta(days=7)
    services_last_7_days_count = ServiceAssignment.query.filter(
        ServiceAssignment.student_id.in_(student_ids_in_battalion),
        ServiceAssignment.service_date > seven_days_ago,
        ServiceAssignment.service_date <= get_localized_now().date(),
    ).count()

    services_today_breakdown = dict(
        db.session.query(
            ServiceAssignment.service_type, func.count(ServiceAssignment.id)
        )
        .filter(
            ServiceAssignment.student_id.in_(student_ids_in_battalion),
            ServiceAssignment.start_datetime <= today_end,
            ServiceAssignment.end_datetime >= today_start,
        )
        .group_by(ServiceAssignment.service_type)
        .all()
    )

    companies_data_roll_call = {}
    companies_data_now = {}  # New dictionary for current data
    companies_in_battalion = sorted(
        list(set(s.companie for s in students_in_battalion_all if s.companie))
    )

    # Handle students without a platoon at the battalion level
    students_without_platoon_in_battalion = [
        s for s in students_in_battalion_all if not s.pluton
    ]
    if students_without_platoon_in_battalion:
        category_name = "Gradati / Personal neîncadrat Batalion"
        companies_data_roll_call[category_name] = _calculate_presence_data(
            students_without_platoon_in_battalion, roll_call_datetime
        )
        companies_data_now[category_name] = _calculate_presence_data(
            students_without_platoon_in_battalion, now_localized_b
        )

    for company_id_str_loop in companies_in_battalion:
        # We only want to process students *with* platoons here, as the others are handled above.
        students_in_company_loop = [
            s
            for s in students_in_battalion_all
            if s.companie == company_id_str_loop and s.pluton
        ]
        if (
            students_in_company_loop
        ):  # Only add company if it has students with platoons
            company_name = f"Compania {company_id_str_loop}"
            companies_data_roll_call[company_name] = _calculate_presence_data(
                students_in_company_loop, roll_call_datetime
            )
            companies_data_now[company_name] = _calculate_presence_data(
                students_in_company_loop, now_localized_b
            )  # Calculate for now

    active_public_codes = (
        PublicViewCode.query.filter_by(
            created_by_user_id=current_user.id,
            is_active=True,
            scope_type="battalion",
            scope_id=battalion_id_str,
        )
        .filter(PublicViewCode.expires_at > get_localized_now())
        .all()
    )

    # Serviciile de azi pentru batalion (listă)
    todays_services_battalion = (
        ServiceAssignment.query.options(joinedload(ServiceAssignment.student))
        .filter(
            ServiceAssignment.student_id.in_(student_ids_in_battalion),
            ServiceAssignment.start_datetime <= today_end,
            ServiceAssignment.end_datetime >= today_start,
        )
        .order_by(ServiceAssignment.start_datetime.asc())
        .all()
    )

    return render_template(
        "battalion_commander_dashboard.html",
        battalion_id=battalion_id_str,
        roll_call_time_str=roll_call_time_str,
        total_battalion_presence=total_battalion_presence_roll_call,
        companies_data=companies_data_roll_call,
        companies_data_now=companies_data_now,  # Pass new data to template
        # Statistici "Astăzi"
        permissions_today_count=permissions_today_battalion_count,
        daily_leaves_today_battalion=daily_leaves_today_battalion_count,  # Nume nou
        weekend_leaves_today_battalion=weekend_leaves_today_battalion_count,  # Nume nou
        total_leaves_today_count=total_leaves_today_battalion_count,
        services_today_count=services_today_battalion_count,
        total_students_battalion=len(student_ids_in_battalion),
        # Statistici "ACUM"
        permissions_active_now_battalion=permissions_active_now_battalion,
        daily_leaves_active_now_battalion=daily_leaves_active_now_battalion,
        weekend_leaves_active_now_battalion=weekend_leaves_active_now_battalion,
        total_on_leave_now_battalion=total_on_leave_now_battalion,
        services_active_now_battalion=services_active_now_battalion,
        current_time_for_display=now_localized_b,
        active_public_codes=active_public_codes,
        # New service stats
        services_last_7_days_count=services_last_7_days_count,
        services_today_breakdown=services_today_breakdown,
        todays_services=todays_services_battalion,
    )


# --- Scoped Access Login and Decorator ---
def scoped_access_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Allow regular logged-in users to pass through without checks
        if current_user.is_authenticated:
            return f(*args, **kwargs)

        if "scoped_access" not in session:
            flash(
                "Trebuie să vă autentificați cu un cod de acces valid.",
                "warning",
            )
            return redirect(url_for("scoped_login"))

        access_info = session["scoped_access"]

        # Check expiry
        if (
            datetime.fromisoformat(access_info["expires_at"])
            < datetime.utcnow()
        ):
            session.pop("scoped_access", None)
            flash("Sesiunea de acces delegat a expirat.", "info")
            return redirect(url_for("scoped_login"))

        # Check permission for the requested page
        allowed_paths = access_info.get("permissions", [])
        if not any(request.path.startswith(path) for path in allowed_paths):
            flash("Nu aveți permisiunea să accesați această pagină.", "danger")
            # Redirect to a safe page, maybe a scoped dashboard if one exists, or just the login
            return redirect(url_for("scoped_login"))

        # Set the 'gradat' user for the context of this request
        g.scoped_user = db.session.get(User, access_info["user_id"])
        if not g.scoped_user:
            session.pop("scoped_access", None)
            flash("Utilizatorul asociat cu acest cod nu mai există.", "danger")
            return redirect(url_for("scoped_login"))

        return f(*args, **kwargs)

    return decorated_function


@app.route("/scoped_login", methods=["GET", "POST"])
def scoped_login():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    if "scoped_access" in session:
        # Re-validate session on page load
        if (
            datetime.fromisoformat(session["scoped_access"].get("expires_at"))
            < datetime.utcnow()
        ):
            session.pop("scoped_access", None)
            flash("Sesiunea de acces delegat a expirat.", "info")
        else:
            # If session is valid, maybe redirect to the first permitted page?
            # For now, let's just show a generic message or redirect to a future scoped dashboard
            flash("Sunteți deja autentificat cu un cod de acces.", "info")
            # This is a bit of a dead end, user should use logout. Or we can create a simple dashboard for them.
            # Let's redirect to the first accessible path.
            first_path = session["scoped_access"].get("permissions", [])[0]
            if first_path:
                # This is tricky as url_for needs an endpoint name, not a path.
                # We'll need a mapping or a simple dashboard.
                # For now, let's just let them logout.
                return render_template(
                    "scoped_login.html"
                )  # Or a simple "you are logged in" page

    if request.method == "POST":
        code_to_check = request.form.get("access_code", "").strip()
        if not code_to_check:
            flash("Vă rugăm introduceți un cod de acces.", "warning")
            return redirect(url_for("scoped_login"))

        now_utc = datetime.utcnow()
        access_code = ScopedAccessCode.query.filter_by(
            code=code_to_check, is_active=True
        ).first()

        if access_code and access_code.expires_at > now_utc:
            # Do not consume the code, it can be used by multiple people until it expires
            # access_code.is_active = False
            # db.session.commit()

            session["scoped_access"] = {
                "user_id": access_code.created_by_user_id,
                "permissions": access_code.get_permissions_list(),
                "expires_at": access_code.expires_at.isoformat(),
                "description": access_code.description,
            }
            session.permanent = True

            flash(
                f"Autentificare reușită! Acces permis pentru: {access_code.description}",
                "success",
            )

            # Redirect to the first available page
            permissions = access_code.get_permissions_list()
            if permissions:
                # This is still tricky. Let's redirect to a known safe endpoint that is likely to be in permissions.
                if "/volunteer" in permissions:
                    return redirect(url_for("volunteer_home"))
                elif "/gradat/students" in permissions:
                    return redirect(url_for("list_students"))

            return redirect(url_for("scoped_login"))  # Fallback
        else:
            flash(
                "Codul de acces este invalid, a expirat sau a fost dezactivat.",
                "danger",
            )
            return redirect(url_for("scoped_login"))

    return render_template("scoped_login.html")


@app.route("/scoped_logout")
def scoped_logout():
    session.pop("scoped_access", None)
    flash("Ați fost deconectat din sesiunea de acces delegat.", "success")
    return redirect(url_for("scoped_login"))


# --- Volunteer Module ---


def get_current_user_or_scoped():
    """Helper to get the active user, either from Flask-Login or the scoped session."""
    if current_user.is_authenticated:
        return current_user
    elif hasattr(g, "scoped_user"):
        return g.scoped_user
    return None


@app.route("/volunteer", methods=["GET"])
@scoped_access_required
def volunteer_home():
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat la modulul de voluntariat.", "danger")
        return redirect(url_for("dashboard"))

    # GET request logic
    activities = (
        VolunteerActivity.query.filter_by(created_by_user_id=user.id)
        .order_by(VolunteerActivity.activity_date.desc())
        .all()
    )
    students_with_points = (
        Student.query.filter_by(created_by_user_id=user.id)
        .order_by(Student.volunteer_points.desc(), Student.nume)
        .all()
    )

    return render_template(
        "volunteer_home.html",
        activities=activities,
        students_with_points=students_with_points,
    )


@app.route("/volunteer/assign_multiple_activities", methods=["GET", "POST"])
@scoped_access_required
def assign_multiple_activities():
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        student_ids = request.form.getlist("student_ids")
        activity_id = request.form.get("activity_id")
        points_to_award = request.form.get("points_to_award", 0, type=int)

        if not student_ids or not activity_id:
            flash(
                "Trebuie să selectați cel puțin un student și o activitate.",
                "warning",
            )
            return redirect(url_for("assign_multiple_activities"))

        activity = VolunteerActivity.query.get_or_404(activity_id)
        if activity.created_by_user_id != user.id:
            flash("Acces neautorizat la această activitate.", "danger")
            return redirect(url_for("volunteer_home"))

        students_to_assign = Student.query.filter(
            Student.id.in_(student_ids), Student.created_by_user_id == user.id
        ).all()

        added_count = 0
        skipped_count = 0
        for student in students_to_assign:
            # Verifică dacă studentul este deja participant
            existing_participant = ActivityParticipant.query.filter_by(
                activity_id=activity.id, student_id=student.id
            ).first()
            if not existing_participant:
                new_participant = ActivityParticipant(
                    activity_id=activity.id,
                    student_id=student.id,
                    points_awarded=points_to_award,
                )
                db.session.add(new_participant)

                # Actualizează punctajul total al studentului
                student.volunteer_points = (
                    student.volunteer_points or 0
                ) + points_to_award
                added_count += 1
            else:
                skipped_count += 1

        if added_count > 0:
            try:
                db.session.commit()
                flash(
                    f'{added_count} studenți au fost adăugați cu succes la activitatea "{activity.name}".',
                    "success",
                )
            except Exception as e:
                db.session.rollback()
                flash(f"Eroare la asignarea studenților: {str(e)}", "danger")

        if skipped_count > 0:
            flash(
                f"{skipped_count} studenți erau deja participanți și au fost omiși.",
                "info",
            )

        return redirect(
            url_for("volunteer_activity_details", activity_id=activity.id)
        )

    # GET request
    students_managed = (
        Student.query.filter_by(created_by_user_id=user.id)
        .order_by(Student.nume)
        .all()
    )
    activities = (
        VolunteerActivity.query.filter_by(created_by_user_id=user.id)
        .order_by(VolunteerActivity.activity_date.desc())
        .all()
    )

    return render_template(
        "assign_multiple_activities.html",
        students_managed=students_managed,
        activities=activities,
    )


@app.route("/volunteer/activity/create", methods=["GET", "POST"])
@scoped_access_required
def volunteer_activity_create():
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("volunteer_home"))

    if request.method == "POST":
        activity_name = request.form.get("activity_name", "").strip()
        activity_description = request.form.get(
            "activity_description", ""
        ).strip()
        activity_date_str = request.form.get("activity_date")

        if not activity_name or not activity_date_str:
            flash("Numele activității și data sunt obligatorii.", "warning")
            today_str_for_form = (
                get_localized_now().date().strftime("%Y-%m-%d")
            )
            return render_template(
                "volunteer_activity_create.html", today_str=today_str_for_form
            )

        try:
            activity_date_obj = datetime.strptime(
                activity_date_str, "%Y-%m-%d"
            ).date()
        except ValueError:
            flash("Format dată invalid pentru activitate.", "danger")
            today_str_for_form = (
                get_localized_now().date().strftime("%Y-%m-%d")
            )
            return render_template(
                "volunteer_activity_create.html", today_str=today_str_for_form
            )

        new_activity = VolunteerActivity(
            name=activity_name,
            description=activity_description,
            activity_date=activity_date_obj,
            created_by_user_id=user.id,
        )
        db.session.add(new_activity)
        try:
            db.session.commit()
            flash(
                f'Activitatea de voluntariat "{activity_name}" a fost creată cu succes.',
                "success",
            )
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la crearea activității: {str(e)}", "danger")
        return redirect(url_for("volunteer_home"))

    # GET request
    today_str_for_form = get_localized_now().date().strftime("%Y-%m-%d")
    return render_template(
        "volunteer_activity_create.html", today_str=today_str_for_form
    )


@app.route("/volunteer/activity/<int:activity_id>", methods=["GET", "POST"])
@scoped_access_required
def volunteer_activity_details(activity_id):
    user = get_current_user_or_scoped()
    activity = VolunteerActivity.query.get_or_404(activity_id)
    if (
        not user
        or user.role != "gradat"
        or activity.created_by_user_id != user.id
    ):
        flash(
            "Acces neautorizat la această activitate de voluntariat.", "danger"
        )
        return redirect(url_for("volunteer_home"))

    if request.method == "POST":
        action = request.form.get("action")

        if action == "update_participants":
            selected_student_ids = set(
                request.form.getlist("participant_ids[]", type=int)
            )

            # Remove participants not selected anymore
            current_participants_in_activity = (
                ActivityParticipant.query.filter_by(
                    activity_id=activity.id
                ).all()
            )
            for ap in current_participants_in_activity:
                if ap.student_id not in selected_student_ids:
                    # Before deleting, if points were awarded, subtract them from student's total.
                    # This logic might need refinement: what if points were awarded for *this specific* participation?
                    # For now, if a student is removed, their points from this activity are effectively "lost" from this AP record.
                    # The Student.volunteer_points should ideally be a sum from all their AP records.
                    student_obj = db.session.get(Student, ap.student_id)
                    if student_obj:  # Subtract points if student exists
                        student_obj.volunteer_points = max(
                            0, student_obj.volunteer_points - ap.points_awarded
                        )  # Ensure not negative
                    db.session.delete(ap)

            # Add new participants
            for student_id_to_add in selected_student_ids:
                exists = ActivityParticipant.query.filter_by(
                    activity_id=activity.id, student_id=student_id_to_add
                ).first()
                if not exists:
                    # Ensure the student belongs to the current gradat
                    student_check = Student.query.filter_by(
                        id=student_id_to_add, created_by_user_id=user.id
                    ).first()
                    if student_check:
                        new_participant = ActivityParticipant(
                            activity_id=activity.id,
                            student_id=student_id_to_add,
                            points_awarded=0,
                        )
                        db.session.add(new_participant)
                    else:
                        flash(
                            f"Studentul cu ID {student_id_to_add} nu a putut fi adăugat (nu este gestionat de dvs).",
                            "warning",
                        )

            try:
                db.session.commit()
                # Recalculate total points for all affected students
                all_involved_student_ids = selected_student_ids.union(
                    set(
                        ap.student_id
                        for ap in current_participants_in_activity
                    )
                )
                for s_id in all_involved_student_ids:
                    stud = db.session.get(Student, s_id)
                    if stud:
                        stud.volunteer_points = (
                            db.session.query(
                                func.sum(ActivityParticipant.points_awarded)
                            )
                            .filter_by(student_id=s_id)
                            .scalar()
                            or 0
                        )
                db.session.commit()
                flash("Lista de participanți a fost actualizată.", "success")
            except Exception as e:
                db.session.rollback()
                flash(
                    f"Eroare la actualizarea participanților: {str(e)}",
                    "danger",
                )

        elif action == "award_points":
            points_to_award_val = request.form.get("points_to_award", type=int)
            participant_ids_for_points = set(
                request.form.getlist("points_participant_ids[]", type=int)
            )

            if points_to_award_val is None or points_to_award_val < 0:
                flash("Numărul de puncte de acordat este invalid.", "warning")
            else:
                updated_count = 0
                for student_id_for_points in participant_ids_for_points:
                    participant_record = ActivityParticipant.query.filter_by(
                        activity_id=activity.id,
                        student_id=student_id_for_points,
                    ).first()
                    if participant_record:
                        # Option 1: Add to existing points for this activity
                        # participant_record.points_awarded += points_to_award_val
                        # Option 2: Set points for this activity (if points are per activity, not cumulative for it)
                        participant_record.points_awarded = points_to_award_val
                        updated_count += 1
                    else:
                        flash(
                            f"Studentul cu ID {student_id_for_points} nu este participant la această activitate pentru a primi puncte.",
                            "warning",
                        )

                if updated_count > 0:
                    try:
                        db.session.commit()
                        # Recalculate total points for all students who received points
                        for s_id in participant_ids_for_points:
                            stud = db.session.get(Student, s_id)
                            if stud:  # Check if student still exists
                                total_points = (
                                    db.session.query(
                                        func.sum(
                                            ActivityParticipant.points_awarded
                                        )
                                    )
                                    .filter_by(student_id=s_id)
                                    .scalar()
                                )
                                stud.volunteer_points = (
                                    total_points
                                    if total_points is not None
                                    else 0
                                )
                        db.session.commit()
                        flash(
                            f"{points_to_award_val} puncte acordate pentru {updated_count} participanți selectați.",
                            "success",
                        )
                    except Exception as e:
                        db.session.rollback()
                        flash(
                            f"Eroare la acordarea punctelor: {str(e)}",
                            "danger",
                        )
        else:
            flash("Acțiune necunoscută.", "danger")

        return redirect(
            url_for("volunteer_activity_details", activity_id=activity.id)
        )

    # GET request
    students_managed = (
        Student.query.filter_by(created_by_user_id=user.id)
        .order_by(Student.nume)
        .all()
    )

    # Get current participant student IDs for this activity
    current_participant_ids = [
        ap.student_id
        for ap in ActivityParticipant.query.filter_by(
            activity_id=activity.id
        ).all()
    ]

    # Get detailed participant info (ActivityParticipant object + Student object)
    activity_participants_detailed_query = (
        db.session.query(ActivityParticipant, Student)
        .join(Student, ActivityParticipant.student_id == Student.id)
        .filter(ActivityParticipant.activity_id == activity.id)
        .all()
    )

    # activity_participants_detailed will be a list of (ActivityParticipant, Student) tuples
    # The template already uses this structure in the loop: {% for participant, student_detail in activity_participants_detailed %}

    return render_template(
        "volunteer_activity_details.html",
        activity=activity,
        students_managed=students_managed,
        current_participant_ids=current_participant_ids,
        activity_participants_detailed=activity_participants_detailed_query,
    )


def _generate_eligible_volunteers(
    gradat_id,
    num_to_generate,
    exclude_girls=False,
    activity_date=None,
    exclude_student_ids=None,
):
    """
    Helper function to generate a list of eligible students for volunteering.
    """
    students_query = Student.query.filter_by(created_by_user_id=gradat_id)

    if exclude_girls:
        students_query = students_query.filter(Student.gender != "F")

    # Exclude students with medical exemptions
    students_query = students_query.filter(Student.is_smt == False)
    students_query = students_query.filter(
        or_(Student.exemption_details == None, Student.exemption_details == "")
    )

    # Exclude students already provided in the exclusion list (e.g., from a session)
    if exclude_student_ids:
        students_query = students_query.filter(
            Student.id.notin_(exclude_student_ids)
        )

    # Exclude students with leaves on the specified activity date
    if activity_date:
        ids_on_leave = set()
        activity_day_start_aware = EUROPE_BUCHAREST.localize(
            datetime.combine(activity_date, time.min)
        )
        activity_day_end_aware = EUROPE_BUCHAREST.localize(
            datetime.combine(activity_date, time.max)
        )

    all_managed_student_ids = [
        sid
        for (sid,) in Student.query.filter_by(created_by_user_id=gradat_id)
        .with_entities(Student.id)
        .all()
    ]

    # Permissions
    act_day_start_naive = activity_day_start_aware.replace(tzinfo=None)
    act_day_end_naive = activity_day_end_aware.replace(tzinfo=None)
    permissions = (
        Permission.query.filter(
            Permission.student_id.in_(all_managed_student_ids),
            Permission.status == "Aprobată",
            Permission.start_datetime < act_day_end_naive,
            Permission.end_datetime > act_day_start_naive,
        )
        .with_entities(Permission.student_id)
        .all()
    )
    for p_id in permissions:
        ids_on_leave.add(p_id[0])

    # Daily Leaves
    daily_leaves = DailyLeave.query.filter(
        DailyLeave.student_id.in_(all_managed_student_ids),
        DailyLeave.status == "Aprobată",
        DailyLeave.leave_date == activity_date,
    ).all()
    for dl in daily_leaves:
        if (
            dl.start_datetime < act_day_end_naive
            and dl.end_datetime > act_day_start_naive
        ):
            ids_on_leave.add(dl.student_id)

    # Weekend Leaves
    weekend_leaves = WeekendLeave.query.filter(
        WeekendLeave.student_id.in_(all_managed_student_ids),
        WeekendLeave.status == "Aprobată",
        WeekendLeave.weekend_start_date <= activity_date,
        WeekendLeave.weekend_start_date >= activity_date - timedelta(days=3),
    ).all()
    for wl in weekend_leaves:
        for interval in wl.get_intervals():
            if (
                interval["start"] < activity_day_end_aware
                and interval["end"] > activity_day_start_aware
            ):
                ids_on_leave.add(wl.student_id)
                break

    if ids_on_leave:
        students_query = students_query.filter(
            Student.id.notin_(list(ids_on_leave))
        )

    return (
        students_query.order_by(
            Student.volunteer_points.asc(), Student.nume.asc()
        )
        .limit(num_to_generate)
        .all()
    )


@app.route("/volunteer/generate", methods=["GET", "POST"])
@scoped_access_required
def volunteer_generate_students():
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    generated_students_list = None
    form_data = {
        "num_students": request.form.get("num_students", 5),
        "exclude_girls": "exclude_girls" in request.form,
        "activity_date_for_check": request.form.get(
            "activity_date_for_check", get_localized_now().date().isoformat()
        ),
    }

    if request.method == "POST":
        try:
            num_to_generate = int(form_data["num_students"])
            if num_to_generate <= 0:
                flash(
                    "Numărul de studenți necesari trebuie să fie pozitiv.",
                    "warning",
                )
                num_to_generate = 5
        except (ValueError, TypeError):
            flash("Număr de studenți invalid.", "warning")
            num_to_generate = 5

        try:
            activity_date = datetime.strptime(
                form_data["activity_date_for_check"], "%Y-%m-%d"
            ).date()
        except (ValueError, TypeError):
            flash(
                "Format dată activitate invalid. Se folosește data curentă.",
                "warning",
            )
            activity_date = get_localized_now().date()
            form_data["activity_date_for_check"] = activity_date.isoformat()

        generated_students_list = _generate_eligible_volunteers(
            gradat_id=user.id,
            num_to_generate=num_to_generate,
            exclude_girls=form_data["exclude_girls"],
            activity_date=activity_date,
        )

        if not generated_students_list:
            flash(
                "Nu s-au găsit studenți eligibili conform criteriilor.", "info"
            )

    return render_template(
        "volunteer_generate_students.html",
        generated_students=generated_students_list,
        form_data=form_data,
    )


@app.route(
    "/volunteer/save_session",
    methods=["POST"],
    endpoint="save_volunteer_session",
)
@scoped_access_required
def save_volunteer_session():
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    session_name = request.form.get("session_name", "").strip()
    student_ids = request.form.getlist("student_ids[]", type=int)

    if not session_name:
        flash("Numele listei de voluntari este obligatoriu.", "warning")
        # This redirect is not ideal as it loses the context of the generated list.
        # A more advanced implementation might use session to store and repopulate the previous page.
        # For now, redirecting to the generation page is a simple fallback.
        return redirect(url_for("volunteer_generate_students"))

    if not student_ids:
        flash(
            "Niciun student nu a fost selectat pentru a fi salvat în listă.",
            "warning",
        )
        return redirect(url_for("volunteer_generate_students"))

    # Create the new session
    new_session = VolunteerSession(
        name=session_name, created_by_user_id=user.id
    )

    # Find and associate students
    students_to_add = Student.query.filter(
        Student.id.in_(student_ids), Student.created_by_user_id == user.id
    ).all()

    if len(students_to_add) != len(student_ids):
        flash(
            "Avertisment: Unii studenți nu au putut fi găsiți sau nu vă aparțin și nu au fost adăugați la listă.",
            "warning",
        )

    new_session.students.extend(students_to_add)

    try:
        db.session.add(new_session)
        db.session.commit()
        flash(
            f'Lista de voluntari "{session_name}" a fost salvată cu succes.',
            "success",
        )
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la salvarea listei de voluntari: {str(e)}", "danger")
        app.logger.error(
            f"Error saving volunteer session for user {user.id}: {str(e)}"
        )

    # Redirect to the list of saved sessions
    return redirect(url_for("volunteer_sessions_list"))


@app.route(
    "/volunteer/sessions", methods=["GET"], endpoint="volunteer_sessions_list"
)
@scoped_access_required
def volunteer_sessions_list():
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    sessions = (
        VolunteerSession.query.filter_by(created_by_user_id=user.id)
        .order_by(VolunteerSession.created_at.desc())
        .all()
    )

    return render_template("volunteer_sessions_list.html", sessions=sessions)


@app.route(
    "/volunteer/session/<int:session_id>",
    methods=["GET", "POST"],
    endpoint="volunteer_session_details",
)
@scoped_access_required
def volunteer_session_details(session_id):
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    session = VolunteerSession.query.get_or_404(session_id)
    if session.created_by_user_id != user.id:
        flash("Acces neautorizat la această listă.", "danger")
        return redirect(url_for("volunteer_sessions_list"))

    if request.method == "POST":
        action = request.form.get("action")
        if action == "add_students":
            student_ids_to_add = request.form.getlist(
                "student_ids_to_add", type=int
            )
            if not student_ids_to_add:
                flash(
                    "Nu ați selectat niciun student pentru a adăuga.",
                    "warning",
                )
                return redirect(
                    url_for("volunteer_session_details", session_id=session_id)
                )

            students_to_add = Student.query.filter(
                Student.id.in_(student_ids_to_add),
                Student.created_by_user_id == user.id,
            ).all()

            added_count = 0
            for student in students_to_add:
                if student not in session.students:
                    session.students.append(student)
                    added_count += 1

            if added_count > 0:
                try:
                    db.session.commit()
                    flash(
                        f"{added_count} studenți au fost adăugați la listă.",
                        "success",
                    )
                except Exception as e:
                    db.session.rollback()
                    flash(
                        f"Eroare la adăugarea studenților: {str(e)}", "danger"
                    )

            return redirect(
                url_for("volunteer_session_details", session_id=session_id)
            )

        elif action == "award_points":
            points_to_award = request.form.get(
                "points_to_award_direct", type=int
            )
            if points_to_award is None or points_to_award < 0:
                flash("Numărul de puncte de acordat este invalid.", "warning")
                return redirect(
                    url_for("volunteer_session_details", session_id=session_id)
                )

            students_in_session = session.students.all()
            if not students_in_session:
                flash(
                    "Nu există studenți în această listă pentru a le acorda puncte.",
                    "info",
                )
                return redirect(
                    url_for("volunteer_session_details", session_id=session_id)
                )

            updated_count = 0
            for student in students_in_session:
                student.volunteer_points = (
                    student.volunteer_points or 0
                ) + points_to_award
                updated_count += 1

            if updated_count > 0:
                try:
                    db.session.commit()
                    flash(
                        f"{points_to_award} puncte au fost acordate pentru {updated_count} studenți.",
                        "success",
                    )
                except Exception as e:
                    db.session.rollback()
                    flash(f"Eroare la acordarea punctelor: {str(e)}", "danger")

            return redirect(
                url_for("volunteer_session_details", session_id=session_id)
            )

    # GET request or if POST action is not 'add_students'
    students_in_session = session.students.all()
    all_managed_students = Student.query.filter_by(
        created_by_user_id=user.id
    ).all()

    # Get activities created by the user to populate the dropdown
    available_activities = (
        VolunteerActivity.query.filter_by(created_by_user_id=user.id)
        .order_by(VolunteerActivity.activity_date.desc())
        .all()
    )

    return render_template(
        "volunteer_session_details.html",
        session=session,
        students=students_in_session,
        all_managed_students=all_managed_students,
        available_activities=available_activities,
    )


@app.route(
    "/volunteer/session/<int:session_id>/assign",
    methods=["POST"],
    endpoint="assign_session_to_activity",
)
@scoped_access_required
def assign_session_to_activity(session_id):
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    session = VolunteerSession.query.get_or_404(session_id)
    if session.created_by_user_id != user.id:
        flash("Acces neautorizat la această listă.", "danger")
        return redirect(url_for("volunteer_sessions_list"))

    activity_id = request.form.get("activity_id", type=int)
    points_to_award = request.form.get("points_to_award", type=int, default=0)

    if not activity_id:
        flash("Trebuie să selectați o activitate.", "warning")
        return redirect(
            url_for("volunteer_session_details", session_id=session_id)
        )

    activity = VolunteerActivity.query.get_or_404(activity_id)
    if activity.created_by_user_id != user.id:
        flash("Acces neautorizat la activitatea selectată.", "danger")
        return redirect(
            url_for("volunteer_session_details", session_id=session_id)
        )

    students_in_session = session.students.all()
    current_participant_ids = {p.student_id for p in activity.participants}

    added_count = 0
    skipped_count = 0

    for student in students_in_session:
        if student.id not in current_participant_ids:
            new_participant = ActivityParticipant(
                activity_id=activity.id,
                student_id=student.id,
                points_awarded=points_to_award,
            )
            db.session.add(new_participant)
            # If points are awarded, update the student's total
            if points_to_award > 0:
                student.volunteer_points += points_to_award
            added_count += 1
        else:
            skipped_count += 1

    try:
        db.session.commit()
        if added_count > 0:
            flash(
                f'{added_count} studenți au fost adăugați ca participanți la activitatea "{activity.name}".',
                "success",
            )
        if skipped_count > 0:
            flash(
                f"{skipped_count} studenți erau deja participanți și au fost omiși.",
                "info",
            )
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la asignarea participanților: {str(e)}", "danger")
        app.logger.error(
            f"Error assigning volunteer session {session_id} to activity {activity_id}: {str(e)}"
        )

    return redirect(
        url_for("volunteer_activity_details", activity_id=activity.id)
    )


@app.route(
    "/volunteer/session/<int:session_id>/generate_and_add",
    methods=["POST"],
    endpoint="generate_and_add_to_session",
)
@scoped_access_required
def generate_and_add_to_session(session_id):
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    session = VolunteerSession.query.get_or_404(session_id)
    if session.created_by_user_id != user.id:
        flash("Acces neautorizat la această listă.", "danger")
        return redirect(url_for("volunteer_sessions_list"))

    num_students_to_generate = request.form.get(
        "num_students_to_generate", type=int, default=5
    )
    exclude_girls = "exclude_girls_generate" in request.form

    # Re-use the logic from volunteer_generate_students, but adapt it
    students_query = Student.query.filter_by(created_by_user_id=user.id)

    if exclude_girls:
        students_query = students_query.filter(Student.gender != "F")

    # Exclude students already in the session
    student_ids_in_session = [s.id for s in session.students]
    if student_ids_in_session:
        students_query = students_query.filter(
            Student.id.notin_(student_ids_in_session)
        )

    # Simple generation logic (lowest points first)
    generated_students = (
        students_query.order_by(
            Student.volunteer_points.asc(), Student.nume.asc()
        )
        .limit(num_students_to_generate)
        .all()
    )

    if not generated_students:
        flash("Nu s-au găsit studenți eligibili pentru a fi adăugați.", "info")
        return redirect(
            url_for("volunteer_session_details", session_id=session_id)
        )

    for student in generated_students:
        session.students.append(student)

    try:
        db.session.commit()
        flash(
            f"{len(generated_students)} studenți au fost generați și adăugați la listă.",
            "success",
        )
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la adăugarea studenților generați: {str(e)}", "danger")

    return redirect(
        url_for("volunteer_session_details", session_id=session_id)
    )


@app.route(
    "/volunteer/session/delete/<int:session_id>",
    methods=["POST"],
    endpoint="delete_volunteer_session",
)
@scoped_access_required
def delete_volunteer_session(session_id):
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    session_to_delete = VolunteerSession.query.get_or_404(session_id)

    if session_to_delete.created_by_user_id != user.id:
        flash("Nu aveți permisiunea să ștergeți această listă.", "danger")
        return redirect(url_for("volunteer_sessions_list"))

    try:
        # The relationship is many-to-many, so deleting the session
        # should just delete the entries in the association table, not the students themselves.
        db.session.delete(session_to_delete)
        db.session.commit()
        flash(
            f'Lista "{session_to_delete.name}" a fost ștearsă cu succes.',
            "success",
        )
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la ștergerea listei: {str(e)}", "danger")
        app.logger.error(
            f"Error deleting volunteer session {session_id} for user {user.id}: {str(e)}"
        )

    return redirect(url_for("volunteer_sessions_list"))


@app.route(
    "/volunteer/activity/delete/<int:activity_id>",
    methods=["POST"],
    endpoint="delete_volunteer_activity",
)
@scoped_access_required
def delete_volunteer_activity(activity_id):
    user = get_current_user_or_scoped()
    if not user or user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("volunteer_home"))

    activity_to_delete = VolunteerActivity.query.get_or_404(activity_id)

    if activity_to_delete.created_by_user_id != user.id:
        flash(
            "Nu aveți permisiunea să ștergeți această activitate de voluntariat.",
            "danger",
        )
        return redirect(url_for("volunteer_home"))

    activity_name_for_log = activity_to_delete.name
    details_before_activity_delete = model_to_dict(activity_to_delete)
    participants_details_before = []
    student_points_adjustments = (
        {}
    )  # student_id: {old_total: X, points_removed: Y, new_total: Z}

    try:
        # Iterează prin participanți pentru a ajusta punctele și a colecta detalii pentru log
        for (
            participant
        ) in (
            activity_to_delete.participants
        ):  # activity.participants este o interogare lazy='dynamic'
            student = participant.student  # Ar trebui să existe datorită FK
            if student:
                points_to_remove = participant.points_awarded

                participants_details_before.append(
                    {
                        "participant_id": participant.id,
                        "student_id": student.id,
                        "student_name": f"{student.nume} {student.prenume}",
                        "points_awarded_in_activity": points_to_remove,
                        "student_total_points_before_adj": student.volunteer_points,
                    }
                )

                student.volunteer_points = max(
                    0, student.volunteer_points - points_to_remove
                )

                student_points_adjustments[student.id] = {
                    "old_total": participants_details_before[-1][
                        "student_total_points_before_adj"
                    ],  # Ultima valoare adăugată
                    "points_removed": points_to_remove,
                    "new_total": student.volunteer_points,
                }

            # Nu este nevoie să ștergem explicit participantul dacă cascade="all, delete-orphan" este setat
            # pe VolunteerActivity.participants și funcționează corect.
            # Dar dacă vrem să fim siguri sau dacă cascada nu e setată/funcțională:
            # db.session.delete(participant) # Acest lucru va fi gestionat de cascadă

        # Șterge activitatea (participanții ar trebui șterși prin cascadă)
        db.session.delete(activity_to_delete)

        # Log principal pentru ștergerea activității
        log_description = (
            f"User {user.username} deleted volunteer activity '{activity_name_for_log}' (ID: {activity_id}). "
            f"{len(participants_details_before)} participant records were associated. Points adjusted for students."
        )
        # Adăugăm și detaliile despre ajustarea punctelor la descriere sau ca un JSON separat
        # Poate fi prea mult pentru description, mai bine în details_before/after dacă e cazul

        log_action(
            "DELETE_VOLUNTEER_ACTIVITY_SUCCESS",
            target_model_name="VolunteerActivity",
            target_id=activity_id,
            details_before_dict={
                "activity": details_before_activity_delete,
                "participants_before_delete": participants_details_before,
                "student_points_adjustments": student_points_adjustments,
            },
            description=log_description,
        )

        db.session.commit()
        flash(
            f'Activitatea de voluntariat "{activity_name_for_log}" și punctele asociate au fost șterse cu succes.',
            "success",
        )

    except Exception as e:
        db.session.rollback()
        flash(
            f"Eroare la ștergerea activității de voluntariat: {str(e)}",
            "danger",
        )
        log_action(
            "DELETE_VOLUNTEER_ACTIVITY_FAIL",
            target_model_name="VolunteerActivity",
            target_id=activity_id,
            details_before_dict=details_before_activity_delete,  # Log ce s-a încercat să se șteargă
            description=f"User {user.username} failed to delete volunteer activity '{activity_name_for_log}'. Error: {str(e)}",
        )
        db.session.commit()  # Commit log-ul de eroare

    return redirect(url_for("volunteer_home"))


# --- Management Studenți ---
@app.route("/gradat/students")
@app.route("/admin/students")
@scoped_access_required
def list_students():
    user = get_current_user_or_scoped()
    # If no user from either login method, deny access.
    if not user:
        flash("Acces neautorizat.", "danger")
        return redirect(
            url_for("user_login")
        )  # Or scoped_login if we can determine context

    is_admin_view = user.role == "admin" and request.path.startswith("/admin/")
    page = request.args.get("page", 1, type=int)
    per_page = 15

    students_query = Student.query.options(joinedload(Student.creator))
    # For populating filter dropdowns - consider optimizing if it becomes slow
    # These filters are primarily for admin view.
    batalioane, companii, plutoane = [], [], []
    if is_admin_view or user.role in [
        "comandant_companie",
        "comandant_batalion",
    ]:
        all_students_for_filters_q = Student.query
        if user.role == "comandant_companie":
            company_id = _get_commander_unit_id(user.username, "CmdC")
            if company_id:
                all_students_for_filters_q = all_students_for_filters_q.filter(
                    Student.companie == company_id
                )
        elif user.role == "comandant_batalion":
            battalion_id = _get_commander_unit_id(user.username, "CmdB")
            if battalion_id:
                all_students_for_filters_q = all_students_for_filters_q.filter(
                    Student.batalion == battalion_id
                )

        all_students_for_filters = (
            all_students_for_filters_q.with_entities(
                Student.batalion, Student.companie, Student.pluton
            )
            .distinct()
            .all()
        )
        batalioane = sorted(
            list(
                set(s.batalion for s in all_students_for_filters if s.batalion)
            )
        )
        companii = sorted(
            list(
                set(s.companie for s in all_students_for_filters if s.companie)
            )
        )
        plutoane = sorted(
            list(set(s.pluton for s in all_students_for_filters if s.pluton))
        )

    search_term = request.args.get("search", "").strip()
    filter_batalion = request.args.get("batalion", "").strip()
    filter_companie = request.args.get("companie", "").strip()
    filter_pluton = request.args.get("pluton", "").strip()

    if is_admin_view:
        if filter_batalion:
            students_query = students_query.filter(
                Student.batalion == filter_batalion
            )
        if filter_companie:
            students_query = students_query.filter(
                Student.companie == filter_companie
            )
        if filter_pluton:
            students_query = students_query.filter(
                Student.pluton == filter_pluton
            )
        students_query = students_query.order_by(
            Student.batalion,
            Student.companie,
            Student.pluton,
            Student.nume,
            Student.prenume,
        )
    elif user.role == "gradat":
        students_query = students_query.filter_by(created_by_user_id=user.id)
        # Gradat might not need sub-filters for platoon/company as they manage specific students
        students_query = students_query.order_by(Student.nume, Student.prenume)
    elif user.role == "comandant_companie":
        company_id = _get_commander_unit_id(user.username, "CmdC")
        if company_id:
            students_query = students_query.filter(
                Student.companie == company_id
            )
            if filter_pluton:
                students_query = students_query.filter(
                    Student.pluton == filter_pluton
                )  # Allow CmdC to filter by platoon
            students_query = students_query.order_by(
                Student.pluton, Student.nume, Student.prenume
            )
        else:
            flash("ID Companie nevalid pentru comandant.", "danger")
            students_query = students_query.filter(
                Student.id == -1
            )  # No results
    elif user.role == "comandant_batalion":
        battalion_id = _get_commander_unit_id(user.username, "CmdB")
        if battalion_id:
            students_query = students_query.filter(
                Student.batalion == battalion_id
            )
            if filter_companie:
                students_query = students_query.filter(
                    Student.companie == filter_companie
                )  # Allow CmdB to filter by company
            if filter_pluton:
                students_query = students_query.filter(
                    Student.pluton == filter_pluton
                )  # And platoon
            students_query = students_query.order_by(
                Student.companie, Student.pluton, Student.nume, Student.prenume
            )
        else:
            flash("ID Batalion nevalid pentru comandant.", "danger")
            students_query = students_query.filter(
                Student.id == -1
            )  # No results
    else:  # Should not happen due to @login_required and role checks in other views
        flash(
            "Rol utilizator necunoscut pentru listarea studenților.", "danger"
        )
        return redirect(url_for("dashboard"))

    if search_term:
        processed_search_term = unidecode(search_term.lower())
        search_pattern = f"%{processed_search_term}%"
        students_query = students_query.filter(
            or_(
                func.lower(Student.nume).ilike(search_pattern),
                func.lower(Student.prenume).ilike(search_pattern),
                func.lower(Student.id_unic_student).ilike(search_pattern),
            )
        )

    students_pagination = students_query.paginate(
        page=page, per_page=per_page, error_out=False
    )
    students_list = students_pagination.items

    # Determine title based on role
    view_title = "Listă Studenți"
    if is_admin_view:
        view_title = "Listă Generală Studenți (Admin)"
    elif user.role == "gradat":
        view_title = "Listă Studenți Gestionați"
    elif user.role == "comandant_companie":
        view_title = f"Listă Studenți Compania {_get_commander_unit_id(user.username, 'CmdC') or 'N/A'}"
    elif user.role == "comandant_batalion":
        view_title = f"Listă Studenți Batalionul {_get_commander_unit_id(user.username, 'CmdB') or 'N/A'}"

    # Fetch leave templates for the batch action modal
    leave_templates = []
    if user.role == "gradat":
        leave_templates = (
            LeaveTemplate.query.filter_by(created_by_user_id=user.id)
            .order_by(LeaveTemplate.name)
            .all()
        )

    return render_template(
        "list_students.html",
        students=students_list,
        students_pagination=students_pagination,
        is_admin_view=is_admin_view,  # Keep for template logic (e.g. showing creator)
        search_term=search_term,
        # Filters are available based on role now
        filter_batalion=filter_batalion,
        filter_companie=filter_companie,
        filter_pluton=filter_pluton,
        batalioane=batalioane,
        companii=companii,
        plutoane=plutoane,
        title=view_title,
        leave_templates=leave_templates,
        service_types=SERVICE_TYPES,
    )


@app.route("/gradat/students/batch_action", methods=["POST"])
@login_required
def batch_action_students():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_ids = request.form.getlist("student_ids[]")
    action_type = request.form.get("action_type")

    if not student_ids or not action_type:
        flash("Lipsesc ID-urile studenților sau tipul acțiunii.", "warning")
        return redirect(url_for("list_students"))

    added_count = 0
    skipped_count = 0
    conflict_details = []

    # This is a single transaction. If one fails, the whole batch is rolled back.
    try:
        for student_id_str in student_ids:
            student_id = int(student_id_str)
            student = db.session.get(Student, student_id)
            if not student or student.created_by_user_id != current_user.id:
                conflict_details.append(
                    f"Studentul cu ID {student_id_str} nu este valid sau nu vă aparține."
                )
                skipped_count += 1
                continue

            try:
                if action_type == "permission":
                    start_dt = datetime.strptime(
                        request.form.get("start_datetime"), "%Y-%m-%dT%H:%M"
                    )
                    end_dt = datetime.strptime(
                        request.form.get("end_datetime"), "%Y-%m-%dT%H:%M"
                    )
                    if end_dt <= start_dt:
                        raise ValueError(
                            "Data de sfârșit trebuie să fie după cea de început."
                        )

                    conflict = check_leave_conflict(
                        student_id, start_dt, end_dt, leave_type="permission"
                    )
                    if conflict:
                        raise ValueError(f"Conflict: {conflict}")

                    new_perm = Permission(
                        student_id=student_id,
                        start_datetime=start_dt,
                        end_datetime=end_dt,
                        destination=request.form.get("destination"),
                        transport_mode=request.form.get("transport_mode"),
                        reason=request.form.get("reason"),
                        created_by_user_id=current_user.id,
                        status="Aprobată",
                    )
                    db.session.add(new_perm)
                    log_student_action(
                        student_id,
                        "PERMISSION_CREATED_BATCH",
                        f"Permisie adăugată: {start_dt.strftime('%d.%m %H:%M')} - {end_dt.strftime('%d.%m %H:%M')}.",
                    )

                elif action_type == "daily_leave":
                    leave_date = datetime.strptime(
                        request.form.get("leave_date"), "%Y-%m-%d"
                    ).date()
                    start_time = datetime.strptime(
                        request.form.get("start_time"), "%H:%M"
                    ).time()
                    end_time = datetime.strptime(
                        request.form.get("end_time"), "%H:%M"
                    ).time()

                    start_dt = datetime.combine(leave_date, start_time)
                    end_dt = datetime.combine(leave_date, end_time)
                    if end_time < start_time:
                        end_dt += timedelta(days=1)

                    conflict = check_leave_conflict(
                        student_id, start_dt, end_dt, leave_type="daily_leave"
                    )
                    if conflict:
                        raise ValueError(f"Conflict: {conflict}")

                    new_leave = DailyLeave(
                        student_id=student_id,
                        leave_date=leave_date,
                        start_time=start_time,
                        end_time=end_time,
                        created_by_user_id=current_user.id,
                        status="Aprobată",
                    )
                    db.session.add(new_leave)
                    log_student_action(
                        student_id,
                        "DAILY_LEAVE_CREATED_BATCH",
                        f"Învoire zilnică adăugată: {leave_date.strftime('%d.%m.%Y')} {start_time.strftime('%H:%M')}-{end_time.strftime('%H:%M')}.",
                    )

                elif action_type == "service":
                    service_type = request.form.get("service_type")
                    service_date = datetime.strptime(
                        request.form.get("service_date"), "%Y-%m-%d"
                    ).date()
                    start_time = datetime.strptime(
                        request.form.get("start_time"), "%H:%M"
                    ).time()
                    end_time = datetime.strptime(
                        request.form.get("end_time"), "%H:%M"
                    ).time()

                    start_dt = datetime.combine(service_date, start_time)
                    end_dt = datetime.combine(service_date, end_time)
                    if end_time < start_time or (
                        service_type == "GSS" and end_time == start_time
                    ):
                        end_dt += timedelta(days=1)

                    conflict = check_service_conflict_for_student(
                        student_id, start_dt, end_dt, service_type
                    )
                    if conflict:
                        raise ValueError(f"Conflict: {conflict}")

                    new_service = ServiceAssignment(
                        student_id=student_id,
                        service_type=service_type,
                        service_date=service_date,
                        start_datetime=start_dt,
                        end_datetime=end_dt,
                        created_by_user_id=current_user.id,
                    )
                    db.session.add(new_service)
                    log_student_action(
                        student_id,
                        "SERVICE_CREATED_BATCH",
                        f"Serviciu '{service_type}' adăugat: {start_dt.strftime('%d.%m %H:%M')} - {end_dt.strftime('%d.%m %H:%M')}.",
                    )

                else:
                    raise ValueError(f"Tip de acțiune invalid: {action_type}")

                added_count += 1

            except ValueError as e:
                skipped_count += 1
                conflict_details.append(
                    f"{student.nume} {student.prenume}: {e}"
                )
                continue  # Continue to the next student

        if added_count > 0:
            flash(
                f"{added_count} acțiuni au fost aplicate cu succes.", "success"
            )

        if skipped_count > 0:
            flash(
                f"{skipped_count} acțiuni nu au putut fi aplicate din cauza erorilor sau conflictelor.",
                "danger",
            )
            for detail in conflict_details:
                flash(detail, "warning")

        db.session.commit()

    except Exception as e:
        db.session.rollback()
        flash(
            f"Operațiunea a eșuat. O eroare neașteptată a avut loc: {e}",
            "danger",
        )

    return redirect(url_for("list_students"))


@app.route("/gradat/student/add", methods=["GET", "POST"])
@login_required
def add_student():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("home"))

    if request.method == "POST":
        form = request.form
        nume = form.get("nume", "").strip()
        prenume = form.get("prenume", "").strip()
        grad_militar = form.get("grad_militar", "").strip()
        id_unic_student_form = form.get("id_unic_student", "").strip() or None
        gender = form.get("gender")
        pluton = form.get("pluton", "").strip()
        companie = form.get("companie", "").strip()
        batalion = form.get("batalion", "").strip()
        is_platoon_graded_duty_val = "is_platoon_graded_duty" in request.form
        is_smt_val = "is_smt" in request.form  # Read SMT status
        exemption_details_val = (
            form.get("exemption_details", "").strip() or None
        )
        assigned_graded_platoon_val = (
            form.get("assigned_graded_platoon", "").strip() or None
        )

        if not all([nume, prenume, grad_militar, gender, companie, batalion]):
            flash(
                "Toate câmpurile marcate cu * (cu excepția plutonului pentru gradați) sunt obligatorii.",
                "warning",
            )
            return render_template(
                "add_edit_student.html",
                form_title="Adăugare Student Nou",
                student=None,
                genders=GENDERS,
                form_data=request.form,
            )

        if (
            id_unic_student_form
            and Student.query.filter_by(
                id_unic_student=id_unic_student_form
            ).first()
        ):
            flash(
                f"ID unic student '{id_unic_student_form}' există deja.",
                "warning",
            )
            return render_template(
                "add_edit_student.html",
                form_title="Adăugare Student Nou",
                student=None,
                genders=GENDERS,
                form_data=request.form,
            )

        if gender not in GENDERS:
            flash("Valoare invalidă pentru gen.", "warning")
            return render_template(
                "add_edit_student.html",
                form_title="Adăugare Student Nou",
                student=None,
                genders=GENDERS,
                form_data=request.form,
            )

        new_student = Student(
            nume=nume,
            prenume=prenume,
            grad_militar=grad_militar,
            id_unic_student=id_unic_student_form,
            gender=gender,
            pluton=pluton,
            companie=companie,
            batalion=batalion,
            is_platoon_graded_duty=is_platoon_graded_duty_val,
            assigned_graded_platoon=assigned_graded_platoon_val,
            is_smt=is_smt_val,
            exemption_details=exemption_details_val,
            created_by_user_id=current_user.id,
        )
        db.session.add(new_student)
        try:
            db.session.flush()  # Assign ID to new_student for logging
            log_action(
                "CREATE_STUDENT_SUCCESS",
                target_model_name="Student",
                target_id=new_student.id,
                details_after_dict=model_to_dict(new_student),
                description=f"User {current_user.username} added student {new_student.grad_militar} {new_student.nume} {new_student.prenume}.",
            )
            log_student_action(
                new_student.id,
                "STUDENT_CREATED",
                f"Studentul {new_student.grad_militar} {new_student.nume} {new_student.prenume} a fost adăugat în sistem.",
            )
            db.session.commit()
            flash(
                f"Studentul {new_student.grad_militar} {new_student.nume} {new_student.prenume} a fost adăugat!",
                "success",
            )
            return redirect(url_for("list_students"))
        except Exception as e:
            db.session.rollback()
            flash_msg = f"Eroare la salvarea studentului: {str(e)}"
            flash(flash_msg, "danger")
            try:
                # Construct dict of attempted data for logging
                attempted_data = {
                    "nume": nume,
                    "prenume": prenume,
                    "grad_militar": grad_militar,
                    "id_unic_student": id_unic_student_form,
                    "gender": gender,
                    "pluton": pluton,
                    "companie": companie,
                    "batalion": batalion,
                    "is_platoon_graded_duty": is_platoon_graded_duty_val,
                    "created_by_user_id": current_user.id,
                }
                log_action(
                    "CREATE_STUDENT_FAIL",
                    target_model_name="Student",
                    description=f"User {current_user.username} failed to add student. Error: {str(e)}",
                    details_after_dict=attempted_data,
                )  # Log attempted data in details_after
                db.session.commit()
            except Exception as log_e:
                app.logger.error(
                    f"CRITICAL: Failed to commit failure log for CREATE_STUDENT_FAIL: {str(log_e)}"
                )
            return render_template(
                "add_edit_student.html",
                form_title="Adăugare Student Nou",
                student=None,
                genders=GENDERS,
                form_data=request.form,
            )

    return render_template(
        "add_edit_student.html",
        form_title="Adăugare Student Nou",
        student=None,
        genders=GENDERS,
        form_data=None,
    )


@app.route("/gradat/students/edit/<int:student_id>", methods=["GET", "POST"])
@login_required
def edit_student(student_id):
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("home"))

    s_edit = Student.query.filter_by(
        id=student_id, created_by_user_id=current_user.id
    ).first_or_404()
    details_before_edit = model_to_dict(
        s_edit
    )  # Capture state before any changes

    if request.method == "POST":
        form = request.form
        s_edit.nume = form.get("nume", "").strip()
        s_edit.prenume = form.get("prenume", "").strip()
        s_edit.grad_militar = form.get("grad_militar", "").strip()
        s_edit.pluton = form.get("pluton", "").strip()
        s_edit.companie = form.get("companie", "").strip()
        s_edit.batalion = form.get("batalion", "").strip()
        s_edit.gender = form.get("gender")
        s_edit.is_platoon_graded_duty = (
            "is_platoon_graded_duty" in request.form
        )
        s_edit.assigned_graded_platoon = (
            form.get("assigned_graded_platoon", "").strip() or None
        )
        s_edit.is_smt = "is_smt" in request.form
        s_edit.exemption_details = (
            form.get("exemption_details", "").strip() or None
        )
        new_id_unic = form.get("id_unic_student", "").strip() or None

        if not all(
            [
                s_edit.nume,
                s_edit.prenume,
                s_edit.grad_militar,
                s_edit.companie,
                s_edit.batalion,
                s_edit.gender,
            ]
        ):
            flash(
                "Toate câmpurile marcate cu * (cu excepția plutonului pentru gradați) sunt obligatorii.",
                "warning",
            )
            return render_template(
                "add_edit_student.html",
                form_title=f"Editare Student: {details_before_edit.get('grad_militar','')} {details_before_edit.get('nume','')}",
                student=s_edit,
                genders=GENDERS,
                form_data=request.form,
            )

        if s_edit.gender not in GENDERS:
            flash("Valoare invalidă pentru gen.", "warning")
            return render_template(
                "add_edit_student.html",
                form_title=f"Editare Student: {details_before_edit.get('grad_militar','')} {details_before_edit.get('nume','')}",
                student=s_edit,
                genders=GENDERS,
                form_data=request.form,
            )

        if (
            new_id_unic
            and new_id_unic != s_edit.id_unic_student
            and Student.query.filter(
                Student.id_unic_student == new_id_unic, Student.id != s_edit.id
            ).first()
        ):
            flash(
                f"Alt student cu ID unic '{new_id_unic}' există deja.",
                "warning",
            )
            return render_template(
                "add_edit_student.html",
                form_title=f"Editare Student: {details_before_edit.get('grad_militar','')} {details_before_edit.get('nume','')}",
                student=s_edit,
                genders=GENDERS,
                form_data=request.form,
            )

        s_edit.id_unic_student = new_id_unic

        try:
            details_after_edit = model_to_dict(s_edit)
            # Logic to find what changed for the audit log
            changed_fields = []
            for key, new_value in details_after_edit.items():
                old_value = details_before_edit.get(key)
                if new_value != old_value:
                    changed_fields.append(
                        f"'{key}' from '{old_value}' to '{new_value}'"
                    )
            details_str = (
                "Date student actualizate. " + ", ".join(changed_fields)
                if changed_fields
                else "Nicio modificare detectată."
            )

            log_action(
                "UPDATE_STUDENT_SUCCESS",
                target_model_name="Student",
                target_id=s_edit.id,
                details_before_dict=details_before_edit,
                details_after_dict=details_after_edit,
                description=f"User {current_user.username} updated student {s_edit.grad_militar} {s_edit.nume} {s_edit.prenume}.",
            )
            log_student_action(s_edit.id, "STUDENT_UPDATED", details_str)
            db.session.commit()
            flash(
                f"Studentul {s_edit.grad_militar} {s_edit.nume} {s_edit.prenume} a fost actualizat!",
                "success",
            )
            return redirect(url_for("list_students"))
        except Exception as e:
            db.session.rollback()
            flash_msg = f"Eroare la actualizarea studentului: {str(e)}"
            flash(flash_msg, "danger")
            try:
                log_action(
                    "UPDATE_STUDENT_FAIL",
                    target_model_name="Student",
                    target_id=student_id,
                    details_before_dict=details_before_edit,  # Log original state
                    description=f"User {current_user.username} failed to update student ID {student_id}. Error: {str(e)}",
                )
                db.session.commit()
            except Exception as log_e:
                app.logger.error(
                    f"CRITICAL: Failed to commit failure log for UPDATE_STUDENT_FAIL: {str(log_e)}"
                )
            return render_template(
                "add_edit_student.html",
                form_title=f"Editare Student: {details_before_edit.get('grad_militar','')} {details_before_edit.get('nume','')} {details_before_edit.get('prenume','')}",
                student=s_edit,
                genders=GENDERS,
                form_data=request.form,
            )

    return render_template(
        "add_edit_student.html",
        form_title=f"Editare Student: {s_edit.grad_militar} {s_edit.nume} {s_edit.prenume}",
        student=s_edit,
        genders=GENDERS,
        form_data=s_edit,
    )


@app.route(
    "/gradat/students/bulk_import_page",
    methods=["GET", "POST"],
    endpoint="gradat_page_bulk_import_students",
)
@login_required
def gradat_page_bulk_import_students_func():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("list_students"))

    form_data_to_repopulate = None  # Used to repopulate textarea on error
    error_details_list_for_template = []
    processed_added_count = 0
    processed_error_count = 0

    if request.method == "POST":
        student_bulk_data = request.form.get("student_bulk_data", "").strip()
        form_data_to_repopulate = (
            request.form
        )  # Save form data for repopulation

        if not student_bulk_data:
            flash("Nu au fost furnizate date pentru import.", "warning")
            return render_template(
                "gradat_bulk_import_students_page.html",
                form_data=form_data_to_repopulate,
                error_details_list=None,
                added_count=0,
                error_count=0,
            )

        lines = student_bulk_data.splitlines()

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            parts = line.split()
            if (
                len(parts) < 7
            ):  # Grad Nume Prenume Gen Pluton Companie Batalion
                error_details_list_for_template.append(
                    f"Linia {i+1} ('{line[:50]}...'): Format incorect - prea puține câmpuri."
                )
                processed_error_count += 1
                continue

            try:
                batalion = parts[-1]
                companie = parts[-2]
                pluton = parts[-3]
                gender_input_original = parts[-4]
                gender_input_upper = gender_input_original.upper()

                gender_db_val = None
                if gender_input_upper == "M":
                    gender_db_val = "M"
                elif gender_input_upper == "F":
                    gender_db_val = "F"
                elif gender_input_upper in [g.upper() for g in GENDERS]:
                    gender_db_val = next(
                        g_val
                        for g_val in GENDERS
                        if g_val.upper() == gender_input_upper
                    )
                else:
                    error_details_list_for_template.append(
                        f"Linia {i+1} ('{line[:50]}...'): Gen '{gender_input_original}' invalid. Folosiți M, F sau Nespecificat."
                    )
                    processed_error_count += 1
                    continue

                name_rank_parts = parts[:-4]
                if (
                    len(name_rank_parts) < 3
                ):  # Need at least Grad, Nume, Prenume
                    error_details_list_for_template.append(
                        f"Linia {i+1} ('{line[:50]}...'): Format insuficient pentru Grad, Nume, Prenume."
                    )
                    processed_error_count += 1
                    continue

                prenume = name_rank_parts[-1]
                nume = name_rank_parts[-2]
                grad_militar = " ".join(name_rank_parts[:-2])

                if not all(
                    [grad_militar, nume, prenume, pluton, companie, batalion]
                ):
                    error_details_list_for_template.append(
                        f"Linia {i+1} ('{line[:50]}...'): Unul sau mai multe câmpuri obligatorii lipsesc după parsare."
                    )
                    processed_error_count += 1
                    continue

                existing_student_check = Student.query.filter_by(
                    nume=nume,
                    prenume=prenume,
                    grad_militar=grad_militar,
                    pluton=pluton,
                    companie=companie,
                    batalion=batalion,
                    created_by_user_id=current_user.id,
                ).first()

                if existing_student_check:
                    error_details_list_for_template.append(
                        f"Linia {i+1} ('{line[:50]}...'): Student similar există deja."
                    )
                    processed_error_count += 1
                    continue

                new_student = Student(
                    grad_militar=grad_militar,
                    nume=nume,
                    prenume=prenume,
                    gender=gender_db_val,
                    pluton=pluton,
                    companie=companie,
                    batalion=batalion,
                    created_by_user_id=current_user.id,
                    is_platoon_graded_duty=False,
                )
                db.session.add(new_student)
                processed_added_count += 1

            except IndexError:
                error_details_list_for_template.append(
                    f"Linia {i+1} ('{line[:50]}...'): Format incorect - eroare la extragerea câmpurilor."
                )
                processed_error_count += 1
                continue
            except Exception as e:
                error_details_list_for_template.append(
                    f"Linia {i+1} ('{line[:50]}...'): Eroare neașteptată - {str(e)}."
                )
                processed_error_count += 1
                db.session.rollback()
                continue

        if processed_added_count > 0:
            try:
                db.session.commit()
                flash(
                    f"{processed_added_count} studenți importați cu succes.",
                    "success",
                )
                # Log action for successful part
                log_action(
                    "BULK_STUDENT_IMPORT_PAGE_PARTIAL_SUCCESS",
                    description=f"User {current_user.username} bulk imported {processed_added_count} students. Errors: {processed_error_count}.",
                    details_after_dict={
                        "added": processed_added_count,
                        "errors": processed_error_count,
                    },
                )
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                flash(
                    f"Eroare la salvarea studenților în baza de date: {str(e)}",
                    "danger",
                )
                processed_error_count += processed_added_count
                processed_added_count = 0
                # Log action for DB failure
                log_action(
                    "BULK_STUDENT_IMPORT_PAGE_DB_FAIL",
                    description=f"User {current_user.username} bulk student import DB commit failed. Error: {str(e)}.",
                    details_after_dict={
                        "attempted_add": processed_added_count,
                        "initial_errors": processed_error_count,
                        "db_error": str(e),
                    },
                )
                db.session.commit()

        if processed_error_count > 0:
            # Don't flash individual errors here if we are re-rendering the page with error_details_list
            # flash(f'{processed_error_count} linii nu au putut fi procesate. Verificați detaliile.', 'warning')
            # The template will display error_details_list
            return render_template(
                "gradat_bulk_import_students_page.html",
                form_data=form_data_to_repopulate,
                error_details_list=error_details_list_for_template,
                added_count=processed_added_count,  # Show how many were added before errors
                error_count=processed_error_count,
            )

        # If no errors and some were added, redirect to list_students
        if processed_added_count > 0 and processed_error_count == 0:
            return redirect(url_for("list_students"))
        elif (
            processed_added_count == 0
            and processed_error_count == 0
            and student_bulk_data
        ):  # No data processed, no errors, but data was submitted
            flash(
                "Nicio linie validă de importat nu a fost găsită în datele furnizate.",
                "info",
            )
            # Fall through to render the page again, possibly with empty form_data if it was cleared

    # GET request or if POST had issues and needs re-render without specific error list for template
    return render_template(
        "gradat_bulk_import_students_page.html",
        form_data=form_data_to_repopulate,  # None for GET, or form data if POST failed early
        error_details_list=(
            error_details_list_for_template
            if processed_error_count > 0
            else None
        ),
        added_count=processed_added_count,
        error_count=processed_error_count,
    )


# Funcționalitatea 'Gradat Companie' a fost eliminată. Am șters ruta și funcția admin_toggle_company_grader_status.


@app.route("/gradat/delete_student/<int:student_id>", methods=["POST"])
@login_required
def delete_student(student_id):
    if current_user.role not in ["admin", "gradat"]:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("home"))
    student_to_delete = Student.query.get_or_404(student_id)
    if (
        current_user.role == "gradat"
        and student_to_delete.created_by_user_id != current_user.id
    ):
        flash("Nu puteți șterge studenți care nu vă sunt arondați.", "danger")
        return redirect(url_for("list_students"))

    if (
        current_user.role == "admin"
        and hasattr(student_to_delete, "creator")
        and student_to_delete.creator
        and student_to_delete.creator.username != current_user.username
    ):
        flash(
            f"Atenție: Ștergeți un student ({student_to_delete.nume} {student_to_delete.prenume}) care aparține gradatului {student_to_delete.creator.username}.",
            "warning",
        )

    details_before_delete = model_to_dict(student_to_delete)
    student_name_for_flash = f"{student_to_delete.grad_militar} {student_to_delete.nume} {student_to_delete.prenume}"
    student_id_for_log = student_to_delete.id  # Capture before deletion

    try:
        # Student-specific log must be created before the student is deleted, but committed after.
        log_student_action(
            student_id_for_log,
            "STUDENT_DELETED",
            f"Studentul {student_name_for_flash} a fost șters din sistem.",
        )
        db.session.delete(student_to_delete)
        log_action(
            "DELETE_STUDENT_SUCCESS",
            target_model_name="Student",
            target_id=student_id,  # Use student_id from param as object is deleted
            details_before_dict=details_before_delete,
            description=f"User {current_user.username} deleted student {student_name_for_flash} (ID: {student_id}).",
        )
        db.session.commit()
        flash(
            f"Studentul {student_name_for_flash} și toate datele asociate au fost șterse.",
            "success",
        )
    except Exception as e:
        db.session.rollback()
        flash_msg = f"Eroare la ștergerea studentului: {str(e)}"
        flash(flash_msg, "danger")
        try:
            log_action(
                "DELETE_STUDENT_FAIL",
                target_model_name="Student",
                target_id=student_id,
                details_before_dict=details_before_delete,
                description=f"User {current_user.username} failed to delete student {student_name_for_flash} (ID: {student_id}). Error: {str(e)}",
            )
            db.session.commit()
        except Exception as log_e:
            app.logger.error(
                f"CRITICAL: Failed to commit failure log for DELETE_STUDENT_FAIL: {str(log_e)}"
            )

    return redirect(url_for("list_students"))


# --- Rute pentru Permisii ---
@app.route("/gradat/permissions")
@login_required
def list_permissions():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    student_id_tuples = (
        db.session.query(Student.id)
        .filter_by(created_by_user_id=current_user.id)
        .all()
    )
    student_ids_managed_by_gradat = [sid for (sid,) in student_id_tuples]
    if not student_ids_managed_by_gradat:
        return render_template(
            "list_permissions.html",
            active_permissions=[],
            upcoming_permissions=[],
            past_permissions=[],
            title="Listă Permisii",
        )

    now = get_localized_now()  # Folosim ora localizată

    # Optimizare: O singură interogare către baza de date
    all_permissions = (
        Permission.query.options(
            joinedload(Permission.student), joinedload(Permission.creator)
        )
        .filter(Permission.student_id.in_(student_ids_managed_by_gradat))
        .order_by(Permission.start_datetime.desc())
        .all()
    )

    active_permissions = []
    upcoming_permissions = []
    past_permissions = []

    for p in all_permissions:
        if (
            p.status == "Aprobată"
            and p.end_datetime >= now
            and p.start_datetime <= now
        ):
            active_permissions.append(p)
        elif p.status == "Aprobată" and p.start_datetime > now:
            upcoming_permissions.append(p)
        else:
            past_permissions.append(p)

    # Sortarea se face în Python
    active_permissions.sort(key=lambda x: x.start_datetime)
    upcoming_permissions.sort(key=lambda x: x.start_datetime)
    past_permissions.sort(key=lambda x: x.end_datetime, reverse=True)

    # Limitarea permisiunilor trecute la 30
    past_permissions = past_permissions[:30]

    return render_template(
        "list_permissions.html",
        active_permissions=active_permissions,
        upcoming_permissions=upcoming_permissions,
        past_permissions=past_permissions,
        title="Listă Permisii",
    )


@app.route("/gradat/permission/add", methods=["GET", "POST"])
@app.route(
    "/gradat/permission/edit/<int:permission_id>", methods=["GET", "POST"]
)
@login_required
def add_edit_permission(permission_id=None):
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    form_title = "Adaugă Permisie Nouă"
    permission = None
    if permission_id:
        permission = Permission.query.get_or_404(permission_id)
        student_of_permission = Student.query.get(permission.student_id)
        if (
            not student_of_permission
            or student_of_permission.created_by_user_id != current_user.id
        ):
            flash("Acces neautorizat la această permisie.", "danger")
            return redirect(url_for("list_permissions"))
        form_title = f"Editare Permisie: {student_of_permission.grad_militar} {student_of_permission.nume} {student_of_permission.prenume}"
    students_managed = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .order_by(Student.nume)
        .all()
    )
    if request.method == "POST":
        student_id = request.form.get("student_id")
        start_datetime_str = request.form.get("start_datetime")
        end_datetime_str = request.form.get("end_datetime")
        destination = request.form.get("destination", "").strip()
        transport_mode = request.form.get("transport_mode", "").strip()
        reason = request.form.get("reason", "").strip()

        current_form_data_post = request.form  # For repopulating form on error

        if not student_id or not start_datetime_str or not end_datetime_str:
            flash(
                "Studentul, data de început și data de sfârșit sunt obligatorii.",
                "warning",
            )
            return render_template(
                "add_edit_permission.html",
                form_title=form_title,
                permission=permission,
                students=students_managed,
                form_data=current_form_data_post,
            )
        try:
            start_dt = datetime.strptime(start_datetime_str, "%Y-%m-%dT%H:%M")
            end_dt = datetime.strptime(end_datetime_str, "%Y-%m-%dT%H:%M")
        except ValueError:
            flash("Format dată/oră invalid.", "danger")
            return render_template(
                "add_edit_permission.html",
                form_title=form_title,
                permission=permission,
                students=students_managed,
                form_data=current_form_data_post,
            )

        if end_dt <= start_dt:
            flash(
                "Data de sfârșit trebuie să fie după data de început.",
                "warning",
            )
            return render_template(
                "add_edit_permission.html",
                form_title=form_title,
                permission=permission,
                students=students_managed,
                form_data=current_form_data_post,
            )

        student_to_check = db.session.get(
            Student, int(student_id)
        )  # Use db.session.get for clarity
        if (
            not student_to_check
            or student_to_check.created_by_user_id != current_user.id
        ):
            flash("Student invalid sau nu vă aparține.", "danger")
            return render_template(
                "add_edit_permission.html",
                form_title=form_title,
                permission=None,
                students=students_managed,
                form_data=current_form_data_post,
            )

        # Conflict checking (remains largely the same)
        conflicting_service = ServiceAssignment.query.filter(
            ServiceAssignment.student_id == student_id,
            ServiceAssignment.service_type == "Intervenție",
            ServiceAssignment.start_datetime < end_dt,
            ServiceAssignment.end_datetime > start_dt,
        ).first()
        if conflicting_service:
            flash(
                f'Studentul {student_to_check.nume} {student_to_check.prenume} este în serviciu de "Intervenție" și nu poate primi permisie.',
                "danger",
            )
            return render_template(
                "add_edit_permission.html",
                form_title=form_title,
                permission=permission,
                students=students_managed,
                form_data=current_form_data_post,
            )

        general_conflict_msg = check_leave_conflict(
            student_id,
            start_dt,
            end_dt,
            "permission",
            permission.id if permission else None,
        )
        if general_conflict_msg:
            flash(
                f"Conflict detectat: Studentul are deja {general_conflict_msg}.",
                "danger",
            )
            return render_template(
                "add_edit_permission.html",
                form_title=form_title,
                permission=permission,
                students=students_managed,
                form_data=current_form_data_post,
            )

        action_description_prefix = f"User {current_user.username}"
        log_details_before = None
        original_student_name_for_log = (
            student_to_check.nume + " " + student_to_check.prenume
        )

        if permission:  # Editing existing permission
            log_details_before = model_to_dict(permission)
            permission.student_id = int(student_id)
            permission.start_datetime = start_dt
            permission.end_datetime = end_dt
            permission.destination = destination
            permission.transport_mode = transport_mode
            permission.reason = reason
            action_type = "UPDATE_PERMISSION_SUCCESS"
            flash_msg_text = "Permisie actualizată cu succes!"
            log_description = f"{action_description_prefix} updated permission for {original_student_name_for_log}."
            log_student_action(
                permission.student_id,
                "PERMISSION_UPDATED",
                f"Permisie actualizată: {start_dt.strftime('%d.%m %H:%M')} - {end_dt.strftime('%d.%m %H:%M')}. Motiv: {reason or 'N/A'}",
            )
        else:  # Adding new permission
            permission = Permission(
                student_id=int(student_id),
                start_datetime=start_dt,
                end_datetime=end_dt,
                destination=destination,
                transport_mode=transport_mode,
                reason=reason,
                status="Aprobată",
                created_by_user_id=current_user.id,
            )
            db.session.add(permission)
            action_type = "CREATE_PERMISSION_SUCCESS"
            flash_msg_text = "Permisie adăugată cu succes!"
            log_description = f"{action_description_prefix} created permission for {original_student_name_for_log}."
            log_student_action(
                permission.student_id,
                "PERMISSION_CREATED",
                f"Permisie adăugată: {start_dt.strftime('%d.%m %H:%M')} - {end_dt.strftime('%d.%m %H:%M')}. Motiv: {reason or 'N/A'}",
            )

        try:
            db.session.flush()  # Ensure permission object has ID if new, and updates are in session
            log_details_after = model_to_dict(permission)
            log_action(
                action_type,
                target_model_name="Permission",
                target_id=permission.id,
                details_before_dict=log_details_before,
                details_after_dict=log_details_after,
                description=log_description,
            )
            db.session.commit()
            flash(flash_msg_text, "success")
        except Exception as e:
            db.session.rollback()
            flash_msg_fail = f"Eroare la salvarea permisiei: {str(e)}"
            flash(flash_msg_fail, "danger")
            try:
                fail_action_type = (
                    "UPDATE_PERMISSION_FAIL"
                    if permission_id
                    else "CREATE_PERMISSION_FAIL"
                )
                # For create fail, permission object might not have an ID yet if flush failed.
                # If it was an edit, permission.id is valid.
                target_id_for_fail_log = (
                    permission.id if permission and permission.id else None
                )
                # Log attempted data for create fail
                attempted_data_on_fail = (
                    model_to_dict(permission)
                    if permission
                    else current_form_data_post
                )

                log_action(
                    fail_action_type,
                    target_model_name="Permission",
                    target_id=target_id_for_fail_log,
                    details_before_dict=(
                        log_details_before if permission_id else None
                    ),
                    details_after_dict=(
                        attempted_data_on_fail
                        if not permission_id
                        else model_to_dict(permission)
                    ),  # Log current state of 'permission' on edit fail
                    description=f"{action_description_prefix} failed to {action_type.split('_')[0].lower()} permission for {original_student_name_for_log}. Error: {str(e)}",
                )
                db.session.commit()
            except Exception as log_e:
                app.logger.error(
                    f"CRITICAL: Failed to commit failure log for {fail_action_type}: {str(log_e)}"
                )
            return render_template(
                "add_edit_permission.html",
                form_title=form_title,
                permission=permission if permission_id else None,
                students=students_managed,
                form_data=current_form_data_post,
            )
        return redirect(url_for("list_permissions"))

    # GET request handling
    form_data_on_get = (
        {}
    )  # Renamed from data_to_populate_form_with for clarity
    if permission:  # Editing existing permission
        form_data_on_get = {
            "student_id": str(permission.student_id),
            "start_datetime": (
                permission.start_datetime.strftime("%Y-%m-%dT%H:%M")
                if permission.start_datetime
                else ""
            ),
            "end_datetime": (
                permission.end_datetime.strftime("%Y-%m-%dT%H:%M")
                if permission.end_datetime
                else ""
            ),
            "destination": permission.destination or "",
            "transport_mode": permission.transport_mode or "",
            "reason": permission.reason or "",
        }
    else:
        # Pre-fill from query parameters to allow quick \"alte scutiri\" creation from student management
        prefill_student_id = request.args.get("student_id", type=int)
        prefill_reason = request.args.get("prefill_reason", type=str)
        if prefill_student_id:
            form_data_on_get["student_id"] = str(prefill_student_id)
        if prefill_reason:
            form_data_on_get["reason"] = prefill_reason

    # If it's a POST request that failed validation and re-rendered, current_form_data_post (passed as form_data) will be used by template.
    # If it's a fresh GET for 'add', form_data_on_get may contain prefilled values from query params.

    return render_template(
        "add_edit_permission.html",
        form_title=form_title,
        permission=permission,  # Pass the permission object itself for the template
        students=students_managed,
        form_data=(
            form_data_on_get
            if request.method == "GET"
            else request.form if request.method == "POST" else {}
        ),
    )


def find_student_for_bulk_import(name_line, students_or_user_id):
    """
    Optimized helper function to find a student.
    It can accept a pre-fetched list of students OR a user_id to fetch them.
    """
    if isinstance(students_or_user_id, int):
        students_managed = Student.query.filter_by(
            created_by_user_id=students_or_user_id
        ).all()
    else:
        students_managed = students_or_user_id
    name_line_norm = unidecode(name_line.lower().strip())
    if not name_line_norm:
        return None, "Linie nume goală."

    if not students_managed:
        return None, "Lista de studenți pre-încărcată este goală."

    # Try to extract rank for more precise matching
    parsed_grad_bulk = None
    student_name_str_bulk = (
        name_line  # Default to full line if no rank pattern matches
    )
    for pattern in KNOWN_RANK_PATTERNS:  # KNOWN_RANK_PATTERNS is global
        match = pattern.match(name_line)
        if match:
            parsed_grad_bulk = match.group(0).strip()
            student_name_str_bulk = pattern.sub("", name_line).strip()
            break

    # Clean up trailing time-like patterns from student_name_str_bulk
    student_name_str_bulk = re.sub(
        r"\s+\d{1,2}:\d{2}(-\d{1,2}:\d{2})?$", "", student_name_str_bulk
    ).strip()

    normalized_search_name_bulk = unidecode(student_name_str_bulk.lower())

    matched_students = []
    # Pass 1: Exact match on normalized name and rank (if rank parsed)
    for s in students_managed:
        s_fullname_norm = unidecode(f"{s.nume} {s.prenume}".lower())
        s_grad_norm = unidecode(s.grad_militar.lower())

        if parsed_grad_bulk:
            parsed_grad_bulk_norm = unidecode(parsed_grad_bulk.lower())
            if (
                normalized_search_name_bulk == s_fullname_norm
                and parsed_grad_bulk_norm == s_grad_norm
            ):
                matched_students.append(s)
        else:  # No rank in input, try to match only by name
            if normalized_search_name_bulk == s_fullname_norm:
                matched_students.append(s)

    if len(matched_students) == 1:
        return matched_students[0], None
    if len(matched_students) > 1:
        return None, f"Potriviri multiple exacte pentru '{name_line}'"

    # Pass 2: More lenient matching (e.g., input name is part of DB name or vice-versa, rank is similar)
    # This is more complex and might lead to false positives, keeping it simpler for now.
    # For now, if exact match failed, we rely on the user providing very precise names.

    # Pass 2: Lenient match on name parts (Nume Prenume vs Prenume Nume), rank similarity
    # This pass tries to find a unique student if the exact match (Pass 1) failed.
    if not matched_students:  # Only if Pass 1 found nothing
        potential_matches = []
        for s in students_managed:
            s_fullname_norm = unidecode(f"{s.nume} {s.prenume}".lower())
            s_fullname_reversed_norm = unidecode(
                f"{s.prenume} {s.nume}".lower()
            )
            s_grad_norm = unidecode(s.grad_militar.lower())

            # Check if normalized_search_name_bulk (name from input without rank)
            # is a substring of student's full name in either order
            name_match_direct = normalized_search_name_bulk in s_fullname_norm
            name_match_reversed = (
                normalized_search_name_bulk in s_fullname_reversed_norm
            )

            if name_match_direct or name_match_reversed:
                if parsed_grad_bulk:  # If rank was parsed from input
                    parsed_grad_bulk_norm = unidecode(parsed_grad_bulk.lower())
                    # Check for rank similarity (e.g., "sdt" vs "sdt.", "cap" vs "cap.")
                    # A simple check: if parsed rank is a substring of DB rank or vice-versa
                    if (
                        parsed_grad_bulk_norm in s_grad_norm
                        or s_grad_norm in parsed_grad_bulk_norm
                    ):
                        potential_matches.append(s)
                    # else: name matches, but rank doesn't, so it's not a good potential match
                else:  # No rank in input, name match is enough to be a potential candidate
                    potential_matches.append(s)

        if len(potential_matches) == 1:
            return (
                potential_matches[0],
                None,
            )  # Found a single, reasonably good match
        elif len(potential_matches) > 1:
            # Construct a more informative error message for multiple lenient matches
            student_names_found = [
                f"{s.grad_militar} {s.nume} {s.prenume}"
                for s in potential_matches
            ]
            return (
                None,
                f"Potriviri multiple pentru '{name_line}': {', '.join(student_names_found)}. Clarificați gradul sau numele.",
            )

    # If still no single match after both passes
    if (
        not matched_students and not potential_matches
    ):  # Check potential_matches from Pass 2
        return (
            None,
            f"Studentul '{name_line}' nu a fost găsit. Verificați numele și gradul.",
        )
    elif len(matched_students) > 1:  # From Pass 1
        return (
            None,
            f"Potriviri multiple exacte pentru '{name_line}'. Clarificați.",
        )

    # Fallback if logic is somehow bypassed, though it shouldn't be.
    return (
        None,
        f"Studentul '{name_line}' nu a fost găsit sau potrivirea este ambiguă (situație neașteptată).",
    )


@app.route(
    "/gradat/permissions/bulk_import",
    methods=["POST"],
    endpoint="gradat_bulk_import_permissions",
)
@login_required
def bulk_import_permissions():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("list_permissions"))

    permission_bulk_data = request.form.get("permission_bulk_data", "").strip()
    if not permission_bulk_data:
        flash("Nu au fost furnizate date pentru import.", "warning")
        return redirect(url_for("list_permissions"))

    lines = permission_bulk_data.splitlines()
    added_count = 0
    error_count = 0
    error_details = []

    # Optimization: Fetch all managed students once
    students_managed = Student.query.filter_by(
        created_by_user_id=current_user.id
    ).all()

    i = 0
    while i < len(lines):
        name_line = lines[i].strip()
        if not name_line:  # Skip empty lines used as separators
            i += 1
            continue

        current_block_lines = []
        temp_i = i
        while temp_i < len(lines) and lines[temp_i].strip():
            current_block_lines.append(lines[temp_i].strip())
            temp_i += 1

        num_actual_lines_for_entry = len(current_block_lines)

        i = temp_i
        if temp_i < len(lines) and not lines[temp_i].strip():
            i += 1

        if num_actual_lines_for_entry < 3:
            if num_actual_lines_for_entry > 0:
                error_details.append(
                    f"Intrare incompletă începând cu '{current_block_lines[0]}'. Necesită cel puțin Nume, Interval, Destinație."
                )
                error_count += 1
            continue

        name_line = current_block_lines[0]
        datetime_line = current_block_lines[1]
        destination_line = current_block_lines[2]
        transport_mode_line = (
            current_block_lines[3] if num_actual_lines_for_entry > 3 else ""
        )
        reason_car_plate_line = (
            current_block_lines[4] if num_actual_lines_for_entry > 4 else ""
        )

        # Use the pre-fetched list of students
        student_obj, student_error = find_student_for_bulk_import(
            name_line, students_managed
        )
        if student_error:
            error_details.append(f"Linia '{name_line}': {student_error}")
            error_count += 1
            continue

        try:
            # More robust regex for date/time parsing
            dt_match = re.search(
                r"(\d{1,2}[./-]\d{1,2}[./-]\d{4})\s+(\d{1,2}:\d{2})\s*-\s*(?:(\d{1,2}[./-]\d{1,2}[./-]\d{4})\s+)?(\d{1,2}:\d{2})",
                datetime_line,
            )
            if not dt_match:
                raise ValueError(
                    "Formatul liniei de dată/timp este invalid. Folosiți 'DD.MM.YYYY HH:MM - DD.MM.YYYY HH:MM' sau 'DD.MM.YYYY HH:MM - HH:MM'."
                )

            start_date_str, start_time_str, end_date_str_opt, end_time_str = (
                dt_match.groups()
            )

            # Normalize date separator
            start_date_str = re.sub(r"[./-]", ".", start_date_str)
            if end_date_str_opt:
                end_date_str_opt = re.sub(r"[./-]", ".", end_date_str_opt)

            start_dt = datetime.strptime(
                f"{start_date_str} {start_time_str}", "%d.%m.%Y %H:%M"
            )

            if end_date_str_opt:
                end_dt = datetime.strptime(
                    f"{end_date_str_opt} {end_time_str}", "%d.%m.%Y %H:%M"
                )
            else:
                end_time_obj_parsed = datetime.strptime(
                    end_time_str, "%H:%M"
                ).time()
                end_date_assumed = start_dt.date()
                if end_time_obj_parsed < start_dt.time():
                    end_date_assumed += timedelta(days=1)
                end_dt = datetime.combine(
                    end_date_assumed, end_time_obj_parsed
                )

            if end_dt <= start_dt:
                raise ValueError(
                    "Data de sfârșit trebuie să fie după data de început."
                )

        except ValueError as ve:
            error_details.append(
                f"Student '{name_line}': Eroare format dată/oră în '{datetime_line}' - {str(ve)}"
            )
            error_count += 1
            continue

        parsed_destination = destination_line.strip()
        parsed_transport_mode = (
            transport_mode_line.strip() if transport_mode_line else None
        )
        parsed_reason = (
            reason_car_plate_line.strip() if reason_car_plate_line else None
        )

        conflict = check_leave_conflict(
            student_obj.id, start_dt, end_dt, leave_type="permission"
        )
        if conflict:
            error_details.append(
                f"Student '{name_line}': Conflict - {conflict}."
            )
            error_count += 1
            continue

        new_permission = Permission(
            student_id=student_obj.id,
            start_datetime=start_dt,
            end_datetime=end_dt,
            destination=parsed_destination,
            transport_mode=parsed_transport_mode,
            reason=parsed_reason,
            status="Aprobată",
            created_by_user_id=current_user.id,
        )
        db.session.add(new_permission)
        log_student_action(
            student_obj.id,
            "PERMISSION_CREATED_BULK",
            f"Permisie adăugată prin import text: {start_dt.strftime('%d.%m %H:%M')} - {end_dt.strftime('%d.%m %H:%M')}.",
        )
        added_count += 1

    if added_count > 0:
        try:
            db.session.commit()
            flash(
                f"{added_count} permisii au fost adăugate cu succes.",
                "success",
            )
        except Exception as e:
            db.session.rollback()
            flash(
                f"Eroare la salvarea permisiilor în baza de date: {str(e)}",
                "danger",
            )
            error_count += added_count
            added_count = 0

    if error_count > 0:
        flash(
            f"{error_count} intrări nu au putut fi procesate sau au generat erori.",
            "danger",
        )
        # Create a more detailed flash message
        error_summary = "Detalii erori:<br>" + "<br>".join(
            error_details[:5]
        )  # Show first 5 errors
        if len(error_details) > 5:
            error_summary += "<br>...și altele."
        flash(error_summary, "warning")

    if added_count > 0 or error_count > 0:
        log_action(
            "BULK_IMPORT_PERMISSIONS_COMPLETED",
            description=f"User {current_user.username} ran bulk permission import. Added: {added_count}, Errors: {error_count}.",
            details_after_dict={
                "added_count": added_count,
                "error_count": error_count,
                "error_details": error_details[:5],
            },
        )
        db.session.commit()

    return redirect(url_for("list_permissions"))


@app.route(
    "/gradat/permissions/export_word",
    endpoint="gradat_export_permissions_word",
)
@login_required
def export_permissions_word():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_id_tuples = (
        db.session.query(Student.id)
        .filter_by(created_by_user_id=current_user.id)
        .all()
    )
    student_ids_managed_by_gradat = [s[0] for s in student_id_tuples]

    if not student_ids_managed_by_gradat:
        flash("Nu aveți studenți pentru a exporta permisii.", "info")
        return redirect(url_for("list_permissions"))

    now = get_localized_now()  # Folosim ora localizată
    # Fetch active and upcoming permissions only for the export
    permissions_to_export = (
        Permission.query.options(joinedload(Permission.student))
        .filter(
            Permission.student_id.in_(student_ids_managed_by_gradat),
            Permission.status == "Aprobată",
            Permission.end_datetime >= now,  # Active or upcoming
        )
        .join(Student)
        .order_by(Permission.start_datetime, Student.nume, Student.prenume)
        .all()
    )  # Initial sort by start_datetime

    if not permissions_to_export:
        flash("Nicio permisie activă sau viitoare de exportat.", "info")
        return redirect(url_for("list_permissions"))

    # Group permissions by period (start_datetime, end_datetime)
    # Using naive datetime for grouping keys as they come from DB
    grouped_permissions = {}
    for p in permissions_to_export:
        period_key = (p.start_datetime, p.end_datetime)
        if period_key not in grouped_permissions:
            grouped_permissions[period_key] = []
        grouped_permissions[period_key].append(p)

    # Helper function to get day name in Romanian
    def get_day_name_ro(date_obj):
        days = [
            "Luni",
            "Marți",
            "Miercuri",
            "Joi",
            "Vineri",
            "Sâmbătă",
            "Duminică",
        ]
        return days[date_obj.weekday()]

    # Custom sorting for periods
    def get_period_sort_key(period_item):
        period_key, _ = period_item  # period_key is (start_dt, end_dt)
        start_dt = period_key[0]
        end_dt = period_key[1]

        # Make datetimes timezone-aware for correct weekday calculation if they are naive
        # Assuming EUROPE_BUCHAREST as the reference timezone
        start_dt_aware = (
            EUROPE_BUCHAREST.localize(start_dt)
            if start_dt.tzinfo is None
            else start_dt.astimezone(EUROPE_BUCHAREST)
        )
        end_dt_aware = (
            EUROPE_BUCHAREST.localize(end_dt)
            if end_dt.tzinfo is None
            else end_dt.astimezone(EUROPE_BUCHAREST)
        )

        start_day_ro = get_day_name_ro(
            start_dt_aware
        )  # Luni, Marti, ..., Duminica
        end_day_ro = get_day_name_ro(end_dt_aware)

        # Define sort order values
        # Lower value means earlier in sort
        # Joi (Thursday) = 3, Vineri (Friday) = 4
        # Duminica (Sunday) = 6, Luni (Monday) = 0 (in weekday())

        # Priority 1: Thursday starts
        if start_day_ro == "Joi":
            if end_day_ro == "Duminică":
                return (1, start_dt)  # Joi - Duminica
            if end_day_ro == "Luni":
                return (2, start_dt)  # Joi - Luni
            return (3, start_dt)  # Other Joi starts (fallback)
        # Priority 2: Friday starts
        elif start_day_ro == "Vineri":
            if end_day_ro == "Duminică":
                return (4, start_dt)  # Vineri - Duminica
            if end_day_ro == "Luni":
                return (5, start_dt)  # Vineri - Luni
            return (6, start_dt)  # Other Vineri starts (fallback)
        # Fallback for other start days (sort by start_datetime)
        return (7, start_dt)

    sorted_grouped_permissions = sorted(
        grouped_permissions.items(), key=get_period_sort_key
    )

    document = Document()
    # General document heading (optional, could be removed if each table has a full title)
    # General document heading (optional, could be removed if each table has a full title)
    # document.add_heading('Raport Permisii Studenți', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # User and date info
    user_info_text = f"Raport generat de: {current_user.username}\nData generării: {get_localized_now().strftime('%d-%m-%Y %H:%M')}"
    p_user_info = document.add_paragraph()
    p_user_info.add_run(user_info_text).italic = True
    p_user_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()  # Spacer

    # Column titles for each table (Period column is removed)
    column_titles = [
        "Nr. crt.",
        "Grad",
        "Nume și Prenume",
        "Grupa",
        "Localitate",
        "Transport",
    ]
    # New column widths for 6 columns
    new_widths = {
        0: Inches(0.4),  # Nr. crt.
        1: Inches(0.8),  # Grad
        2: Inches(2.0),  # Nume și Prenume
        3: Inches(0.8),  # Grupa
        4: Inches(1.5),  # Localitate
        5: Inches(1.5),  # Transport
    }  # Total width approx 7.0 inches

    for period_key, permissions_in_period in sorted_grouped_permissions:
        start_dt_period = (
            EUROPE_BUCHAREST.localize(period_key[0])
            if period_key[0].tzinfo is None
            else period_key[0].astimezone(EUROPE_BUCHAREST)
        )
        end_dt_period = (
            EUROPE_BUCHAREST.localize(period_key[1])
            if period_key[1].tzinfo is None
            else period_key[1].astimezone(EUROPE_BUCHAREST)
        )

        # Format period string for the title
        # Example: Joi, 25.07.2024 (14:00) - Duminică, 28.07.2024 (22:00)
        period_title_str = (
            f"{get_day_name_ro(start_dt_period)}, {start_dt_period.strftime('%d.%m.%Y (%H:%M)')} - "
            f"{get_day_name_ro(end_dt_period)}, {end_dt_period.strftime('%d.%m.%Y (%H:%M)')}"
        )

        document.add_heading(
            f"Tabel nominal cu studenții care pleacă in permisie în perioada {period_title_str}",
            level=2,
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.add_table(
            rows=1, cols=len(column_titles)
        )  # Create table with 6 columns
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        for i, title in enumerate(column_titles):
            hdr_cells[i].text = title
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sort permissions within this period group by student name for consistency
        permissions_in_period.sort(
            key=lambda p: (p.student.nume, p.student.prenume)
        )

        for idx, p_item in enumerate(permissions_in_period):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx + 1)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[1].text = p_item.student.grad_militar
            row_cells[2].text = (
                f"{p_item.student.nume} {p_item.student.prenume}"  # Combined name
            )

            row_cells[3].text = p_item.student.pluton  # Grupa/Pluton
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[4].text = (
                p_item.destination if p_item.destination else "-"
            )
            row_cells[5].text = (
                p_item.transport_mode if p_item.transport_mode else "-"
            )

        # Apply column widths to the current table
        for col_idx, width_val in new_widths.items():
            for row in table.rows:  # Apply to all rows including header
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = width_val

        document.add_paragraph()  # Add some space between tables

    # Change font for the whole document (optional)
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    f = io.BytesIO()
    document.save(f)
    f.seek(0)

    filename = f"Raport_Permisii_{current_user.username}_{date.today().strftime('%Y%m%d')}.docx"

    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/gradat/permission/cancel/<int:permission_id>", methods=["POST"])
@login_required
def cancel_permission(permission_id):
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    permission = Permission.query.get_or_404(permission_id)
    student_of_permission = Student.query.get(permission.student_id)
    if (
        not student_of_permission
        or student_of_permission.created_by_user_id != current_user.id
    ):
        flash("Nu aveți permisiunea să anulați această permisie.", "danger")
        return redirect(url_for("list_permissions"))
    if permission.status == "Aprobată":
        permission.status = "Anulată"
        try:
            db.session.commit()
            flash(
                f"Permisia pentru {student_of_permission.nume} {student_of_permission.prenume} a fost anulată.",
                "success",
            )
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la anularea permisiei: {str(e)}", "danger")
    else:
        flash(
            'Această permisie nu poate fi anulată (statusul curent nu este "Aprobată").',
            "warning",
        )
    return redirect(url_for("list_permissions"))


@app.route("/gradat/permissions/delete/<int:permission_id>", methods=["POST"])
@app.route("/admin/permissions/delete/<int:permission_id>", methods=["POST"])
@login_required
def delete_permission(permission_id):
    permission_to_delete = db.session.get(Permission, permission_id)
    if not permission_to_delete:
        flash("Permisia nu a fost găsită.", "danger")
        return redirect(
            url_for("list_permissions")
            if current_user.role == "gradat"
            else url_for("admin_dashboard_route")
        )  # Sau o pagină admin relevantă

    student_owner = db.session.get(Student, permission_to_delete.student_id)

    if current_user.role == "gradat":
        if (
            not student_owner
            or student_owner.created_by_user_id != current_user.id
        ):
            flash(
                "Nu aveți permisiunea să ștergeți această permisie.", "danger"
            )
            return redirect(url_for("list_permissions"))
        redirect_url = url_for("list_permissions")
    elif current_user.role == "admin":
        # Admin poate șterge orice permisie, dar poate afișăm un warning dacă aparține altui gradat
        if (
            student_owner
            and student_owner.creator
            and student_owner.creator.username != current_user.username
        ):  # Presupunând că admin nu e creator direct
            flash(
                f"Atenție: Ștergeți o permisie pentru studentul {student_owner.nume} {student_owner.prenume}, gestionat de {student_owner.creator.username}.",
                "warning",
            )
        redirect_url = request.referrer or url_for(
            "admin_dashboard_route"
        )  # sau o listă de permisii admin, dacă există
    else:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_name_for_flash = (
        f"{student_owner.grad_militar} {student_owner.nume} {student_owner.prenume}"
        if student_owner
        else "N/A"
    )
    permission_details_for_flash = f"din {permission_to_delete.start_datetime.strftime('%d.%m.%Y %H:%M')} până în {permission_to_delete.end_datetime.strftime('%d.%m.%Y %H:%M')}"
    student_id_for_log = permission_to_delete.student_id

    details_before_delete = model_to_dict(permission_to_delete)
    try:
        log_student_action(
            student_id_for_log,
            "PERMISSION_DELETED",
            f"Permisia ({permission_details_for_flash}) a fost ștearsă.",
        )
        db.session.delete(permission_to_delete)
        log_action(
            "DELETE_PERMISSION_SUCCESS",
            target_model_name="Permission",
            target_id=permission_id,
            details_before_dict=details_before_delete,
            description=f"User {current_user.username} deleted permission for student {student_name_for_flash} (ID: {permission_id}) details: {permission_details_for_flash}.",
        )
        db.session.commit()
        flash(
            f"Permisia pentru {student_name_for_flash} ({permission_details_for_flash}) a fost ștearsă.",
            "success",
        )
    except Exception as e:
        db.session.rollback()
        flash_msg = f"Eroare la ștergerea permisiei: {str(e)}"
        flash(flash_msg, "danger")
        try:
            log_action(
                "DELETE_PERMISSION_FAIL",
                target_model_name="Permission",
                target_id=permission_id,
                details_before_dict=details_before_delete,
                description=f"User {current_user.username} failed to delete permission for {student_name_for_flash} (ID: {permission_id}). Error: {str(e)}",
            )
            db.session.commit()
        except Exception as log_e:
            app.logger.error(
                f"CRITICAL: Failed to commit failure log for DELETE_PERMISSION_FAIL: {str(log_e)}"
            )
    return redirect(redirect_url)


# --- Rute pentru Învoiri Zilnice ---
@app.route("/gradat/daily_leaves")
@login_required
def list_daily_leaves():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    student_id_tuples = (
        db.session.query(Student.id)
        .filter_by(created_by_user_id=current_user.id)
        .all()
    )
    student_ids = [s[0] for s in student_id_tuples]
    today_string_for_form = (
        get_localized_now().date().strftime("%Y-%m-%d")
    )  # Folosim data localizată
    if not student_ids:
        return render_template(
            "list_daily_leaves.html",
            active_leaves=[],
            upcoming_leaves=[],
            past_leaves=[],
            title="Listă Învoiri Zilnice",
            today_str=today_string_for_form,
        )

    all_relevant_leaves = (
        DailyLeave.query.options(
            joinedload(DailyLeave.student), joinedload(DailyLeave.creator)
        )
        .filter(DailyLeave.student_id.in_(student_ids))
        .order_by(DailyLeave.leave_date.desc(), DailyLeave.start_time.desc())
        .all()
    )
    active_leaves = []
    upcoming_leaves = []
    past_leaves = []
    for leave in all_relevant_leaves:
        if leave.status == "Anulată":
            past_leaves.append(leave)
        elif leave.is_active:
            active_leaves.append(leave)
        elif leave.is_upcoming:
            upcoming_leaves.append(leave)
        elif leave.is_past:
            past_leaves.append(leave)
    active_leaves.sort(key=lambda x: (x.leave_date, x.start_time))
    upcoming_leaves.sort(key=lambda x: (x.leave_date, x.start_time))
    past_leaves = past_leaves[:50]
    return render_template(
        "list_daily_leaves.html",
        active_leaves=active_leaves,
        upcoming_leaves=upcoming_leaves,
        past_leaves=past_leaves,
        title="Listă Învoiri Zilnice",
        today_str=today_string_for_form,
    )


@app.route("/gradat/daily_leave/add", methods=["GET", "POST"])
@app.route("/gradat/daily_leave/edit/<int:leave_id>", methods=["GET", "POST"])
@login_required
def add_edit_daily_leave(leave_id=None):
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    form_title = "Adaugă Învoire Zilnică"
    daily_leave = None
    today_string = (
        get_localized_now().date().strftime("%Y-%m-%d")
    )  # Folosim data localizată
    if leave_id:
        daily_leave = DailyLeave.query.get_or_404(leave_id)
        student_of_leave = Student.query.get(daily_leave.student_id)
        if (
            not student_of_leave
            or student_of_leave.created_by_user_id != current_user.id
        ):
            flash("Acces neautorizat la această învoire.", "danger")
            return redirect(url_for("list_daily_leaves"))
        form_title = f"Editare Învoire Zilnică: {student_of_leave.grad_militar} {student_of_leave.nume} {student_of_leave.prenume}"
    students_managed = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .order_by(Student.nume)
        .all()
    )
    if request.method == "POST":
        student_id = request.form.get("student_id")
        leave_date_str = request.form.get("leave_date")
        start_time_str = request.form.get("start_time")
        end_time_str = request.form.get("end_time")
        reason = request.form.get("reason", "").strip()
        current_form_data_post = request.form
        if not all([student_id, leave_date_str, start_time_str, end_time_str]):
            flash(
                "Toate câmpurile (student, dată, oră început, oră sfârșit) sunt obligatorii.",
                "warning",
            )
            return render_template(
                "add_edit_daily_leave.html",
                form_title=form_title,
                daily_leave=daily_leave,
                students=students_managed,
                today_str=today_string,
                form_data=current_form_data_post,
            )
        try:
            leave_date_obj = datetime.strptime(
                leave_date_str, "%Y-%m-%d"
            ).date()
            start_time_obj = datetime.strptime(start_time_str, "%H:%M").time()
            end_time_obj = datetime.strptime(end_time_str, "%H:%M").time()
        except ValueError:
            flash("Format dată sau oră invalid.", "danger")
            return render_template(
                "add_edit_daily_leave.html",
                form_title=form_title,
                daily_leave=daily_leave,
                students=students_managed,
                today_str=today_string,
                form_data=current_form_data_post,
            )
        is_valid_day, day_message = validate_daily_leave_times(
            start_time_obj, end_time_obj, leave_date_obj
        )
        if not is_valid_day:
            flash(day_message, "danger")
            return render_template(
                "add_edit_daily_leave.html",
                form_title=form_title,
                daily_leave=daily_leave,
                students=students_managed,
                today_str=today_string,
                form_data=current_form_data_post,
            )

        # Removed the restrictive is_in_program and is_out_program checks.
        # The primary validation will be that end_datetime is after start_datetime.

        start_dt = datetime.combine(leave_date_obj, start_time_obj)
        effective_end_date = leave_date_obj
        # Determine if end_time implies the next day
        if (
            end_time_obj < start_time_obj
        ):  # This condition means it spans midnight
            effective_end_date += timedelta(days=1)
        end_dt = datetime.combine(effective_end_date, end_time_obj)

        if end_dt <= start_dt:
            flash(
                "Data/ora de sfârșit trebuie să fie după data/ora de început, chiar și când trece în ziua următoare.",
                "warning",
            )
            return render_template(
                "add_edit_daily_leave.html",
                form_title=form_title,
                daily_leave=daily_leave,
                students=students_managed,
                today_str=today_string,
                form_data=current_form_data_post,
            )
        student_to_check = Student.query.get(student_id)
        if student_to_check:
            conflict_msg_intervention = check_leave_conflict(
                student_id, start_dt, end_dt, "daily_leave", leave_id
            )
            if (
                conflict_msg_intervention
                and "serviciu (Intervenție)" in conflict_msg_intervention
            ):
                flash(
                    f'Studentul {student_to_check.nume} {student_to_check.prenume} este în serviciu de "Intervenție" și nu poate primi învoire zilnică în acest interval.',
                    "danger",
                )
                return render_template(
                    "add_edit_daily_leave.html",
                    form_title=form_title,
                    daily_leave=daily_leave,
                    students=students_managed,
                    today_str=today_string,
                    form_data=current_form_data_post,
                )

        general_conflict_msg = check_leave_conflict(
            student_id,
            start_dt,
            end_dt,
            "daily_leave",
            leave_id if daily_leave else None,
        )
        if general_conflict_msg:
            flash(
                f"Conflict detectat: Studentul are deja {general_conflict_msg}.",
                "danger",
            )
            return render_template(
                "add_edit_daily_leave.html",
                form_title=form_title,
                daily_leave=daily_leave,
                students=students_managed,
                today_str=today_string,
                form_data=current_form_data_post,
            )

        if daily_leave:
            daily_leave.student_id = int(student_id)
            daily_leave.leave_date = leave_date_obj
            daily_leave.start_time = start_time_obj
            daily_leave.end_time = end_time_obj
            daily_leave.reason = reason
            log_student_action(
                daily_leave.student_id,
                "DAILY_LEAVE_UPDATED",
                f"Învoire zilnică actualizată: {leave_date_obj.strftime('%d.%m.%Y')} de la {start_time_str} la {end_time_str}.",
            )
            flash("Învoire zilnică actualizată!", "success")
        else:
            new_leave = DailyLeave(
                student_id=int(student_id),
                leave_date=leave_date_obj,
                start_time=start_time_obj,
                end_time=end_time_obj,
                reason=reason,
                status="Aprobată",
                created_by_user_id=current_user.id,
            )
            log_student_action(
                new_leave.student_id,
                "DAILY_LEAVE_CREATED",
                f"Învoire zilnică adăugată: {leave_date_obj.strftime('%d.%m.%Y')} de la {start_time_str} la {end_time_str}.",
            )
            db.session.add(new_leave)
            flash("Învoire zilnică adăugată!", "success")
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la salvarea învoirii: {str(e)}", "danger")
            return render_template(
                "add_edit_daily_leave.html",
                form_title=form_title,
                daily_leave=daily_leave,
                students=students_managed,
                today_str=today_string,
                form_data=current_form_data_post,
            )
        return redirect(url_for("list_daily_leaves"))

    data_to_populate_form_with = {}
    if request.method == "POST":
        data_to_populate_form_with = request.form
    elif daily_leave:
        data_to_populate_form_with = {
            "student_id": str(daily_leave.student_id),
            "leave_date": daily_leave.leave_date.strftime("%Y-%m-%d"),
            "start_time": daily_leave.start_time.strftime("%H:%M"),
            "end_time": daily_leave.end_time.strftime("%H:%M"),
            "reason": daily_leave.reason or "",
        }
    return render_template(
        "add_edit_daily_leave.html",
        form_title=form_title,
        daily_leave=daily_leave,
        students=students_managed,
        today_str=today_string,
        form_data=data_to_populate_form_with,
    )


@app.route("/gradat/daily_leave/cancel/<int:leave_id>", methods=["POST"])
@login_required
def cancel_daily_leave(leave_id):
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    leave = DailyLeave.query.get_or_404(leave_id)
    student_of_leave = Student.query.get(leave.student_id)
    if (
        not student_of_leave
        or student_of_leave.created_by_user_id != current_user.id
    ):
        flash("Nu aveți permisiunea să anulați această învoire.", "danger")
        return redirect(url_for("list_daily_leaves"))
    if leave.status == "Aprobată":
        leave.status = "Anulată"
        try:
            db.session.commit()
            flash(
                f'Învoirea zilnică pentru {student_of_leave.nume} {student_of_leave.prenume} din data {leave.leave_date.strftime("%d.%m.%Y")} a fost anulată.',
                "success",
            )
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la anularea învoirii: {str(e)}", "danger")
    else:
        flash(
            'Această învoire nu poate fi anulată (statusul curent nu este "Aprobată").',
            "warning",
        )
    return redirect(url_for("list_daily_leaves"))


def parse_leave_line(
    line_text,
):  # Renamed from parse_leave_line_new for replacement
    parts = line_text.strip().split()
    grad = None
    parsed_start_time_obj = None
    parsed_end_time_obj = None
    normalized_name_search = None

    if not parts:
        return None, None, None, None

    name_parts = list(parts)  # Make a mutable copy

    # Try to parse time range HH:MM-HH:MM from the end
    if len(name_parts) > 0:
        time_range_match = re.fullmatch(
            r"(\d{1,2}:\d{2})-(\d{1,2}:\d{2})", name_parts[-1]
        )
        if time_range_match:
            try:
                parsed_start_time_obj = datetime.strptime(
                    time_range_match.group(1), "%H:%M"
                ).time()
                parsed_end_time_obj = datetime.strptime(
                    time_range_match.group(2), "%H:%M"
                ).time()
                name_parts.pop()  # Remove the time string from name_parts
            except ValueError:
                # Invalid time format in range, reset times, it will be treated as part of name
                parsed_start_time_obj = None
                parsed_end_time_obj = None
                # name_parts remains as is, time string is part of the name

    if not name_parts:  # If only time was provided, or all parts were consumed
        return None, None, None, None

    student_name_str = " ".join(name_parts)
    # Attempt to extract military rank
    for pattern in KNOWN_RANK_PATTERNS:
        match = pattern.match(student_name_str)
        if match:
            grad = match.group(0).strip()
            student_name_str = pattern.sub("", student_name_str).strip()
            break

    if student_name_str:  # If there's any name left after stripping rank
        normalized_name_search = unidecode(student_name_str.lower())
    else:  # Only rank was found, or empty string
        return (
            None,
            grad,
            parsed_start_time_obj,
            parsed_end_time_obj,
        )  # Might be only rank + time

    return (
        normalized_name_search,
        grad,
        parsed_start_time_obj,
        parsed_end_time_obj,
    )


def parse_weekend_leave_line(line_text_raw):
    line_text = line_text_raw.strip()
    is_biserica_requested = False

    # 1. Check for "biserica" keyword and remove it
    # Handles "biserica" at the end, possibly preceded by a comma.
    biserica_keyword = "biserica"
    if line_text.lower().endswith(f" {biserica_keyword}"):
        is_biserica_requested = True
        line_text = line_text[: -len(f" {biserica_keyword}")].strip()
    elif line_text.lower().endswith(f",{biserica_keyword}"):
        is_biserica_requested = True
        line_text = line_text[: -len(f",{biserica_keyword}")].strip()
    elif line_text.lower() == biserica_keyword:  # Line is ONLY "biserica"
        return (
            "",
            [],
            True,
            "Linia conține doar 'biserica', fără student sau intervale.",
        )

    # 2. Extract all date-time interval strings first using a regex pattern
    # Pattern: DD.MM.YYYY HH:MM-HH:MM (flexible spacing around hyphen)
    interval_pattern = re.compile(
        r"(\d{1,2}\.\d{1,2}\.\d{4})\s+(\d{1,2}:\d{2})\s*-\s*(\d{1,2}:\d{2})"
    )

    raw_interval_parts = []  # Store (date_str, start_str, end_str)

    # Find all matches and store their string parts
    for match in interval_pattern.finditer(line_text):
        raw_interval_parts.append(
            match.groups()
        )  # (date_str, start_str, end_str)

    if not raw_interval_parts:
        student_name_part_if_no_intervals = line_text.replace(",", "").strip()
        if (
            not student_name_part_if_no_intervals and not is_biserica_requested
        ):  # Empty line
            app.logger.debug(
                f"parse_weekend_leave_line: Skipping empty or biserica-only line: '{line_text_raw}'"
            )
            return (
                None,
                [],
                False,
                None,
            )  # Signal to skip empty line processing
        app.logger.warning(
            f"parse_weekend_leave_line: No valid datetime intervals found in line: '{line_text_raw}'. Student part considered: '{student_name_part_if_no_intervals}'. Biserica: {is_biserica_requested}"
        )
        return (
            student_name_part_if_no_intervals,
            [],
            is_biserica_requested,
            "Niciun interval de timp valid (DD.MM.YYYY HH:MM-HH:MM) găsit.",
        )

    # Attempt to parse these string parts into datetime objects
    parsed_intervals = []
    for date_str, start_str, end_str in raw_interval_parts:
        try:
            date_obj = datetime.strptime(date_str, "%d.%m.%Y").date()
            start_time_obj = datetime.strptime(start_str, "%H:%M").time()
            end_time_obj = datetime.strptime(end_str, "%H:%M").time()

            if start_time_obj == end_time_obj:
                return (
                    line_text,
                    [],
                    is_biserica_requested,
                    f"Interval orar invalid (început=sfârșit) în '{date_str} {start_str}-{end_str}'.",
                )

            parsed_intervals.append(
                {
                    "date_obj": date_obj,
                    "start_time_obj": start_time_obj,
                    "end_time_obj": end_time_obj,
                    "raw_match": f"{date_str} {start_str}-{end_str}",  # Store the raw match for removal later
                }
            )
        except ValueError:
            return (
                line_text,
                [],
                is_biserica_requested,
                f"Format dată/oră invalid în intervalul '{date_str} {start_str}-{end_str}'.",
            )

    # 3. Determine student name part by removing interval strings from the original line
    student_name_part = line_text
    # Iterate over a sorted list of raw matches by length (desc) to remove longer matches first (more specific)
    # This helps if one raw match string could be a substring of another, though less likely with this pattern.
    # For safety, or just iterate as found.
    raw_matches_to_remove = [
        item["raw_match"] for item in parsed_intervals
    ]  # Get all parts that were successfully parsed

    # Reconstruct the full matched strings as they appeared, to remove them accurately
    # The regex match objects themselves would be better for this if we captured their span.
    # For now, let's try replacing based on the reconstructed raw_match.
    # This part is tricky if formatting varies wildly (e.g. "01.01.2025 10:00 - 12:00" vs "01.01.2025 10:00-12:00")
    # The regex `\s*-\s*` handles spaces around hyphen.
    # We need to reconstruct the string that the regex actually matched.
    # Let's refine interval_pattern to capture the whole interval string for easier removal.

    # Re-defining interval_pattern to capture the whole thing for easier removal:
    full_interval_pattern = re.compile(
        r"(\d{1,2}\.\d{1,2}\.\d{4}\s+\d{1,2}:\d{2}\s*-\s*\d{1,2}:\d{2})"
    )
    actual_matched_interval_strings = full_interval_pattern.findall(line_text)

    for matched_str in actual_matched_interval_strings:
        student_name_part = student_name_part.replace(matched_str, "")

    student_name_part = student_name_part.replace(
        ",", ""
    ).strip()  # Remove commas and strip whitespace

    if not student_name_part:
        # This could happen if the line ONLY contained intervals (and maybe "biserica")
        return (
            "",
            parsed_intervals,
            is_biserica_requested,
            (
                "Numele studentului lipsește (linia conține doar intervale/biserica)."
                if parsed_intervals
                else "Linie invalidă."
            ),
        )

    # Sort intervals by date and start time (already parsed)
    parsed_intervals.sort(key=lambda x: (x["date_obj"], x["start_time_obj"]))

    return student_name_part, parsed_intervals, is_biserica_requested, None


# Endpoint vechi pentru procesarea textului din modal, înlocuit de pagina dedicată gradat_page_import_weekend_leaves
# @app.route('/gradat/weekend_leaves/process_text', methods=['POST'], endpoint='gradat_process_weekend_leaves_text')
# @login_required
# def process_weekend_leaves_text():
#     # ... (codul vechi aici)
#     pass


@app.route(
    "/gradat/daily_leaves/process_text",
    methods=["POST"],
    endpoint="gradat_process_daily_leaves_text",
)
@login_required
def process_daily_leaves_text():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    leave_list_text = request.form.get("leave_list_text")
    apply_date_str = request.form.get("apply_date")
    if not leave_list_text or not apply_date_str:
        flash(
            "Lista de învoiri și data de aplicare sunt obligatorii.", "warning"
        )
        return redirect(url_for("list_daily_leaves"))
    try:
        apply_date_obj = datetime.strptime(apply_date_str, "%Y-%m-%d").date()
    except ValueError:
        flash("Format dată aplicare invalid.", "danger")
        return redirect(url_for("list_daily_leaves"))
    if apply_date_obj.weekday() > 3:
        flash(
            "Învoirile din text pot fi aplicate doar pentru zile de Luni până Joi.",
            "warning",
        )
        return redirect(url_for("list_daily_leaves"))

    lines = leave_list_text.strip().splitlines()
    students_managed = Student.query.filter_by(
        created_by_user_id=current_user.id
    ).all()

    default_start_time_obj = time(15, 0)
    default_end_time_obj = time(19, 0)

    processed_count, error_count, already_exists_count = 0, 0, 0
    error_details_import_dl = []

    for line_raw in lines:
        line_for_student_find = line_raw.strip()
        if not line_for_student_find:
            continue

        current_iter_start_time = default_start_time_obj
        current_iter_end_time = default_end_time_obj

        # Accept either a range HH:MM-HH:MM OR a single end time HH:MM at the end of the line
        time_range_match = re.search(
            r"(\d{1,2}:\d{2})-(\d{1,2}:\d{2})$", line_for_student_find
        )
        single_time_match = None if time_range_match else re.search(
            r"(\d{1,2}:\d{2})$", line_for_student_find
        )

        student_name_grad_part = line_for_student_find

        if time_range_match:
            student_name_grad_part = line_for_student_find[
                : time_range_match.start()
            ].strip()
            try:
                parsed_line_start_time = datetime.strptime(
                    time_range_match.group(1), "%H:%M"
                ).time()
                parsed_line_end_time = datetime.strptime(
                    time_range_match.group(2), "%H:%M"
                ).time()
                current_iter_start_time = parsed_line_start_time
                current_iter_end_time = parsed_line_end_time
            except ValueError:
                error_details_import_dl.append(
                    f"Linia '{line_raw}': Format orar invalid, s-au folosit orele implicite."
                )
        elif single_time_match:
            # If only one time is provided, interpret it as the END time and keep default start (15:00)
            student_name_grad_part = line_for_student_find[
                : single_time_match.start()
            ].strip()
            try:
                parsed_single_end_time = datetime.strptime(
                    single_time_match.group(1), "%H:%M"
                ).time()
                current_iter_end_time = parsed_single_end_time
            except ValueError:
                error_details_import_dl.append(
                    f"Linia '{line_raw}': Oră invalidă, s-au folosit orele implicite."
                )

        if not student_name_grad_part:
            error_details_import_dl.append(
                f"Linia '{line_raw}': Nume student lipsă."
            )
            error_count += 1
            continue

        found_student, student_error = find_student_for_bulk_import(
            student_name_grad_part, students_managed
        )

        if student_error:
            error_details_import_dl.append(
                f"Linia '{line_raw}': {student_error}"
            )
            error_count += 1
            continue

        valid_schedule, validation_message = validate_daily_leave_times(
            current_iter_start_time, current_iter_end_time, apply_date_obj
        )
        if not valid_schedule:
            error_details_import_dl.append(
                f"Linia '{line_raw}' ({found_student.nume}): {validation_message}."
            )
            error_count += 1
            continue

        start_dt_bulk = datetime.combine(
            apply_date_obj, current_iter_start_time
        )
        effective_end_date_bulk = apply_date_obj
        if current_iter_end_time < current_iter_start_time:
            effective_end_date_bulk += timedelta(days=1)
        end_dt_bulk = datetime.combine(
            effective_end_date_bulk, current_iter_end_time
        )

        conflict_msg = check_leave_conflict(
            found_student.id, start_dt_bulk, end_dt_bulk, "daily_leave"
        )
        if conflict_msg:
            error_details_import_dl.append(
                f"Linia '{line_raw}' ({found_student.nume}): Conflict - {conflict_msg}."
            )
            error_count += 1
            continue

        # Check for duplicates before adding to session
        # This is a simplified check. A more robust check might consider overlaps.
        existing_leave = DailyLeave.query.filter_by(
            student_id=found_student.id,
            leave_date=apply_date_obj,
            start_time=current_iter_start_time,
            end_time=current_iter_end_time,
            status="Aprobată",
        ).first()
        if existing_leave:
            already_exists_count += 1
            continue

        new_leave = DailyLeave(
            student_id=found_student.id,
            leave_date=apply_date_obj,
            start_time=current_iter_start_time,
            end_time=current_iter_end_time,
            status="Aprobată",
            created_by_user_id=current_user.id,
            reason=f"Procesare text: {line_raw}",
        )
        db.session.add(new_leave)
        log_student_action(
            found_student.id,
            "DAILY_LEAVE_CREATED_BULK",
            f"Învoire zilnică adăugată prin import text: {apply_date_obj.strftime('%d.%m.%Y')} de la {current_iter_start_time.strftime('%H:%M')} la {current_iter_end_time.strftime('%H:%M')}.",
        )
        processed_count += 1

    try:
        db.session.commit()
        if processed_count > 0:
            flash(
                f"{processed_count} învoiri procesate și adăugate.", "success"
            )

        final_error_messages = []
        if error_count > 0:
            final_error_messages.append(
                f"{error_count} linii nu au putut fi procesate."
            )
            # Flash detailed errors
            error_summary = "Detalii erori:<br>" + "<br>".join(
                error_details_import_dl[:5]
            )
            if len(error_details_import_dl) > 5:
                error_summary += "<br>...și altele."
            flash(error_summary, "warning")
        if already_exists_count > 0:
            final_error_messages.append(
                f"{already_exists_count} învoiri identice existau deja și au fost ignorate."
            )

        if final_error_messages:
            flash(" | ".join(final_error_messages), "info")

    except Exception as e:
        db.session.rollback()
        flash(
            f"Eroare majoră la salvarea învoirilor din text: {str(e)}",
            "danger",
        )

    log_action(
        "BULK_IMPORT_DAILY_LEAVES_COMPLETED",
        description=f"User {current_user.username} ran bulk daily leave import. Added: {processed_count}, Errors: {error_count}, Duplicates: {already_exists_count}.",
        details_after_dict={
            "added": processed_count,
            "errors": error_count,
            "duplicates": already_exists_count,
            "error_details": error_details_import_dl[:5],
        },
    )
    db.session.commit()

    return redirect(url_for("list_daily_leaves"))


@app.route("/gradat/daily_leaves/delete/<int:leave_id>", methods=["POST"])
@app.route("/admin/daily_leaves/delete/<int:leave_id>", methods=["POST"])
@login_required
def delete_daily_leave(leave_id):
    leave_to_delete = db.session.get(DailyLeave, leave_id)
    if not leave_to_delete:
        flash("Învoirea zilnică nu a fost găsită.", "danger")
        return redirect(
            url_for("list_daily_leaves")
            if current_user.role == "gradat"
            else url_for("admin_dashboard_route")
        )

    student_owner = db.session.get(Student, leave_to_delete.student_id)

    if current_user.role == "gradat":
        if (
            not student_owner
            or student_owner.created_by_user_id != current_user.id
        ):
            flash(
                "Nu aveți permisiunea să ștergeți această învoire zilnică.",
                "danger",
            )
            return redirect(url_for("list_daily_leaves"))
        redirect_url = url_for("list_daily_leaves")
    elif current_user.role == "admin":
        if (
            student_owner
            and student_owner.creator
            and student_owner.creator.username != current_user.username
        ):
            flash(
                f"Atenție: Ștergeți o învoire zilnică pentru studentul {student_owner.nume} {student_owner.prenume}, gestionat de {student_owner.creator.username}.",
                "warning",
            )
        redirect_url = request.referrer or url_for("admin_dashboard_route")
    else:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_name_for_flash = (
        f"{student_owner.grad_militar} {student_owner.nume} {student_owner.prenume}"
        if student_owner
        else "N/A"
    )
    leave_details_for_flash = f"din {leave_to_delete.leave_date.strftime('%d.%m.%Y')} ({leave_to_delete.start_time.strftime('%H:%M')}-{leave_to_delete.end_time.strftime('%H:%M')})"
    student_id_for_log = leave_to_delete.student_id
    try:
        log_student_action(
            student_id_for_log,
            "DAILY_LEAVE_DELETED",
            f"Învoirea zilnică ({leave_details_for_flash}) a fost ștearsă.",
        )
        db.session.delete(leave_to_delete)
        db.session.commit()
        flash(
            f"Învoirea zilnică pentru {student_name_for_flash} {leave_details_for_flash} a fost ștearsă.",
            "success",
        )
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la ștergerea învoirii zilnice: {str(e)}", "danger")

    return redirect(redirect_url)


# --- Rute pentru Învoiri Weekend ---
@app.route("/gradat/weekend_leaves")
@login_required
def list_weekend_leaves():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    student_id_tuples = (
        db.session.query(Student.id)
        .filter_by(created_by_user_id=current_user.id)
        .all()
    )
    student_ids = [s[0] for s in student_id_tuples]
    if not student_ids:
        return render_template(
            "list_weekend_leaves.html",
            active_or_upcoming_leaves=[],
            past_leaves=[],
            title="Listă Învoiri Weekend",
        )

    all_relevant_leaves = (
        WeekendLeave.query.options(
            joinedload(WeekendLeave.student), joinedload(WeekendLeave.creator)
        )
        .filter(WeekendLeave.student_id.in_(student_ids))
        .order_by(WeekendLeave.weekend_start_date.desc())
        .all()
    )
    active_or_upcoming_leaves = []
    past_leaves = []
    for leave in all_relevant_leaves:
        if leave.status == "Anulată":
            past_leaves.append(leave)
        elif leave.is_overall_active_or_upcoming:
            active_or_upcoming_leaves.append(leave)
        else:
            past_leaves.append(leave)
    active_or_upcoming_leaves.sort(key=lambda x: x.weekend_start_date)
    past_leaves = past_leaves[:50]
    return render_template(
        "list_weekend_leaves.html",
        active_or_upcoming_leaves=active_or_upcoming_leaves,
        past_leaves=past_leaves,
        title="Listă Învoiri Weekend",
    )


@app.route("/gradat/weekend_leave/add", methods=["GET", "POST"])
@app.route(
    "/gradat/weekend_leave/edit/<int:leave_id>", methods=["GET", "POST"]
)
@login_required
def add_edit_weekend_leave(leave_id=None):
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    form_title = "Adaugă Învoire Weekend"
    weekend_leave = None
    form_data_on_get = {}
    if leave_id:
        weekend_leave = WeekendLeave.query.get_or_404(leave_id)
        student_of_leave = Student.query.get(weekend_leave.student_id)
        if (
            not student_of_leave
            or student_of_leave.created_by_user_id != current_user.id
        ):
            flash("Acces neautorizat la această învoire de weekend.", "danger")
            return redirect(url_for("list_weekend_leaves"))
        form_title = f"Editare Învoire Weekend: {student_of_leave.grad_militar} {student_of_leave.nume} {student_of_leave.prenume}"
        form_data_on_get["student_id"] = str(weekend_leave.student_id)
        form_data_on_get["weekend_start_date"] = (
            weekend_leave.weekend_start_date.strftime("%Y-%m-%d")
        )
        form_data_on_get["reason"] = weekend_leave.reason
        form_data_on_get["duminica_biserica"] = (
            weekend_leave.duminica_biserica
        )  # Populate church checkbox state
        selected_days_from_db = []
        # Helper to populate form_data_on_get for existing leave days
        day_fields_map = {
            "day1": (
                weekend_leave.day1_date,
                weekend_leave.day1_start_time,
                weekend_leave.day1_end_time,
                weekend_leave.day1_selected,
            ),
            "day2": (
                weekend_leave.day2_date,
                weekend_leave.day2_start_time,
                weekend_leave.day2_end_time,
                weekend_leave.day2_selected,
            ),
            "day3": (
                weekend_leave.day3_date,
                weekend_leave.day3_start_time,
                weekend_leave.day3_end_time,
                weekend_leave.day3_selected,
            ),
        }
        day_names_ro_map = {
            0: "Luni",
            1: "Marti",
            2: "Miercuri",
            3: "Joi",
            4: "Vineri",
            5: "Sambata",
            6: "Duminica",
        }

        for _field_prefix, (
            d_date,
            s_time,
            e_time,
            d_name_selected,
        ) in day_fields_map.items():
            if (
                d_date and d_name_selected
            ):  # d_name_selected is the actual day name like "Vineri"
                # day_name_template_key = day_names_ro_map.get(d_date.weekday(), "Nespecificat").lower() # This was problematic if d_name_selected is the source of truth
                day_name_template_key = (
                    d_name_selected.lower()
                )  # Use the stored day name directly
                if (
                    day_name_template_key not in selected_days_from_db
                ):  # Ensure unique day names for selection
                    selected_days_from_db.append(
                        d_name_selected
                    )  # Use original casing for selected_days[] list

                form_data_on_get[f"{day_name_template_key}_start_time"] = (
                    s_time.strftime("%H:%M") if s_time else ""
                )
                form_data_on_get[f"{day_name_template_key}_end_time"] = (
                    e_time.strftime("%H:%M") if e_time else ""
                )

        form_data_on_get["selected_days[]"] = (
            selected_days_from_db  # This will be used by template to check checkboxes
        )
    students_managed = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .order_by(Student.nume)
        .all()
    )
    upcoming_fridays_list = get_upcoming_fridays()
    if request.method == "POST":
        student_id = request.form.get("student_id")
        weekend_start_date_str = request.form.get("weekend_start_date")
        selected_days = request.form.getlist("selected_days[]")
        reason = request.form.get("reason", "").strip()
        current_form_data_post = (
            request.form
        )  # Used to repopulate form on error
        if not student_id or not weekend_start_date_str:
            flash(
                "Studentul și data de început a weekendului (Vineri) sunt obligatorii.",
                "warning",
            )
            return render_template(
                "add_edit_weekend_leave.html",
                form_title=form_title,
                weekend_leave=weekend_leave,
                students=students_managed,
                upcoming_weekends=upcoming_fridays_list,
                form_data=current_form_data_post,
            )
        if (
            not selected_days
            or len(selected_days) == 0
            or len(selected_days) > 3
        ):  # Allow 1 to 3 days
            flash(
                "Trebuie să selectați între 1 și 3 zile din weekend.",
                "warning",
            )
            return render_template(
                "add_edit_weekend_leave.html",
                form_title=form_title,
                weekend_leave=weekend_leave,
                students=students_managed,
                upcoming_weekends=upcoming_fridays_list,
                form_data=current_form_data_post,
            )
        try:
            friday_date_obj = datetime.strptime(
                weekend_start_date_str, "%Y-%m-%d"
            ).date()
        except ValueError:
            flash("Format dată weekend invalid.", "danger")
            return render_template(
                "add_edit_weekend_leave.html",
                form_title=form_title,
                weekend_leave=weekend_leave,
                students=students_managed,
                upcoming_weekends=upcoming_fridays_list,
                form_data=current_form_data_post,
            )

        day_data = []  # To store processed day information before saving
        for day_name_selected in selected_days:
            start_time_str = request.form.get(
                f"{day_name_selected.lower()}_start_time"
            )
            end_time_str = request.form.get(
                f"{day_name_selected.lower()}_end_time"
            )
            if not start_time_str or not end_time_str:
                flash(
                    f"Orele de început și sfârșit sunt obligatorii pentru {day_name_selected}.",
                    "warning",
                )
                return render_template(
                    "add_edit_weekend_leave.html",
                    form_title=form_title,
                    weekend_leave=weekend_leave,
                    students=students_managed,
                    upcoming_weekends=upcoming_fridays_list,
                    form_data=current_form_data_post,
                )
            try:
                start_time_obj = datetime.strptime(
                    start_time_str, "%H:%M"
                ).time()
                end_time_obj = datetime.strptime(end_time_str, "%H:%M").time()
            except ValueError:
                flash(
                    f"Format oră invalid pentru {day_name_selected}.", "danger"
                )
                return render_template(
                    "add_edit_weekend_leave.html",
                    form_title=form_title,
                    weekend_leave=weekend_leave,
                    students=students_managed,
                    upcoming_weekends=upcoming_fridays_list,
                    form_data=current_form_data_post,
                )
            if end_time_obj == start_time_obj:
                flash(
                    f"Ora de început și sfârșit nu pot fi identice pentru {day_name_selected}.",
                    "warning",
                )
                return render_template(
                    "add_edit_weekend_leave.html",
                    form_title=form_title,
                    weekend_leave=weekend_leave,
                    students=students_managed,
                    upcoming_weekends=upcoming_fridays_list,
                    form_data=current_form_data_post,
                )
            day_offset_map = {"Vineri": 0, "Sambata": 1, "Duminica": 2}
            day_offset = day_offset_map.get(day_name_selected)
            if day_offset is None:
                flash(f"Nume zi invalid: {day_name_selected}", "danger")
                return render_template(
                    "add_edit_weekend_leave.html",
                    form_title=form_title,
                    weekend_leave=weekend_leave,
                    students=students_managed,
                    upcoming_weekends=upcoming_fridays_list,
                    form_data=current_form_data_post,
                )
            actual_date_obj = friday_date_obj + timedelta(days=day_offset)
            current_interval_start_dt = datetime.combine(
                actual_date_obj, start_time_obj
            )
            effective_end_date_for_interval = actual_date_obj
            if end_time_obj < start_time_obj:
                effective_end_date_for_interval += timedelta(days=1)
            current_interval_end_dt = datetime.combine(
                effective_end_date_for_interval, end_time_obj
            )
            if current_interval_end_dt <= current_interval_start_dt:
                flash(
                    f"Interval orar invalid pentru {day_name_selected}.",
                    "warning",
                )
                return render_template(
                    "add_edit_weekend_leave.html",
                    form_title=form_title,
                    weekend_leave=weekend_leave,
                    students=students_managed,
                    upcoming_weekends=upcoming_fridays_list,
                    form_data=current_form_data_post,
                )
            day_data.append(
                {
                    "name": day_name_selected,
                    "date": actual_date_obj,
                    "start": start_time_obj,
                    "end": end_time_obj,
                    "start_dt": current_interval_start_dt,
                    "end_dt": current_interval_end_dt,
                }
            )
        day_data.sort(key=lambda x: x["start_dt"])
        student_to_check = Student.query.get(student_id)
        if student_to_check:
            for interval_data in day_data:
                active_intervention_service = ServiceAssignment.query.filter(
                    ServiceAssignment.student_id == student_id,
                    ServiceAssignment.service_type == "Intervenție",
                    ServiceAssignment.start_datetime < interval_data["end_dt"],
                    ServiceAssignment.end_datetime > interval_data["start_dt"],
                ).first()
                if active_intervention_service:
                    flash(
                        f'Studentul {student_to_check.nume} este în "Intervenție" pe {interval_data["name"]} și nu poate primi învoire.',
                        "danger",
                    )
                    return render_template(
                        "add_edit_weekend_leave.html",
                        form_title=form_title,
                        weekend_leave=weekend_leave,
                        students=students_managed,
                        upcoming_weekends=upcoming_fridays_list,
                        form_data=current_form_data_post,
                    )

        # Clear previous day data before setting new, especially for edits
        if weekend_leave:
            target_leave = weekend_leave
            target_leave.day1_selected = None
            target_leave.day1_date = None
            target_leave.day1_start_time = None
            target_leave.day1_end_time = None
            target_leave.day2_selected = None
            target_leave.day2_date = None
            target_leave.day2_start_time = None
            target_leave.day2_end_time = None
            target_leave.day3_selected = None
            target_leave.day3_date = None
            target_leave.day3_start_time = None
            target_leave.day3_end_time = None
            flash_msg = "Învoire Weekend actualizată!"
        else:
            target_leave = WeekendLeave(
                created_by_user_id=current_user.id, status="Aprobată"
            )
            flash_msg = "Învoire Weekend adăugată!"

        target_leave.student_id = int(student_id)
        target_leave.weekend_start_date = friday_date_obj
        target_leave.reason = reason
        target_leave.duminica_biserica = (
            request.form.get("duminica_biserica") == "true"
        )

        # Assign processed day_data to the model fields
        if len(day_data) >= 1:
            target_leave.day1_selected = day_data[0]["name"]
            target_leave.day1_date = day_data[0]["date"]
            target_leave.day1_start_time = day_data[0]["start"]
            target_leave.day1_end_time = day_data[0]["end"]
        if len(day_data) >= 2:
            target_leave.day2_selected = day_data[1]["name"]
            target_leave.day2_date = day_data[1]["date"]
            target_leave.day2_start_time = day_data[1]["start"]
            target_leave.day2_end_time = day_data[1]["end"]
        if len(day_data) >= 3:  # Assign third day if present
            target_leave.day3_selected = day_data[2]["name"]
            target_leave.day3_date = day_data[2]["date"]
            target_leave.day3_start_time = day_data[2]["start"]
            target_leave.day3_end_time = day_data[2]["end"]

        if not weekend_leave:  # If it's a new leave, add to session
            db.session.add(target_leave)

        try:
            log_student_action(
                target_leave.student_id,
                (
                    "WEEKEND_LEAVE_CREATED"
                    if not weekend_leave
                    else "WEEKEND_LEAVE_UPDATED"
                ),
                f"Învoire weekend salvată pentru {friday_date_obj.strftime('%d.%m')}. Zile: {', '.join(selected_days)}.",
            )
            db.session.commit()
            if len(selected_days) == 3:
                flash(
                    flash_msg + " Ați selectat 3 zile pentru învoire.",
                    "success",
                )  # Add warning if 3 days selected
            else:
                flash(flash_msg, "success")
        except Exception as e:
            db.session.rollback()
            flash(
                f"Eroare la salvarea învoirii de weekend: {str(e)}", "danger"
            )
            return render_template(
                "add_edit_weekend_leave.html",
                form_title=form_title,
                weekend_leave=weekend_leave,
                students=students_managed,
                upcoming_weekends=upcoming_fridays_list,
                form_data=current_form_data_post,
            )
        return redirect(url_for("list_weekend_leaves"))

    data_to_populate_form_with = {}
    if request.method == "POST":
        data_to_populate_form_with = request.form
    elif weekend_leave:
        data_to_populate_form_with = form_data_on_get

    return render_template(
        "add_edit_weekend_leave.html",
        form_title=form_title,
        weekend_leave=weekend_leave,
        students=students_managed,
        upcoming_weekends=upcoming_fridays_list,
        form_data=data_to_populate_form_with,
    )


@app.route("/gradat/weekend_leave/cancel/<int:leave_id>", methods=["POST"])
@login_required
def cancel_weekend_leave(leave_id):
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    leave = WeekendLeave.query.get_or_404(leave_id)
    student_of_leave = Student.query.get(leave.student_id)
    if (
        not student_of_leave
        or student_of_leave.created_by_user_id != current_user.id
    ):
        flash(
            "Nu aveți permisiunea să anulați această învoire de weekend.",
            "danger",
        )
        return redirect(url_for("list_weekend_leaves"))
    if leave.status == "Aprobată":
        leave.status = "Anulată"
        try:
            db.session.commit()
            flash(
                f'Învoirea de weekend pentru {student_of_leave.nume} {student_of_leave.prenume} (începând cu {leave.weekend_start_date.strftime("%d.%m")}) a fost anulată.',
                "success",
            )
        except Exception as e:
            db.session.rollback()
            flash(
                f"Eroare la anularea învoirii de weekend: {str(e)}", "danger"
            )
    else:
        flash(
            'Această învoire de weekend nu poate fi anulată (statusul curent nu este "Aprobată").',
            "warning",
        )
    return redirect(url_for("list_weekend_leaves"))


@app.route(
    "/gradat/weekend_leaves/process_text",
    methods=["POST"],
    endpoint="gradat_process_weekend_leaves_text",
)
@login_required
def process_weekend_leaves_text():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    students_managed = Student.query.filter_by(
        created_by_user_id=current_user.id
    ).all()
    leave_list_text = request.form.get(
        "weekend_leave_list_text", ""
    ).strip()  # Assuming this will be the form field name
    if not leave_list_text:
        flash("Lista de învoiri este goală.", "warning")
        return redirect(url_for("list_weekend_leaves"))

    lines = leave_list_text.strip().splitlines()
    processed_count = 0
    error_count = 0
    error_details_list = []  # More structured than just a list of strings

    day_names_map_ro_to_en = {
        "vineri": "Friday",
        "sambata": "Saturday",
        "duminica": "Sunday",
    }  # For model's dayX_selected
    day_name_to_weekday = {
        "vineri": 4,
        "sambata": 5,
        "duminica": 6,
    }  # Monday is 0, Sunday is 6

    for line_raw in lines:
        line_content = line_raw.strip()
        if not line_content:
            continue

        student_name_str, parsed_intervals, is_biserica_req, error_msg = (
            parse_weekend_leave_line(line_content)
        )

        if (
            error_msg and student_name_str is None and not parsed_intervals
        ):  # Skip empty line signal
            continue
        if error_msg:
            error_details_list.append(
                {"line": line_content, "error": error_msg}
            )
            error_count += 1
            continue

        if (
            not parsed_intervals
        ):  # Should have been caught by error_msg, but as a safeguard
            error_details_list.append(
                {
                    "line": line_content,
                    "error": "Niciun interval valid de procesat.",
                }
            )
            error_count += 1
            continue

        student_obj, student_error = find_student_for_bulk_import(
            student_name_str, students_managed
        )
        if student_error:
            error_details_list.append(
                {
                    "line": line_content,
                    "error": f"Student '{student_name_str}': {student_error}",
                }
            )
            error_count += 1
            continue

        # --- Logic to determine weekend_start_date and map intervals to day1/2/3 ---
        if not parsed_intervals:  # Should not happen if error_msg was None
            error_details_list.append(
                {
                    "line": line_content,
                    "student": student_obj.nume,
                    "error": "Eroare internă: Intervale goale după parsare reușită.",
                }
            )
            error_count += 1
            continue

        # Determine the primary weekend: find the Friday of the week of the first interval's date.
        first_interval_date = parsed_intervals[0]["date_obj"]
        days_to_subtract_for_friday = (first_interval_date.weekday() - 4) % 7
        # (weekday - 4) gives offset from Friday. If it's negative, add 7.
        # E.g. Mon(0): (0-4)%7 = -4%7 = 3 -> Mon - 3 days is Fri of PREVIOUS week. This is not what we want.
        # We want Friday of the week *containing* the first_interval_date.
        # If first_interval_date is Mon, its Friday is +4 days. If Sun, its Friday is -2 days.
        weekend_start_date_obj = (
            first_interval_date
            - timedelta(days=first_interval_date.weekday())
            + timedelta(days=4)
        )

        # Validate that all intervals fall within this determined weekend (Fri, Sat, Sun)
        # and map them to day1_ (Friday), day2_ (Saturday), day3_ (Sunday)

        current_weekend_leave_data = {  # Temp storage for this student's leave
            "day1_date": None,
            "day1_start_time": None,
            "day1_end_time": None,
            "day1_selected": None,
            "day2_date": None,
            "day2_start_time": None,
            "day2_end_time": None,
            "day2_selected": None,
            "day3_date": None,
            "day3_start_time": None,
            "day3_end_time": None,
            "day3_selected": None,
            "intervals_for_conflict_check": [],
        }

        distinct_days_processed = set()

        for interval in parsed_intervals:
            interval_date = interval["date_obj"]

            # Check if interval_date is Fri, Sat, or Sun of the determined weekend_start_date_obj
            delta_days = (interval_date - weekend_start_date_obj).days
            day_slot_key = None  # Will be 'day1', 'day2', or 'day3'
            day_name_ro = None

            if delta_days == 0 and interval_date.weekday() == 4:  # Friday
                day_slot_key = "day1"
                day_name_ro = "Vineri"
            elif delta_days == 1 and interval_date.weekday() == 5:  # Saturday
                day_slot_key = "day2"
                day_name_ro = "Sambata"
            elif delta_days == 2 and interval_date.weekday() == 6:  # Sunday
                day_slot_key = "day3"
                day_name_ro = "Duminica"
            else:
                error_details_list.append(
                    {
                        "line": line_content,
                        "student": student_obj.nume,
                        "error": f"Data {interval_date.strftime('%d.%m.%Y')} nu corespunde weekendului definit de prima dată ({weekend_start_date_obj.strftime('%d.%m.%Y')}).",
                    }
                )
                error_count += 1
                break  # Break from this student's intervals

            if (
                day_slot_key in distinct_days_processed
            ):  # Already have an interval for this day slot
                error_details_list.append(
                    {
                        "line": line_content,
                        "student": student_obj.nume,
                        "error": f"Intervale multiple specificate pentru aceeași zi ({day_name_ro}). Doar primul va fi considerat.",
                    }
                )
                # error_count +=1 # Optionally count this as an error or just a warning
                continue  # Skip this additional interval for the same day

            distinct_days_processed.add(day_slot_key)

            current_weekend_leave_data[f"{day_slot_key}_date"] = interval_date
            current_weekend_leave_data[f"{day_slot_key}_start_time"] = (
                interval["start_time_obj"]
            )
            current_weekend_leave_data[f"{day_slot_key}_end_time"] = interval[
                "end_time_obj"
            ]
            current_weekend_leave_data[f"{day_slot_key}_selected"] = (
                day_name_ro  # Store "Vineri", "Sambata", "Duminica"
            )

            # For conflict checking:
            start_dt = datetime.combine(
                interval_date, interval["start_time_obj"]
            )
            effective_end_date = interval_date
            if interval["end_time_obj"] < interval["start_time_obj"]:
                effective_end_date += timedelta(days=1)
            end_dt = datetime.combine(
                effective_end_date, interval["end_time_obj"]
            )
            current_weekend_leave_data["intervals_for_conflict_check"].append(
                {"start": start_dt, "end": end_dt, "day_name": day_name_ro}
            )

        if (
            error_count > 0 and error_details_list[-1]["line"] == line_content
        ):  # If error occurred for this student, skip to next line
            continue

        if not distinct_days_processed:
            error_details_list.append(
                {
                    "line": line_content,
                    "student": student_obj.nume,
                    "error": "Niciun interval valid mapat la zilele weekendului.",
                }
            )
            error_count += 1
            continue

        # --- Conflict Checking ---
        conflict_found_for_student = False
        for interval_to_check in current_weekend_leave_data[
            "intervals_for_conflict_check"
        ]:
            conflict = check_leave_conflict(
                student_obj.id,
                interval_to_check["start"],
                interval_to_check["end"],
                leave_type="weekend_leave",
            )
            if conflict:
                error_details_list.append(
                    {
                        "line": line_content,
                        "student": student_obj.nume,
                        "error": f"Conflict pentru {interval_to_check['day_name']}: {conflict}.",
                    }
                )
                error_count += 1
                conflict_found_for_student = True
                break
        if conflict_found_for_student:
            continue

        # --- Create WeekendLeave object ---
        new_wl = WeekendLeave(
            student_id=student_obj.id,
            weekend_start_date=weekend_start_date_obj,
            day1_selected=current_weekend_leave_data["day1_selected"],
            day1_date=current_weekend_leave_data["day1_date"],
            day1_start_time=current_weekend_leave_data["day1_start_time"],
            day1_end_time=current_weekend_leave_data["day1_end_time"],
            day2_selected=current_weekend_leave_data["day2_selected"],
            day2_date=current_weekend_leave_data["day2_date"],
            day2_start_time=current_weekend_leave_data["day2_start_time"],
            day2_end_time=current_weekend_leave_data["day2_end_time"],
            day3_selected=current_weekend_leave_data["day3_selected"],
            day3_date=current_weekend_leave_data["day3_date"],
            day3_start_time=current_weekend_leave_data["day3_start_time"],
            day3_end_time=current_weekend_leave_data["day3_end_time"],
            duminica_biserica=(
                is_biserica_req
                and current_weekend_leave_data["day3_selected"] == "Duminica"
            ),  # Only if Sunday is actually selected
            status="Aprobată",
            created_by_user_id=current_user.id,
            reason=f"Procesare text: {line_content[:100]}",  # Truncate reason if line is too long
        )
        db.session.add(new_wl)
        log_student_action(
            student_obj.id,
            "WEEKEND_LEAVE_CREATED_BULK",
            f"Învoire weekend adăugată prin import text pentru {weekend_start_date_obj.strftime('%d.%m')}.",
        )
        processed_count += 1

    # --- Commit and Flash Messages ---
    try:
        if processed_count > 0:  # Only commit if there are successful items
            db.session.commit()
            flash(
                f"{processed_count} învoiri de weekend procesate și adăugate cu succes.",
                "success",
            )
        elif (
            error_count == 0 and processed_count == 0
        ):  # No lines or only empty lines
            flash("Nu au fost furnizate date de procesat.", "info")

        if error_count > 0:
            flash(
                f"{error_count} linii nu au putut fi procesate sau au generat erori.",
                "danger",
            )
            # Construct detailed error message for flash
            error_flash_message = "Detalii erori:<br>"
            for err_detail in error_details_list[:5]:  # Show first 5 errors
                error_flash_message += f"- Linia: '{err_detail['line'][:60]}...' Student: {err_detail.get('student','N/A')} Eroare: {err_detail['error']}<br>"
            if len(error_details_list) > 5:
                error_flash_message += (
                    f"... și încă {len(error_details_list) - 5} erori."
                )
            flash(error_flash_message, "warning")

        # Log the bulk operation attempt
        log_action(
            "BULK_IMPORT_WEEKEND_LEAVES_COMPLETED",
            description=f"User {current_user.username} ran bulk weekend leave import. Added: {processed_count}, Errors: {error_count}. Line count: {len(lines)}",
            details_after_dict={
                "added_count": processed_count,
                "error_count": error_count,
                "total_lines_input": len(lines),
                "first_few_error_details": error_details_list[:3],
            },
        )
        db.session.commit()  # Commit the log

    except Exception as e:
        db.session.rollback()
        flash(
            f"Eroare majoră la salvarea învoirilor de weekend din text: {str(e)}",
            "danger",
        )
        try:
            log_action(
                "BULK_IMPORT_WEEKEND_LEAVES_FAIL_MAJOR",
                description=f"User {current_user.username} bulk weekend leave import failed critically. Error: {str(e)}",
                details_after_dict={
                    "added_count": processed_count,
                    "error_count": error_count,
                    "exception": str(e),
                },
            )
            db.session.commit()
        except Exception as log_e:
            app.logger.error(
                f"CRITICAL: Failed to commit failure log for BULK_IMPORT_WEEKEND_LEAVES_FAIL_MAJOR: {str(log_e)}"
            )

    return redirect(url_for("list_weekend_leaves"))


@app.route(
    "/gradat/weekend_leaves/export_word",
    endpoint="gradat_export_weekend_leaves_word",
)
@login_required
def export_weekend_leaves_word():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_id_tuples = (
        db.session.query(Student.id)
        .filter_by(created_by_user_id=current_user.id)
        .all()
    )
    student_ids = [s[0] for s in student_id_tuples]

    if not student_ids:
        flash("Nu aveți studenți pentru a exporta învoiri de weekend.", "info")
        return redirect(url_for("list_weekend_leaves"))

    leaves_to_export = (
        WeekendLeave.query.options(joinedload(WeekendLeave.student))
        .filter(
            WeekendLeave.student_id.in_(student_ids),
            WeekendLeave.status == "Aprobată",
        )
        .join(Student)
        .order_by(
            Student.nume, Student.prenume, WeekendLeave.weekend_start_date
        )
        .all()
    )

    # păstrăm doar cele active sau viitoare
    leaves_to_export = [
        leave
        for leave in leaves_to_export
        if leave.is_overall_active_or_upcoming
    ]

    if not leaves_to_export:
        flash(
            "Nicio învoire de weekend activă sau viitoare de exportat.", "info"
        )
        return redirect(url_for("list_weekend_leaves"))

    document = Document()
    document.add_heading("Raport Învoiri Weekend", level=1).alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    user_info_text = f"Raport generat de: {current_user.username}\nData generării: {get_localized_now().strftime('%d-%m-%Y %H:%M')}"
    p_user = document.add_paragraph()
    p_user.add_run(user_info_text).italic = True
    p_user.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    # Agregăm intervalele pe student (un rând per student cu toate perioadele listate)
    from collections import defaultdict

    student_map = {}
    periods_map = defaultdict(list)

    for leave in leaves_to_export:
        st = leave.student
        student_map[st.id] = st
        for iv in leave.get_intervals():
            # includem doar intervale care încă nu au trecut complet
            if iv["end"] >= get_localized_now():
                start_local = iv["start"].astimezone(EUROPE_BUCHAREST)
                end_local = iv["end"].astimezone(EUROPE_BUCHAREST)
                day_label = iv.get("day_name") or start_local.strftime("%A")
                date_label = start_local.strftime("%d-%m")
                periods_map[st.id].append(
                    (
                        start_local,
                        f"{day_label} ({date_label}): {start_local.strftime('%H:%M')} - {end_local.strftime('%H:%M')}",
                    )
                )

    # Tabel: Nr., Grad, Nume și Prenume, Plutonul, Perioade
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = [
        "Nr. crt.",
        "Grad",
        "Nume și Prenume",
        "Plutonul (Grupa)",
        "Perioade",
    ]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ordonăm după nume
    ordered_students = sorted(
        periods_map.keys(),
        key=lambda sid: (student_map[sid].nume, student_map[sid].prenume),
    )
    for idx, sid in enumerate(ordered_students, start=1):
        st = student_map[sid]
        # sortăm perioadele per student după început
        periods_sorted = [
            p for _, p in sorted(periods_map[sid], key=lambda t: t[0])
        ]
        periods_str = "; ".join(periods_sorted)
        row = table.add_row().cells
        row[0].text = str(idx)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = st.grad_militar
        row[2].text = f"{st.nume} {st.prenume}"
        row[3].text = st.pluton
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[4].text = periods_str

    # Lățimi coloane
    new_widths = {
        0: Inches(0.5),  # Nr
        1: Inches(0.8),  # Grad
        2: Inches(2.2),  # Nume Prenume
        3: Inches(0.9),  # Pluton
        4: Inches(3.0),  # Perioade
    }
    for col_idx, width_val in new_widths.items():
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = width_val

    document.add_paragraph()  # spațiu

    # --- Tabel separat: participă la Biserică (Duminică) ---
    church_attendees = []
    for leave in leaves_to_export:
        if leave.duminica_biserica:
            if any(
                iv["day_name"] == "Duminica" for iv in leave.get_intervals()
            ):
                church_attendees.append(leave.student)

    if church_attendees:
        document.add_heading(
            "Studenți care participă la Biserică (Duminică 09:00-11:00)",
            level=2,
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        church_table = document.add_table(rows=1, cols=4)
        church_table.style = "Table Grid"
        church_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ch_hdr = church_table.rows[0].cells
        for i, title in enumerate(
            ["Nr. crt.", "Grad", "Nume și Prenume", "Plutonul (Grupa)"]
        ):
            ch_hdr[i].text = title
            ch_hdr[i].paragraphs[0].runs[0].font.bold = True
            ch_hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        church_attendees = sorted(
            {s.id: s for s in church_attendees}.values(),
            key=lambda s: (s.nume, s.prenume),
        )
        for idx, s in enumerate(church_attendees, start=1):
            r = church_table.add_row().cells
            r[0].text = str(idx)
            r[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r[1].text = s.grad_militar
            r[2].text = f"{s.nume} {s.prenume}"
            r[3].text = s.pluton
            r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        widths = {
            0: Inches(0.5),
            1: Inches(0.8),
            2: Inches(2.5),
            3: Inches(1.0),
        }
        for ci, w in widths.items():
            for row in church_table.rows:
                if ci < len(row.cells):
                    row.cells[ci].width = w

    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    f = io.BytesIO()
    document.save(f)
    f.seek(0)

    filename = f"Raport_Invoiri_Weekend_{current_user.username}_{date.today().strftime('%Y%m%d')}.docx"
    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/gradat/weekend_leaves/delete/<int:leave_id>", methods=["POST"])
@app.route("/admin/weekend_leaves/delete/<int:leave_id>", methods=["POST"])
@login_required
def delete_weekend_leave(leave_id):
    leave_to_delete = db.session.get(WeekendLeave, leave_id)
    if not leave_to_delete:
        flash("Învoirea de weekend nu a fost găsită.", "danger")
        return redirect(
            url_for("list_weekend_leaves")
            if current_user.role == "gradat"
            else url_for("admin_dashboard_route")
        )

    student_owner = db.session.get(Student, leave_to_delete.student_id)

    if current_user.role == "gradat":
        if (
            not student_owner
            or student_owner.created_by_user_id != current_user.id
        ):
            flash(
                "Nu aveți permisiunea să ștergeți această învoire de weekend.",
                "danger",
            )
            return redirect(url_for("list_weekend_leaves"))
        redirect_url = url_for("list_weekend_leaves")
    elif current_user.role == "admin":
        if (
            student_owner
            and student_owner.creator
            and student_owner.creator.username != current_user.username
        ):
            flash(
                f"Atenție: Ștergeți o învoire de weekend pentru studentul {student_owner.nume} {student_owner.prenume}, gestionat de {student_owner.creator.username}.",
                "warning",
            )
        redirect_url = request.referrer or url_for("admin_dashboard_route")
    else:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_name_for_flash = (
        f"{student_owner.grad_militar} {student_owner.nume} {student_owner.prenume}"
        if student_owner
        else "N/A"
    )
    leave_details_for_flash = f"din weekend-ul {leave_to_delete.weekend_start_date.strftime('%d.%m.%Y')}"
    student_id_for_log = leave_to_delete.student_id

    try:
        log_student_action(
            student_id_for_log,
            "WEEKEND_LEAVE_DELETED",
            f"Învoirea de weekend ({leave_details_for_flash}) a fost ștearsă.",
        )
        db.session.delete(leave_to_delete)
        db.session.commit()
        flash(
            f"Învoirea de weekend pentru {student_name_for_flash} {leave_details_for_flash} a fost ștearsă.",
            "success",
        )
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la ștergerea învoirii de weekend: {str(e)}", "danger")

    return redirect(redirect_url)


# --- Rute pentru Servicii ---
@app.route("/gradat/services")
@login_required
def list_services():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    students_managed_by_gradat = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .with_entities(Student.id)
        .all()
    )
    student_ids = [sid for (sid,) in students_managed_by_gradat]

    if not student_ids:
        return render_template(
            "list_services.html",
            upcoming_services=[],
            past_services=[],
            title="Management Servicii",
        )

    now = get_localized_now()  # Folosim ora localizată

    all_services = (
        ServiceAssignment.query.options(
            joinedload(ServiceAssignment.student),
            joinedload(ServiceAssignment.creator),
        )
        .filter(ServiceAssignment.student_id.in_(student_ids))
        .order_by(ServiceAssignment.start_datetime.desc())
        .all()
    )

    upcoming_services = []
    past_services = []

    for s in all_services:
        # Folosim proprietatea .is_past care deja compară cu ora localizată
        if s.is_past:
            past_services.append(s)
        else:
            upcoming_services.append(s)

    # Sortarea finală
    upcoming_services.sort(key=lambda x: x.start_datetime)
    # past_services este deja sortat descendent de la query

    # Limitarea serviciilor trecute
    past_services = past_services[:50]

    return render_template(
        "list_services.html",
        upcoming_services=upcoming_services,
        past_services=past_services,
        title="Management Servicii",
    )


@app.route("/gradat/services/assign", methods=["GET", "POST"])
@app.route(
    "/gradat/services/edit/<int:assignment_id>", methods=["GET", "POST"]
)
@login_required
def assign_service(assignment_id=None):
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    form_title = "Asignează Serviciu Nou"
    service_assignment = None
    form_data_for_template = {}

    if assignment_id:
        service_assignment = ServiceAssignment.query.get_or_404(assignment_id)
        student_of_service = Student.query.get(service_assignment.student_id)
        if (
            not student_of_service
            or student_of_service.created_by_user_id != current_user.id
        ):
            flash("Acces neautorizat la acest serviciu.", "danger")
            return redirect(url_for("list_services"))
        form_title = f"Editare Serviciu: {student_of_service.grad_militar} {student_of_service.nume} ({service_assignment.service_type})"
        form_data_for_template = {
            "student_id": str(service_assignment.student_id),
            "service_type": service_assignment.service_type,
            "service_date": service_assignment.service_date.strftime(
                "%Y-%m-%d"
            ),
            "start_time": service_assignment.start_datetime.strftime("%H:%M"),
            "end_time": service_assignment.end_datetime.strftime("%H:%M"),
            "participates_in_roll_call": (
                "true" if service_assignment.participates_in_roll_call else ""
            ),
            "notes": service_assignment.notes or "",
        }

    students = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .order_by(Student.nume)
        .all()
    )
    if not students and not assignment_id:
        flash("Nu aveți studenți pentru a le asigna servicii.", "warning")
        return redirect(url_for("list_students"))

    default_times_for_js = (
        {  # Acestea sunt stringuri, nu necesită localizare directă aici
            "GSS": ("07:00", "07:00"),
            "SVM": ("05:50", "20:00"),
            "Intervenție": ("20:00", "00:00"),
            "Planton 1": ("22:00", "00:00"),
            "Planton 2": ("00:00", "02:00"),
            "Planton 3": ("02:00", "04:00"),
            "Altul": ("", ""),
        }
    )
    today_iso_str = (
        get_localized_now().date().isoformat()
    )  # Pentru valoarea default a datei

    if request.method == "POST":
        student_id = request.form.get("student_id")
        service_type = request.form.get("service_type")
        service_date_str = request.form.get("service_date")
        start_time_str = request.form.get("start_time")
        end_time_str = request.form.get("end_time")
        participates = "participates_in_roll_call" in request.form
        notes = request.form.get("notes", "").strip()

        current_form_data = request.form

        if not all(
            [
                student_id,
                service_type,
                service_date_str,
                start_time_str,
                end_time_str,
            ]
        ):
            flash(
                "Toate câmpurile marcate cu * (student, tip, dată, ore) sunt obligatorii.",
                "warning",
            )
            return render_template(
                "assign_service.html",
                form_title=form_title,
                service_assignment=service_assignment,
                students=students,
                service_types=SERVICE_TYPES,
                default_times=default_times_for_js,
                today_str=today_iso_str,
                form_data=current_form_data,
            )

        try:
            service_date_obj = datetime.strptime(
                service_date_str, "%Y-%m-%d"
            ).date()
            start_time_obj = datetime.strptime(start_time_str, "%H:%M").time()
            end_time_obj = datetime.strptime(end_time_str, "%H:%M").time()
        except ValueError:
            flash("Format dată sau oră invalid.", "danger")
            return render_template(
                "assign_service.html",
                form_title=form_title,
                service_assignment=service_assignment,
                students=students,
                service_types=SERVICE_TYPES,
                default_times=default_times_for_js,
                today_str=today_iso_str,
                form_data=current_form_data,
            )

        # Datele și orele sunt deja locale din formular, le combinăm
        start_dt_obj = datetime.combine(
            service_date_obj, start_time_obj
        )  # Naive, local
        effective_end_date = service_date_obj
        if end_time_obj < start_time_obj:
            effective_end_date += timedelta(days=1)
        elif (
            service_type == "GSS" and end_time_obj == start_time_obj
        ):  # GSS special case for 24h
            effective_end_date += timedelta(days=1)
        end_dt_obj = datetime.combine(effective_end_date, end_time_obj)

        if (
            end_dt_obj <= start_dt_obj
        ):  # This check should now be generally correct
            flash(
                "Intervalul orar al serviciului este invalid (sfârșitul trebuie să fie după început).",
                "danger",
            )
            return render_template(
                "assign_service.html",
                form_title=form_title,
                service_assignment=service_assignment,
                students=students,
                service_types=SERVICE_TYPES,
                default_times=default_times_for_js,
                today_str=today_iso_str,
                form_data=current_form_data,
            )

        stud = Student.query.filter_by(
            id=student_id, created_by_user_id=current_user.id
        ).first()
        if not stud:
            flash("Student selectat invalid sau nu vă aparține.", "danger")
            return render_template(
                "assign_service.html",
                form_title=form_title,
                service_assignment=service_assignment,
                students=students,
                service_types=SERVICE_TYPES,
                default_times=default_times_for_js,
                today_str=today_iso_str,
                form_data=current_form_data,
            )

        # check_service_conflict_for_student se așteaptă la datetime-uri naive locale
        conflict_msg = check_service_conflict_for_student(
            stud.id, start_dt_obj, end_dt_obj, service_type, assignment_id
        )
        if conflict_msg:
            flash(f"Conflict: studentul are deja {conflict_msg}", "danger")
            return render_template(
                "assign_service.html",
                form_title=form_title,
                service_assignment=service_assignment,
                students=students,
                service_types=SERVICE_TYPES,
                default_times=default_times_for_js,
                today_str=today_iso_str,
                form_data=current_form_data,
            )

        if service_assignment:
            service_assignment.student_id = stud.id
            service_assignment.service_type = service_type
            service_assignment.service_date = service_date_obj
            service_assignment.start_datetime = start_dt_obj
            service_assignment.end_datetime = end_dt_obj
            service_assignment.participates_in_roll_call = participates
            service_assignment.notes = notes
            log_student_action(
                stud.id,
                "SERVICE_UPDATED",
                f"Serviciu '{service_type}' actualizat: {start_dt_obj.strftime('%d.%m %H:%M')} - {end_dt_obj.strftime('%d.%m %H:%M')}.",
            )
            flash_msg = f"Serviciul {service_type} pentru {stud.nume} {stud.prenume} a fost actualizat!"
        else:
            new_assignment = ServiceAssignment(
                student_id=stud.id,
                service_type=service_type,
                service_date=service_date_obj,
                start_datetime=start_dt_obj,
                end_datetime=end_dt_obj,
                participates_in_roll_call=participates,
                notes=notes,
                created_by_user_id=current_user.id,
            )
            db.session.add(new_assignment)
            log_student_action(
                stud.id,
                "SERVICE_CREATED",
                f"Serviciu '{service_type}' adăugat: {start_dt_obj.strftime('%d.%m %H:%M')} - {end_dt_obj.strftime('%d.%m %H:%M')}.",
            )
            flash_msg = f"Serviciul {service_type} a fost asignat lui {stud.nume} {stud.prenume}."

        try:
            db.session.commit()
            flash(flash_msg, "success")
            return redirect(url_for("list_services"))
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la salvarea serviciului: {str(e)}", "danger")
            return render_template(
                "assign_service.html",
                form_title=form_title,
                service_assignment=service_assignment,
                students=students,
                service_types=SERVICE_TYPES,
                default_times=default_times_for_js,
                today_str=today_iso_str,
                form_data=current_form_data,
            )

    data_to_populate_form_with = {}
    if request.method == "POST":
        data_to_populate_form_with = request.form
    elif service_assignment:  # GET pentru editare
        data_to_populate_form_with = form_data_for_template

    return render_template(
        "assign_service.html",
        form_title=form_title,
        service_assignment=service_assignment,
        students=students,
        service_types=SERVICE_TYPES,
        default_times=default_times_for_js,
        today_str=today_iso_str,  # Folosim variabila actualizată
        form_data=data_to_populate_form_with,
    )


@app.route("/gradat/services/assign_multiple", methods=["GET", "POST"])
@login_required
def assign_multiple_services():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    # POST request handles the submission from the details form (Step 2)
    if request.method == "POST":
        student_ids_to_process = [
            key.split("_")[2]
            for key in request.form
            if key.startswith("student_id_")
        ]
        added_count = 0
        error_count = 0
        conflict_details_messages = []

        for student_id_str in student_ids_to_process:
            student_id = int(student_id_str)

            # Retrieve all data for this student from the form
            service_type = request.form.get(f"service_type_{student_id}")
            service_date_str = request.form.get(f"service_date_{student_id}")
            start_time_str = request.form.get(f"start_time_{student_id}")
            end_time_str = request.form.get(f"end_time_{student_id}")
            participates = f"participates_{student_id}" in request.form
            notes = request.form.get(f"notes_{student_id}", "").strip()

            student_obj = db.session.get(Student, student_id)
            if (
                not student_obj
                or student_obj.created_by_user_id != current_user.id
            ):
                error_count += 1
                conflict_details_messages.append(
                    f"Studentul cu ID {student_id} este invalid sau nu vă aparține."
                )
                continue

            if not all(
                [service_type, service_date_str, start_time_str, end_time_str]
            ):
                error_count += 1
                conflict_details_messages.append(
                    f"Datele de serviciu sunt incomplete pentru {student_obj.nume}."
                )
                continue

            try:
                service_date_obj = datetime.strptime(
                    service_date_str, "%Y-%m-%d"
                ).date()
                start_time_obj = datetime.strptime(
                    start_time_str, "%H:%M"
                ).time()
                end_time_obj = datetime.strptime(end_time_str, "%H:%M").time()

                start_dt = datetime.combine(service_date_obj, start_time_obj)
                effective_end_date = service_date_obj
                if end_time_obj < start_time_obj or (
                    service_type == "GSS" and end_time_obj == start_time_obj
                ):
                    effective_end_date += timedelta(days=1)
                end_dt = datetime.combine(effective_end_date, end_time_obj)

                if end_dt <= start_dt:
                    raise ValueError(
                        "Data de sfârșit trebuie să fie după data de început."
                    )
            except ValueError as e:
                error_count += 1
                conflict_details_messages.append(
                    f"Format dată/oră invalid pentru {student_obj.nume}: {e}"
                )
                continue

            conflict_msg = check_service_conflict_for_student(
                student_id, start_dt, end_dt, service_type, None
            )
            if conflict_msg:
                error_count += 1
                conflict_details_messages.append(
                    f"{student_obj.nume} {student_obj.prenume}: Conflict ({conflict_msg})."
                )
                continue

            new_assignment = ServiceAssignment(
                student_id=student_id,
                service_type=service_type,
                service_date=service_date_obj,
                start_datetime=start_dt,
                end_datetime=end_dt,
                participates_in_roll_call=participates,
                notes=notes,
                created_by_user_id=current_user.id,
            )
            db.session.add(new_assignment)
            log_student_action(
                student_id,
                "SERVICE_CREATED_BULK",
                f"Serviciu '{service_type}' adăugat prin formularul multiplu: {start_dt.strftime('%d.%m %H:%M')} - {end_dt.strftime('%d.%m %H:%M')}.",
            )
            added_count += 1

        if added_count > 0:
            try:
                db.session.commit()
                flash(
                    f"{added_count} servicii au fost adăugate cu succes.",
                    "success",
                )
            except Exception as e:
                db.session.rollback()
                added_count = 0
                error_count = len(student_ids_to_process)
                flash(
                    f"Eroare majoră la salvarea serviciilor: {str(e)}",
                    "danger",
                )

        if error_count > 0:
            flash(
                f"{error_count} servicii nu au putut fi adăugate din cauza erorilor sau conflictelor.",
                "danger",
            )
            for msg in conflict_details_messages:
                flash(msg, "warning")

        return redirect(url_for("list_services"))

    # GET request handles both Step 1 (selection) and Step 2 (details form)
    student_ids_selected = request.args.getlist("student_ids", type=int)
    students_to_prepare = None

    if student_ids_selected:
        # This is Step 2: Prepare the details table
        students_to_prepare = (
            Student.query.filter(
                Student.id.in_(student_ids_selected),
                Student.created_by_user_id == current_user.id,
            )
            .order_by(Student.pluton, Student.nume)
            .all()
        )

        if len(students_to_prepare) != len(student_ids_selected):
            flash(
                "Unii studenți selectați nu au putut fi găsiți sau nu vă aparțin.",
                "warning",
            )
            return redirect(url_for("assign_multiple_services"))

    # This is Step 1: Show the student selection list
    all_students_managed = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .order_by(Student.pluton, Student.nume)
        .all()
    )
    if not all_students_managed:
        flash(
            "Nu aveți studenți în evidență pentru a le asigna servicii.",
            "warning",
        )
        return redirect(url_for("list_services"))

    default_times_for_js = {
        "GSS": ("07:00", "07:00"),
        "SVM": ("05:50", "20:00"),
        "Intervenție": ("20:00", "00:00"),
        "Planton 1": ("22:00", "00:00"),
        "Planton 2": ("00:00", "02:00"),
        "Planton 3": ("02:00", "04:00"),
        "Altul": ("", ""),
    }

    return render_template(
        "assign_multiple_services.html",
        students=all_students_managed,
        students_to_prepare=students_to_prepare,
        service_types=SERVICE_TYPES,
        default_times_json=json.dumps(default_times_for_js),
        today_str=get_localized_now().date().isoformat(),
    )


@app.route("/gradat/services/delete/<int:assignment_id>", methods=["POST"])
@login_required
def delete_service_assignment(assignment_id):
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    assign_del = ServiceAssignment.query.get_or_404(assignment_id)
    student_owner = Student.query.filter_by(
        id=assign_del.student_id, created_by_user_id=current_user.id
    ).first()

    if not student_owner:
        flash("Acces neautorizat la acest serviciu pentru ștergere.", "danger")
        return redirect(url_for("list_services"))

    student_name = (
        assign_del.student.nume + " " + assign_del.student.prenume
        if assign_del.student
        else "N/A"
    )
    service_type_deleted = assign_del.service_type
    student_id_for_log = assign_del.student_id

    try:
        log_student_action(
            student_id_for_log,
            "SERVICE_DELETED",
            f"Serviciul '{service_type_deleted}' din {assign_del.service_date.strftime('%d.%m.%Y')} a fost șters.",
        )
        db.session.delete(assign_del)
        db.session.commit()
        flash(
            f"Serviciul ({service_type_deleted}) pentru {student_name} a fost șters.",
            "success",
        )
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la ștergerea serviciului: {str(e)}", "danger")

    return redirect(url_for("list_services"))


# --- Rapoarte ---
@app.route("/company_commander/report/text", methods=["GET"])
@login_required
def text_report_display_company():
    if current_user.role != "comandant_companie":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    company_id_str = _get_commander_unit_id(current_user.username, "CmdC")
    if not company_id_str:
        flash("ID-ul companiei nu a putut fi determinat.", "danger")
        return redirect(url_for("dashboard"))

    roll_call_datetime = get_standard_roll_call_datetime()
    report_datetime_str = roll_call_datetime.strftime("%d.%m.%Y, %H:%M")

    students_in_company = Student.query.filter_by(
        companie=company_id_str
    ).all()
    if not students_in_company:
        flash(
            f"Niciun student în compania {company_id_str} pentru a genera raportul.",
            "info",
        )
        # Poate afișa un template gol sau un mesaj specific
        return render_template(
            "text_report_display.html",
            report_title=f"Raport Compania {company_id_str}",
            report_content="Niciun student în unitate.",
            report_datetime_str=report_datetime_str,
        )

    company_presence_data = _calculate_presence_data(
        students_in_company, roll_call_datetime
    )

    report_lines = []
    report_lines.append(f"RAPORT OPERATIV - COMPANIA {company_id_str}")
    report_lines.append(f"Data și ora raportului: {report_datetime_str}")
    report_lines.append("-" * 30)
    report_lines.append(
        f"Efectiv control (Ec): {company_presence_data['efectiv_control']}"
    )
    report_lines.append(
        f"Efectiv prezent (Ep): {company_presence_data['efectiv_prezent_total']}"
    )
    report_lines.append(
        f"  - În formație: {company_presence_data['in_formation_count']}"
    )
    report_lines.append(
        f"  - La Servicii: {company_presence_data['on_duty_count']}"
    )  # Changed label
    report_lines.append(
        f"  - Gradat Pluton (prezent): {company_presence_data['platoon_graded_duty_count']}"
    )
    report_lines.append(
        f"Efectiv absent (Ea): {company_presence_data['efectiv_absent_total']}"
    )
    report_lines.append("-" * 30)

    if company_presence_data["in_formation_students_details"]:
        report_lines.append("\nPREZENȚI ÎN FORMAȚIE:")
        for detail in company_presence_data["in_formation_students_details"]:
            report_lines.append(f"  - {detail}")

    if company_presence_data["on_duty_students_details"]:
        report_lines.append("\nLA SERVICII:")  # Changed label
        for detail in company_presence_data["on_duty_students_details"]:
            report_lines.append(f"  - {detail}")

    if company_presence_data["platoon_graded_duty_students_details"]:
        report_lines.append("\nGRADAȚI (în afara formației):")
        for detail in company_presence_data[
            "platoon_graded_duty_students_details"
        ]:
            report_lines.append(f"  - {detail}")
    if company_presence_data.get("present_exempt_not_in_formation_details"):
        report_lines.append("\nSCUTIȚI TEMPORAR (prezenți):")
        for detail in company_presence_data[
            "present_exempt_not_in_formation_details"
        ]:
            report_lines.append(f"  - {detail}")

    if company_presence_data["absent_students_details"]:
        report_lines.append("\nABSENȚI MOTIVAT:")
        for detail in company_presence_data["absent_students_details"]:
            report_lines.append(f"  - {detail}")

    report_lines.append("\n" + "-" * 30)
    report_lines.append("Raport generat de sistem.")

    final_report_content = "\n".join(report_lines)

    return render_template(
        "text_report_display.html",
        report_title=f"Raport Text Compania {company_id_str}",
        report_content=final_report_content,
        report_datetime_str=report_datetime_str,
    )


@app.route("/battalion_commander/report/text", methods=["GET"])
@login_required
def text_report_display_battalion():
    if current_user.role != "comandant_batalion":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    battalion_id_str = _get_commander_unit_id(current_user.username, "CmdB")
    if not battalion_id_str:
        flash("ID-ul batalionului nu a putut fi determinat.", "danger")
        return redirect(url_for("dashboard"))

    roll_call_datetime = get_standard_roll_call_datetime()
    report_datetime_str = roll_call_datetime.strftime("%d.%m.%Y, %H:%M")

    students_in_battalion = Student.query.filter_by(
        batalion=battalion_id_str
    ).all()
    if not students_in_battalion:
        flash(
            f"Niciun student în batalionul {battalion_id_str} pentru a genera raportul.",
            "info",
        )
        return render_template(
            "text_report_display.html",
            report_title=f"Raport Batalionul {battalion_id_str}",
            report_content="Niciun student în unitate.",
            report_datetime_str=report_datetime_str,
        )

    total_battalion_presence = _calculate_presence_data(
        students_in_battalion, roll_call_datetime
    )

    report_lines = []
    report_lines.append(f"RAPORT OPERATIV - BATALIONUL {battalion_id_str}")
    report_lines.append(f"Data și ora raportului: {report_datetime_str}")
    report_lines.append("=" * 40)
    report_lines.append("SITUAȚIE CENTRALIZATOARE BATALION:")
    report_lines.append(
        f"  Efectiv control (Ec): {total_battalion_presence['efectiv_control']}"
    )
    report_lines.append(
        f"  Efectiv prezent (Ep): {total_battalion_presence['efectiv_prezent_total']}"
    )
    report_lines.append(
        f"    - În formație: {total_battalion_presence['in_formation_count']}"
    )
    report_lines.append(
        f"    - La Servicii: {total_battalion_presence['on_duty_count']}"
    )  # Changed label
    report_lines.append(
        f"    - Gradat Pluton (prezent): {total_battalion_presence['platoon_graded_duty_count']}"
    )
    report_lines.append(
        f"  Efectiv absent (Ea): {total_battalion_presence['efectiv_absent_total']}"
    )
    report_lines.append("=" * 40)

    companies_in_battalion = sorted(
        list(set(s.companie for s in students_in_battalion if s.companie))
    )
    for company_id_loop in companies_in_battalion:
        students_in_company_loop = [
            s for s in students_in_battalion if s.companie == company_id_loop
        ]
        company_presence_data = _calculate_presence_data(
            students_in_company_loop, roll_call_datetime
        )

        report_lines.append(f"\nSITUAȚIE COMPANIA {company_id_loop}:")
        report_lines.append(
            f"  Ec: {company_presence_data['efectiv_control']}, Ep: {company_presence_data['efectiv_prezent_total']}, Ea: {company_presence_data['efectiv_absent_total']}"
        )
        report_lines.append(
            f"    În formație: {company_presence_data['in_formation_count']}"
        )
        report_lines.append(
            f"    La Servicii: {company_presence_data['on_duty_count']}"
        )  # Changed label
        report_lines.append(
            f"    Gradați (în afara formației): {company_presence_data['platoon_graded_duty_count']}"
        )

        if company_presence_data.get("platoon_graded_duty_students_details"):
            report_lines.append("    Gradați (în afara formației) - detalii:")
            for detail in company_presence_data[
                "platoon_graded_duty_students_details"
            ]:
                report_lines.append(f"      - {detail}")
        if company_presence_data.get(
            "present_exempt_not_in_formation_details"
        ):
            report_lines.append("    Scutiți temporar (prezenți):")
            for detail in company_presence_data[
                "present_exempt_not_in_formation_details"
            ]:
                report_lines.append(f"      - {detail}")
        if company_presence_data["absent_students_details"]:
            report_lines.append("    Absenți motivat:")
            for detail in company_presence_data["absent_students_details"]:
                report_lines.append(f"      - {detail}")
        report_lines.append("-" * 30)

    report_lines.append("\n" + "=" * 40)
    report_lines.append("DETALII ABSENȚE LA NIVEL DE BATALION (dacă există):")
    if total_battalion_presence["absent_students_details"]:
        for detail in total_battalion_presence["absent_students_details"]:
            report_lines.append(f"  - {detail}")
    else:
        report_lines.append(
            "  Nicio absență înregistrată la nivel de batalion."
        )

    report_lines.append("\n" + "=" * 40)
    report_lines.append("Raport generat de sistem.")
    final_report_content = "\n".join(report_lines)

    return render_template(
        "text_report_display.html",
        report_title=f"Raport Text Batalionul {battalion_id_str}",
        report_content=final_report_content,
        report_datetime_str=report_datetime_str,
    )


# --- Admin Action Log Viewer ---
@app.route("/admin/action_logs", endpoint="admin_action_logs")
@login_required
def admin_action_logs():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    page = request.args.get("page", 1, type=int)
    per_page = 20  # Or a configurable value

    # Basic query
    logs_query = ActionLog.query.options(joinedload(ActionLog.user)).order_by(
        ActionLog.timestamp.desc()
    )

    # TODO: Add filtering based on request.args
    # filter_user_id = request.args.get('user_id')
    # filter_action_type = request.args.get('action_type')
    # filter_target_model = request.args.get('target_model')
    # filter_date_from = request.args.get('date_from')
    filter_user_id_str = request.args.get("user_id_filter_val", "").strip()
    filter_action_type = request.args.get("action_type_filter_val", "").strip()
    filter_target_model = request.args.get(
        "target_model_filter_val", ""
    ).strip()
    filter_date_from_str = request.args.get("filter_date_from", "").strip()
    filter_date_to_str = request.args.get("filter_date_to", "").strip()

    if filter_user_id_str:
        try:
            filter_user_id = int(filter_user_id_str)
            logs_query = logs_query.filter(ActionLog.user_id == filter_user_id)
        except ValueError:
            flash("ID Utilizator invalid pentru filtrare.", "warning")

    if filter_action_type:
        logs_query = logs_query.filter(
            ActionLog.action_type.ilike(f"%{filter_action_type}%")
        )

    if filter_target_model:
        logs_query = logs_query.filter(
            ActionLog.target_model.ilike(f"%{filter_target_model}%")
        )

    if filter_date_from_str:
        try:
            date_from = datetime.strptime(filter_date_from_str, "%Y-%m-%d")
            # logs_query = logs_query.filter(ActionLog.timestamp >= date_from)
            # To handle timestamp comparison correctly, especially if timestamp is timezone-aware (UTC)
            # and date_from is naive, it's better to define date_from as start of day in UTC.
            # However, ActionLog.timestamp is default=datetime.utcnow, so it's already UTC.
            # We can make date_from timezone-aware UTC or ensure comparison handles this.
            # For simplicity, if timestamps are stored UTC, compare with UTC datetimes.
            # If date_from is just a date, it implies start of that day.
            logs_query = logs_query.filter(
                ActionLog.timestamp
                >= datetime.combine(date_from, time.min).replace(
                    tzinfo=pytz.UTC
                )
            )
        except ValueError:
            flash(
                "Format dată 'De la' invalid. Folosiți YYYY-MM-DD.", "warning"
            )

    if filter_date_to_str:
        try:
            date_to = datetime.strptime(filter_date_to_str, "%Y-%m-%d")
            # To include the whole day, filter up to end of day.
            # logs_query = logs_query.filter(ActionLog.timestamp <= datetime.combine(date_to, time.max))
            # Similar to date_from, ensure timezone consistency.
            logs_query = logs_query.filter(
                ActionLog.timestamp
                <= datetime.combine(date_to, time.max).replace(tzinfo=pytz.UTC)
            )
        except ValueError:
            flash(
                "Format dată 'Până la' invalid. Folosiți YYYY-MM-DD.",
                "warning",
            )

    logs_pagination = logs_query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    # Convert iter_pages generator to a list for use with |length in template
    page_iterator_list = list(
        logs_pagination.iter_pages(
            left_edge=1, right_edge=1, left_current=2, right_current=3
        )
    )

    return render_template(
        "admin_action_logs.html",
        logs_pagination=logs_pagination,
        page_iterator_list=page_iterator_list,  # Pass the list
        title="Jurnal Acțiuni Sistem",
        # Pass filter values back to template
        user_id_filter_val=filter_user_id_str,
        action_type_filter_val=filter_action_type,
        target_model_filter_val=filter_target_model,
        filter_date_from=filter_date_from_str,
        filter_date_to=filter_date_to_str,
    )


@app.route(
    "/admin/profile/change_password",
    methods=["GET", "POST"],
    endpoint="admin_change_self_password",
)
@login_required
def admin_change_self_password():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        current_password = request.form.get("current_password")
        new_password = request.form.get("new_password")
        confirm_new_password = request.form.get("confirm_new_password")

        if not current_user.check_password(current_password):
            flash("Parola curentă este incorectă.", "danger")
            return redirect(url_for("admin_change_self_password"))

        if (
            not new_password or len(new_password) < 6
        ):  # Exemplu de politică simplă de parolă
            flash("Parola nouă trebuie să aibă minim 6 caractere.", "warning")
            return redirect(url_for("admin_change_self_password"))

        if new_password != confirm_new_password:
            flash("Parolele noi nu se potrivesc.", "warning")
            return redirect(url_for("admin_change_self_password"))

        details_before = {
            "user_id": current_user.id,
            "username": current_user.username,
            "action": "Attempt change own password",
        }
        current_user.set_password(new_password)
        try:
            log_action(
                "ADMIN_CHANGE_OWN_PASSWORD_SUCCESS",
                target_model_name="User",
                target_id=current_user.id,
                details_before_dict=details_before,  # details_after nu arată parola
                description=f"Admin {current_user.username} changed their own password successfully.",
            )
            db.session.commit()
            flash("Parola a fost schimbată cu succes!", "success")
            # Opțional, se poate face logout adminului după schimbarea parolei pentru a forța re-autentificare
            # logout_user()
            # flash('Parola a fost schimbată cu succes! Te rugăm să te autentifici din nou.', 'success')
            # return redirect(url_for('admin_login'))
            return redirect(url_for("admin_dashboard_route"))
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la schimbarea parolei: {str(e)}", "danger")
            log_action(
                "ADMIN_CHANGE_OWN_PASSWORD_FAIL",
                target_model_name="User",
                target_id=current_user.id,
                details_before_dict=details_before,
                description=f"Admin {current_user.username} failed to change their own password. Error: {str(e)}",
            )
            db.session.commit()  # Commit log-ul de eroare
            return redirect(url_for("admin_change_self_password"))

    return render_template("admin_change_password.html")


@app.route(
    "/admin/user/<int:user_id>/set_personal_code",
    methods=["GET", "POST"],
    endpoint="admin_set_user_personal_code",
)
@login_required
def admin_set_user_personal_code(user_id):
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    target_user = db.session.get(User, user_id)
    if not target_user:
        flash("Utilizatorul specificat nu a fost găsit.", "danger")
        return redirect(url_for("admin_dashboard_route"))

    if target_user.role == "admin":
        flash(
            "Codul personal nu poate fi setat pentru un alt administrator prin această metodă.",
            "warning",
        )
        return redirect(url_for("admin_dashboard_route"))

    if request.method == "POST":
        new_code = request.form.get("new_personal_code")
        confirm_new_code = request.form.get("confirm_new_personal_code")

        if not new_code or len(new_code) < 4:
            flash(
                "Noul cod personal trebuie să aibă minim 4 caractere.",
                "warning",
            )
            return render_template(
                "admin_set_user_personal_code.html", target_user=target_user
            )

        if new_code != confirm_new_code:
            flash("Codurile personale introduse nu se potrivesc.", "warning")
            return render_template(
                "admin_set_user_personal_code.html", target_user=target_user
            )

        # Security fix: Check if the personal code is already in use by another user
        other_users = User.query.filter(
            User.id != target_user.id,
            User.role != "admin",
            User.personal_code_hash.isnot(None),
        ).all()
        for other_user in other_users:
            if other_user.check_personal_code(new_code):
                flash(
                    "Acest cod personal este deja utilizat de altcineva. Vă rugăm alegeți un alt cod.",
                    "danger",
                )
                return render_template(
                    "admin_set_user_personal_code.html",
                    target_user=target_user,
                )

        details_before = model_to_dict(
            target_user,
            exclude_fields=[
                "password_hash",
                "unique_code",
                "personal_code_hash",
            ],
        )
        details_before["personal_code_was_set"] = (
            target_user.personal_code_hash is not None
        )
        details_before["was_first_login"] = target_user.is_first_login

        target_user.set_personal_code(
            new_code
        )  # Aceasta setează hash-ul și is_first_login = False

        try:
            details_after = model_to_dict(
                target_user,
                exclude_fields=[
                    "password_hash",
                    "unique_code",
                    "personal_code_hash",
                ],
            )
            details_after["personal_code_is_set"] = True
            details_after["is_first_login"] = False

            log_action(
                "ADMIN_SET_USER_PERSONAL_CODE_SUCCESS",
                target_model_name="User",
                target_id=target_user.id,
                details_before_dict=details_before,
                details_after_dict=details_after,
                description=f"Admin {current_user.username} set new personal code for user {target_user.username} (ID: {target_user.id}).",
            )
            db.session.commit()
            flash(
                f"Noul cod personal pentru utilizatorul {target_user.username} a fost setat cu succes.",
                "success",
            )
            return redirect(url_for("admin_dashboard_route"))
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la setarea codului personal: {str(e)}", "danger")
            # Log failure
            log_action(
                "ADMIN_SET_USER_PERSONAL_CODE_FAIL",
                target_model_name="User",
                target_id=target_user.id,
                details_before_dict=details_before,  # Log state before attempted change
                description=f"Admin {current_user.username} failed to set new personal code for {target_user.username}. Error: {str(e)}",
            )
            db.session.commit()
            return render_template(
                "admin_set_user_personal_code.html", target_user=target_user
            )

    return render_template(
        "admin_set_user_personal_code.html", target_user=target_user
    )


# --- Admin UpdateTopic (Announcements) Management ---
AVAILABLE_STATUS_COLORS = [
    "primary",
    "secondary",
    "success",
    "danger",
    "warning",
    "info",
    "light",
    "dark",
]


@app.route("/admin/updates", methods=["GET"], endpoint="admin_list_updates")
@login_required
def admin_list_updates():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    page = request.args.get("page", 1, type=int)
    per_page = 15  # Or your preferred number
    updates_pagination = (
        UpdateTopic.query.join(User, UpdateTopic.user_id == User.id)
        .options(joinedload(UpdateTopic.author))
        .order_by(UpdateTopic.is_pinned.desc(), UpdateTopic.updated_at.desc())
        .paginate(page=page, per_page=per_page, error_out=False)
    )

    return render_template(
        "admin_list_updates.html",
        updates_pagination=updates_pagination,
        title="Management Anunțuri",
    )


@app.route(
    "/admin/updates/create",
    methods=["GET", "POST"],
    endpoint="admin_create_update",
)
@login_required
def admin_create_update():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("admin_list_updates"))

    form_data_to_pass = request.form if request.method == "POST" else {}

    if request.method == "POST":
        title = request.form.get("title", "").strip()
        content = request.form.get("content", "").strip()
        status_color = request.form.get("status_color")
        is_pinned = "is_pinned" in request.form
        is_visible = "is_visible" in request.form

        if not title or not content:
            flash("Titlul și conținutul sunt obligatorii.", "warning")
            return render_template(
                "admin_edit_update_topic.html",
                title="Creează Anunț Nou",
                topic=None,  # No topic object for create
                available_colors=AVAILABLE_STATUS_COLORS,
                form_data=request.form,
            )  # Repopulate with current form data

        if status_color == "None" or not status_color:
            status_color = None  # Store as None if "Niciuna" selected

        new_topic = UpdateTopic(
            title=title,
            content=content,
            user_id=current_user.id,
            status_color=status_color,
            is_pinned=is_pinned,
            is_visible=is_visible,
        )
        db.session.add(new_topic)
        try:
            db.session.commit()
            flash(f'Anunțul "{title}" a fost creat cu succes.', "success")
            return redirect(url_for("admin_list_updates"))
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la crearea anunțului: {str(e)}", "danger")
            # form_data already contains current form input

    return render_template(
        "admin_edit_update_topic.html",
        title="Creează Anunț Nou",
        topic=None,  # Explicitly pass None for topic
        available_colors=AVAILABLE_STATUS_COLORS,
        form_data=form_data_to_pass,
    )  # Empty for GET, or form data for failed POST


@app.route(
    "/admin/updates/edit/<int:topic_id>",
    methods=["GET", "POST"],
    endpoint="admin_edit_update",
)
@login_required
def admin_edit_update(topic_id):
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("admin_list_updates"))

    topic_to_edit = db.session.get(UpdateTopic, topic_id)
    if not topic_to_edit:
        flash("Anunțul nu a fost găsit.", "danger")
        return redirect(url_for("admin_list_updates"))

    form_data_on_get = (
        {  # Populate with existing topic data for GET
            "title": topic_to_edit.title,
            "content": topic_to_edit.content,
            "status_color": topic_to_edit.status_color or "None",
            "is_pinned": topic_to_edit.is_pinned,
            "is_visible": topic_to_edit.is_visible,
        }
        if request.method == "GET"
        else request.form
    )  # Use request.form for POST repopulation

    if request.method == "POST":
        topic_to_edit.title = request.form.get("title", "").strip()
        topic_to_edit.content = request.form.get("content", "").strip()
        status_color_form = request.form.get("status_color")
        topic_to_edit.is_pinned = "is_pinned" in request.form
        topic_to_edit.is_visible = "is_visible" in request.form

        if not topic_to_edit.title or not topic_to_edit.content:
            flash("Titlul și conținutul sunt obligatorii.", "warning")
            # Pass current (failed) form data back to template
            return render_template(
                "admin_edit_update_topic.html",
                title=f"Editare Anunț: {topic_to_edit.title[:30]}...",
                topic=topic_to_edit,
                available_colors=AVAILABLE_STATUS_COLORS,
                form_data=request.form,
            )

        topic_to_edit.status_color = (
            None
            if status_color_form == "None" or not status_color_form
            else status_color_form
        )
        topic_to_edit.updated_at = (
            datetime.utcnow()
        )  # Manually update 'updated_at'

        try:
            db.session.commit()
            flash(
                f'Anunțul "{topic_to_edit.title}" a fost actualizat.',
                "success",
            )
            return redirect(url_for("admin_list_updates"))
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la actualizarea anunțului: {str(e)}", "danger")
            # Pass current (failed) form data back

    # For GET, or for POST if validation failed and we re-render
    return render_template(
        "admin_edit_update_topic.html",
        title=f"Editare Anunț: {topic_to_edit.title[:30]}...",
        topic=topic_to_edit,
        available_colors=AVAILABLE_STATUS_COLORS,
        form_data=form_data_on_get,
    )


@app.route(
    "/admin/updates/delete/<int:topic_id>",
    methods=["POST"],
    endpoint="admin_delete_update",
)
@login_required
def admin_delete_update(topic_id):
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("admin_list_updates"))
    topic = db.session.get(UpdateTopic, topic_id)
    if topic:
        try:
            db.session.delete(topic)
            db.session.commit()
            flash(f'Anunțul "{topic.title}" a fost șters.', "success")
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la ștergerea anunțului: {str(e)}", "danger")
    else:
        flash("Anunțul nu a fost găsit.", "warning")
    return redirect(url_for("admin_list_updates"))


@app.route(
    "/admin/updates/toggle_pin/<int:topic_id>",
    methods=["POST"],
    endpoint="admin_toggle_pin_update",
)
@login_required
def admin_toggle_pin_update(topic_id):
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("admin_list_updates"))
    topic = db.session.get(UpdateTopic, topic_id)
    if topic:
        topic.is_pinned = not topic.is_pinned
        topic.updated_at = datetime.utcnow()
        try:
            db.session.commit()
            flash(
                f'Statusul "Fixat" pentru anunțul "{topic.title}" a fost schimbat.',
                "info",
            )
        except Exception as e:
            db.session.rollback()
            flash(
                f'Eroare la schimbarea statusului "Fixat": {str(e)}', "danger"
            )
    else:
        flash("Anunțul nu a fost găsit.", "warning")
    return redirect(url_for("admin_list_updates"))


@app.route(
    "/admin/updates/toggle_visibility/<int:topic_id>",
    methods=["POST"],
    endpoint="admin_toggle_visibility_update",
)
@login_required
def admin_toggle_visibility_update(topic_id):
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("admin_list_updates"))
    topic = db.session.get(UpdateTopic, topic_id)
    if topic:
        topic.is_visible = not topic.is_visible
        topic.updated_at = datetime.utcnow()  # Consider this an update
        try:
            db.session.commit()
            flash(
                f'Vizibilitatea pentru anunțul "{topic.title}" a fost schimbată.',
                "info",
            )
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la schimbarea vizibilității: {str(e)}", "danger")
    else:
        flash("Anunțul nu a fost găsit.", "warning")
    return redirect(url_for("admin_list_updates"))


@app.route(
    "/gradat/weekend_leave/bulk_add",
    methods=["GET", "POST"],
    endpoint="gradat_bulk_add_weekend_leave",
)
@login_required
def gradat_bulk_add_weekend_leave():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    students_managed = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .order_by(Student.pluton, Student.nume)
        .all()
    )
    upcoming_fridays_list = get_upcoming_fridays()  # Funcție existentă

    if request.method == "POST":
        student_ids_selected = request.form.getlist("student_ids")
        weekend_start_date_str = request.form.get("weekend_start_date")
        selected_days_names = request.form.getlist(
            "selected_days"
        )  # Numele zilelor: 'Vineri', 'Sambata', 'Duminica'
        reason_common = request.form.get("reason", "").strip()

        if not student_ids_selected:
            flash("Nu ați selectat niciun student.", "warning")
            return render_template(
                "bulk_add_weekend_leave.html",
                students=students_managed,
                upcoming_fridays=upcoming_fridays_list,
                form_data=request.form,
            )

        if not weekend_start_date_str:
            flash(
                "Data de început a weekendului (Vineri) este obligatorie.",
                "warning",
            )
            return render_template(
                "bulk_add_weekend_leave.html",
                students=students_managed,
                upcoming_fridays=upcoming_fridays_list,
                form_data=request.form,
            )

        if not selected_days_names:
            flash("Nu ați selectat nicio zi din weekend.", "warning")
            return render_template(
                "bulk_add_weekend_leave.html",
                students=students_managed,
                upcoming_fridays=upcoming_fridays_list,
                form_data=request.form,
            )

        try:
            friday_date_obj = datetime.strptime(
                weekend_start_date_str, "%Y-%m-%d"
            ).date()
            if friday_date_obj.weekday() != 4:  # Vineri
                flash(
                    "Data de început a weekendului selectată nu este o zi de Vineri.",
                    "warning",
                )
                return render_template(
                    "bulk_add_weekend_leave.html",
                    students=students_managed,
                    upcoming_fridays=upcoming_fridays_list,
                    form_data=request.form,
                )
        except ValueError:
            flash("Format dată weekend invalid.", "danger")
            return render_template(
                "bulk_add_weekend_leave.html",
                students=students_managed,
                upcoming_fridays=upcoming_fridays_list,
                form_data=request.form,
            )

        day_inputs_from_form = []
        for day_name_form in selected_days_names:
            start_time_str = request.form.get(
                f"bulk_{day_name_form.lower()}_start_time"
            )
            end_time_str = request.form.get(
                f"bulk_{day_name_form.lower()}_end_time"
            )

            if not start_time_str or not end_time_str:
                flash(
                    f"Orele de început și sfârșit sunt obligatorii pentru {day_name_form} în formularul de adăugare rapidă.",
                    "warning",
                )
                return render_template(
                    "bulk_add_weekend_leave.html",
                    students=students_managed,
                    upcoming_fridays=upcoming_fridays_list,
                    form_data=request.form,
                )
            try:
                start_time_obj = datetime.strptime(
                    start_time_str, "%H:%M"
                ).time()
                end_time_obj = datetime.strptime(end_time_str, "%H:%M").time()
            except ValueError:
                flash(f"Format oră invalid pentru {day_name_form}.", "danger")
                return render_template(
                    "bulk_add_weekend_leave.html",
                    students=students_managed,
                    upcoming_fridays=upcoming_fridays_list,
                    form_data=request.form,
                )

            if end_time_obj == start_time_obj:
                flash(
                    f"Ora de început și sfârșit nu pot fi identice pentru {day_name_form}.",
                    "warning",
                )
                return render_template(
                    "bulk_add_weekend_leave.html",
                    students=students_managed,
                    upcoming_fridays=upcoming_fridays_list,
                    form_data=request.form,
                )

            day_offset_map = {"Vineri": 0, "Sambata": 1, "Duminica": 2}
            actual_date_for_day = friday_date_obj + timedelta(
                days=day_offset_map[day_name_form]
            )

            # Validare interval (similar cu formularul individual)
            current_interval_start_dt = datetime.combine(
                actual_date_for_day, start_time_obj
            )
            effective_end_date_for_interval = actual_date_for_day
            if end_time_obj < start_time_obj:  # Trece în ziua următoare
                effective_end_date_for_interval += timedelta(days=1)
            current_interval_end_dt = datetime.combine(
                effective_end_date_for_interval, end_time_obj
            )

            if current_interval_end_dt <= current_interval_start_dt:
                flash(
                    f"Interval orar invalid pentru {day_name_form} (sfârșitul trebuie să fie după început).",
                    "warning",
                )
                return render_template(
                    "bulk_add_weekend_leave.html",
                    students=students_managed,
                    upcoming_fridays=upcoming_fridays_list,
                    form_data=request.form,
                )

            day_inputs_from_form.append(
                {
                    "name": day_name_form,
                    "date": actual_date_for_day,
                    "start_time": start_time_obj,
                    "end_time": end_time_obj,
                }
            )

        day_inputs_from_form.sort(
            key=lambda x: x["date"]
        )  # Sortează după dată

        added_count = 0
        skipped_due_to_conflict_count = 0
        error_details_conflict = []

        duminica_biserica_selected = "duminica_biserica" in request.form

        for student_id_str in student_ids_selected:
            student_id = int(student_id_str)
            # Verificare conflict pentru acest student cu intervalele selectate
            conflict_for_this_student = False
            for day_data_input in day_inputs_from_form:
                start_dt_check = datetime.combine(
                    day_data_input["date"], day_data_input["start_time"]
                )
                end_dt_check = datetime.combine(
                    day_data_input["date"], day_data_input["end_time"]
                )
                if (
                    day_data_input["end_time"] < day_data_input["start_time"]
                ):  # trece in ziua urmatoare
                    end_dt_check += timedelta(days=1)

                # Simplificare: check_leave_conflict verifică și servicii. Pentru bulk, putem fi mai permisivi sau loga conflictele
                # Aici vom face o verificare de bază, dar ideal ar fi o strategie mai complexă.
                # Momentan, vom sări peste student dacă există *orice* conflict pe *oricare* din zilele selectate.
                conflict_reason = check_leave_conflict(
                    student_id,
                    start_dt_check,
                    end_dt_check,
                    leave_type="weekend_leave",
                    existing_leave_id=None,
                )
                if conflict_reason:
                    student_obj = db.session.get(Student, student_id)
                    error_details_conflict.append(
                        f"Studentul {student_obj.nume} {student_obj.prenume}: conflict pe {day_data_input['name']} ({conflict_reason}). Învoirea nu a fost adăugată."
                    )
                    conflict_for_this_student = True
                    break

            if conflict_for_this_student:
                skipped_due_to_conflict_count += 1
                continue

            # Creează învoirea
            new_leave = WeekendLeave(
                student_id=student_id,
                weekend_start_date=friday_date_obj,
                reason=reason_common,
                status="Aprobată",
                created_by_user_id=current_user.id,
                duminica_biserica=duminica_biserica_selected,
            )

            # Atribuie zilele și orele
            if len(day_inputs_from_form) >= 1:
                new_leave.day1_selected = day_inputs_from_form[0]["name"]
                new_leave.day1_date = day_inputs_from_form[0]["date"]
                new_leave.day1_start_time = day_inputs_from_form[0][
                    "start_time"
                ]
                new_leave.day1_end_time = day_inputs_from_form[0]["end_time"]
            if len(day_inputs_from_form) >= 2:
                new_leave.day2_selected = day_inputs_from_form[1]["name"]
                new_leave.day2_date = day_inputs_from_form[1]["date"]
                new_leave.day2_start_time = day_inputs_from_form[1][
                    "start_time"
                ]
                new_leave.day2_end_time = day_inputs_from_form[1]["end_time"]
            if len(day_inputs_from_form) >= 3:
                new_leave.day3_selected = day_inputs_from_form[2]["name"]
                new_leave.day3_date = day_inputs_from_form[2]["date"]
                new_leave.day3_start_time = day_inputs_from_form[2][
                    "start_time"
                ]
                new_leave.day3_end_time = day_inputs_from_form[2]["end_time"]

            db.session.add(new_leave)
            log_student_action(
                student_id,
                "WEEKEND_LEAVE_CREATED_BULK",
                f"Învoire weekend adăugată prin formular multiplu pentru {friday_date_obj.strftime('%d.%m')}.",
            )
            added_count += 1

        try:
            db.session.commit()
            if added_count > 0:
                flash(
                    f"{added_count} învoiri de weekend au fost adăugate cu succes.",
                    "success",
                )
            if skipped_due_to_conflict_count > 0:
                flash(
                    f"{skipped_due_to_conflict_count} învoiri au fost omise din cauza conflictelor existente. Detalii mai jos.",
                    "warning",
                )
                for err_detail in error_details_conflict:
                    flash(err_detail, "info")  # Afișează fiecare conflict
            if (
                added_count == 0
                and skipped_due_to_conflict_count == 0
                and len(student_ids_selected) > 0
            ):
                flash(
                    "Nicio învoire nu a fost adăugată. Verificați selecțiile și încercați din nou.",
                    "info",
                )

            # Logare acțiune bulk
            log_action(
                "BULK_ADD_WEEKEND_LEAVE",
                description=f"User {current_user.username} attempted bulk weekend leave. Added: {added_count}, Skipped (conflict): {skipped_due_to_conflict_count} for weekend starting {friday_date_obj.isoformat()}.",
                details_after_dict={
                    "students_selected_count": len(student_ids_selected),
                    "days_selected": selected_days_names,
                    "conflicts_details": error_details_conflict[:5],
                },
            )  # Log first 5 conflicts
            db.session.commit()  # Commit log-ul separat

            return redirect(url_for("list_weekend_leaves"))
        except Exception as e:
            db.session.rollback()
            flash(
                f"Eroare la salvarea învoirilor de weekend: {str(e)}", "danger"
            )
            log_action(
                "BULK_ADD_WEEKEND_LEAVE_FAIL",
                description=f"Bulk add weekend leave failed for user {current_user.username}. Error: {str(e)}",
            )
            db.session.commit()

    return render_template(
        "bulk_add_weekend_leave.html",
        students=students_managed,
        upcoming_fridays=upcoming_fridays_list,
        form_data=None,
    )


@app.route(
    "/gradat/permission/bulk_add",
    methods=["GET", "POST"],
    endpoint="gradat_bulk_add_permission",
)
@login_required
def gradat_bulk_add_permission():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        # This is Step 2: Processing the details form
        student_ids_to_process = [
            key.split("_")[2]
            for key in request.form
            if key.startswith("student_id_")
        ]

        added_count = 0
        error_count = 0
        conflict_details_messages = []

        for student_id_str in student_ids_to_process:
            student_id = int(student_id_str)
            start_dt_str = request.form.get(f"start_datetime_{student_id}")
            end_dt_str = request.form.get(f"end_datetime_{student_id}")
            destination = request.form.get(
                f"destination_{student_id}", ""
            ).strip()
            transport_mode = request.form.get(
                f"transport_mode_{student_id}", ""
            ).strip()
            reason = request.form.get(f"reason_{student_id}", "").strip()

            student_obj = db.session.get(Student, student_id)
            if (
                not student_obj
                or student_obj.created_by_user_id != current_user.id
            ):
                error_count += 1
                conflict_details_messages.append(
                    f"Student ID {student_id} invalid sau nu vă aparține."
                )
                continue

            if not start_dt_str or not end_dt_str:
                error_count += 1
                conflict_details_messages.append(
                    f"Datele de început/sfârșit lipsesc pentru {student_obj.nume}."
                )
                continue

            try:
                start_dt = datetime.strptime(start_dt_str, "%Y-%m-%dT%H:%M")
                end_dt = datetime.strptime(end_dt_str, "%Y-%m-%dT%H:%M")
                if end_dt <= start_dt:
                    raise ValueError(
                        "Data de sfârșit trebuie să fie după data de început."
                    )
            except ValueError as e:
                error_count += 1
                conflict_details_messages.append(
                    f"Format dată invalid pentru {student_obj.nume}: {e}"
                )
                continue

            conflict_msg = check_leave_conflict(
                student_id, start_dt, end_dt, "permission", None
            )
            if conflict_msg:
                error_count += 1
                conflict_details_messages.append(
                    f"{student_obj.nume} {student_obj.prenume}: Conflict ({conflict_msg})."
                )
                continue

            new_permission = Permission(
                student_id=student_id,
                start_datetime=start_dt,
                end_datetime=end_dt,
                destination=destination,
                transport_mode=transport_mode,
                reason=reason,
                status="Aprobată",
                created_by_user_id=current_user.id,
            )
            db.session.add(new_permission)
            log_student_action(
                student_id,
                "PERMISSION_CREATED_BULK",
                f"Permisie adăugată prin formular multiplu: {start_dt.strftime('%d.%m %H:%M')} - {end_dt.strftime('%d.%m %H:%M')}.",
            )
            added_count += 1

        if added_count > 0:
            try:
                db.session.commit()
                flash(
                    f"{added_count} permisii au fost adăugate cu succes.",
                    "success",
                )
            except Exception as e:
                db.session.rollback()
                added_count = 0
                error_count = len(student_ids_to_process)
                flash(
                    f"Eroare majoră la salvarea permisiilor: {str(e)}",
                    "danger",
                )

        if error_count > 0:
            flash(
                f"{error_count} permisii nu au putut fi adăugate din cauza erorilor sau conflictelor.",
                "danger",
            )
            for msg in conflict_details_messages:
                flash(msg, "warning")

        return redirect(url_for("list_permissions"))

    # This is Step 1: Displaying the student selection or the details form
    student_ids_selected = request.args.getlist("student_ids", type=int)
    students_to_prepare = None

    if student_ids_selected:
        # Step 2 View: Prepare the details table
        students_to_prepare = (
            Student.query.filter(
                Student.id.in_(student_ids_selected),
                Student.created_by_user_id == current_user.id,
            )
            .order_by(Student.pluton, Student.nume)
            .all()
        )
        # Verify all requested students were found and belong to the user
        if len(students_to_prepare) != len(student_ids_selected):
            flash(
                "Unii studenți selectați nu au putut fi găsiți sau nu vă aparțin.",
                "warning",
            )
            return redirect(url_for("gradat_bulk_add_permission"))

    # Step 1 View: Show the student selection list
    all_students_managed = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .order_by(Student.pluton, Student.nume)
        .all()
    )

    return render_template(
        "bulk_add_permission.html",
        students=all_students_managed,
        students_to_prepare=students_to_prepare,
    )


# --- Istoric Învoiri Gradat ---
@app.route(
    "/gradat/invoiri/istoric",
    methods=["GET"],
    endpoint="gradat_invoiri_istoric",
)
@login_required
def gradat_invoiri_istoric():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    students_managed_by_gradat = (
        Student.query.filter_by(created_by_user_id=current_user.id)
        .with_entities(Student.id)
        .all()
    )
    student_ids = [sid for (sid,) in students_managed_by_gradat]

    if not student_ids:
        return render_template(
            "invoiri_istoric.html",
            leaves_history=[],
            title="Istoric Învoiri Pluton",
            form_data=request.args,
        )

    # Parsare filtre din request
    perioada = request.args.get(
        "perioada", "ultimele_7_zile"
    )  # Default la ultimele 7 zile
    data_start_custom_str = request.args.get("data_start_custom")
    data_sfarsit_custom_str = request.args.get("data_sfarsit_custom")

    end_date = (
        get_localized_now().date()
    )  # Data de sfârșit a intervalului implicit
    start_date = None

    if perioada == "ieri":
        start_date = end_date - timedelta(days=1)
        end_date = start_date  # Pentru 'ieri', intervalul este o singură zi
    elif perioada == "ultimele_2_zile":
        start_date = end_date - timedelta(
            days=1
        )  # Ultimele 2 zile înseamnă ieri și azi
    elif perioada == "ultimele_7_zile":
        start_date = end_date - timedelta(days=6)
    elif (
        perioada == "custom"
        and data_start_custom_str
        and data_sfarsit_custom_str
    ):
        try:
            start_date = datetime.strptime(
                data_start_custom_str, "%Y-%m-%d"
            ).date()
            end_date = datetime.strptime(
                data_sfarsit_custom_str, "%Y-%m-%d"
            ).date()
            if start_date > end_date:
                flash(
                    "Data de început custom nu poate fi după data de sfârșit.",
                    "warning",
                )
                start_date = end_date  # Sau resetează la un default valid
        except ValueError:
            flash(
                "Format dată custom invalid. Se afișează ultimele 7 zile.",
                "warning",
            )
            perioada = "ultimele_7_zile"  # Revert la default
            start_date = get_localized_now().date() - timedelta(days=6)
            end_date = get_localized_now().date()
    elif perioada == "toate":
        start_date = None  # Fără filtru de start
        end_date = None  # Fără filtru de end
    else:  # Default la ultimele 7 zile dacă perioada e invalidă
        start_date = get_localized_now().date() - timedelta(days=6)
        end_date = get_localized_now().date()

    leaves_history = []

    # Procesare Daily Leaves
    daily_leaves_query = DailyLeave.query.options(
        joinedload(DailyLeave.student)
    ).filter(DailyLeave.student_id.in_(student_ids))
    if start_date and end_date:  # Aplică filtru de dată doar dacă nu e "toate"
        daily_leaves_query = daily_leaves_query.filter(
            DailyLeave.leave_date >= start_date,
            DailyLeave.leave_date <= end_date,
        )

    for dl in daily_leaves_query.order_by(
        DailyLeave.leave_date.desc(), DailyLeave.start_time.desc()
    ).all():
        # These are times, not datetimes, but if comparison is made elsewhere, ensure awareness.
        leaves_history.append(
            {
                "student_name": f"{dl.student.grad_militar} {dl.student.nume} {dl.student.prenume}",
                "tip": "Zilnică",
                "data_start": dl.leave_date,
                "ora_start": dl.start_time,
                "ora_sfarsit": dl.end_time,
                "detalii": f"{dl.leave_type_display}",
                "motiv": dl.reason or "-",
                "status": dl.status,
            }
        )

    # Procesare Weekend Leaves
    weekend_leaves_query = WeekendLeave.query.options(
        joinedload(WeekendLeave.student)
    ).filter(WeekendLeave.student_id.in_(student_ids))
    # Pentru weekend leaves, filtrul de dată e mai complex.
    # Vom prelua toate și vom filtra în Python dacă e necesar, sau adaptăm query-ul.
    # Deocamdată, preluăm toate și filtrăm manual dacă nu e 'toate'.

    all_wl_gradat = weekend_leaves_query.order_by(
        WeekendLeave.weekend_start_date.desc()
    ).all()

    for wl in all_wl_gradat:
        # Verificăm dacă vreun interval din învoirea de weekend se încadrează în perioada filtrată
        relevant_for_period = False
        if perioada == "toate":
            relevant_for_period = True
        else:  # start_date și end_date sunt definite
            for (
                interval
            ) in (
                wl.get_intervals()
            ):  # get_intervals returnează datetimes aware
                # Convertim start_date/end_date (naive date) la datetime naive pentru comparație corectă cu datele intervalelor
                filter_start_dt_aware = EUROPE_BUCHAREST.localize(
                    datetime.combine(start_date, time.min)
                )
                filter_end_dt_aware = EUROPE_BUCHAREST.localize(
                    datetime.combine(end_date, time.max)
                )

                interval_start_aware = interval["start"].astimezone(
                    EUROPE_BUCHAREST
                )
                interval_end_aware = interval["end"].astimezone(
                    EUROPE_BUCHAREST
                )

                # Verificare intersecție intervale: (StartA <= EndB) and (EndA >= StartB)
                if (
                    interval_start_aware <= filter_end_dt_aware
                    and interval_end_aware >= filter_start_dt_aware
                ):
                    relevant_for_period = True
                    break

        if relevant_for_period:
            intervals_display = []
            for interval_data in wl.get_intervals():
                intervals_display.append(
                    f"{interval_data['day_name']} ({interval_data['start'].strftime('%d.%m')}) "
                    f"{interval_data['start'].strftime('%H:%M')}-{interval_data['end'].strftime('%H:%M')}"
                )

            leaves_history.append(
                {
                    "student_name": f"{wl.student.grad_militar} {wl.student.nume} {wl.student.prenume}",
                    "tip": "Weekend",
                    "data_start": wl.weekend_start_date,  # Data de vineri a weekendului
                    "ora_start": None,  # Nu e direct aplicabil, detaliile sunt în 'detalii'
                    "ora_sfarsit": None,
                    "detalii": "; ".join(intervals_display)
                    + (
                        f", Biserica Duminică"
                        if wl.duminica_biserica
                        and any(
                            d["day_name"] == "Duminica"
                            for d in wl.get_intervals()
                        )
                        else ""
                    ),
                    "motiv": wl.reason or "-",
                    "status": wl.status,
                }
            )

    # Sortare finală a listei combinate după data de început (descendent)
    leaves_history.sort(key=lambda x: x["data_start"], reverse=True)

    # Paginare manuală simplă (dacă se dorește) sau se poate folosi paginate() pe query-uri separate și apoi combina
    # Pentru simplitate, momentan fără paginare complexă pe lista combinată.

    return render_template(
        "invoiri_istoric.html",
        leaves_history=leaves_history,
        title="Istoric Învoiri Pluton",
        form_data=request.args,  # Pentru a repopula filtrele
        selected_period=perioada,
        selected_start_custom=data_start_custom_str,
        selected_end_custom=data_sfarsit_custom_str,
    )


# --- Admin Word Export Routes ---
@app.route(
    "/admin/permissions/export_word", endpoint="admin_export_permissions_word"
)
@login_required
def admin_export_permissions_word():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    now = get_localized_now()
    permissions_to_export = (
        Permission.query.options(joinedload(Permission.student))
        .join(Student, Permission.student_id == Student.id)
        .filter(
            Permission.status == "Aprobată",
            Permission.end_datetime >= now,  # Active or upcoming
        )
        .order_by(Student.nume, Student.prenume, Permission.start_datetime)
        .all()
    )

    if not permissions_to_export:
        flash(
            "Nicio permisie activă sau viitoare de exportat în sistem.", "info"
        )
        return redirect(url_for("admin_dashboard_route"))

    document = Document()
    document.add_heading(
        "Raport General Permisii (Admin)", level=1
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    user_info_text = f"Raport generat de: {current_user.username} (Admin)\nData generării: {datetime.now().strftime('%d-%m-%Y %H:%M')}"
    p_user = document.add_paragraph()
    p_user.add_run(user_info_text).italic = True
    p_user.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    table = document.add_table(
        rows=1, cols=7
    )  # Nr.crt, Grad, Nume și Prenume, Perioada, Grupa, Localitate, Transport
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    column_titles = [
        "Nr. crt.",
        "Grad",
        "Nume și Prenume",
        "Perioada",
        "Grupa",
        "Localitate",
        "Transport",
    ]
    for i, title in enumerate(column_titles):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, p in enumerate(permissions_to_export):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_cells[1].text = p.student.grad_militar
        row_cells[2].text = f"{p.student.nume} {p.student.prenume}"

        start_dt_local = (
            EUROPE_BUCHAREST.localize(p.start_datetime)
            if p.start_datetime.tzinfo is None
            else p.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        end_dt_local = (
            EUROPE_BUCHAREST.localize(p.end_datetime)
            if p.end_datetime.tzinfo is None
            else p.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        if start_dt_local.date() == end_dt_local.date():
            period_str = f"{start_dt_local.strftime('%d.%m.%Y %H:%M')} - {end_dt_local.strftime('%H:%M')}"
        else:
            period_str = f"{start_dt_local.strftime('%d.%m.%Y %H:%M')} - {end_dt_local.strftime('%d.%m.%Y %H:%M')}"
        row_cells[3].text = period_str
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_cells[4].text = p.student.pluton
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[5].text = p.destination if p.destination else "-"
        row_cells[6].text = p.transport_mode if p.transport_mode else "-"

    new_widths = {
        0: Inches(0.4),
        1: Inches(0.7),
        2: Inches(1.8),
        3: Inches(2.5),
        4: Inches(0.8),
        5: Inches(1.2),
        6: Inches(1.1),
    }
    for col_idx, width_val in new_widths.items():
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = width_val

    document.add_paragraph()  # Spacer

    # --- Separate table for church attendees (Admin view) ---
    church_attendees_admin = []
    for (
        leave
    ) in (
        leaves_to_export
    ):  # leaves_to_export is already filtered for active/upcoming
        if leave.duminica_biserica:
            is_sunday_selected_for_leave = False
            for interval in leave.get_intervals():
                if interval["day_name"] == "Duminica":
                    is_sunday_selected_for_leave = True
                    break
            if is_sunday_selected_for_leave:
                church_attendees_admin.append(leave.student)

    if church_attendees_admin:
        document.add_heading(
            "Studenți care participă la Biserică (Duminică 09:00-11:00)",
            level=2,
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        church_table_admin = document.add_table(
            rows=1, cols=4
        )  # Nr.crt, Grad, Nume și Prenume, Pluton
        church_table_admin.style = "Table Grid"
        church_table_admin.alignment = WD_TABLE_ALIGNMENT.CENTER

        church_hdr_cells_admin = church_table_admin.rows[0].cells
        church_col_titles_admin = [
            "Nr. crt.",
            "Grad",
            "Nume și Prenume",
            "Plutonul (Grupa)",
        ]
        for i, title in enumerate(church_col_titles_admin):
            church_hdr_cells_admin[i].text = title
            church_hdr_cells_admin[i].paragraphs[0].runs[0].font.bold = True
            church_hdr_cells_admin[i].paragraphs[
                0
            ].alignment = WD_ALIGN_PARAGRAPH.CENTER

        church_attendees_admin.sort(
            key=lambda s: (s.batalion, s.companie, s.pluton, s.nume, s.prenume)
        )  # Sort for admin view

        for idx, student in enumerate(church_attendees_admin):
            row_cells_admin = church_table_admin.add_row().cells
            row_cells_admin[0].text = str(idx + 1)
            row_cells_admin[0].paragraphs[
                0
            ].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells_admin[1].text = student.grad_militar
            row_cells_admin[2].text = f"{student.nume} {student.prenume}"
            row_cells_admin[3].text = (
                student.pluton
            )  # Admin might want to see Company/Battalion too, but Pluton is consistent
            row_cells_admin[3].paragraphs[
                0
            ].alignment = WD_ALIGN_PARAGRAPH.CENTER

        church_table_widths_admin = {
            0: Inches(0.5),
            1: Inches(0.8),
            2: Inches(2.5),
            3: Inches(1.0),
        }
        for col_idx, width_val in church_table_widths_admin.items():
            for row in church_table_admin.rows:
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = width_val
        document.add_paragraph()

    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    filename = (
        f"Raport_General_Permisii_Admin_{date.today().strftime('%Y%m%d')}.docx"
    )
    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route(
    "/admin/weekend_leaves/export_word",
    endpoint="admin_export_weekend_leaves_word",
)
@login_required
def admin_export_weekend_leaves_word():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    leaves_to_export = (
        WeekendLeave.query.options(joinedload(WeekendLeave.student))
        .join(Student, WeekendLeave.student_id == Student.id)
        .filter(WeekendLeave.status == "Aprobată")
        .order_by(
            Student.nume, Student.prenume, WeekendLeave.weekend_start_date
        )
        .all()
    )

    leaves_to_export = [
        leave
        for leave in leaves_to_export
        if leave.is_overall_active_or_upcoming
    ]

    if not leaves_to_export:
        flash(
            "Nicio învoire de weekend activă sau viitoare de exportat în sistem.",
            "info",
        )
        return redirect(url_for("admin_dashboard_route"))

    document = Document()
    document.add_heading(
        "Raport General Învoiri Weekend (Admin)", level=1
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    user_info_text = f"Raport generat de: {current_user.username} (Admin)\nData generării: {get_localized_now().strftime('%d-%m-%Y %H:%M')}"
    p_user = document.add_paragraph()
    p_user.add_run(user_info_text).italic = True
    p_user.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    # Agregare pe student
    from collections import defaultdict

    student_map = {}
    periods_map = defaultdict(list)
    for leave in leaves_to_export:
        st = leave.student
        student_map[st.id] = st
        for iv in leave.get_intervals():
            if iv["end"] >= get_localized_now():
                s_loc = iv["start"].astimezone(EUROPE_BUCHAREST)
                e_loc = iv["end"].astimezone(EUROPE_BUCHAREST)
                dname = iv.get("day_name") or s_loc.strftime("%A")
                periods_map[st.id].append(
                    (
                        s_loc,
                        f"{dname} ({s_loc.strftime('%d-%m')}): {s_loc.strftime('%H:%M')} - {e_loc.strftime('%H:%M')}",
                    )
                )

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, t in enumerate(
        ["Nr. crt.", "Grad", "Nume și Prenume", "Plutonul (Grupa)", "Perioade"]
    ):
        hdr[i].text = t
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    ordered = sorted(
        periods_map.keys(),
        key=lambda sid: (student_map[sid].nume, student_map[sid].prenume),
    )
    for idx, sid in enumerate(ordered, start=1):
        st = student_map[sid]
        periods_sorted = [
            p for _, p in sorted(periods_map[sid], key=lambda t: t[0])
        ]
        row = table.add_row().cells
        row[0].text = str(idx)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = st.grad_militar
        row[2].text = f"{st.nume} {st.prenume}"
        row[3].text = st.pluton
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[4].text = "; ".join(periods_sorted)

    widths = {
        0: Inches(0.5),
        1: Inches(0.8),
        2: Inches(2.2),
        3: Inches(0.9),
        4: Inches(3.0),
    }
    for ci, w in widths.items():
        for r in table.rows:
            if ci < len(r.cells):
                r.cells[ci].width = w

    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    filename = (
        f"Raport_General_Weekend_Admin_{date.today().strftime('%Y%m%d')}.docx"
    )
    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# --- Company Commander Word Export Routes ---
@app.route(
    "/company_commander/permissions/export_word",
    endpoint="company_commander_export_permissions_word",
)
@login_required
def company_commander_export_permissions_word():
    if current_user.role != "comandant_companie":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    company_id_str = _get_commander_unit_id(current_user.username, "CmdC")
    if not company_id_str:
        flash("ID-ul companiei nu a putut fi determinat.", "warning")
        return redirect(url_for("company_commander_dashboard"))

    student_ids_in_company = [
        sid
        for (sid,) in Student.query.filter_by(companie=company_id_str)
        .with_entities(Student.id)
        .all()
    ]
    if not student_ids_in_company:
        flash(
            f"Niciun student în compania {company_id_str} pentru a exporta permisii.",
            "info",
        )
        return redirect(url_for("company_commander_dashboard"))

    now = get_localized_now()
    permissions_to_export = (
        Permission.query.options(joinedload(Permission.student))
        .join(Student, Permission.student_id == Student.id)
        .filter(
            Permission.student_id.in_(student_ids_in_company),
            Permission.status == "Aprobată",
            Permission.end_datetime >= now,
        )
        .order_by(Student.nume, Student.prenume, Permission.start_datetime)
        .all()
    )

    if not permissions_to_export:
        flash(
            f"Nicio permisie activă sau viitoare de exportat pentru compania {company_id_str}.",
            "info",
        )
        return redirect(url_for("company_commander_dashboard"))

    document = Document()
    document.add_heading(
        f"Raport Permisii Compania {company_id_str}", level=1
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    user_info_text = f"Raport generat de: {current_user.username}\nData generării: {datetime.now().strftime('%d-%m-%Y %H:%M')}"
    p_user = document.add_paragraph()
    p_user.add_run(user_info_text).italic = True
    p_user.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    table = document.add_table(
        rows=1, cols=7
    )  # Nr.crt, Grad, Nume și Prenume, Perioada, Grupa, Localitate, Transport
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    column_titles = [
        "Nr. crt.",
        "Grad",
        "Nume și Prenume",
        "Perioada",
        "Grupa",
        "Localitate",
        "Transport",
    ]
    for i, title in enumerate(column_titles):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, p in enumerate(permissions_to_export):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_cells[1].text = p.student.grad_militar
        row_cells[2].text = f"{p.student.nume} {p.student.prenume}"

        start_dt_local = (
            EUROPE_BUCHAREST.localize(p.start_datetime)
            if p.start_datetime.tzinfo is None
            else p.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        end_dt_local = (
            EUROPE_BUCHAREST.localize(p.end_datetime)
            if p.end_datetime.tzinfo is None
            else p.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        if start_dt_local.date() == end_dt_local.date():
            period_str = f"{start_dt_local.strftime('%d.%m.%Y %H:%M')} - {end_dt_local.strftime('%H:%M')}"
        else:
            period_str = f"{start_dt_local.strftime('%d.%m.%Y %H:%M')} - {end_dt_local.strftime('%d.%m.%Y %H:%M')}"
        row_cells[3].text = period_str
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_cells[4].text = p.student.pluton
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[5].text = p.destination if p.destination else "-"
        row_cells[6].text = p.transport_mode if p.transport_mode else "-"

    new_widths = {
        0: Inches(0.4),
        1: Inches(0.7),
        2: Inches(1.8),
        3: Inches(2.5),
        4: Inches(0.8),
        5: Inches(1.2),
        6: Inches(1.1),
    }
    for col_idx, width_val in new_widths.items():
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = width_val

    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    filename = f"Raport_Permisii_Compania_{company_id_str}_{date.today().strftime('%Y%m%d')}.docx"
    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route(
    "/company_commander/weekend_leaves/export_word",
    endpoint="company_commander_export_weekend_leaves_word",
)
@login_required
def company_commander_export_weekend_leaves_word():
    if current_user.role != "comandant_companie":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    company_id_str = _get_commander_unit_id(current_user.username, "CmdC")
    if not company_id_str:
        flash("ID-ul companiei nu a putut fi determinat.", "warning")
        return redirect(url_for("company_commander_dashboard"))

    student_ids_in_company = [
        sid
        for (sid,) in Student.query.filter_by(companie=company_id_str)
        .with_entities(Student.id)
        .all()
    ]
    if not student_ids_in_company:
        flash(
            f"Niciun student în compania {company_id_str} pentru a exporta învoiri de weekend.",
            "info",
        )
        return redirect(url_for("company_commander_dashboard"))

    leaves_to_export = (
        WeekendLeave.query.options(joinedload(WeekendLeave.student))
        .join(Student, WeekendLeave.student_id == Student.id)
        .filter(
            WeekendLeave.student_id.in_(student_ids_in_company),
            WeekendLeave.status == "Aprobată",
        )
        .order_by(
            Student.nume, Student.prenume, WeekendLeave.weekend_start_date
        )
        .all()
    )

    leaves_to_export = [
        leave
        for leave in leaves_to_export
        if leave.is_overall_active_or_upcoming
    ]

    if not leaves_to_export:
        flash(
            f"Nicio învoire de weekend activă sau viitoare de exportat pentru compania {company_id_str}.",
            "info",
        )
        return redirect(url_for("company_commander_dashboard"))

    document = Document()
    document.add_heading(
        f"Raport Învoiri Weekend Compania {company_id_str}", level=1
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    user_info_text = f"Raport generat de: {current_user.username}\nData generării: {get_localized_now().strftime('%d-%m-%Y %H:%M')}"
    p_user = document.add_paragraph()
    p_user.add_run(user_info_text).italic = True
    p_user.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    # Agregare perioade pe student
    from collections import defaultdict

    student_map = {}
    periods_map = defaultdict(list)
    for leave in leaves_to_export:
        st = leave.student
        student_map[st.id] = st
        for iv in leave.get_intervals():
            if iv["end"] >= get_localized_now():
                s_loc = iv["start"].astimezone(EUROPE_BUCHAREST)
                e_loc = iv["end"].astimezone(EUROPE_BUCHAREST)
                dname = iv.get("day_name") or s_loc.strftime("%A")
                periods_map[st.id].append(
                    (
                        s_loc,
                        f"{dname} ({s_loc.strftime('%d-%m')}): {s_loc.strftime('%H:%M')} - {e_loc.strftime('%H:%M')}",
                    )
                )

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, t in enumerate(
        ["Nr. crt.", "Grad", "Nume și Prenume", "Plutonul (Grupa)", "Perioade"]
    ):
        hdr[i].text = t
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    ordered = sorted(
        periods_map.keys(),
        key=lambda sid: (student_map[sid].nume, student_map[sid].prenume),
    )
    for idx, sid in enumerate(ordered, start=1):
        st = student_map[sid]
        per_sorted = [
            p for _, p in sorted(periods_map[sid], key=lambda t: t[0])
        ]
        r = table.add_row().cells
        r[0].text = str(idx)
        r[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r[1].text = st.grad_militar
        r[2].text = f"{st.nume} {st.prenume}"
        r[3].text = st.pluton
        r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r[4].text = "; ".join(per_sorted)

    widths = {
        0: Inches(0.5),
        1: Inches(0.8),
        2: Inches(2.2),
        3: Inches(0.9),
        4: Inches(3.0),
    }
    for ci, w in widths.items():
        for row in table.rows:
            if ci < len(row.cells):
                row.cells[ci].width = w

    document.add_paragraph()  # Spacer

    # --- Biserică (duminică) ---
    church_attendees_company = []
    for leave in leaves_to_export:
        if leave.duminica_biserica and any(
            iv["day_name"] == "Duminica" for iv in leave.get_intervals()
        ):
            church_attendees_company.append(leave.student)

    if church_attendees_company:
        document.add_heading(
            "Studenți care participă la Biserică (Duminică 09:00-11:00)",
            level=2,
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        church_table_company = document.add_table(rows=1, cols=4)
        church_table_company.style = "Table Grid"
        church_table_company.alignment = WD_TABLE_ALIGNMENT.CENTER

        ch_hdr = church_table_company.rows[0].cells
        for i, t in enumerate(
            ["Nr. crt.", "Grad", "Nume și Prenume", "Plutonul (Grupa)"]
        ):
            ch_hdr[i].text = t
            ch_hdr[i].paragraphs[0].runs[0].font.bold = True
            ch_hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        church_attendees_company = sorted(
            {s.id: s for s in church_attendees_company}.values(),
            key=lambda s: (s.pluton, s.nume, s.prenume),
        )
        for idx, s in enumerate(church_attendees_company, start=1):
            row = church_table_company.add_row().cells
            row[0].text = str(idx)
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[1].text = s.grad_militar
            row[2].text = f"{s.nume} {s.prenume}"
            row[3].text = s.pluton
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        ch_w = {0: Inches(0.5), 1: Inches(0.8), 2: Inches(2.5), 3: Inches(1.0)}
        for ci, w in ch_w.items():
            for row in church_table_company.rows:
                if ci < len(row.cells):
                    row.cells[ci].width = w

    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    filename = f"Raport_Weekend_Compania_{company_id_str}_{date.today().strftime('%Y%m%d')}.docx"
    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# --- Battalion Commander Word Export Routes ---
@app.route(
    "/battalion_commander/permissions/export_word",
    endpoint="battalion_commander_export_permissions_word",
)
@login_required
def battalion_commander_export_permissions_word():
    if current_user.role != "comandant_batalion":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    battalion_id_str = _get_commander_unit_id(current_user.username, "CmdB")
    if not battalion_id_str:
        flash("ID-ul batalionului nu a putut fi determinat.", "warning")
        return redirect(url_for("battalion_commander_dashboard"))

    student_ids_in_battalion = [
        sid
        for (sid,) in Student.query.filter_by(batalion=battalion_id_str)
        .with_entities(Student.id)
        .all()
    ]
    if not student_ids_in_battalion:
        flash(
            f"Niciun student în batalionul {battalion_id_str} pentru a exporta permisii.",
            "info",
        )
        return redirect(url_for("battalion_commander_dashboard"))

    now = get_localized_now()
    permissions_to_export = (
        Permission.query.options(joinedload(Permission.student))
        .join(Student, Permission.student_id == Student.id)
        .filter(
            Permission.student_id.in_(student_ids_in_battalion),
            Permission.status == "Aprobată",
            Permission.end_datetime >= now,
        )
        .order_by(Student.nume, Student.prenume, Permission.start_datetime)
        .all()
    )

    if not permissions_to_export:
        flash(
            f"Nicio permisie activă sau viitoare de exportat pentru batalionul {battalion_id_str}.",
            "info",
        )
        return redirect(url_for("battalion_commander_dashboard"))

    document = Document()
    document.add_heading(
        f"Raport Permisii Batalionul {battalion_id_str}", level=1
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    user_info_text = f"Raport generat de: {current_user.username}\nData generării: {datetime.now().strftime('%d-%m-%Y %H:%M')}"
    p_user = document.add_paragraph()
    p_user.add_run(user_info_text).italic = True
    p_user.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    table = document.add_table(
        rows=1, cols=7
    )  # Nr.crt, Grad, Nume și Prenume, Perioada, Grupa, Localitate, Transport
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    column_titles = [
        "Nr. crt.",
        "Grad",
        "Nume și Prenume",
        "Perioada",
        "Grupa",
        "Localitate",
        "Transport",
    ]
    for i, title in enumerate(column_titles):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, p in enumerate(permissions_to_export):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_cells[1].text = p.student.grad_militar
        row_cells[2].text = f"{p.student.nume} {p.student.prenume}"

        start_dt_local = (
            EUROPE_BUCHAREST.localize(p.start_datetime)
            if p.start_datetime.tzinfo is None
            else p.start_datetime.astimezone(EUROPE_BUCHAREST)
        )
        end_dt_local = (
            EUROPE_BUCHAREST.localize(p.end_datetime)
            if p.end_datetime.tzinfo is None
            else p.end_datetime.astimezone(EUROPE_BUCHAREST)
        )
        if start_dt_local.date() == end_dt_local.date():
            period_str = f"{start_dt_local.strftime('%d.%m.%Y %H:%M')} - {end_dt_local.strftime('%H:%M')}"
        else:
            period_str = f"{start_dt_local.strftime('%d.%m.%Y %H:%M')} - {end_dt_local.strftime('%d.%m.%Y %H:%M')}"
        row_cells[3].text = period_str
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_cells[4].text = p.student.pluton
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[5].text = p.destination if p.destination else "-"
        row_cells[6].text = p.transport_mode if p.transport_mode else "-"

    new_widths = {
        0: Inches(0.4),
        1: Inches(0.7),
        2: Inches(1.8),
        3: Inches(2.5),
        4: Inches(0.8),
        5: Inches(1.2),
        6: Inches(1.1),
    }
    for col_idx, width_val in new_widths.items():
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = width_val

    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    filename = f"Raport_Permisii_Batalion_{battalion_id_str}_{date.today().strftime('%Y%m%d')}.docx"
    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route(
    "/battalion_commander/weekend_leaves/export_word",
    endpoint="battalion_commander_export_weekend_leaves_word",
)
@login_required
def battalion_commander_export_weekend_leaves_word():
    if current_user.role != "comandant_batalion":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    battalion_id_str = _get_commander_unit_id(current_user.username, "CmdB")
    if not battalion_id_str:
        flash("ID-ul batalionului nu a putut fi determinat.", "warning")
        return redirect(url_for("battalion_commander_dashboard"))

    student_ids_in_battalion = [
        sid
        for (sid,) in Student.query.filter_by(batalion=battalion_id_str)
        .with_entities(Student.id)
        .all()
    ]
    if not student_ids_in_battalion:
        flash(
            f"Niciun student în batalionul {battalion_id_str} pentru a exporta învoiri de weekend.",
            "info",
        )
        return redirect(url_for("battalion_commander_dashboard"))

    leaves_to_export = (
        WeekendLeave.query.options(joinedload(WeekendLeave.student))
        .join(Student, WeekendLeave.student_id == Student.id)
        .filter(
            WeekendLeave.student_id.in_(student_ids_in_battalion),
            WeekendLeave.status == "Aprobată",
        )
        .order_by(
            Student.nume, Student.prenume, WeekendLeave.weekend_start_date
        )
        .all()
    )

    leaves_to_export = [
        leave
        for leave in leaves_to_export
        if leave.is_overall_active_or_upcoming
    ]

    if not leaves_to_export:
        flash(
            f"Nicio învoire de weekend activă sau viitoare de exportat pentru batalionul {battalion_id_str}.",
            "info",
        )
        return redirect(url_for("battalion_commander_dashboard"))

    document = Document()
    document.add_heading(
        f"Raport Învoiri Weekend Batalionul {battalion_id_str}", level=1
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    user_info_text = f"Raport generat de: {current_user.username}\nData generării: {get_localized_now().strftime('%d-%m-%Y %H:%M')}"
    p_user = document.add_paragraph()
    p_user.add_run(user_info_text).italic = True
    p_user.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    # Agregare pe student
    from collections import defaultdict

    student_map = {}
    periods_map = defaultdict(list)
    for leave in leaves_to_export:
        st = leave.student
        student_map[st.id] = st
        for iv in leave.get_intervals():
            if iv["end"] >= get_localized_now():
                s_loc = iv["start"].astimezone(EUROPE_BUCHAREST)
                e_loc = iv["end"].astimezone(EUROPE_BUCHAREST)
                dname = iv.get("day_name") or s_loc.strftime("%A")
                periods_map[st.id].append(
                    (
                        s_loc,
                        f"{dname} ({s_loc.strftime('%d-%m')}): {s_loc.strftime('%H:%M')} - {e_loc.strftime('%H:%M')}",
                    )
                )

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, t in enumerate(
        ["Nr. crt.", "Grad", "Nume și Prenume", "Plutonul (Grupa)", "Perioade"]
    ):
        hdr[i].text = t
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    ordered = sorted(
        periods_map.keys(),
        key=lambda sid: (student_map[sid].nume, student_map[sid].prenume),
    )
    for idx, sid in enumerate(ordered, start=1):
        st = student_map[sid]
        per_sorted = [
            p for _, p in sorted(periods_map[sid], key=lambda t: t[0])
        ]
        r = table.add_row().cells
        r[0].text = str(idx)
        r[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r[1].text = st.grad_militar
        r[2].text = f"{st.nume} {st.prenume}"
        r[3].text = st.pluton
        r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r[4].text = "; ".join(per_sorted)

    widths = {
        0: Inches(0.5),
        1: Inches(0.8),
        2: Inches(2.2),
        3: Inches(0.9),
        4: Inches(3.0),
    }
    for ci, w in widths.items():
        for row in table.rows:
            if ci < len(row.cells):
                row.cells[ci].width = w

    document.add_paragraph()  # Spacer

    # --- Biserică (duminică) ---
    church_attendees_battalion = []
    for leave in leaves_to_export:
        if leave.duminica_biserica and any(
            iv["day_name"] == "Duminica" for iv in leave.get_intervals()
        ):
            church_attendees_battalion.append(leave.student)

    if church_attendees_battalion:
        document.add_heading(
            "Studenți care participă la Biserică (Duminică 09:00-11:00)",
            level=2,
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        church_table_battalion = document.add_table(rows=1, cols=5)
        church_table_battalion.style = "Table Grid"
        church_table_battalion.alignment = WD_TABLE_ALIGNMENT.CENTER

        ch_hdr = church_table_battalion.rows[0].cells
        for i, t in enumerate(
            ["Nr. crt.", "Grad", "Nume și Prenume", "Compania", "Plutonul"]
        ):
            ch_hdr[i].text = t
            ch_hdr[i].paragraphs[0].runs[0].font.bold = True
            ch_hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        church_attendees_battalion = sorted(
            {s.id: s for s in church_attendees_battalion}.values(),
            key=lambda s: (s.companie, s.pluton, s.nume, s.prenume),
        )
        for idx, s in enumerate(church_attendees_battalion, start=1):
            row = church_table_battalion.add_row().cells
            row[0].text = str(idx)
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[1].text = s.grad_militar
            row[2].text = f"{s.nume} {s.prenume}"
            row[3].text = s.companie
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[4].text = s.pluton
            row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        ch_w = {
            0: Inches(0.4),
            1: Inches(0.7),
            2: Inches(2.2),
            3: Inches(0.8),
            4: Inches(0.8),
        }
        for ci, w in ch_w.items():
            for row in church_table_battalion.rows:
                if ci < len(row.cells):
                    row.cells[ci].width = w

    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    filename = f"Raport_Weekend_Batalion_{battalion_id_str}_{date.today().strftime('%Y%m%d')}.docx"
    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# --- Gradat Daily Leaves Word Export ---
@app.route(
    "/gradat/daily_leaves/export_word",
    endpoint="gradat_export_daily_leaves_word",
)
@login_required
def gradat_export_daily_leaves_word():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_id_tuples = (
        db.session.query(Student.id)
        .filter_by(created_by_user_id=current_user.id)
        .all()
    )
    student_ids = [s[0] for s in student_id_tuples]

    if not student_ids:
        flash("Nu aveți studenți pentru a exporta învoiri zilnice.", "info")
        return redirect(url_for("list_daily_leaves"))

    # Fetch all approved daily leaves for the gradat's students
    # For export, typically all relevant (approved) leaves are included, or based on a filter.
    # Here, we'll fetch all approved ones and sort them.
    leaves_to_export = (
        DailyLeave.query.options(joinedload(DailyLeave.student))
        .filter(
            DailyLeave.student_id.in_(student_ids),
            DailyLeave.status == "Aprobată",
        )
        .join(Student)
        .order_by(
            DailyLeave.leave_date.asc(),  # Sort by date first
            Student.nume.asc(),  # Then by student name
            Student.prenume.asc(),
            DailyLeave.start_time.asc(),
        )
        .all()
    )

    if not leaves_to_export:
        flash("Nicio învoire zilnică aprobată de exportat.", "info")
        return redirect(url_for("list_daily_leaves"))

    document = Document()
    document.add_heading("Raport Învoiri Zilnice", level=1).alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    user_info_text = f"Raport generat de: {current_user.username}\nData generării: {get_localized_now().strftime('%d-%m-%Y %H:%M')}"
    p_user = document.add_paragraph()
    p_user.add_run(user_info_text).italic = True
    p_user.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()  # Spacer

    table = document.add_table(
        rows=1, cols=5
    )  # Nr. crt, Grad, Nume și Prenume, Pluton(Grupa), Data
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    column_titles = [
        "Nr. crt.",
        "Grad",
        "Nume și Prenume",
        "Plutonul (Grupa)",
        "Data",
    ]
    for i, title in enumerate(column_titles):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, leave in enumerate(leaves_to_export):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = leave.student.grad_militar
        row_cells[2].text = (
            f"{leave.student.nume} {leave.student.prenume}"  # Combined Name
        )
        row_cells[3].text = leave.student.pluton
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].text = leave.leave_date.strftime("%d.%m.%Y")
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths
    # NrCrt(0.5), Grad(0.8), Nume și Prenume(2.5), Pluton(1.0), Data(1.0)
    widths = {
        0: Inches(0.5),
        1: Inches(0.8),
        2: Inches(2.5),  # Nume și Prenume
        3: Inches(1.0),
        4: Inches(1.0),
    }
    for col_idx, width_val in widths.items():
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = width_val

    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    f = io.BytesIO()
    document.save(f)
    f.seek(0)

    filename = f"Raport_Invoiri_Zilnice_{current_user.username}_{get_localized_now().date().strftime('%Y%m%d')}.docx"

    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# START JULES BLOCK - MORNING INVIGORATION REPORT (PROBLEM 5)
@app.route("/reports/morning_invigoration", methods=["GET", "POST"])
@login_required
def morning_invigoration_report():
    if current_user.role not in [
        "gradat",
        "admin",
        "comandant_companie",
        "comandant_batalion",
    ]:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    form_data = {
        "report_date_str": get_localized_now().date().isoformat(),
        "platoon_3_participates": False,  # Default for the toggle
    }
    report_data_to_render = None

    if request.method == "POST":
        form_data["report_date_str"] = request.form.get(
            "report_date", get_localized_now().date().isoformat()
        )
        # Checkbox value is 'on' if checked, otherwise not present in form
        form_data["platoon_3_participates"] = (
            request.form.get("platoon_3_participates") == "on"
        )

        try:
            report_date = datetime.strptime(
                form_data["report_date_str"], "%Y-%m-%d"
            ).date()
        except ValueError:
            flash("Format dată invalid. Se folosește data curentă.", "warning")
            report_date = get_localized_now().date()
            form_data["report_date_str"] = report_date.isoformat()

        invigoration_time = time(6, 5)  # Standard invigoration time
        # Create aware datetime for the check
        datetime_to_check = EUROPE_BUCHAREST.localize(
            datetime.combine(report_date, invigoration_time)
        )
        report_title_detail = f"Înviorare Dimineață ({datetime_to_check.strftime('%d.%m.%Y %H:%M')})"

        students_for_report_query = Student.query
        report_base_title = "Raport Prezență Înviorare"

        # Scope students based on user role
        if current_user.role == "gradat":
            students_for_report_query = students_for_report_query.filter(
                Student.created_by_user_id == current_user.id
            )
            # If gradat is for Platoon 3, the toggle on their view might be confusing.
            # However, the current logic means they see their own students; if they are Platoon 3,
            # and the toggle (if shown to them, though not planned) is off, their students would be "non-participating".
            # For simplicity, a gradat's report is just for their students.
            # The platoon_3_participates toggle is primarily for commanders/admin.
        elif current_user.role == "comandant_companie":
            company_id = _get_commander_unit_id(current_user.username, "CmdC")
            if company_id:
                students_for_report_query = students_for_report_query.filter(
                    Student.companie == company_id
                )
                report_base_title = f"Raport Înviorare Compania {company_id}"
            else:
                flash("ID Companie invalid.", "danger")
                return redirect(url_for("dashboard"))
        elif current_user.role == "comandant_batalion":
            battalion_id = _get_commander_unit_id(
                current_user.username, "CmdB"
            )
            if battalion_id:
                students_for_report_query = students_for_report_query.filter(
                    Student.batalion == battalion_id
                )
                report_base_title = (
                    f"Raport Înviorare Batalionul {battalion_id}"
                )
            else:
                flash("ID Batalion invalid.", "danger")
                return redirect(url_for("dashboard"))
        elif current_user.role == "admin":
            report_base_title = "Raport Înviorare General (Admin)"
            # No specific unit filter for admin, they see all by default for this report type.

        all_students_in_scope = students_for_report_query.all()

        students_to_calculate_presence_for = []
        platoon_3_non_participants_details = (
            []
        )  # Students from Platoon 3 not participating

        if not form_data["platoon_3_participates"]:
            for s in all_students_in_scope:
                # Assuming platoon is stored as a string, e.g., '3', '11', '21'
                if s.pluton == "3":
                    platoon_3_non_participants_details.append(
                        f"{s.grad_militar} {s.nume} {s.prenume} - Pluton 3 (Neparticipant)"
                    )
                else:
                    students_to_calculate_presence_for.append(s)
        else:  # Platoon 3 participates
            students_to_calculate_presence_for = all_students_in_scope

        if not all_students_in_scope:
            flash("Niciun student în evidență pentru acest raport.", "info")
            # Still render the template but report_data_calculated will be based on empty list

        report_data_calculated = _calculate_presence_data(
            students_to_calculate_presence_for, datetime_to_check
        )

        # Adjust final counts for the report display
        final_efectiv_control = len(all_students_in_scope)
        # Absentees = those calculated as absent from participating students + non-participating Platoon 3 + SMT from Platoon 3 (if any)
        # _calculate_presence_data already includes SMT in its absent_total for the list it processed.
        # So, if Platoon 3 students were in students_to_calculate_presence_for, their SMT status is handled.
        # If Platoon 3 students were moved to platoon_3_non_participants_details, their SMT status is not in report_data_calculated.smt_count.

        # Let's ensure SMT students from a non-participating Platoon 3 are also listed in the SMT section of the main report
        # if they are SMT, rather than just "Pluton 3 (Neparticipant)".
        # The current _calculate_presence_data handles SMT first.
        # So, if Platoon 3 *is* participating, their SMTs are caught.
        # If Platoon 3 is *not* participating, they are entirely in platoon_3_non_participants_details.
        # This seems acceptable: if Platoon 3 is marked as non-participating, all its members are listed as such for this specific report.

        final_efectiv_absent = report_data_calculated[
            "efectiv_absent_total"
        ] + len(platoon_3_non_participants_details)
        final_efectiv_prezent = report_data_calculated[
            "efectiv_prezent_total"
        ]  # Based on those who were supposed to participate

        report_data_to_render = {
            **report_data_calculated,
            "efectiv_control": final_efectiv_control,
            "efectiv_prezent_total": final_efectiv_prezent,
            "efectiv_absent_total": final_efectiv_absent,
            "platoon_3_non_participants_details": sorted(
                platoon_3_non_participants_details
            ),
            "platoon_3_non_participants_count": len(
                platoon_3_non_participants_details
            ),
            "platoon_3_participated_fully": form_data[
                "platoon_3_participates"
            ],
            "title": f"{report_base_title} - {report_title_detail}",
            "datetime_checked": datetime_to_check.strftime(
                "%d %B %Y, %H:%M:%S"
            ),
        }

    return render_template(
        "morning_invigoration_report.html",
        form_data=form_data,  # Pass current form data for repopulation
        report_data=report_data_to_render,
        current_user_role=current_user.role,
    )  # For conditional display of toggle in template


# END JULES BLOCK - MORNING INVIGORATION REPORT (PROBLEM 5)


# START JULES BLOCK - NEW AND MODIFIED FUNCTIONS (PROBLEMS 2 & 4)


# --- Commander Scoped List Views ---
@app.route("/commander/permissions", endpoint="view_scoped_permissions")
@login_required
def view_scoped_permissions():
    if current_user.role not in ["comandant_companie", "comandant_batalion"]:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_ids_in_scope = []
    unit_id_str = ""
    report_title_detail = ""
    return_dashboard_url = "dashboard"

    if current_user.role == "comandant_companie":
        unit_id_str = _get_commander_unit_id(current_user.username, "CmdC")
        if unit_id_str:
            students_in_unit = (
                Student.query.filter_by(companie=unit_id_str)
                .with_entities(Student.id)
                .all()
            )
            student_ids_in_scope = [sid for (sid,) in students_in_unit]
            report_title_detail = f"Permisii Compania {unit_id_str}"
            return_dashboard_url = "company_commander_dashboard"
        else:
            flash("ID Companie invalid.", "danger")
            return redirect(url_for("dashboard"))
    elif current_user.role == "comandant_batalion":
        unit_id_str = _get_commander_unit_id(current_user.username, "CmdB")
        if unit_id_str:
            students_in_unit = (
                Student.query.filter_by(batalion=unit_id_str)
                .with_entities(Student.id)
                .all()
            )
            student_ids_in_scope = [sid for (sid,) in students_in_unit]
            report_title_detail = f"Permisii Batalionul {unit_id_str}"
            return_dashboard_url = "battalion_commander_dashboard"
        else:
            flash("ID Batalion invalid.", "danger")
            return redirect(url_for("dashboard"))

    from sqlalchemy import false

    page = request.args.get("page", 1, type=int)
    per_page = 20
    query_filter = (
        Permission.student_id.in_(student_ids_in_scope)
        if student_ids_in_scope
        else false()
    )
    query = Permission.query.options(
        joinedload(Permission.student), joinedload(Permission.creator)
    ).filter(query_filter)

    search_student_name = request.args.get("search_student_name", "").strip()
    filter_status = request.args.get("filter_status", "").strip()
    filter_date_type = request.args.get(
        "filter_date_type", "active_today"
    ).strip()

    if search_student_name:
        search_pattern = f"%{unidecode(search_student_name.lower())}%"
        query = query.join(Permission.student).filter(
            or_(
                func.lower(unidecode(Student.nume)).like(search_pattern),
                func.lower(unidecode(Student.prenume)).like(search_pattern),
            )
        )
    if filter_status:
        query = query.filter(Permission.status == filter_status)

    now_localized = get_localized_now()
    today_start_naive = datetime.combine(now_localized.date(), time.min)
    today_end_naive = datetime.combine(now_localized.date(), time.max)
    now_naive_for_db_compare = now_localized.replace(tzinfo=None)

    if filter_date_type == "active_today":
        query = query.filter(
            Permission.start_datetime <= today_end_naive,
            Permission.end_datetime >= today_start_naive,
            Permission.status == "Aprobată",
        )
    elif filter_date_type == "active_now":
        query = query.filter(
            Permission.start_datetime <= now_naive_for_db_compare,
            Permission.end_datetime >= now_naive_for_db_compare,
            Permission.status == "Aprobată",
        )
    elif filter_date_type == "upcoming":
        query = query.filter(
            Permission.start_datetime > now_naive_for_db_compare,
            Permission.status == "Aprobată",
        )
    elif filter_date_type == "past_week":
        seven_days_ago_naive = now_localized.date() - timedelta(days=7)
        query = query.filter(
            Permission.end_datetime
            >= datetime.combine(seven_days_ago_naive, time.min)
        )

    query = query.order_by(Permission.start_datetime.desc())
    permissions_pagination = query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    return render_template(
        "scoped_list_permissions.html",
        permissions_pagination=permissions_pagination,
        title=report_title_detail,
        return_dashboard_url=return_dashboard_url,
    )


@app.route("/commander/daily_leaves", endpoint="view_scoped_daily_leaves")
@login_required
def view_scoped_daily_leaves():
    if current_user.role not in ["comandant_companie", "comandant_batalion"]:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_ids_in_scope = []
    unit_id_str = ""
    view_title = ""
    return_dashboard_url = "dashboard"

    if current_user.role == "comandant_companie":
        unit_id_str = _get_commander_unit_id(current_user.username, "CmdC")
        if unit_id_str:
            students_in_unit = (
                Student.query.filter_by(companie=unit_id_str)
                .with_entities(Student.id)
                .all()
            )
            student_ids_in_scope = [sid for (sid,) in students_in_unit]
            view_title = f"Învoiri Zilnice Compania {unit_id_str}"
            return_dashboard_url = "company_commander_dashboard"
        else:
            flash("ID Companie invalid.", "danger")
            return redirect(url_for("dashboard"))
    elif current_user.role == "comandant_batalion":
        unit_id_str = _get_commander_unit_id(current_user.username, "CmdB")
        if unit_id_str:
            students_in_unit = (
                Student.query.filter_by(batalion=unit_id_str)
                .with_entities(Student.id)
                .all()
            )
            student_ids_in_scope = [sid for (sid,) in students_in_unit]
            view_title = f"Învoiri Zilnice Batalionul {unit_id_str}"
            return_dashboard_url = "battalion_commander_dashboard"
        else:
            flash("ID Batalion invalid.", "danger")
            return redirect(url_for("dashboard"))

    from sqlalchemy import false

    page = request.args.get("page", 1, type=int)
    per_page = 20
    query_filter = (
        DailyLeave.student_id.in_(student_ids_in_scope)
        if student_ids_in_scope
        else false()
    )
    query = DailyLeave.query.options(
        joinedload(DailyLeave.student), joinedload(DailyLeave.creator)
    ).filter(query_filter)

    search_student_name = request.args.get("search_student_name", "").strip()
    filter_status = request.args.get("filter_status", "").strip()
    filter_date_str = request.args.get("filter_date", "").strip()

    if search_student_name:
        search_pattern = f"%{unidecode(search_student_name.lower())}%"
        query = query.join(DailyLeave.student).filter(
            or_(
                func.lower(unidecode(Student.nume)).like(search_pattern),
                func.lower(unidecode(Student.prenume)).like(search_pattern),
            )
        )
    if filter_status:
        query = query.filter(DailyLeave.status == filter_status)
    if filter_date_str:
        try:
            filter_date_obj = datetime.strptime(
                filter_date_str, "%Y-%m-%d"
            ).date()
            query = query.filter(DailyLeave.leave_date == filter_date_obj)
        except ValueError:
            flash(
                "Format dată invalid pentru filtrare. Folosiți YYYY-MM-DD.",
                "warning",
            )

    query = query.order_by(
        DailyLeave.leave_date.desc(), DailyLeave.start_time.desc()
    )
    daily_leaves_pagination = query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    return render_template(
        "scoped_list_daily_leaves.html",
        daily_leaves_pagination=daily_leaves_pagination,
        title=view_title,
        return_dashboard_url=return_dashboard_url,
    )


@app.route("/commander/weekend_leaves", endpoint="view_scoped_weekend_leaves")
@login_required
def view_scoped_weekend_leaves():
    if current_user.role not in ["comandant_companie", "comandant_batalion"]:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_ids_in_scope = []
    unit_id_str = ""
    view_title = ""
    return_dashboard_url = "dashboard"

    if current_user.role == "comandant_companie":
        unit_id_str = _get_commander_unit_id(current_user.username, "CmdC")
        if unit_id_str:
            students_in_unit = (
                Student.query.filter_by(companie=unit_id_str)
                .with_entities(Student.id)
                .all()
            )
            student_ids_in_scope = [sid for (sid,) in students_in_unit]
            view_title = f"Învoiri Weekend Compania {unit_id_str}"
            return_dashboard_url = "company_commander_dashboard"
        else:
            flash("ID Companie invalid.", "danger")
            return redirect(url_for("dashboard"))
    elif current_user.role == "comandant_batalion":
        unit_id_str = _get_commander_unit_id(current_user.username, "CmdB")
        if unit_id_str:
            students_in_unit = (
                Student.query.filter_by(batalion=unit_id_str)
                .with_entities(Student.id)
                .all()
            )
            student_ids_in_scope = [sid for (sid,) in students_in_unit]
            view_title = f"Învoiri Weekend Batalionul {unit_id_str}"
            return_dashboard_url = "battalion_commander_dashboard"
        else:
            flash("ID Batalion invalid.", "danger")
            return redirect(url_for("dashboard"))

    from sqlalchemy import false

    page = request.args.get("page", 1, type=int)
    per_page = 20
    query_filter = (
        WeekendLeave.student_id.in_(student_ids_in_scope)
        if student_ids_in_scope
        else false()
    )
    query = WeekendLeave.query.options(
        joinedload(WeekendLeave.student), joinedload(WeekendLeave.creator)
    ).filter(query_filter)

    search_student_name = request.args.get("search_student_name", "").strip()
    filter_status = request.args.get("filter_status", "").strip()
    filter_weekend_start_date_str = request.args.get(
        "filter_weekend_start_date", ""
    ).strip()

    if search_student_name:
        search_pattern = f"%{unidecode(search_student_name.lower())}%"
        query = query.join(WeekendLeave.student).filter(
            or_(
                func.lower(unidecode(Student.nume)).like(search_pattern),
                func.lower(unidecode(Student.prenume)).like(search_pattern),
            )
        )
    if filter_status:
        query = query.filter(WeekendLeave.status == filter_status)
    if filter_weekend_start_date_str:
        try:
            filter_date_obj = datetime.strptime(
                filter_weekend_start_date_str, "%Y-%m-%d"
            ).date()
            query = query.filter(
                WeekendLeave.weekend_start_date == filter_date_obj
            )
        except ValueError:
            flash(
                "Format dată invalid pentru filtrare (Vineri). Folosiți YYYY-MM-DD.",
                "warning",
            )

    query = query.order_by(
        WeekendLeave.weekend_start_date.desc(), WeekendLeave.student_id
    )
    weekend_leaves_pagination = query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    return render_template(
        "scoped_list_weekend_leaves.html",
        weekend_leaves_pagination=weekend_leaves_pagination,
        title=view_title,
        return_dashboard_url=return_dashboard_url,
    )


@app.route("/commander/services", endpoint="view_scoped_services")
@login_required
def view_scoped_services():
    if current_user.role not in ["comandant_companie", "comandant_batalion"]:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    student_ids_in_scope = []
    unit_id_str = ""
    view_title = ""
    return_dashboard_url = "dashboard"

    if current_user.role == "comandant_companie":
        unit_id_str = _get_commander_unit_id(current_user.username, "CmdC")
        if unit_id_str:
            students_in_unit = (
                Student.query.filter_by(companie=unit_id_str)
                .with_entities(Student.id)
                .all()
            )
            student_ids_in_scope = [sid for (sid,) in students_in_unit]
            view_title = f"Servicii Compania {unit_id_str}"
            return_dashboard_url = "company_commander_dashboard"
        else:
            flash("ID Companie invalid.", "danger")
            return redirect(url_for("dashboard"))
    elif current_user.role == "comandant_batalion":
        unit_id_str = _get_commander_unit_id(current_user.username, "CmdB")
        if unit_id_str:
            students_in_unit = (
                Student.query.filter_by(batalion=unit_id_str)
                .with_entities(Student.id)
                .all()
            )
            student_ids_in_scope = [sid for (sid,) in students_in_unit]
            view_title = f"Servicii Batalionul {unit_id_str}"
            return_dashboard_url = "battalion_commander_dashboard"
        else:
            flash("ID Batalion invalid.", "danger")
            return redirect(url_for("dashboard"))

    from sqlalchemy import false

    page = request.args.get("page", 1, type=int)
    per_page = 20
    query_filter = (
        ServiceAssignment.student_id.in_(student_ids_in_scope)
        if student_ids_in_scope
        else false()
    )
    query = ServiceAssignment.query.options(
        joinedload(ServiceAssignment.student),
        joinedload(ServiceAssignment.creator),
    ).filter(query_filter)

    search_student_name = request.args.get("search_student_name", "").strip()
    filter_service_type = request.args.get("filter_service_type", "").strip()
    filter_service_date_str = request.args.get(
        "filter_service_date", ""
    ).strip()

    if search_student_name:
        search_pattern = f"%{unidecode(search_student_name.lower())}%"
        query = query.join(ServiceAssignment.student).filter(
            or_(
                func.lower(unidecode(Student.nume)).like(search_pattern),
                func.lower(unidecode(Student.prenume)).like(search_pattern),
            )
        )
    if filter_service_type:
        query = query.filter(
            ServiceAssignment.service_type == filter_service_type
        )
    if filter_service_date_str:
        try:
            filter_date_obj = datetime.strptime(
                filter_service_date_str, "%Y-%m-%d"
            ).date()
            query = query.filter(
                ServiceAssignment.service_date == filter_date_obj
            )
        except ValueError:
            flash(
                "Format dată invalid pentru filtrare. Folosiți YYYY-MM-DD.",
                "warning",
            )

    service_types_for_filter = SERVICE_TYPES

    query = query.order_by(ServiceAssignment.start_datetime.desc())
    services_pagination = query.paginate(
        page=page, per_page=per_page, error_out=False
    )

    return render_template(
        "scoped_list_services.html",
        services_pagination=services_pagination,
        title=view_title,
        return_dashboard_url=return_dashboard_url,
        service_types_for_filter=service_types_for_filter,
    )


# --- Gradat New Bulk Import Pages ---
@app.route(
    "/gradat/import/permissions",
    methods=["GET", "POST"],
    endpoint="gradat_page_import_permissions",
)
@login_required
def gradat_page_import_permissions():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        permission_bulk_data = request.form.get(
            "permission_bulk_data", ""
        ).strip()
        if not permission_bulk_data:
            flash("Nu au fost furnizate date pentru import.", "warning")
            return render_template(
                "gradat_import_permissions.html",
                permission_bulk_data=permission_bulk_data,
            )

        lines = permission_bulk_data.splitlines()
        added_count = 0
        error_count = 0
        error_details = []
        processed_students_details = []

        students_managed = Student.query.filter_by(
            created_by_user_id=current_user.id
        ).all()

        i = 0
        while i < len(lines):
            name_line_from_block = lines[i].strip()
            if not name_line_from_block:
                i += 1
                continue

            current_block_lines = []
            temp_i = i
            while temp_i < len(lines) and lines[temp_i].strip():
                current_block_lines.append(lines[temp_i].strip())
                temp_i += 1

            i = temp_i
            if temp_i < len(lines) and not lines[temp_i].strip():
                i += 1

            if len(current_block_lines) < 3:
                if len(current_block_lines) > 0:
                    err_msg = f"Intrare incompletă începând cu '{current_block_lines[0]}'. Necesită Nume, Interval, Destinație."
                    app.logger.warning(
                        f"Bulk Permission Page Import: {err_msg}"
                    )
                    error_details.append(
                        {
                            "line_content": current_block_lines[0],
                            "error_message": err_msg,
                        }
                    )
                    error_count += 1
                continue

            name_line_parsed = current_block_lines[0]
            datetime_line = current_block_lines[1]
            destination_line = current_block_lines[2]
            transport_mode_line = (
                current_block_lines[3] if len(current_block_lines) > 3 else ""
            )
            reason_car_plate_line = (
                current_block_lines[4] if len(current_block_lines) > 4 else ""
            )

            student_obj, student_error = find_student_for_bulk_import(
                name_line_parsed, students_managed
            )
            if student_error:
                app.logger.warning(
                    f"Bulk Permission Page Import: Student find error for '{name_line_parsed}': {student_error}"
                )
                error_details.append(
                    {
                        "line_content": name_line_parsed,
                        "error_message": student_error,
                    }
                )
                error_count += 1
                continue

            try:
                dt_match = re.search(
                    r"(\d{1,2}\.\d{1,2}\.\d{4})\s+(\d{1,2}:\d{2})\s*-\s*(?:(\d{1,2}\.\d{1,2}\.\d{4})\s+)?(\d{1,2}:\d{2})",
                    datetime_line,
                )
                if not dt_match:
                    app.logger.warning(
                        f"Bulk Permission Page Import: Invalid datetime format for '{name_line_parsed}'. Input: '{datetime_line}'."
                    )
                    raise ValueError("Format interval datetime invalid.")

                (
                    start_date_str,
                    start_time_str,
                    end_date_str_opt,
                    end_time_str,
                ) = dt_match.groups()
                start_dt = datetime.strptime(
                    f"{start_date_str} {start_time_str}", "%d.%m.%Y %H:%M"
                )
                end_time_obj_parsed = datetime.strptime(
                    end_time_str, "%H:%M"
                ).time()

                if end_date_str_opt:
                    end_date_assumed = datetime.strptime(
                        end_date_str_opt, "%d.%m.%Y"
                    ).date()
                else:
                    end_date_assumed = start_dt.date()

                if (
                    not end_date_str_opt
                    and end_time_obj_parsed < start_dt.time()
                ):
                    end_date_assumed += timedelta(days=1)
                end_dt = datetime.combine(
                    end_date_assumed, end_time_obj_parsed
                )

                if end_dt <= start_dt:
                    app.logger.warning(
                        f"Bulk Permission Page Import: End datetime not after start for '{name_line_parsed}'. Start: {start_dt}, End: {end_dt}"
                    )
                    raise ValueError(
                        "Data/ora de sfârșit trebuie să fie după data/ora de început."
                    )

            except ValueError as ve:
                error_details.append(
                    {
                        "line_content": f"{name_line_parsed}, {datetime_line}",
                        "error_message": f"Eroare format dată/oră: {str(ve)}",
                    }
                )
                error_count += 1
                continue

            parsed_destination = destination_line.strip()
            parsed_transport_mode = (
                transport_mode_line.strip() if transport_mode_line else None
            )
            parsed_reason = (
                reason_car_plate_line.strip()
                if reason_car_plate_line
                else None
            )

            if not student_obj:
                app.logger.error(
                    f"Bulk Permission Page Import: student_obj is None for '{name_line_parsed}' after student_error was None."
                )
                error_details.append(
                    {
                        "line_content": name_line_parsed,
                        "error_message": "Eroare internă la procesarea studentului.",
                    }
                )
                error_count += 1
                continue

            conflict = check_leave_conflict(
                student_obj.id, start_dt, end_dt, leave_type="permission"
            )
            if conflict:
                app.logger.info(
                    f"Bulk Permission Page Import: Conflict for student '{name_line_parsed}': {conflict}"
                )
                error_details.append(
                    {
                        "line_content": name_line_parsed,
                        "error_message": f"Conflict - {conflict}.",
                    }
                )
                error_count += 1
                continue

            new_permission = Permission(
                student_id=student_obj.id,
                start_datetime=start_dt,
                end_datetime=end_dt,
                destination=parsed_destination,
                transport_mode=parsed_transport_mode,
                reason=parsed_reason,
                status="Aprobată",
                created_by_user_id=current_user.id,
            )
            db.session.add(new_permission)
            log_student_action(
                student_obj.id,
                "PERMISSION_CREATED_BULK_PAGE",
                f"Permisie adăugată prin import text (pagină nouă): {start_dt.strftime('%d.%m %H:%M')} - {end_dt.strftime('%d.%m %H:%M')}.",
            )
            added_count += 1
            processed_students_details.append(
                f"Permisie pt {student_obj.nume} {student_obj.prenume}: {start_dt.strftime('%d.%m %H:%M')} - {end_dt.strftime('%d.%m %H:%M')}"
            )

        if added_count > 0:
            try:
                db.session.commit()
                flash(
                    f"{added_count} permisii au fost adăugate cu succes.",
                    "success",
                )
                log_action(
                    "BULK_IMPORT_PERMISSIONS_SUCCESS_PAGE",  # Changed action type
                    description=f"User {current_user.username} bulk imported {added_count} permissions via page.",
                    details_after_dict={
                        "added_count": added_count,
                        "details": processed_students_details[:5],
                    },
                )
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                flash(
                    f"Eroare la salvarea permisiilor în baza de date: {str(e)}",
                    "danger",
                )
                app.logger.error(
                    f"Bulk Permission Page Import: DB commit error: {str(e)}"
                )
                error_count += added_count
                added_count = 0
                log_action(
                    "BULK_IMPORT_PERMISSIONS_FAIL_DB_PAGE",  # Changed action type
                    description=f"User {current_user.username} bulk permission import (page) DB commit failed. Error: {str(e)}",
                    details_after_dict={
                        "attempted_add_count": added_count,
                        "error_count_at_fail": error_count,
                    },
                )
                db.session.commit()

        if error_count > 0:
            flash(
                f"{error_count} intrări nu au putut fi procesate sau au generat erori. Verificați detaliile afișate.",
                "danger",
            )
            return render_template(
                "gradat_import_permissions.html",
                permission_bulk_data=permission_bulk_data,
                error_details=error_details,
            )

        if added_count == 0 and error_count == 0 and permission_bulk_data:
            flash(
                "Nicio permisie validă de importat din datele furnizate. Verificați formatul.",
                "info",
            )
            return render_template(
                "gradat_import_permissions.html",
                permission_bulk_data=permission_bulk_data,
            )

        return redirect(url_for("list_permissions"))

    return render_template(
        "gradat_import_permissions.html", permission_bulk_data=""
    )


@app.route(
    "/gradat/import/weekend_leaves",
    methods=["GET", "POST"],
    endpoint="gradat_page_import_weekend_leaves",
)
@login_required
def gradat_page_import_weekend_leaves():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        leave_list_text = request.form.get(
            "weekend_leave_bulk_data", ""
        ).strip()
        if not leave_list_text:
            flash("Lista de învoiri este goală.", "warning")
            return render_template(
                "gradat_import_weekend_leaves.html",
                weekend_leave_bulk_data=leave_list_text,
                error_details=None,
            )

        lines = leave_list_text.strip().splitlines()
        processed_count = 0
        error_count = 0
        error_details_list = []
        success_details_list = []

        students_managed = Student.query.filter_by(
            created_by_user_id=current_user.id
        ).all()

        for line_raw in lines:
            line_content = line_raw.strip()
            if not line_content:
                continue

            student_name_str, parsed_intervals, is_biserica_req, error_msg = (
                parse_weekend_leave_line(line_content)
            )

            if error_msg and student_name_str is None and not parsed_intervals:
                continue
            if error_msg:
                error_details_list.append(
                    {"line": line_content, "error": error_msg}
                )
                error_count += 1
                continue
            if not parsed_intervals:
                error_details_list.append(
                    {
                        "line": line_content,
                        "error": "Niciun interval valid de procesat.",
                    }
                )
                error_count += 1
                continue

            student_obj, student_error = find_student_for_bulk_import(
                student_name_str, students_managed
            )
            if student_error:
                error_details_list.append(
                    {
                        "line": line_content,
                        "error": f"Student '{student_name_str}': {student_error}",
                    }
                )
                error_count += 1
                continue

            first_interval_date = parsed_intervals[0]["date_obj"]
            weekend_start_date_obj = (
                first_interval_date
                - timedelta(days=first_interval_date.weekday())
                + timedelta(days=4)
            )

            current_weekend_leave_data = {
                "day1_date": None,
                "day1_start_time": None,
                "day1_end_time": None,
                "day1_selected": None,
                "day2_date": None,
                "day2_start_time": None,
                "day2_end_time": None,
                "day2_selected": None,
                "day3_date": None,
                "day3_start_time": None,
                "day3_end_time": None,
                "day3_selected": None,
                "intervals_for_conflict_check": [],
            }
            distinct_days_processed = set()
            valid_intervals_for_this_student = True

            for interval in parsed_intervals:
                interval_date = interval["date_obj"]
                delta_days = (interval_date - weekend_start_date_obj).days
                day_slot_key = None
                day_name_ro = None

                if delta_days == 0 and interval_date.weekday() == 4:
                    day_slot_key, day_name_ro = "day1", "Vineri"
                elif delta_days == 1 and interval_date.weekday() == 5:
                    day_slot_key, day_name_ro = "day2", "Sambata"
                elif delta_days == 2 and interval_date.weekday() == 6:
                    day_slot_key, day_name_ro = "day3", "Duminica"
                else:
                    error_details_list.append(
                        {
                            "line": line_content,
                            "student": student_obj.nume,
                            "error": f"Data {interval_date.strftime('%d.%m.%Y')} nu corespunde weekendului.",
                        }
                    )
                    error_count += 1
                    valid_intervals_for_this_student = False
                    break

                if day_slot_key in distinct_days_processed:
                    error_details_list.append(
                        {
                            "line": line_content,
                            "student": student_obj.nume,
                            "error": f"Intervale multiple pentru {day_name_ro}.",
                        }
                    )
                    valid_intervals_for_this_student = False
                    break
                distinct_days_processed.add(day_slot_key)

                current_weekend_leave_data[f"{day_slot_key}_date"] = (
                    interval_date
                )
                current_weekend_leave_data[f"{day_slot_key}_start_time"] = (
                    interval["start_time_obj"]
                )
                current_weekend_leave_data[f"{day_slot_key}_end_time"] = (
                    interval["end_time_obj"]
                )
                current_weekend_leave_data[f"{day_slot_key}_selected"] = (
                    day_name_ro
                )

                start_dt = datetime.combine(
                    interval_date, interval["start_time_obj"]
                )
                effective_end_date = interval_date
                if interval["end_time_obj"] < interval["start_time_obj"]:
                    effective_end_date += timedelta(days=1)
                end_dt = datetime.combine(
                    effective_end_date, interval["end_time_obj"]
                )
                current_weekend_leave_data[
                    "intervals_for_conflict_check"
                ].append(
                    {"start": start_dt, "end": end_dt, "day_name": day_name_ro}
                )

            if not valid_intervals_for_this_student:
                continue
            if not distinct_days_processed:
                error_details_list.append(
                    {
                        "line": line_content,
                        "student": student_obj.nume,
                        "error": "Niciun interval valid mapat.",
                    }
                )
                error_count += 1
                continue

            conflict_found_for_student = False
            for interval_to_check in current_weekend_leave_data[
                "intervals_for_conflict_check"
            ]:
                conflict = check_leave_conflict(
                    student_obj.id,
                    interval_to_check["start"],
                    interval_to_check["end"],
                    leave_type="weekend_leave",
                )
                if conflict:
                    error_details_list.append(
                        {
                            "line": line_content,
                            "student": student_obj.nume,
                            "error": f"Conflict {interval_to_check['day_name']}: {conflict}.",
                        }
                    )
                    error_count += 1
                    conflict_found_for_student = True
                    break
            if conflict_found_for_student:
                continue

            new_wl = WeekendLeave(
                student_id=student_obj.id,
                weekend_start_date=weekend_start_date_obj,
                day1_selected=current_weekend_leave_data["day1_selected"],
                day1_date=current_weekend_leave_data["day1_date"],
                day1_start_time=current_weekend_leave_data["day1_start_time"],
                day1_end_time=current_weekend_leave_data["day1_end_time"],
                day2_selected=current_weekend_leave_data["day2_selected"],
                day2_date=current_weekend_leave_data["day2_date"],
                day2_start_time=current_weekend_leave_data["day2_start_time"],
                day2_end_time=current_weekend_leave_data["day2_end_time"],
                day3_selected=current_weekend_leave_data["day3_selected"],
                day3_date=current_weekend_leave_data["day3_date"],
                day3_start_time=current_weekend_leave_data["day3_start_time"],
                day3_end_time=current_weekend_leave_data["day3_end_time"],
                duminica_biserica=(
                    is_biserica_req
                    and current_weekend_leave_data["day3_selected"]
                    == "Duminica"
                ),
                status="Aprobată",
                created_by_user_id=current_user.id,
                reason=f"Import text: {line_content[:100]}",
            )
            db.session.add(new_wl)
            log_student_action(
                student_obj.id,
                "WEEKEND_LEAVE_CREATED_BULK_PAGE",
                f"Învoire weekend adăugată prin import text (pagină nouă) pentru {weekend_start_date_obj.strftime('%d.%m')}.",
            )
            processed_count += 1
            success_details_list.append(
                f"Învoire weekend pentru {student_obj.nume} {student_obj.prenume} ({weekend_start_date_obj.strftime('%d.%m')}) adăugată."
            )

        if processed_count > 0:
            try:
                db.session.commit()
                flash(
                    f"{processed_count} învoiri de weekend procesate și adăugate.",
                    "success",
                )
                for detail in success_details_list:
                    flash(detail, "info")
                log_action(
                    "BULK_IMPORT_WEEKEND_LEAVES_SUCCESS_PAGE",
                    description=f"User {current_user.username} bulk imported {processed_count} weekend leaves via page.",
                    details_after_dict={
                        "added_count": processed_count,
                        "success_details": success_details_list[:5],
                    },
                )
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                flash(
                    f"Eroare la salvarea învoirilor de weekend: {str(e)}",
                    "danger",
                )
                error_count += processed_count
                processed_count = 0
                log_action(
                    "BULK_IMPORT_WEEKEND_LEAVES_FAIL_DB_PAGE",
                    description=f"User {current_user.username} bulk weekend leave (page) DB commit failed. Error: {str(e)}",
                    details_after_dict={
                        "attempted_add_count": processed_count,
                        "error_count_at_fail": error_count,
                    },
                )
                db.session.commit()

        if error_count > 0:
            flash(
                f"{error_count} linii/intrări nu au putut fi procesate sau au generat erori.",
                "danger",
            )
            return render_template(
                "gradat_import_weekend_leaves.html",
                weekend_leave_bulk_data=leave_list_text,
                error_details=error_details_list,
            )

        if processed_count == 0 and error_count == 0 and leave_list_text:
            flash(
                "Nicio învoire validă de importat. Verificați formatul.",
                "info",
            )

        return redirect(url_for("list_weekend_leaves"))

    return render_template(
        "gradat_import_weekend_leaves.html", weekend_leave_bulk_data=""
    )


# --- Presence Report Route ---
@app.route("/reports/presence", methods=["GET", "POST"])
@login_required
def presence_report():
    if current_user.role not in [
        "gradat",
        "admin",
        "comandant_companie",
        "comandant_batalion",
    ]:
        flash("Acces neautorizat pentru rolul dumneavoastră.", "danger")
        return redirect(url_for("dashboard"))

    current_dt_str_for_form = get_localized_now().strftime("%Y-%m-%dT%H:%M")
    report_data_to_render = None

    if request.method == "POST":
        report_type = request.form.get("report_type")
        custom_datetime_str = request.form.get("custom_datetime")
        datetime_to_check = None
        report_title_detail = ""

        if report_type == "current":
            datetime_to_check = get_localized_now()
            report_title_detail = "Prezență Curentă"
        elif report_type == "evening_roll_call":
            naive_dt = get_standard_roll_call_datetime()
            datetime_to_check = EUROPE_BUCHAREST.localize(naive_dt)
            report_title_detail = (
                f"Apel de Seară ({datetime_to_check.strftime('%H:%M')})"
            )
        elif report_type == "company_report":
            naive_dt = datetime.combine(
                get_localized_now().date(), time(14, 20)
            )
            datetime_to_check = EUROPE_BUCHAREST.localize(naive_dt)
            report_title_detail = "Raport Companie (14:20)"
        elif report_type == "morning_check":
            target_d = get_localized_now().date()
            if custom_datetime_str:
                try:
                    target_d = datetime.strptime(
                        custom_datetime_str, "%Y-%m-%dT%H:%M"
                    ).date()
                except (ValueError, TypeError):
                    flash(
                        "Data custom specificată era invalidă, s-a folosit data curentă pentru raportul de dimineață.",
                        "warning",
                    )
            naive_dt = datetime.combine(target_d, time(7, 0))
            datetime_to_check = EUROPE_BUCHAREST.localize(naive_dt)
            report_title_detail = (
                f"Prezență Dimineață ({target_d.strftime('%d.%m.%Y')} 07:00)"
            )
        elif report_type == "custom":
            try:
                naive_dt = datetime.strptime(
                    custom_datetime_str, "%Y-%m-%dT%H:%M"
                )
                datetime_to_check = EUROPE_BUCHAREST.localize(naive_dt)
                report_title_detail = f"Dată Specifică ({datetime_to_check.strftime('%d.%m.%Y %H:%M')})"
            except (ValueError, TypeError):
                flash(
                    "Format dată și oră custom invalid. Folosiți formatul corect.",
                    "danger",
                )
                return render_template(
                    "presence_report.html",
                    current_datetime_str=current_dt_str_for_form,
                    report_data=None,
                )
        else:
            flash("Tip de raport invalid selectat.", "danger")
            return render_template(
                "presence_report.html",
                current_datetime_str=current_dt_str_for_form,
                report_data=None,
            )

        students_for_report = []
        report_base_title = "Raport Prezență"

        if current_user.role == "gradat":
            students_for_report = Student.query.filter_by(
                created_by_user_id=current_user.id
            ).all()
            gradat_pluton = (
                students_for_report[0].pluton if students_for_report else "N/A"
            )
            report_base_title = f"Raport Prezență Plutonul {gradat_pluton}"
        elif current_user.role == "comandant_companie":
            company_id = _get_commander_unit_id(current_user.username, "CmdC")
            if company_id:
                students_for_report = Student.query.filter_by(
                    companie=company_id
                ).all()
                report_base_title = f"Raport Prezență Compania {company_id}"
            else:
                flash("Nu s-a putut determina ID-ul companiei.", "danger")
        elif current_user.role == "comandant_batalion":
            battalion_id = _get_commander_unit_id(
                current_user.username, "CmdB"
            )
            if battalion_id:
                students_for_report = Student.query.filter_by(
                    batalion=battalion_id
                ).all()
                report_base_title = (
                    f"Raport Prezență Batalionul {battalion_id}"
                )
            else:
                flash("Nu s-a putut determina ID-ul batalionului.", "danger")
        elif current_user.role == "admin":  # Admin can also generate reports
            # For admin, maybe they select a unit or see all?
            # For now, let's assume admin sees all students if they access this.
            # This part can be refined with a selection UI for admin.
            students_for_report = Student.query.all()
            report_base_title = "Raport Prezență General (Admin)"

        if not students_for_report and current_user.role == "gradat":
            flash(
                "Nu aveți studenți în evidență pentru a genera raportul.",
                "info",
            )

        if students_for_report:
            report_data_calculated = _calculate_presence_data(
                students_for_report, datetime_to_check
            )
            report_data_to_render = {
                **report_data_calculated,
                "title": f"{report_base_title} - {report_title_detail}",
                "datetime_checked": datetime_to_check.strftime(
                    "%d %B %Y, %H:%M:%S"
                ),
            }
        elif (
            not students_for_report
            and current_user.role != "gradat"
            and not request.form.get("suppress_no_students_flash")
        ):
            flash(
                f"Niciun student găsit pentru {current_user.role} {current_user.username} pentru a genera raportul.",
                "info",
            )

    return render_template(
        "presence_report.html",
        current_datetime_str=current_dt_str_for_form,
        report_data=report_data_to_render,
    )


# --- Admin Data Management (Import/Export) ---
@app.route("/admin/data_management", endpoint="admin_data_management")
@login_required
def admin_data_management():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))
    return render_template("admin_data_management.html")


@app.route(
    "/admin/data/export", methods=["POST"], endpoint="admin_export_data"
)
@login_required
def admin_export_data():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    # Log the export action
    log_action(
        "ADMIN_EXPORT_DATA_START",
        description=f"Admin {current_user.username} initiated a full data export.",
    )
    db.session.commit()

    # A dictionary to hold all the data
    full_data = {}

    # List of all models to export
    # Order doesn't matter for export, but good to be consistent
    models_to_export = [
        User,
        Student,
        Permission,
        DailyLeave,
        WeekendLeave,
        ServiceAssignment,
        VolunteerActivity,
        ActivityParticipant,
        VolunteerSession,
        ActionLog,
        UpdateTopic,
        SiteSetting,
        PublicViewCode,
    ]

    # Using the existing model_to_dict utility
    # We might need to adjust it if it doesn't handle all cases well (e.g., relationships)
    # For a simple JSON dump, excluding relationships and just storing IDs is often safest.
    # The existing model_to_dict excludes _sa_instance_state, which is good.

    for model in models_to_export:
        model_name = model.__tablename__
        records = model.query.all()
        full_data[model_name] = [model_to_dict(r) for r in records]

    # Also export association table data if any (e.g., volunteer_session_participants)
    # This is a bit more manual as it's not a model with a __dict__
    try:
        session_participants = db.session.execute(
            text("SELECT * FROM volunteer_session_participants")
        ).fetchall()
        full_data["volunteer_session_participants"] = [
            dict(row._mapping) for row in session_participants
        ]
    except Exception as e:
        app.logger.error(
            f"Could not export volunteer_session_participants association table: {e}"
        )
        full_data["volunteer_session_participants"] = []

    # Create a JSON file in memory
    json_string = json.dumps(
        full_data, ensure_ascii=False, indent=4, default=str
    )
    json_bytes = json_string.encode("utf-8")

    f = io.BytesIO(json_bytes)
    f.seek(0)

    filename = (
        f"full_export_{get_localized_now().strftime('%Y-%m-%d_%H-%M')}.json"
    )

    return send_file(
        f,
        download_name=filename,
        as_attachment=True,
        mimetype="application/json",
    )


@app.route(
    "/admin/data/import", methods=["POST"], endpoint="admin_import_data"
)
@login_required
def admin_import_data():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    if "import_file" not in request.files:
        flash("Niciun fișier selectat pentru import.", "warning")
        return redirect(url_for("admin_data_management"))

    file = request.files["import_file"]
    if file.filename == "":
        flash("Niciun fișier selectat pentru import.", "warning")
        return redirect(url_for("admin_data_management"))

    if file and file.filename.endswith(".json"):
        try:
            # Load data from file
            json_data = json.load(
                io.TextIOWrapper(file.stream, encoding="utf-8")
            )

            # --- DELETION PHASE ---
            # Order is crucial to respect foreign key constraints
            # Start with many-to-many tables, then models that depend on others.
            db.session.execute(
                text("DELETE FROM volunteer_session_participants")
            )
            MODELS_TO_DELETE = [
                ActivityParticipant,
                Permission,
                DailyLeave,
                WeekendLeave,
                ServiceAssignment,
                ActionLog,
                PublicViewCode,
                VolunteerActivity,
                VolunteerSession,
                UpdateTopic,
                Student,
                User,
                SiteSetting,
            ]
            for model in MODELS_TO_DELETE:
                # This is a simple but potentially slow way. For large dbs, raw SQL is faster.
                model.query.delete()

            db.session.commit()
            flash(
                "Baza de date a fost ștearsă cu succes. Se începe importul...",
                "info",
            )

            # --- INSERTION PHASE ---
            # Order is the reverse of deletion
            MODELS_TO_INSERT = {
                "sitesetting": SiteSetting,
                "user": User,
                "student": Student,
                "updatetopic": UpdateTopic,
                "volunteer_session": VolunteerSession,
                "volunteer_activity": VolunteerActivity,
                "public_view_code": PublicViewCode,
                "actionlog": ActionLog,
                "permission": Permission,
                "daily_leave": DailyLeave,
                "weekend_leave": WeekendLeave,
                "service_assignment": ServiceAssignment,
                "activity_participant": ActivityParticipant,
            }

            # Disable foreign key checks for SQLite during import
            # This is DB specific. For others like PostgreSQL, this is different.
            if db.engine.dialect.name == "sqlite":
                db.session.execute(text("PRAGMA foreign_keys=OFF"))

            for table_name, model_class in MODELS_TO_INSERT.items():
                if table_name in json_data:
                    # Use bulk_insert_mappings for efficiency
                    db.session.bulk_insert_mappings(
                        model_class, json_data[table_name]
                    )

            # Handle the many-to-many association table separately
            if "volunteer_session_participants" in json_data:
                # Since we use bulk insert, we need to use the table object directly
                participants_table = db.metadata.tables[
                    "volunteer_session_participants"
                ]
                db.session.execute(
                    participants_table.insert(),
                    json_data["volunteer_session_participants"],
                )

            db.session.commit()

            if db.engine.dialect.name == "sqlite":
                db.session.execute(text("PRAGMA foreign_keys=ON"))

            flash("Importul datelor a fost finalizat cu succes!", "success")
            log_action(
                "ADMIN_IMPORT_DATA_SUCCESS",
                description=f"Admin {current_user.username} successfully imported data from {file.filename}.",
            )
            db.session.commit()

        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la importul datelor: {str(e)}", "danger")
            app.logger.error(f"Data import failed: {str(e)}")
            log_action(
                "ADMIN_IMPORT_DATA_FAIL",
                description=f"Admin {current_user.username} failed to import data. Error: {str(e)}",
            )
            db.session.commit()

        return redirect(url_for("admin_data_management"))

    else:
        flash(
            "Format fișier invalid. Vă rugăm încărcați un fișier .json.",
            "warning",
        )
        return redirect(url_for("admin_data_management"))


# END JULES BLOCK - MORNING INVIGORATION REPORT (PROBLEM 5)


# START JULES - ADMIN HOME PAGE SETTINGS
@app.route(
    "/admin/settings/homepage",
    methods=["GET", "POST"],
    endpoint="admin_homepage_settings",
)
@login_required
def admin_homepage_settings():
    if current_user.role != "admin":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    default_title = "UNAP User Panel"
    default_badge_text = "Beta v2.5"  # Original badge text

    if request.method == "POST":
        new_title = request.form.get("home_page_title", default_title).strip()
        new_badge_text = request.form.get(
            "home_page_badge_text", default_badge_text
        ).strip()

        title_setting = SiteSetting.query.filter_by(
            key="home_page_title"
        ).first()
        if not title_setting:
            title_setting = SiteSetting(key="home_page_title", value=new_title)
            db.session.add(title_setting)
        else:
            title_setting.value = new_title

        badge_setting = SiteSetting.query.filter_by(
            key="home_page_badge_text"
        ).first()
        if not badge_setting:
            badge_setting = SiteSetting(
                key="home_page_badge_text", value=new_badge_text
            )
            db.session.add(badge_setting)
        else:
            badge_setting.value = new_badge_text

        try:
            db.session.commit()
            flash(
                "Setările pentru pagina principală au fost actualizate.",
                "success",
            )
            log_action(
                "ADMIN_UPDATE_HOMEPAGE_SETTINGS",
                target_model_name="SiteSetting",
                description=f"Admin {current_user.username} updated homepage settings. Title: '{new_title}', Badge: '{new_badge_text}'.",
            )
            db.session.commit()  # Commit log
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la salvarea setărilor: {str(e)}", "danger")
            log_action(
                "ADMIN_UPDATE_HOMEPAGE_SETTINGS_FAIL",
                target_model_name="SiteSetting",
                description=f"Admin {current_user.username} failed to update homepage settings. Error: {str(e)}",
            )
            db.session.commit()  # Commit log

        return redirect(url_for("admin_homepage_settings"))

    # GET request
    current_title_setting = SiteSetting.query.filter_by(
        key="home_page_title"
    ).first()
    current_badge_setting = SiteSetting.query.filter_by(
        key="home_page_badge_text"
    ).first()

    current_title = (
        current_title_setting.value if current_title_setting else default_title
    )
    current_badge_text = (
        current_badge_setting.value
        if current_badge_setting
        else default_badge_text
    )

    return render_template(
        "admin_homepage_settings.html",
        title="Setări Pagină Principală",
        current_title=current_title,
        current_badge_text=current_badge_text,
    )


# END JULES - ADMIN HOME PAGE SETTINGS


# --- Scoped Access Code Routes ---
@app.route(
    "/scoped_access/generate",
    methods=["POST"],
    endpoint="generate_scoped_access_code",
)
@login_required
def generate_scoped_access_code():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    description = request.form.get("description", "").strip()
    permissions = request.form.getlist("permissions")
    expiry_hours = request.form.get("expiry_hours", type=int, default=8)

    if not description:
        flash("Descrierea este obligatorie.", "warning")
        return redirect(url_for("dashboard"))
    if not permissions:
        flash("Trebuie să selectați cel puțin o permisiune.", "warning")
        return redirect(url_for("dashboard"))

    new_code_str = secrets.token_hex(8)
    while ScopedAccessCode.query.filter_by(code=new_code_str).first():
        new_code_str = secrets.token_hex(8)

    expires_at_dt = datetime.utcnow() + timedelta(hours=expiry_hours)

    new_code = ScopedAccessCode(
        code=new_code_str,
        description=description,
        permissions=json.dumps(permissions),
        expires_at=expires_at_dt,
        created_by_user_id=current_user.id,
    )

    try:
        db.session.add(new_code)
        db.session.commit()
        flash(f"Cod de acces delegat generat: {new_code_str}", "success")
        flash(
            f"Descriere: {description} | Valabil pentru {expiry_hours} ore.",
            "info",
        )
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la generarea codului: {str(e)}", "danger")

    return redirect(url_for("dashboard"))


@app.route(
    "/scoped_access/deactivate/<int:code_id>",
    methods=["POST"],
    endpoint="deactivate_scoped_access_code",
)
@login_required
def deactivate_scoped_access_code(code_id):
    code_to_deactivate = ScopedAccessCode.query.get_or_404(code_id)

    if code_to_deactivate.created_by_user_id != current_user.id:
        flash("Nu aveți permisiunea să dezactivați acest cod.", "danger")
        return redirect(url_for("dashboard"))

    code_to_deactivate.is_active = False
    try:
        db.session.commit()
        flash(f"Codul {code_to_deactivate.code} a fost dezactivat.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la dezactivarea codului: {str(e)}", "danger")

    return redirect(url_for("dashboard"))


# --- Public View Code Routes ---
@app.route(
    "/view_access/generate",
    methods=["POST"],
    endpoint="generate_public_view_code",
)
@login_required
def generate_public_view_code():
    if current_user.role not in [
        "admin",
        "comandant_companie",
        "comandant_batalion",
    ]:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    scope_type = request.form.get("scope_type")
    scope_id = request.form.get("scope_id")
    expiry_hours = request.form.get("expiry_hours", type=int, default=24)

    # Validation
    if (
        not scope_type
        or not scope_id
        or scope_type not in ["company", "battalion"]
    ):
        flash("Tip sau ID unitate invalid.", "danger")
        return redirect(request.referrer or url_for("dashboard"))

    # Security check: ensure commanders can only create codes for their own unit
    if current_user.role == "comandant_companie":
        if scope_type != "company" or scope_id != _get_commander_unit_id(
            current_user.username, "CmdC"
        ):
            flash("Nu puteți genera coduri pentru altă unitate.", "danger")
            return redirect(url_for("company_commander_dashboard"))
    elif current_user.role == "comandant_batalion":
        if scope_type != "battalion" or scope_id != _get_commander_unit_id(
            current_user.username, "CmdB"
        ):
            flash("Nu puteți genera coduri pentru altă unitate.", "danger")
            return redirect(url_for("battalion_commander_dashboard"))

    # Generate unique code
    new_code_str = secrets.token_hex(8)
    while PublicViewCode.query.filter_by(code=new_code_str).first():
        new_code_str = secrets.token_hex(8)

    expires_at_dt = datetime.utcnow() + timedelta(hours=expiry_hours)

    new_code = PublicViewCode(
        code=new_code_str,
        scope_type=scope_type,
        scope_id=scope_id,
        expires_at=expires_at_dt,
        created_by_user_id=current_user.id,
    )

    try:
        db.session.add(new_code)
        db.session.commit()
        # Create the full URL for sharing
        public_url = url_for(
            "public_view_login", code=new_code_str, _external=True
        )

        # Flash messages with the new code and the full URL
        flash(f"Cod de acces public generat: {new_code_str}", "success")
        flash(
            f"Link de partajat (valabil {expiry_hours} ore): {public_url}",
            "info",
        )

    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la generarea codului: {str(e)}", "danger")

    return redirect(request.referrer or url_for("dashboard"))


@app.route(
    "/view_access/deactivate/<int:code_id>",
    methods=["POST"],
    endpoint="deactivate_public_view_code",
)
@login_required
def deactivate_public_view_code(code_id):
    code_to_deactivate = PublicViewCode.query.get_or_404(code_id)

    # Security check
    if (
        current_user.role != "admin"
        and code_to_deactivate.created_by_user_id != current_user.id
    ):
        flash("Nu aveți permisiunea să dezactivați acest cod.", "danger")
        return redirect(request.referrer or url_for("dashboard"))

    code_to_deactivate.is_active = False
    try:
        db.session.commit()
        flash(f"Codul {code_to_deactivate.code} a fost dezactivat.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la dezactivarea codului: {str(e)}", "danger")

    return redirect(request.referrer or url_for("dashboard"))


@app.route(
    "/public_view", methods=["GET", "POST"], endpoint="public_view_login"
)
def public_view_login():
    if "public_view_access" in session:
        # Re-check expiry on page load in case the session is old but not cleared
        if (
            datetime.fromisoformat(
                session["public_view_access"].get("expires_at")
            )
            < datetime.utcnow()
        ):
            session.pop("public_view_access", None)
            flash("Sesiunea de vizualizare a expirat.", "info")
            return redirect(url_for("public_view_login"))
        return redirect(url_for("public_dashboard"))

    code_to_check = None
    if request.method == "POST":
        code_to_check = request.form.get("code", "").strip()
        if not code_to_check:
            flash("Vă rugăm introduceți un cod de acces.", "warning")
            return redirect(url_for("public_view_login"))
    elif request.method == "GET":
        code_to_check = request.args.get("code", "").strip()

    if code_to_check:
        now_utc = datetime.utcnow()
        access_code = PublicViewCode.query.filter_by(
            code=code_to_check, is_active=True
        ).first()

        if access_code and access_code.expires_at > now_utc:
            # Consume the code on first use
            access_code.is_active = False
            try:
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                app.logger.error(
                    f"PublicViewCode consume failed for code {access_code.code}: {e}"
                )
            session["public_view_access"] = {
                "scope_type": access_code.scope_type,
                "scope_id": access_code.scope_id,
                "expires_at": access_code.expires_at.isoformat(),
            }
            session.permanent = (
                True  # Make the session last for PERMANENT_SESSION_LIFETIME
            )
            return redirect(url_for("public_dashboard"))
        else:
            if request.method == "GET" and code_to_check:
                flash(
                    "Codul de acces este invalid, a expirat sau a fost dezactivat.",
                    "danger",
                )
            elif request.method == "POST":
                flash(
                    "Codul de acces este invalid, a expirat sau a fost dezactivat.",
                    "danger",
                )
            return redirect(url_for("public_view_login"))

    # Render the login form for a GET request with no code
    return render_template("public_login.html")


@app.route("/public_dashboard", endpoint="public_dashboard")
def public_dashboard():
    if "public_view_access" not in session:
        return redirect(url_for("public_view_login"))

    access_info = session["public_view_access"]

    # Check expiry on every page load
    if datetime.fromisoformat(access_info["expires_at"]) < datetime.utcnow():
        session.pop("public_view_access", None)
        flash("Sesiunea de vizualizare a expirat.", "info")
        return redirect(url_for("public_view_login"))

    scope_type = access_info["scope_type"]
    scope_id = access_info["scope_id"]

    students_in_scope = []
    if scope_type == "company":
        students_in_scope = Student.query.filter_by(companie=scope_id).all()
    elif scope_type == "battalion":
        students_in_scope = Student.query.filter_by(batalion=scope_id).all()

    if not students_in_scope:
        return render_template(
            "public_dashboard.html",
            access_info=access_info,
            report_data=None,
            error="Niciun student găsit pentru unitatea specificată.",
        )

    now = get_localized_now()
    report_data = _calculate_presence_data(students_in_scope, now)

    return render_template(
        "public_dashboard.html",
        access_info=access_info,
        report_data=report_data,
    )


@app.route("/public_view/logout", endpoint="public_view_logout")
def public_view_logout():
    session.pop("public_view_access", None)
    flash("Ați fost deconectat din modulul de vizualizare publică.", "success")
    return redirect(url_for("public_view_login"))


# Initialize DB (for admin user creation, etc.)
# init_db() # This also needs app context if called outside a request or app setup


# --- Student Profile Page ---
@app.route("/student/profile/<int:student_id>")
@login_required
def student_profile(student_id):
    student = db.session.get(Student, student_id)
    if not student:
        flash("Studentul nu a fost găsit.", "danger")
        return redirect(url_for("dashboard"))

    # Security Check
    can_view = False
    if current_user.role == "admin":
        can_view = True
    elif (
        current_user.role == "gradat"
        and student.created_by_user_id == current_user.id
    ):
        can_view = True
    elif current_user.role == "comandant_companie":
        company_id = _get_commander_unit_id(current_user.username, "CmdC")
        if company_id and student.companie == company_id:
            can_view = True
    elif current_user.role == "comandant_batalion":
        battalion_id = _get_commander_unit_id(current_user.username, "CmdB")
        if battalion_id and student.batalion == battalion_id:
            can_view = True

    if not can_view:
        flash("Acces neautorizat la profilul acestui student.", "danger")
        return redirect(url_for("dashboard"))

    # Fetch all related data
    permissions = (
        Permission.query.options(joinedload(Permission.creator))
        .filter_by(student_id=student_id)
        .order_by(Permission.start_datetime.desc())
        .all()
    )
    daily_leaves = (
        DailyLeave.query.options(joinedload(DailyLeave.creator))
        .filter_by(student_id=student_id)
        .order_by(DailyLeave.leave_date.desc())
        .all()
    )
    weekend_leaves = (
        WeekendLeave.query.options(joinedload(WeekendLeave.creator))
        .filter_by(student_id=student_id)
        .order_by(WeekendLeave.weekend_start_date.desc())
        .all()
    )
    services = (
        ServiceAssignment.query.options(joinedload(ServiceAssignment.creator))
        .filter_by(student_id=student_id)
        .order_by(ServiceAssignment.start_datetime.desc())
        .all()
    )
    volunteer_participations = (
        ActivityParticipant.query.options(
            joinedload(ActivityParticipant.activity).joinedload(
                VolunteerActivity.creator
            )
        )
        .filter_by(student_id=student_id)
        .join(VolunteerActivity)
        .order_by(VolunteerActivity.activity_date.desc())
        .all()
    )

    # Determine current status
    now = get_localized_now()
    # The _calculate_presence_data function expects a list of students
    presence_data_list = _calculate_presence_data([student], now)

    current_status = "Prezent"  # Default status
    if presence_data_list.get("smt_students_details"):
        current_status = "Scutit Medical Total"
    elif presence_data_list.get("exempt_other_students_details"):
        current_status = f"Scutit: {student.exemption_details}"
    elif presence_data_list.get("on_duty_students_details"):
        # Extract more specific service detail if available
        current_status = presence_data_list["on_duty_students_details"][
            0
        ].split(" - ", 1)[1]
    elif presence_data_list.get("absent_students_details"):
        # Extract more specific leave detail if available
        current_status = presence_data_list["absent_students_details"][
            0
        ].split(" - ", 1)[1]
    elif (
        presence_data_list.get("platoon_graded_duty_count", 0) > 0
        and student.is_platoon_graded_duty
    ):
        # Check if this student is the one counted as a present gradat
        # This check is complex; for now, we simplify
        is_present_leader = False
        for detail in presence_data_list.get(
            "platoon_graded_duty_students_details", []
        ):
            if student.nume in detail and "Absent" not in detail:
                is_present_leader = True
                break
        if is_present_leader:
            current_status = "Prezent (Gradat Pluton)"

    return render_template(
        "student_profile.html",
        student=student,
        current_status=current_status,
        permissions=permissions,
        daily_leaves=daily_leaves,
        weekend_leaves=weekend_leaves,
        services=services,
        volunteer_participations=volunteer_participations,
    )


@app.route("/student/<int:student_id>/audit")
@login_required
def student_audit_log(student_id):
    student = db.session.get(Student, student_id)
    if not student:
        flash("Studentul nu a fost găsit.", "danger")
        return redirect(url_for("dashboard"))

    # Security Check (similar to student_profile)
    can_view = False
    if current_user.role == "admin":
        can_view = True
    elif (
        current_user.role == "gradat"
        and student.created_by_user_id == current_user.id
    ):
        can_view = True
    elif current_user.role == "comandant_companie":
        company_id = _get_commander_unit_id(current_user.username, "CmdC")
        if company_id and student.companie == company_id:
            can_view = True
    elif current_user.role == "comandant_batalion":
        battalion_id = _get_commander_unit_id(current_user.username, "CmdB")
        if battalion_id and student.batalion == battalion_id:
            can_view = True

    if not can_view:
        flash("Acces neautorizat la jurnalul acestui student.", "danger")
        return redirect(url_for("dashboard"))

    page = request.args.get("page", 1, type=int)
    per_page = 25

    audit_logs_pagination = (
        AuditLog.query.options(joinedload(AuditLog.user))
        .filter_by(student_id=student_id)
        .order_by(AuditLog.timestamp.desc())
        .paginate(page=page, per_page=per_page, error_out=False)
    )

    return render_template(
        "student_audit.html",
        student=student,
        logs_pagination=audit_logs_pagination,
    )


# --- Calendar View ---
@app.route("/calendar")
@login_required
def calendar_view():
    if current_user.role not in ["gradat", "admin"]:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    return render_template(
        "calendar_view.html", events_url=url_for("api_events")
    )


@app.route("/api/events")
@login_required
def api_events():
    if current_user.role not in ["gradat", "admin"]:
        return jsonify({"error": "Unauthorized"}), 403

    student_ids_in_scope = []
    if current_user.role == "admin":
        student_ids_in_scope = [
            s.id for s in Student.query.with_entities(Student.id).all()
        ]
    elif current_user.role == "gradat":
        student_ids_in_scope = [
            s.id
            for s in Student.query.filter_by(
                created_by_user_id=current_user.id
            )
            .with_entities(Student.id)
            .all()
        ]

    if not student_ids_in_scope:
        return jsonify([])

    events = []

    permissions = (
        Permission.query.options(joinedload(Permission.student))
        .filter(
            Permission.student_id.in_(student_ids_in_scope),
            Permission.status == "Aprobată",
        )
        .all()
    )
    for p in permissions:
        events.append(
            {
                "title": f"Permisie: {p.student.nume} {p.student.prenume}",
                "start": p.start_datetime.isoformat(),
                "end": p.end_datetime.isoformat(),
                "backgroundColor": "#ff9f40",
                "borderColor": "#ff9f40",
                "extendedProps": {
                    "type": "Permisie",
                    "student": f"{p.student.grad_militar} {p.student.nume} {p.student.prenume}",
                    "details": p.reason or "",
                },
            }
        )

    daily_leaves = (
        DailyLeave.query.options(joinedload(DailyLeave.student))
        .filter(
            DailyLeave.student_id.in_(student_ids_in_scope),
            DailyLeave.status == "Aprobată",
        )
        .all()
    )
    for dl in daily_leaves:
        events.append(
            {
                "title": f"Învoire: {dl.student.nume} {dl.student.prenume}",
                "start": dl.start_datetime.isoformat(),
                "end": dl.end_datetime.isoformat(),
                "backgroundColor": "#36a2eb",
                "borderColor": "#36a2eb",
                "extendedProps": {
                    "type": "Învoire Zilnică",
                    "student": f"{dl.student.grad_militar} {dl.student.nume} {dl.student.prenume}",
                    "details": dl.leave_type_display,
                },
            }
        )

    weekend_leaves = (
        WeekendLeave.query.options(joinedload(WeekendLeave.student))
        .filter(
            WeekendLeave.student_id.in_(student_ids_in_scope),
            WeekendLeave.status == "Aprobată",
        )
        .all()
    )
    for wl in weekend_leaves:
        for interval in wl.get_intervals():
            events.append(
                {
                    "title": f"Weekend: {wl.student.nume} {wl.student.prenume}",
                    "start": interval["start"].isoformat(),
                    "end": interval["end"].isoformat(),
                    "backgroundColor": "#4bc0c0",
                    "borderColor": "#4bc0c0",
                    "extendedProps": {
                        "type": "Învoire Weekend",
                        "student": f"{wl.student.grad_militar} {wl.student.nume} {wl.student.prenume}",
                        "details": f"Ziua: {interval['day_name']}",
                    },
                }
            )

    services = (
        ServiceAssignment.query.options(joinedload(ServiceAssignment.student))
        .filter(ServiceAssignment.student_id.in_(student_ids_in_scope))
        .all()
    )
    for s in services:
        events.append(
            {
                "title": f"Serviciu ({s.service_type}): {s.student.nume}",
                "start": s.start_datetime.isoformat(),
                "end": s.end_datetime.isoformat(),
                "backgroundColor": "#ff6384",
                "borderColor": "#ff6384",
                "extendedProps": {
                    "type": "Serviciu",
                    "student": f"{s.student.grad_militar} {s.student.nume} {s.student.prenume}",
                    "details": s.service_type,
                },
            }
        )

    return jsonify(events)


# --- Leave/Service Templates ---
@app.route("/gradat/templates", methods=["GET"])
@login_required
def list_leave_templates():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    templates = (
        LeaveTemplate.query.filter_by(created_by_user_id=current_user.id)
        .order_by(LeaveTemplate.template_type, LeaveTemplate.name)
        .all()
    )
    return render_template("list_leave_templates.html", templates=templates)


@app.route("/gradat/templates/create", methods=["GET", "POST"])
@login_required
def create_leave_template():
    if current_user.role != "gradat":
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        name = request.form.get("name", "").strip()
        template_type = request.form.get("template_type")

        if not name or not template_type:
            flash("Numele și tipul șablonului sunt obligatorii.", "warning")
            return render_template(
                "create_edit_leave_template.html",
                template=None,
                service_types=SERVICE_TYPES,
                form_data=request.form,
            )

        template_data = {}
        if template_type == "daily_leave":
            template_data["start_time"] = request.form.get("start_time")
            template_data["end_time"] = request.form.get("end_time")
            if (
                not template_data["start_time"]
                or not template_data["end_time"]
            ):
                flash(
                    "Orele de început și sfârșit sunt obligatorii pentru învoirile zilnice.",
                    "warning",
                )
                return render_template(
                    "create_edit_leave_template.html",
                    template=None,
                    service_types=SERVICE_TYPES,
                    form_data=request.form,
                )

        elif template_type == "permission":
            template_data["duration_hours"] = request.form.get(
                "duration_hours", type=int
            )
            template_data["destination"] = request.form.get(
                "destination", ""
            ).strip()
            template_data["transport_mode"] = request.form.get(
                "transport_mode", ""
            ).strip()
            if (
                not template_data["duration_hours"]
                or template_data["duration_hours"] <= 0
            ):
                flash(
                    "Durata în ore este obligatorie și trebuie să fie pozitivă pentru permisii.",
                    "warning",
                )
                return render_template(
                    "create_edit_leave_template.html",
                    template=None,
                    service_types=SERVICE_TYPES,
                    form_data=request.form,
                )

        elif template_type == "service":
            template_data["service_type"] = request.form.get("service_type")
            template_data["duration_hours"] = request.form.get(
                "duration_hours", type=int
            )
            if (
                not template_data["service_type"]
                or not template_data["duration_hours"]
                or template_data["duration_hours"] <= 0
            ):
                flash(
                    "Tipul serviciului și durata sunt obligatorii pentru servicii.",
                    "warning",
                )
                return render_template(
                    "create_edit_leave_template.html",
                    template=None,
                    service_types=SERVICE_TYPES,
                    form_data=request.form,
                )

        else:
            flash("Tip de șablon invalid.", "danger")
            return render_template(
                "create_edit_leave_template.html",
                template=None,
                service_types=SERVICE_TYPES,
            )

        new_template = LeaveTemplate(
            name=name,
            template_type=template_type,
            created_by_user_id=current_user.id,
            data=json.dumps(template_data),
        )
        db.session.add(new_template)
        try:
            db.session.commit()
            flash(f'Șablonul "{name}" a fost creat cu succes.', "success")
            return redirect(url_for("list_leave_templates"))
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la crearea șablonului: {e}", "danger")

    return render_template(
        "create_edit_leave_template.html",
        template=None,
        service_types=SERVICE_TYPES,
        form_data={},
    )


@app.route("/gradat/templates/edit/<int:template_id>", methods=["GET", "POST"])
@login_required
def edit_leave_template(template_id):
    template = LeaveTemplate.query.get_or_404(template_id)
    if template.created_by_user_id != current_user.id:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("list_leave_templates"))

    if request.method == "POST":
        name = request.form.get("name", "").strip()
        if not name:
            flash("Numele șablonului este obligatoriu.", "warning")
            return render_template(
                "create_edit_leave_template.html",
                template=template,
                service_types=SERVICE_TYPES,
                form_data=request.form,
            )

        template_data = {}
        if template.template_type == "daily_leave":
            template_data["start_time"] = request.form.get("start_time")
            template_data["end_time"] = request.form.get("end_time")
            if (
                not template_data["start_time"]
                or not template_data["end_time"]
            ):
                flash(
                    "Orele de început și sfârșit sunt obligatorii.", "warning"
                )
                return render_template(
                    "create_edit_leave_template.html",
                    template=template,
                    service_types=SERVICE_TYPES,
                    form_data=request.form,
                )
        elif template.template_type == "permission":
            template_data["duration_hours"] = request.form.get(
                "duration_hours", type=int
            )
            template_data["destination"] = request.form.get(
                "destination", ""
            ).strip()
            template_data["transport_mode"] = request.form.get(
                "transport_mode", ""
            ).strip()
            if (
                not template_data["duration_hours"]
                or template_data["duration_hours"] <= 0
            ):
                flash(
                    "Durata în ore este obligatorie și trebuie să fie pozitivă.",
                    "warning",
                )
                return render_template(
                    "create_edit_leave_template.html",
                    template=template,
                    service_types=SERVICE_TYPES,
                    form_data=request.form,
                )
        elif template.template_type == "service":
            template_data["service_type"] = request.form.get("service_type")
            template_data["duration_hours"] = request.form.get(
                "duration_hours", type=int
            )
            if (
                not template_data["service_type"]
                or not template_data["duration_hours"]
                or template_data["duration_hours"] <= 0
            ):
                flash(
                    "Tipul serviciului și durata sunt obligatorii.", "warning"
                )
                return render_template(
                    "create_edit_leave_template.html",
                    template=template,
                    service_types=SERVICE_TYPES,
                    form_data=request.form,
                )

        template.name = name
        template.data = json.dumps(template_data)
        try:
            db.session.commit()
            flash(f'Șablonul "{name}" a fost actualizat.', "success")
            return redirect(url_for("list_leave_templates"))
        except Exception as e:
            db.session.rollback()
            flash(f"Eroare la actualizarea șablonului: {e}", "danger")

    form_data = template.get_data()
    form_data["name"] = template.name
    form_data["template_type"] = template.template_type

    return render_template(
        "create_edit_leave_template.html",
        template=template,
        service_types=SERVICE_TYPES,
        form_data=form_data,
    )


@app.route("/gradat/templates/delete/<int:template_id>", methods=["POST"])
@login_required
def delete_leave_template(template_id):
    template = LeaveTemplate.query.get_or_404(template_id)
    if template.created_by_user_id != current_user.id:
        flash("Acces neautorizat.", "danger")
        return redirect(url_for("list_leave_templates"))

    try:
        db.session.delete(template)
        db.session.commit()
        flash(f'Șablonul "{template.name}" a fost șters.', "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Eroare la ștergerea șablonului: {e}", "danger")

    return redirect(url_for("list_leave_templates"))


if __name__ == "__main__":
    with app.app_context():
        print("Applying database migrations...")
        try:
            from flask_migrate import upgrade as flask_upgrade

            flask_upgrade()
            print("Database migrations applied successfully.")
        except Exception as e:
            print(f"An error occurred during database migration: {e}")

        print("Initializing database (if needed)...")
        init_db()
        print("Database initialization complete.")

    app.run(host="0.0.0.0", port=5001, debug=True)
